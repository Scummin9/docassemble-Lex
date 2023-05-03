from docassemble.base.util import (
    path_and_mimetype,
    ocr_file_in_background,
    DAFile,
    DAObject,
    DADict,
    DASet,
    DAList,
    get_config,
)
from docassemble.base.functions import word, single_paragraph
import math
import re
import PyPDF2
import docx
import pdfquery
from docx import Document
import PIL
import pytesseract
from kraken import pageseg

# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
__all__ = [
    "quaero_rrfps",
    "quaero_x",
    "populor_bluebook_cites",
    "populabantur_bluebook_cites",
    "getText",
    "cleanup",
    "OCRPX",
    "ELEVEN_INCHES",
    "get_jury_instructions_9cir",
    "get_emails",
    "binarize",
    "BLACKS_DICTIONARY",
    "char_width",
    "find_best_binarizer",
    "show_boxes",
    "show_column_boxes",
    "calculate_line_height",
    "gap_check",
    "process_image",
    "columnizer",
    "refinarize",
]

def detect_text(path):
    """Detects text in the file."""
    from google.cloud import vision
    import io
    client = vision.ImageAnnotatorClient()

    with io.open(path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations
    print('Texts:')

    for text in texts:
        print('\n"{}"'.format(text.description))

        vertices = (['({},{})'.format(vertex.x, vertex.y)
                    for vertex in text.bounding_poly.vertices])

        print('bounds: {}'.format(','.join(vertices)))

    if response.error.message:
        raise Exception(
            '{}\nFor more info on error messages, check: '
            'https://cloud.google.com/apis/design/errors'.format(
                response.error.message))
        
def async_detect_document(gcs_source_uri, gcs_destination_uri):
    """OCR with PDF/TIFF as source files on GCS"""

# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
OCRPX = int(get_config("ocr dpi"))
ELEVEN_INCHES = int(get_config("ocr dpi") * 11)
BLACKS_DICTIONARY = ("county", "parish", "circuit", "court", "district", "for", "the", "case", "dept", "department", "number", "no", "civil", "request", "production", "interrogatory", "admission", "documents", "of", "interrogatories", "attorney", 'plaintiff', 'defendant', 'third-party', 'petitioner', 'respondent', 'intervenor', 'interpleader', 'cross-defendant', 'cross-plaintiff', 'cross-respondent', 'cross-petitioner', 'claimant', 'appellant', 'obligor', 'obligee', 'deceased', 'his', "her", "its", "their", "in", "comes", "now", "corporation", "limited", "liability", "company", "assumed", "business", "name", "individually", "llc", "inc", "by", "and", "through", "pursuant", "must", "produce", "answer", "response", "corp", "attorneys", "defendants", "plaintiffs", "first", "second", "third", "fourth", "fifth", "inspection", "copying", "copy", "copies", "take", "place", "personal", "representative", "estate", "minor", "within", "days", "mail", "allow", "extends", "all", "possession", "instructions", "definitions", "beyond", "offices", "requested", "custody", "control", "seek", "obtain", "specifically", "information", "require")
char_width=25

#Helps find the best threshold for binarization of legal docs. This function assumes a one word scan, which is unlikely. I need to return and figure out how to parse the words into a list of "foundlings."
def find_best_binarizer(img):
    #img = PIL.Image.open("imge")
    output = []
    for thresh_test in range(100, 200, 25):
      foundlings=[]
      my_string=pytesseract.image_to_string(binarize(img, thresh_test))
      #my_string=detect_text(binarize(img, thresh_test))
      my_string=my_string.lower()
      import string
      comparison=''
      for character in my_string:
        if character in string.ascii_lowercase or character=="-":
          comparison=comparison+character
        else:
          if comparison in BLACKS_DICTIONARY:
            foundlings.append(str(comparison))
          comparison=''
      output.append((int(thresh_test), len(foundlings)))
    #return output
    most=0
    for toop in output:
      if toop[1]>most:
        most=toop[1]
        winner=toop[0]
    output = []
    for thresh_test in range(int(winner-12), int(winner+12), 6):
      foundlings=[]
      my_string=pytesseract.image_to_string(binarize(img, thresh_test))
      my_string=my_string.lower()
      import string
      comparison=''
      for character in my_string:
        if character in string.ascii_lowercase or character=="-":
          comparison=comparison+character
        else:
          if comparison in BLACKS_DICTIONARY:
            foundlings.append(str(comparison))
          comparison=''
      output.append((str(thresh_test), len(foundlings)))
    #return output
    for toop in output:
      if toop[1]>most:
        most=toop[1]
        winner=toop[0]
    return (winner, most)
  
def refinarize(winner, img):
    output = []
    for thresh_test in range(int(winner-24), int(winner+24), 24):
      foundlings=[]
      my_string=pytesseract.image_to_string(binarize(img, thresh_test))
      my_string=my_string.lower()
      import string
      comparison=''
      for character in my_string:
        if character in string.ascii_lowercase or character=="-":
          comparison=comparison+character
        else:
          if comparison in BLACKS_DICTIONARY:
            foundlings.append(str(comparison))
          comparison=''
      output.append((str(thresh_test), len(foundlings)))
    #return output
    most=0
    for toop in output:
      if toop[1]>most:
        most=toop[1]
        winner=toop[0]
    return (winner, most)
  
def binarize(image_to_transform, threshold):
  #from PIL import Image
  img = PIL.Image.open(image_to_transform.path())
  #(left, upper, right, lower) = (0, 0, imgc.width, int(imgc.height*.666666666666))
  
  #img = imgc.crop((left, upper, right, lower))
  output_image = img.convert("L")
  for x in range(output_image.width):
    for y in range(output_image.height):
      if output_image.getpixel((x, y))<threshold:
        output_image.putpixel((x, y), 0)
      else:
        output_image.putpixel((x, y), 255)
  return output_image

def show_boxes(binarized_img):
  #Modifies the passed image to show a series of bounding boxes on an image as run by kraken.
  #param binarized_img: A PIL.Image object that has already been binarized
  #return img: The Modified PIL.Image object
  #First bring in the ImageDraw object
  img = Image.open(binarized_img)
  from PIL import ImageDraw
  #Second grab a drawing object to annotate that image
  drawing.object=ImageDraw.Draw(img)
  #Third, we can create a set of boxes using pageseg.segment
  bounding_boxes=pageseg.segment(img.convert("1"))["boxes"]
  #Now lets go through the list of bounding boxes
  for box in bounding_boxes:
    #And Just draw a nice rectangle around it.
    drawing_object.rectangle(box, fill=None, outline="red")
  return img

def show_column_boxes(binarized_img):
  #Modifies the passed image to show a series of bounding boxes on an image as run by kraken.
  #param binarized_img: A PIL.Image object that has already been binarized
  #return img: The Modified PIL.Image object
  #First bring in the ImageDraw object
  img = Image.open(binarized_img)
  from PIL import ImageDraw
  #Second grab a drawing object to annotate that image
  drawing.object=ImageDraw.Draw(img)
  #Third, we can create a set of boxes using pageseg.segment
  bounding_boxes=pageseg.segment(img.convert("1"), black_colseps=True)["boxes"]
  #Now lets go through the list of bounding boxes
  for box in bounding_boxes:
    #And Just draw a nice rectangle around it.
    drawing_object.rectangle(box, fill=None, outline="red")
  return binarized_img

def calculate_line_height(imge):
  #Calculates the average height of a line from a given image
  #param img: A PIL.Image object
  #return: The average line height in pixels
  img = Image.open(imge)
  bounding_boxes=pageseg.segment(img.convert("1"))["boxes"]
  height_accumulator=0
  for box in bounding_boxes:
    height_accumulator=height_accumulator+box[3]-box[1]
    #Because we start counting at the upper left corner in PIL
  return int(height_accumulator/len(bounding_boxes))
                                            
                                            
def gap_check(imge, location):
    img = Image.open(imge)
    line_height=calculate_line_height(img)
    gap_box=(0, 0, char_width, line_height*6)
    #Checks the img in a given (x,y) location to see if it fits the description of a gap_box
    #param img: a PIL.Image file
    #param location: A tuple (x,y) which is a pixel location in that image
    # return: True if it fits the definition of a gap_box, otherwise false
    for x in range(location[0], location[0]+gap_box[2]):
      for y in range(location[1], location[1]+gap_box[3]):
        if x < img.width and y <img.height:
          if img.getpixel((x,y)) != 255:
            return False
    return True

def columnizer(imge, location):
  img = Image.open(imge)
  line_height=calculate_line_height(img)
  gap_box=(0, 0, char_width, line_height*6)
  #Draws a line in img in the middle of the gap discovered at location
  #this doesn't draw the line in location, but draws it at the middle of a gap_box starting at location.
  #param img: A PIL.Image file
  #param location: A tuple(x, y) which is a pixel location in the img.
  drawing_object=ImageDraw.Draw(img)
  #Define a Line to draw by first defining what "in the middle" means (x1) and then how wide it should be (x2). Here I make x2=x1 to create a 1 pixel line.
  x1=location[0]+int(gap_box[2]/2)
  x2=x1
  #First y coordinate should be the y coordinate that was passed in (top of the gap_box), and we want the second to be the bottom of the gap box
  y1=location[1]
  y2=y1+gap_box[3]
  drawing_object.rectangle((x1, y1, x2, y2), fill="black", outline="black")
                                            
def process_image(imge):
  #Takes an image of text and adds black vertical bars to break up columns
   #param img: A PIL.Image file
   #return: img
  img = Image.open(imge)
  for x in range(0, img.width, char_width):
    for y in range(0, img.height, line_height):
      if (gap_check(img, (x,y))):
        columnizer(img, (x,y))
  return img
                                            
# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def get_jury_instructions_9cir(docxfile):
    text = str(docxfile)
    Text = []
    # for para in doc.paragraphs:
    # Text.append(para.text)
    # text= '\n'.join(Text)
    jiRegex = re.compile(
        r"((\d?\d\.\d)  ?([A-Z\s\(\)]+))(?=[A-Z]{1}[a-z]+)(.*?)(?=\s\d?\d\.\d  ?[A-Z\s\(\)]+)",
        re.DOTALL,
    )
    # jiRegex = re.compile(r'(.*?)(\f)', re.DOTALL)
    ji = jiRegex.findall(text)
    return ji


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def getText(filename):
    doc = docx.Document(filename[0].path())
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    text = "\n".join(fullText)
    bluebookRegex = re.compile(
        r"(\,?\s?(\d\d?\d?)\s(Or\.\s?App\.|Or\.|P\.2d|P\.3d)\s(\d\d?\d?)(,?\s?(\d\d?\d?)((\-|–|—)(\d\d?\d?))?)?\,?\s?(\(\d\d\d\d\))?)"
    )
    starpaginationRegex = re.compile(r"(\*\s?\d\d?\d?)")
    orsRegex = re.compile(r"(ORS\s(\d\d\d?\.\d\d\d)(\(\w\))?(\(\w\))?)")
    text = bluebookRegex.sub("", text)
    text = starpaginationRegex.sub("", text)
    text = orsRegex.sub("Oregon Statute", text)
    return text


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def populor_bluebook_cites(document):
    text = getText(document)
    bluebookRegex = re.compile(
        r"(\,?\s?(\d\d?\d?)\s(Or\.\s?App\.|Or\.|P\.2d|P\.3d)\s(\d\d?\d?)(,?\s?(\d\d?\d?)((\-|–|—)(\d\d?\d?))?)?\,?\s?(\(\d\d\d\d\))?)"
    )
    starpaginationRegex = re.compile(r"(\*\d\d?\d?)")
    orsRegex = re.compile(r"(ORS\s(\d\d\d?\.\d\d\d)(\(\w\))?(\(\w\))?)")
    text2 = bluebookRegex.sub("", text)
    text3 = starpaginationRegex.sub("", text2)
    text4 = str(orsRegex.sub("Oregon Statute", text3))
    return text4


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def populabantur_bluebook_cites(document):
    tabula = DAFile("tabula")
    tabula.initialize(filename="tabula.docx")
    bluebookRegex = re.compile(
        r"(\,?\s?(\d\d?\d?)\s(Or\.\s?App\.|Or\.|P\.2d|P\.3d)\s(\d\d?\d?)(,?\s?(\d\d?\d?)((\-|–|—)(\d\d?\d?))?)?\,?\s?(\(\d\d\d\d\))?)"
    )
    starpaginationRegex = re.compile(r"(\*\s?\d\d?\d?)")
    orsRegex = re.compile(r"(ORS\s(\d\d\d?\.\d\d\d)(\(\w\))?(\(\w\))?)")
    text = getText(document)
    text2 = bluebookRegex.sub("", text2)
    text3 = starpaginationRegex.sub("", text3)
    text4 = orsRegex.sub("Oregon Statute", text4)
    # return text
    # tabula.commit()
    # document.save(tabula.path())
    # text.save(tabula.path())
    # word/document.xml.write(text)
    # document.save()
    tabula.write(text4)
    tabula.commit()
    # tabula.retreive()
    # tabula.save()
    return tabula


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def quaero_x(document):
    pdf = pdfquery.PDFQuery(document[0].path())
    pdf.load()
    label = pdf.pq('LTTextLineHorizontal:contains("24 ")')
    if pdf.pq('LTTextLineHorizontal:contains("28 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("28 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("27 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("27 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("26 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("26 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("25 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("25 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("24 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("24 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("23 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("23 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("22 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("22 ")')
    elif pdf.pq('LTTextLineHorizontal:contains("21 ")'):
        label2 = pdf.pq('LTTextLineHorizontal:contains("21 ")')
    else:
        label2 = pdf.pq('LTTextLineHorizontal:contains("20 ")')
    q = float(label2.attr("y0"))
    s = (q / 72) * OCRPX
    if s < 290:
        s = OCRPX
    t = s * 0.97
    a = float(label.attr("x1"))
    b = (a / 72) * OCRPX
    c = b * 1.11
    return (int(c), int(s))

def get_emails(textf):
    text = str(textf)
    Text = []
    emailRegex = re.compile(
        r"([A-Za-z0-9]+[.-_])*[A-Za-z0-9]+@[A-Za-z0-9-]+(\.[A-Z|a-z]{2,})+",
        re.DOTALL,
    )
    elist = emailRegex.findall(textf)
    for theemail in elist:
      Text.append(theemail)
    return Text


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def quaero_rrfps(document):
    text = str(document)
    requests = []
    interrogs = []
    admissions = []
    rRFPRegex = re.compile(r"""
          ((\s?\|?\s?\|?\s?
            ((
                ((T|l|I|\|)N(T|l|I|\|)(E|F|L|R)(R|P|L|S|C)(R|P|L|S|C)(O|o|0)GA(T|l|I|\|)(O|o|0)(P|R)Y|Interrogatory)                                                          # INTERROGATORY
                    |((R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|))|Request)?\s?                                                                                                       # REQUEST/Request
                    ((E|L|F)(Q|O)(R|P|L|S|C|\<\.|\,|-)\s                                                                                                                                                             # FOR
                    (R|P|L|S|C|\<\.|\,|-)(R|P|L|S|C|\<\.|\,|-)(Q|O)DU(R|P|L|S|C)(T|l|I|\|)(T|l|I|\|)(Q|O)N\s)?\s?                                                                         # PRODUCTION
                (N(O|o|0)(\.|\,|\:|;)?|NUMBER)\s?)                                                                                                                                                            # NO./NUMBER
                  (\_?(\d|I|l|\||i)+)((\—|\-|\:|\;|\'|\.|\,|\s)))                                                                                                                                                       # e.g., "22:"
          (.+?)                                                                                                                                                                                                                 # substance of request/interrogatory
          (?=(\s?\|?\s?\|?\s?
          ((((R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|))|Request)?\s?
          ((E|L|F)(Q|O)(R|P|L|S|C|\<\.|\,|-)\s?
          (R|P|L|S|C|\<\.|\,|-)(R|P|L|S|C|\<\.|\,|-)(Q|O)DU(R|P|L|S|C)(T|l|I|\|)(T|l|I|\|)(Q|O)N\s?)\s?
          (N(O|o|0)(\.|\,|\:|;)?|NUMBER)\s?)
          (\_?(\d|I|l|\||i)+)((\—|\-|\:|\;|\'|\.|\,|\s)))
          |(Response)
          |(Answer)
          |(ANSW(E|F|L|R)(R|P|L|S|C))
          |((R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)S(P|R)ONS(E|F|L))|(\s\d?\d\.|Dated|DATED)))
          """, re.VERBOSE | re.DOTALL)
    OtherRFPRegex = re.compile(
        r"(\s((\_?(\d|I|l|\||i){1,2})((\—|\-|\:|\;|\'|\.)))(.+?)(?=(\s(\_?(\d|I|l|\||i){1,2})((\—|\-|\:|\;|\'|\.|\,)))|(\s\d?\d\.|Dated|DATED)))",
        re.DOTALL,
    )
    casenoRegex = re.compile(
        r"(Case\sNo\.\s\d?\:?\d\d\-?(CV|cv)\-?\d\d\d\d\d\-?\D?\D?)"
    )
    text2 = str(casenoRegex.sub("", text))
    reqs = rRFPRegex.findall(text2)
    #The following categories are "out of order" so to speak in terms of the integer that identifies them and the occurence of each type. However, if rfas were to come before rfas, the issue arises that "REQUEST NO. X:" is very often used for both discovery requests. Accordingly, if rfas were analyzed after rfps, rfas would occasionally be miscategorized as rfps.
    for req in reqs:
        if (
            "ADMISSION" in req[1].upper()
            or "ADMIT" in req[1].upper()
            or "ADMIT" in req[38].upper()
            or "REQUEST FOR ADMISSION" in req[1].upper()
            or "RFA" in req[1].upper()
        ):
            requests.append((single_paragraph(cleanup(req[38]).strip()), 2))
        elif (
            "INTERROGATORY" in req[1].upper()
            or "INTER" in req[1].upper()
            or "ATORY" in req[1].upper()
            or "ANSWER" in req[1].upper()
            or "ROG" in req[1].upper()
        ):
            requests.append((single_paragraph(cleanup(req[38]).strip()), 1))
        elif (
            "PRODUCTION" in str(req[1]).upper()
            or "PROD" in str(req[1]).upper()
            or "REQUEST" in str(req[1]).upper()
            or "REQ" in str(req[1]).upper()
            or "RESPONSE" in str(req[1]).upper()
            or "RFP" in str(req[1]).upper()
        ):
            requests.append((single_paragraph(cleanup(req[38]).strip()), 0))
        else:
            requests.append((single_paragraph(cleanup(req[38]).strip()), 3))
    if len(requests) < 1:
        requests.clear()
        reqs = OtherRFPRegex.findall(text2)
        for req in reqs:
            requests.append((single_paragraph(cleanup(req[6]).strip()), 3))
    return requests


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def cleanup(listoftext):
    nobreaksRegex = re.compile(
        r"(\/\s?\/\s?|\\\s?\\\s?|\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?\\\s?)"
    )
    nolines = []
    nopipes = []
    nobreaks = []
    nobrackets = []
    #lineNumRegex = re.compile(r"(\s?(\d|\||I|l)?\s?\d\s)")
    nobracketsRegex = re.compile(r"(}|{)")
    pipesRegex = re.compile(r"\|")
    oneRegex = re.compile(r"(\s1n\s)")
    one = oneRegex.sub(" in ", listoftext)
    twoRegex = re.compile(r"(\s1s\s)")
    two = twoRegex.sub(" is ", one)
    threeRegex = re.compile(r"(\sfonn\s)")
    three = threeRegex.sub(" form ", two)
    fourRegex = re.compile(r"(\sre1mbursement)")
    four = fourRegex.sub(" reimbursement", three)
    fiveRegex = re.compile(r"(\f)")
    five = fiveRegex.sub(" ", four)
    sixRegex = re.compile(r"(liagrams)")
    six = sixRegex.sub("diagrams", five)
    sevenRegex = re.compile(r"(1ot)")
    seven = sevenRegex.sub("not", six)
    eightRegex = re.compile(r"(liagnostic)")
    eight = eightRegex.sub("diagnostic", seven)
    nineRegex = re.compile(r"(dlaintiff)")
    nine = nineRegex.sub("plaintiff", eight)
    tenRegex = re.compile(r"\slaintiff")
    ten = tenRegex.sub(" plaintiff", nine)
    elevenRegex = re.compile(r"nedical")
    eleven = elevenRegex.sub("medical", ten)
    twelveRegex = re.compile(r"vhether")
    twelve = twelveRegex.sub("whether", eleven)
    thirteenRegex = re.compile(r"\serein")
    thirteen = thirteenRegex.sub(" herein", twelve)
    fourteenRegex = re.compile(r"\sompensation")
    fourteen = fourteenRegex.sub(" compensation", thirteen)
    fifteenRegex = re.compile(r"\sny")
    fifteen = fifteenRegex.sub("any", fourteen)
    sixteenRegex = re.compile(r"\somplaint")
    sixteen = sixteenRegex.sub(" complaint", fifteen)
    seventeenRegex = re.compile(r"\s1¢\s")
    seventeen = seventeenRegex.sub(" the ", sixteen)
    eighteenRegex = re.compile(r"\sith\s")
    eighteen = eighteenRegex.sub(" with ", seventeen)
    nineteenRegex = re.compile(r"\(\d\.e\.,\s")
    nineteen = nineteenRegex.sub("(i.e., ", eighteen)
    twentyRegex = re.compile(r"\=")
    twenty = twentyRegex.sub("", nineteen)
    twentyoneRegex = re.compile(r"attarnev")
    twentyone = twentyoneRegex.sub("attorney", twenty)
    twentytwoRegex = re.compile(r"nrodiuct")
    twentytwo = twentytwoRegex.sub("product", twentyone)
    twentythreeRegex = re.compile(r"nrivilegce")
    twentythree = twentythreeRegex.sub("privelege", twentytwo)
    twentyfourRegex = re.compile(r"chlent")
    twentyfour = twentyfourRegex.sub("client", twentythree)
    twentyfiveRegex = re.compile(r"_")
    twentyfive = twentyfiveRegex.sub("", twentyfour)
    twentysixRegex = re.compile(r"~")
    twentysix = twentysixRegex.sub("", twentyfive)
    twentysevenRegex = re.compile(r"\s1S\s")
    twentyseven = twentysevenRegex.sub(" is ", twentysix)
    twentyeightRegex = re.compile(r"\s\—\s")
    twentyeight = twentyeightRegex.sub(" ", twentyseven)
    twentynineRegex = re.compile(r"\sainvolved")
    twentynine = twentynineRegex.sub(" involved", twentyeight)
    thirtyRegex = re.compile(r"\sOr\s")
    thirty = thirtyRegex.sub(" or ", twentynine)
    thirtyoneRegex = re.compile(r"\s\[fa\s")
    thirtyone = thirtyoneRegex.sub(" If a ", thirty)
    thirtytwoRegex = re.compile(r"PAGE\s\d\d?\s\-\sDEFENDANT..")
    thirtytwo = thirtytwoRegex.sub("", thirtyone)
    #nolines = lineNumRegex.sub(" ", thirtytwo)
    nopipes = pipesRegex.sub(" ", thirtytwo)
    nobreaks = nobreaksRegex.sub(" ", nopipes)
    nobrackets = nobracketsRegex.sub("", nobreaks)
    return nobrackets


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def train_rrfp_independent(textt):
    text = str(textt)
    rRFPRegex = re.compile(
        r"(\s?\|?\s?\|?\s?(\s?(N(O|o|0)(\.|\,|\:|;)?|NUMBER)\s?)(\_?\d+)(\s?(\—|\-|\:|\;|\'|\.|\,)?)(.+?)(\s?(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)S(P|R)ONS(E|F|L)|(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|)|Response|Request)(\s?(\—|\-|\:|\;|\'|\.|\,)?)(.+?)(Request|(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|)|DATED|Dated))",
        re.DOTALL,
    )
    nobreaksRegex = re.compile(
        r"(\/\s?\/\s?|\\\s?\\\s?|\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?\\\s?)"
    )
    lineNumRegex = re.compile(r"(\s?(\d|\||I|l)?\s?\d\s)")
    nobracketsRegex = re.compile(r"(}|{)")
    pipesRegex = re.compile(r"\|")
    nobreaksRegex = re.compile(
        r"(\/\s?\/\s?|\\\s?\\\s?|\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?\\\s?)"
    )
    nolines = []
    nopipes = []
    nobreaks = []
    casenoRegex = re.compile(r"Case\sNo\.\s\d\d\-?(CV|cv)\-?\d\d\d\d\d")
    nobrackets = []
    text2 = casenoRegex.sub("", text)
    reqs = rRFPRegex.findall(text2)
    requests = []
    for req in reqs:
        requests.append(req[8])
    one = []
    oneRegex = re.compile(r"(\s1n\s)")
    for item in requests:
        one.append(oneRegex.sub(" in ", item))

    two = []
    twoRegex = re.compile(r"(\s1s\s)")
    for item in one:
        two.append(twoRegex.sub(" is ", item))

    three = []
    threeRegex = re.compile(r"(\sfonn\s)")
    for item in two:
        three.append(threeRegex.sub(" form ", item))

    four = []
    fourRegex = re.compile(r"(\sre1mbursement)")
    for item in three:
        four.append(fourRegex.sub(" reimbursement", item))

    five = []
    fiveRegex = re.compile(r"(\f)")
    for item in four:
        five.append(fiveRegex.sub(" ", item))

    six = []
    sixRegex = re.compile(r"(liagrams)")
    for item in five:
        six.append(sixRegex.sub("diagrams", item))

    seven = []
    sevenRegex = re.compile(r"(1ot)")
    for item in six:
        seven.append(sevenRegex.sub("not", item))

    eight = []
    eightRegex = re.compile(r"(liagnostic)")
    for item in seven:
        eight.append(eightRegex.sub("diagnostic", item))

    nine = []
    nineRegex = re.compile(r"(dlaintiff)")
    for item in eight:
        nine.append(nineRegex.sub("plaintiff", item))

    ten = []
    tenRegex = re.compile(r"\slaintiff")
    for item in nine:
        ten.append(tenRegex.sub(" plaintiff", item))

    eleven = []
    elevenRegex = re.compile(r"nedical")
    for item in ten:
        eleven.append(elevenRegex.sub("medical", item))

    twelve = []
    twelveRegex = re.compile(r"vhether")
    for item in eleven:
        twelve.append(twelveRegex.sub("whether", item))

    thirteen = []
    thirteenRegex = re.compile(r"\serein")
    for item in twelve:
        thirteen.append(thirteenRegex.sub(" herein", item))

    fourteen = []
    fourteenRegex = re.compile(r"\sompensation")
    for item in thirteen:
        fourteen.append(fourteenRegex.sub(" compensation", item))

    fifteen = []
    fifteenRegex = re.compile(r"\sny")
    for item in fourteen:
        fifteen.append(fifteenRegex.sub("any", item))

    sixteen = []
    sixteenRegex = re.compile(r"\somplaint")
    for item in fifteen:
        sixteen.append(sixteenRegex.sub(" complaint", item))

    seventeen = []
    seventeenRegex = re.compile(r"\s1¢\s")
    for item in sixteen:
        seventeen.append(seventeenRegex.sub(" the ", item))

    eighteen = []
    eighteenRegex = re.compile(r"\sith\s")
    for item in seventeen:
        eighteen.append(eighteenRegex.sub(" with ", item))

    nineteen = []
    nineteenRegex = re.compile(r"\(\d\.e\.,\s")
    for item in eighteen:
        nineteen.append(nineteenRegex.sub("(i.e., ", item))

    twenty = []
    twentyRegex = re.compile(r"\=")
    for item in nineteen:
        twenty.append(twentyRegex.sub("", item))

    twentyone = []
    twentyoneRegex = re.compile(r"attarnev")
    for item in twenty:
        twentyone.append(twentyoneRegex.sub("attorney", item))

    twentytwo = []
    twentytwoRegex = re.compile(r"nrodiuct")
    for item in twentyone:
        twentytwo.append(twentytwoRegex.sub("product", item))

    twentythree = []
    twentythreeRegex = re.compile(r"nrivilegce")
    for item in twentytwo:
        twentythree.append(twentythreeRegex.sub("privelege", item))

    twentyfour = []
    twentyfourRegex = re.compile(r"chlent")
    for item in twentythree:
        twentyfour.append(twentyfourRegex.sub("client", item))

    twentyfive = []
    twentyfiveRegex = re.compile(r"_")
    for item in twentyfour:
        twentyfive.append(twentyfiveRegex.sub("", item))

    twentysix = []
    twentysixRegex = re.compile(r"~")
    for item in twentyfive:
        twentysix.append(twentysixRegex.sub("", item))

    twentyseven = []
    twentysevenRegex = re.compile(r"\s1S\s")
    for item in twentysix:
        twentyseven.append(twentysevenRegex.sub(" is ", item))

    twentyeight = []
    twentyeightRegex = re.compile(r"\s\—\s")
    for item in twentyseven:
        twentyeight.append(twentyeightRegex.sub(" ", item))

    twentynine = []
    twentynineRegex = re.compile(r"\sainvolved")
    for item in twentyeight:
        twentynine.append(twentynineRegex.sub(" involved", item))

    thirty = []
    thirtyRegex = re.compile(r"\sOr\s")
    for item in twentynine:
        thirty.append(thirtyRegex.sub(" or ", item))

    thirtyone = []
    thirtyoneRegex = re.compile(r"\s\[fa\s")
    for item in thirty:
        thirtyone.append(thirtyoneRegex.sub(" If a ", item))

    thirtytwo = []
    thirtytwoRegex = re.compile(r"PAGE\s\d\d?\s\-\sDEFENDANT..")
    for item in thirtyone:
        thirtytwo.append(thirtytwoRegex.sub("", item))

    for item in thirtytwo:
        nolines.append(lineNumRegex.sub(" ", item))
    for item in nolines:
        nopipes.append(pipesRegex.sub(" ", item))
    for item in nopipes:
        nobreaks.append(nobreaksRegex.sub(" ", item))
    for item in nobreaks:
        nobrackets.append(nobracketsRegex.sub("", item))
    return nobrackets


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def train_rrfp_dependent(textt):
    text = str(textt)
    rRFPRegex = re.compile(
        r"(\s?\|?\s?\|?\s?(\s?(N(O|o|0)(\.|\,|\:|;)?|NUMBER)\s?)(\_?\d+)(\s?(\—|\-|\:|\;|\'|\.|\,)?)(.+?)(\s?(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)S(P|R)ONS(E|F|L)|(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|)|Response|Request)(\s?(\—|\-|\:|\;|\'|\.|\,)?)(.+?)(Request|(R|P|L|S|C|\<\.|\,|-)?(E|F|L|R)(Q|O)U(E|L|F)S(T|l|I|\|)|DATED|Dated))",
        re.DOTALL,
    )
    casenoRegex = re.compile(
        r"(Case\sNo\.\s\d?\:?\d\d\-?(CV|cv)\-?\d\d\d\d\d\-?\D?\D?)"
    )
    nobreaksRegex = re.compile(
        r"(\/\s?\/\s?|\\\s?\\\s?|\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?|\/\s?\/\s?\/\s?\/\s?\/\s?|\\\s?\\\s?\\\s?\\\s?\\\s?)"
    )
    lineNumRegex = re.compile(r"(\s?(\d|\||I|l)?\s?\d\s)")
    nobracketsRegex = re.compile(r"(}|{)")
    pipesRegex = re.compile(r"\|")
    nolines = []
    nopipes = []
    nobrackets = []
    nobreaks = []
    text2 = casenoRegex.sub("", text)
    reqs = rRFPRegex.findall(text2)
    responses = []
    for req in reqs:
        responses.append(req[21])
    one = []
    oneRegex = re.compile(r"(\s1n\s)")
    for item in requests:
        one.append(oneRegex.sub(" in ", item))

    two = []
    twoRegex = re.compile(r"(\s1s\s)")
    for item in one:
        two.append(twoRegex.sub(" is ", item))

    three = []
    threeRegex = re.compile(r"(\sfonn\s)")
    for item in two:
        three.append(threeRegex.sub(" form ", item))

    four = []
    fourRegex = re.compile(r"(\sre1mbursement)")
    for item in three:
        four.append(fourRegex.sub(" reimbursement", item))

    five = []
    fiveRegex = re.compile(r"(\f)")
    for item in four:
        five.append(fiveRegex.sub(" ", item))

    six = []
    sixRegex = re.compile(r"(liagrams)")
    for item in five:
        six.append(sixRegex.sub("diagrams", item))

    seven = []
    sevenRegex = re.compile(r"(1ot)")
    for item in six:
        seven.append(sevenRegex.sub("not", item))

    eight = []
    eightRegex = re.compile(r"(liagnostic)")
    for item in seven:
        eight.append(eightRegex.sub("diagnostic", item))

    nine = []
    nineRegex = re.compile(r"(dlaintiff)")
    for item in eight:
        nine.append(nineRegex.sub("plaintiff", item))

    ten = []
    tenRegex = re.compile(r"\slaintiff")
    for item in nine:
        ten.append(tenRegex.sub(" plaintiff", item))

    eleven = []
    elevenRegex = re.compile(r"nedical")
    for item in ten:
        eleven.append(elevenRegex.sub("medical", item))

    twelve = []
    twelveRegex = re.compile(r"vhether")
    for item in eleven:
        twelve.append(twelveRegex.sub("whether", item))

    thirteen = []
    thirteenRegex = re.compile(r"\serein")
    for item in twelve:
        thirteen.append(thirteenRegex.sub(" herein", item))

    fourteen = []
    fourteenRegex = re.compile(r"\sompensation")
    for item in thirteen:
        fourteen.append(fourteenRegex.sub(" compensation", item))

    fifteen = []
    fifteenRegex = re.compile(r"\sny")
    for item in fourteen:
        fifteen.append(fifteenRegex.sub("any", item))

    sixteen = []
    sixteenRegex = re.compile(r"\somplaint")
    for item in fifteen:
        sixteen.append(sixteenRegex.sub(" complaint", item))

    seventeen = []
    seventeenRegex = re.compile(r"\s1¢\s")
    for item in sixteen:
        seventeen.append(seventeenRegex.sub(" the ", item))

    eighteen = []
    eighteenRegex = re.compile(r"\sith\s")
    for item in seventeen:
        eighteen.append(eighteenRegex.sub(" with ", item))

    nineteen = []
    nineteenRegex = re.compile(r"\(\d\.e\.,\s")
    for item in eighteen:
        nineteen.append(nineteenRegex.sub("(i.e., ", item))

    twenty = []
    twentyRegex = re.compile(r"\=")
    for item in nineteen:
        twenty.append(twentyRegex.sub("", item))

    twentyone = []
    twentyoneRegex = re.compile(r"attarnev")
    for item in twenty:
        twentyone.append(twentyoneRegex.sub("attorney", item))

    twentytwo = []
    twentytwoRegex = re.compile(r"nrodiuct")
    for item in twentyone:
        twentytwo.append(twentytwoRegex.sub("product", item))

    twentythree = []
    twentythreeRegex = re.compile(r"nrivilegce")
    for item in twentytwo:
        twentythree.append(twentythreeRegex.sub("privelege", item))

    twentyfour = []
    twentyfourRegex = re.compile(r"chlent")
    for item in twentythree:
        twentyfour.append(twentyfourRegex.sub("client", item))

    twentyfive = []
    twentyfiveRegex = re.compile(r"_")
    for item in twentyfour:
        twentyfive.append(twentyfiveRegex.sub("", item))

    twentysix = []
    twentysixRegex = re.compile(r"~")
    for item in twentyfive:
        twentysix.append(twentysixRegex.sub("", item))

    twentyseven = []
    twentysevenRegex = re.compile(r"\s1S\s")
    for item in twentysix:
        twentyseven.append(twentysevenRegex.sub(" is ", item))

    twentyeight = []
    twentyeightRegex = re.compile(r"\s\—\s")
    for item in twentyseven:
        twentyeight.append(twentyeightRegex.sub(" ", item))

    twentynine = []
    twentynineRegex = re.compile(r"\sainvolved")
    for item in twentyeight:
        twentynine.append(twentynineRegex.sub(" involved", item))

    thirty = []
    thirtyRegex = re.compile(r"\sOr\s")
    for item in twentynine:
        thirty.append(thirtyRegex.sub(" or ", item))

    thirtyone = []
    thirtyoneRegex = re.compile(r"\s\[fa\s")
    for item in thirty:
        thirtyone.append(thirtyoneRegex.sub(" If a ", item))

    thirtytwo = []
    thirtytwoRegex = re.compile(r"PAGE\s\d\d?\s\-\sDEFENDANT..")
    for item in thirtyone:
        thirtytwo.append(thirtytwoRegex.sub("", item))

    for item in thirtytwo:
        nolines.append(lineNumRegex.sub(" ", item))
    for item in nolines:
        nopipes.append(pipesRegex.sub(" ", item))
    for item in nopipes:
        nobreaks.append(nobreaksRegex.sub(" ", item))
    for item in nobreaks:
        nobrackets.append(nobracketsRegex.sub("", item))
    return thirtytwo
