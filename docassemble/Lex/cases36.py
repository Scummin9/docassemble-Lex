# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
from docassemble.base.util import (
    Individual,
    Person,
    DAList,
    path_and_mimetype,
    ocr_file_in_background,
    DAOAuth,
    DAWeb,
    DAWebError,
    DAObject,
    log,
    Thing,
    IndividualName,
    Name,
    format_date,
    format_time,
    date_difference,
    date_interval,
    current_datetime,
    Value,
    DAFile,
    url_ask,
    Address,
    DAObjectPlusParameters,
    as_datetime,
    DAGlobal,
    DADateTime,
)
from docassemble.base.functions import (
    this_thread,
    possessify,
    possessify_en,
    indefinite_article,
    noun_singular,
    noun_plural,
    comma_and_list,
    ordinal,
    need,
    capitalize,
    server,
    nodoublequote,
    some,
    indefinite_article,
    force_gather,
    quantity_noun,
    invalidate,
    possessify_long,
    a_preposition_b,
    a_in_the_b,
    its,
    their,
    the,
    this,
    her,
    his,
    these,
    your,
    underscore_to_space,
    nice_number,
    verb_past,
    verb_present,
    salutation,
    salutation_default,
    single_paragraph,
    word,
    fix_punctuation,
    item_label,
    state_name,
    states_list,
    add_separators,
    bold,
    phone_number_in_e164,
    phone_number_formatted,
    currency,
    bold,
    url_action,
    space_to_underscore,
    single_to_double_newlines,
)
from docassemble.base.core import DAList
from ics import Calendar, Event, Attendee, Organizer
import math
import xlrd
import re
import PyPDF2
import docxtpl
from docxtpl import RichText, DocxTemplate
import docx
from docx import Document
from docx.shared import Inches, Length, Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import copy
import random
from docassemble.base.filter import markdown_to_html
from markdownify import markdownify
from .civpro_unified import CIV_PRO
from .gvision import case_name_italicizer, replace_newlines, replace_underscores
from docassemble.base.util import Individual, Person, DAObject
from docassemble.base.sql import alchemy_url, upgrade_db, SQLObject, SQLObjectRelationship
from sqlalchemy import Column, ForeignKey, Integer, String, Boolean, Time, Date, Numeric, create_engine, or_, and_, text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship, sessionmaker
import sys
from sqlalchemy.ext.associationproxy import association_proxy
import copy
import os
import json
from docassemble.base.logger import logmessage
from docassemble.base.error import DAAttributeError
from alembic.config import Config
from alembic import command

# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
__all__ = [
    "SCNonEconomics",
    "SCCase",
    "SCIndividual",
    "SCPerson",
    "SCList",
    "possessify",
    "Counter",
    "SCRequests",
    "SCPartyList",
    "SCParty",
    "SCLawfirmList",
    "SCLawyerList",
    "SCLawfirm",
    "SCLawyer",
    "flatten",
    "SCMeds",
    "SCFutureMeds",
    "SCLostWages",
    "SCLostProfits",
    "SCRehab",
    "SCReducedEarningCapacity",
    "SCWrongfulDeath",
    "SCPecuniaryLosses",
    "SCMortuary",
    "get_row",
    "strb",
    "strbu",
    "strbi",
    "strnull",
    "stri",
    "strbiu",
    "Color",
    "h1strb",
    "h1strbu",
    "h1strbi",
    "h1strnull",
    "h1stri",
    "h1strbiu",
    "h2strb",
    "h2strbu",
    "h2strbi",
    "h2strnull",
    "h2stri",
    "h2strbiu",
    "bld",
    "uline",
    "itx",
    "uline_itx",
    "bld_itx",
    "bld_uline",
    "bld_uline",
    "bld_uline_itx",
    "no_indent_body",
    "nullrun",
    "quotefull",
    "SCMVA",
    "parens",
    "the_states",
    "jurisdiction_docx",
    "bluebook_ordinal",
    "SCUnsortedRequestsList",
    "SCUnsortedRequests",
    "format_date_range",
    "SCCrossParties",
    "SCCrossClaims",
    "SCCrossClaim",
    'Party',
    'Lawyer',
    'Lawfirm',
    'Case',
    'Provider',
    'db_find_all',
    'SCPaymentList',
    'SCPayment',
    "make_event",
    "Designation",
    "Jurisdiction",
    "State",
    "County",
    "District",
    "Division",
    "SCJurisdiction",
    "Claim",
    "BusinessEntity",
    "CaptionType",
    "Type",
    "Gender",
    "SCClaimList",
    "SCDesignationList",
    "contains",
    "SCClaimList",
    "PartyClaim",
    "SCCity",
    "JuryInstruction",
    "SCJuryInstructionList",
]


def contains(list, filter):
    for x in list:
        if filter(x):
            return True
    return False


def make_event(title=None, location=None, description=None, begin_date=None, begin_time=None, end_date=None,
               end_time=None, organizer=None, attendees=None):
    if attendees is None:
        attendees = []
    if title is None:
        raise Exception("make_event: a title parameter is required")
    c = Calendar()
    e = Event()
    if organizer is not None:
        e.organizer = Organizer(common_name=organizer.name.full(), email=organizer.email)
    if len(attendees) > 0:
        e.attendees = [Attendee(common_name=attendee.name.full(), email=attendee.email) for attendee in attendees]
    e.name = str(title)
    e.begin = as_datetime(begin_date).format_datetime('yyyy-MM-dd')
    e.end = as_datetime(end_date).format_datetime('yyyy-MM-dd')
    e.make_all_day()
    if location not in (None, ''):
        e.location = str(location)
    if description not in (None, ''):
        e.description = str(description)
    c.events.add(e)
    c.events  # pylint: disable=pointless-statement
    ics_file = DAFile('ics_file')
    ics_file.set_random_instance_name()
    ics_file.initialize(filename="event.ics", mimetype="text/calendar")
    with open(ics_file.path(), 'w', encoding='utf-8') as f:
        f.write(str(c))
    ics_file.commit()
    return ics_file


class SCPaymentList(DAList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCPayment
        self.complete_attribute = 'pay_complete'
        super().init(*pargs, **kwargs)

class SCPayment(DAObject):
    @property
    def pay_complete(self):
        self.date
        self.amount

Base = declarative_base()

class PartyModel(Base):
    __tablename__ = 'party'
    id = Column(Integer, primary_key=True)
    caption_type_id = Column(Integer, ForeignKey("caption_type.id"))
    name = Column(String(250))
    first_name = Column(String(250))
    middle_name = Column(String(250))
    last_name = Column(String(250))
    address = Column(String(250))
    unit = Column(String(250))
    city = Column(String(250))
    state = Column(String(250))
    zip_code = Column(String(250))
    party_type_id = Column(Integer, ForeignKey("designation.id"))
    business_entity_id = Column(Integer, ForeignKey("business_entity.id"))
    trust = Column(String(250))
    dba = Column(String(250))
    state_of_inc_id = Column(Integer, ForeignKey("state.id"))
    caption_text = Column(String(250))
    gender_id = Column(Integer, ForeignKey("gender.id"))
    is_3dp = Column(Boolean())
    is_cross_p = Column(Boolean())
    is_cross_d = Column(Boolean())

class MedsVisitModel(Base):
    __tablename__ = 'meds_visit'
    id = Column(Integer, primary_key=True)
    chart_note = Column(String(250))
    date = Column(Date())
    clinician = Column(Integer, ForeignKey("provider.id"))
    provider = Column(Integer, ForeignKey("provider.id"))
    amount_billed = Column(Numeric())
    pip = Column(Numeric())
    primary_health_insurance = Column(Numeric())
    secondary_health_insurance = Column(Numeric())
    out_of_pocket = Column(Numeric())
    written_off = Column(Numeric())

class BusinessEntityModel(Base):
    __tablename__ = 'business_entity'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))

class CaptionTypeModel(Base):
    __tablename__ = 'caption_type'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))


class ClaimModel(Base):
    __tablename__ = 'claim'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))


class LawfirmModel(Base):
    __tablename__ = 'lawfirm'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)
    address = Column(String(250))
    unit = Column(String(250))
    city = Column(String(250))
    state = Column(String(250))
    zip = Column(String(250))
    phone_number = Column(String(250))
    fax_number = Column(String(250))


class LawyerModel(Base):
    __tablename__ = 'lawyer'
    id = Column(Integer, primary_key=True)
    lawfirm_id = Column(Integer, ForeignKey("lawfirm.id"))
    first_name = Column(String(250))
    middle_name = Column(String(250))
    last_name = Column(String(250))
    email = Column(String(250), unique=True)
    phone_number = Column(String(250))
    fax_number = Column(String(250))
    gender_id = Column(Integer, ForeignKey("gender.id"))
    state_bar_id = Column(Integer, ForeignKey("state.id"))
    bar_no = Column(String(250))


class GenderModel(Base):
    __tablename__ = 'gender'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))


class TypeModel(Base):
    __tablename__ = 'type'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))

class CityModel(Base):
    __tablename__ = 'city'
    id = Column(Integer, primary_key=True)
    name = Column(String(250))

class CaseModel(Base):
    __tablename__ = 'cases'
    id = Column(Integer, primary_key=True)
    number = Column(String(250))
    casename = Column(String(250), unique=True)
    jurisdiction_id = Column(Integer, ForeignKey("jurisdiction.id"))
    state_id = Column(Integer, ForeignKey("state.id"))
    county_id = Column(Integer, ForeignKey("county.id"))
    district_id = Column(Integer, ForeignKey("district.id"))
    division_id = Column(Integer, ForeignKey("division.id"))
    dol = Column(Date())
    tol = Column(Time())
    sol = Column(Date())
    filed = Column(Date())

class JurisdictionModel(Base):
    __tablename__ = 'jurisdiction'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)

class StateModel(Base):
    __tablename__ = 'state'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)

class CountyModel(Base):
    __tablename__ = 'county'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)
    address = Column(String(250))
    unit = Column(String(250))
    city = Column(String(250))
    state = Column(String(250))
    zip = Column(String(250))
    
class DistrictModel(Base):
    __tablename__ = 'district'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)

class DivisionModel(Base):
    __tablename__ = 'division'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)
    

class ProviderModel(Base):
    __tablename__ = 'provider'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)
    address = Column(String(250))
    unit = Column(String(250))
    city = Column(String(250))
    state = Column(String(250))
    zip = Column(String(250))
    fax_number = Column(String(250))
    phone_number = Column(String(250))
    email = Column(String(250))
    specialty_id = Column(Integer, ForeignKey("specialty.id"))
    
class JuryInstructionModel(Base):
    __tablename__ = 'jury_instruction'
    id = Column(Integer, primary_key=True)
    name = Column(String(), unique=True)
    title = Column(String())
    no = Column(String(250))
    instruction = Column(String())
    comment_title = Column(String())
    comment = Column(String())
    update = Column(Date())


class SpecialtyModel(Base):
    __tablename__ = 'specialty'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)


class DesignationModel(Base):
    __tablename__ = 'designation'
    id = Column(Integer, primary_key=True)
    name = Column(String(250), unique=True)


class CaseTypeModel(Base):
    __tablename__ = 'cases_type'
    id = Column(Integer, primary_key=True)
    cases_id = Column(Integer, ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    type_id = Column(Integer, ForeignKey('type.id', ondelete='CASCADE'), nullable=False)


class PartyLawfirmModel(Base):
    __tablename__ = 'party_lawfirm'
    id = Column(Integer, primary_key=True)
    lawfirm_id = Column(Integer, ForeignKey('lawfirm.id', ondelete='CASCADE'), nullable=False)
    party_id = Column(Integer, ForeignKey('party.id', ondelete='CASCADE'), nullable=False)


class LawyerStateBarModel(Base):
    __tablename__ = 'lawyers_state_bar'
    id = Column(Integer, primary_key=True)
    state_id = Column(Integer, ForeignKey('state.id', ondelete='CASCADE'), nullable=False)
    lawyer_id = Column(Integer, ForeignKey('lawyer.id', ondelete='CASCADE'), nullable=False)
    bar_no = Column(String(250))


class CasePartyModel(Base):
    __tablename__ = 'cases_party'
    id = Column(Integer, primary_key=True)
    cases_id = Column(Integer, ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    party_id = Column(Integer, ForeignKey('party.id', ondelete='CASCADE'), nullable=False)


class CaseLawyerModel(Base):
    __tablename__ = 'cases_lawyer'
    id = Column(Integer, primary_key=True)
    cases_id = Column(Integer, ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    lawyer_id = Column(Integer, ForeignKey('lawyer.id', ondelete='CASCADE'), nullable=False)


class CaseLawfirmModel(Base):
    __tablename__ = 'cases_lawfirm'
    id = Column(Integer, primary_key=True)
    cases_id = Column(Integer, ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    lawfirm_id = Column(Integer, ForeignKey('lawfirm.id', ondelete='CASCADE'), nullable=False)


class CasePartyLawfirmModel(Base):
    __tablename__ = 'lawfirm_party'
    id = Column(Integer, primary_key=True)
    cases_id = Column(Integer, ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    party_lawfirm_id = Column(Integer, ForeignKey('party_lawfirm.id', ondelete='CASCADE'), nullable=False)


class LawfirmLawyerModel(Base):
    __tablename__ = 'lawfirm_lawyer'
    id = Column(Integer, primary_key=True)
    lawfirm_id = Column(Integer, ForeignKey('lawfirm.id', ondelete='CASCADE'), nullable=False)
    lawyer_id = Column(Integer, ForeignKey('lawyer.id', ondelete='CASCADE'), nullable=False)


class PartyDesignationModel(Base):
    __tablename__ = 'party_designation'
    id = Column(Integer, primary_key=True)
    designation_id = Column(Integer, ForeignKey('designation.id', ondelete='CASCADE'), nullable=False)
    party_id = Column(Integer, ForeignKey('party.id', ondelete='CASCADE'), nullable=False)


class PartyClaimModel(Base):
    __tablename__ = 'party_claim'
    id = Column(Integer, primary_key=True)
    claim_id = Column(Integer, ForeignKey('claim.id'))
    party_id = Column(Integer, ForeignKey('party.id'))
    aggreived_party_id = Column(Integer, ForeignKey('party.id'))
    third_party_claim = Column(Boolean())
    cross_claim = Column(Boolean())


url = alchemy_url('demo db')
engine = create_engine(url, pool_pre_ping=True)
Base.metadata.create_all(engine)
Base.metadata.bind = engine
DBSession = sessionmaker(bind=engine)()
upgrade_db(url, __file__, engine)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCList(DAList):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("fx", Color)
        super().init(*pargs, **kwargs)

    def pnameu_possessive(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self).upper()
        if output[-1] == "S":
            output += str(APOSTROPHE)
        else:
            output += str(APOSTROPHE + "s")
        return output

    def nameu_possessive(self):
        output = ""
        output += comma_and_list(self).upper()
        if output[-1] == "S":
            output += str(APOSTROPHE)
        else:
            output += str(APOSTROPHE + "s")
        return output

    def nameu(self):
        output = ""
        output += comma_and_list(self).upper()
        return output

    def indlpname(self):
        output = ""
        listout = []
        for party in self.elements:
            listout.append(party.party.pname())
        output += comma_and_list(listout)
        return output

    def pname_does(self, doesverb):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self)
        output += str(" " + self.does_verb(doesverb))
        return output

    def pname_did(self, didverb):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self)
        output += str(" " + self.did_verb(didverb))
        return output

    def pname(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self)
        return output

    def pnameu(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self).upper()
        return output

    def pnameor_does(self, doesverb):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self, and_string="or")
        output += str(" " + self.does_verb(doesverb))
        return output

    def pnameor_did(self, didverb):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self, and_string="or")
        output += str(" " + self.did_verb(didverb))
        return output

    def pnameor(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += comma_and_list(self, and_string="or")
        return output

    def salute(self):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout)
        return output

    def salute_does(self, doesverb):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout)
        output += str(" " + self.does_verb(doesverb))
        return output

    def salute_did(self, didverb):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout)
        output += str(" " + self.did_verb(didverb))
        return output

    def salute_possessive(self):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += possessify(comma_and_list(listout), "")
        return output

    def saluteor(self):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout, and_string="or")
        return output

    def saluteor_does(self, doesverb):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout, and_string="or")
        output += str(" " + self.does_verb(doesverb))
        return output

    def saluteor_did(self, didverb):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += comma_and_list(listout, and_string="or")
        output += str(" " + self.did_verb(didverb))
        return output

    def saluteor_possessive(self):
        output = ""
        listout = []
        for party in self:
            if isinstance(party, SCIndividual) or party.name.caption_type.name.text not in (
                    "governmental entity",
                    "business entity",
            ):
                listout.append(party.salutation(with_name=True))
            else:
                listout.append(party.name.text)
        output += possessify(comma_and_list(listout, and_string="or"), "")
        return output

    def pname_possessive(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += possessify(comma_and_list(self), "")
        return output

    def pnameor_possessive(self):
        output = ""
        output += str(self.asnoun() + " ")
        output += possessify(comma_and_list(self, and_string="or"), "")
        return output

    def client_of(self, firm):
        output = SCList("client_of", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self:
            # log(str("PARTY " + repr(party))
            for lawfirm in party.party.firms:
                if firm.id == lawfirm.id and party.party not in output.elements:
                    output.append(party.party)
        output.there_is_another = False
        output.gathered = True
        return output

    def the_ps(self):
        log("the_ps .  . . . for party in self: . . . . 1")
        the_ps_output = SCList("the_ps_output", there_are_any=True, auto_gather=False)
        the_ps_output.clear()
        for party in self:
            # if any(designation.name.text in ("Plaintiff", "Petitioner", "Claimant", "Appellant", "Obligee") for designation in party.party.party_types):
            log("the_ps .  . . . for party in self: . . . . 2 . . . party " + repr(party))
            for designation in party.party.party_types:
                log("the_ps .  . . . for party in self: . . . . 3  . . . . designation " + repr(designation))
                if designation.name.text == "Plaintiff":
                    log("the_ps .  . . . for party in self: . . . . 4")
                    the_ps_output.append(party.party)
        the_ps_output.there_is_another = False
        the_ps_output.gathered = True
        return the_ps_output

    def the_pets(self):
        the_pets_output = the_pets_output("the_pets", there_are_any=True, auto_gather=False)
        the_pets_output.clear()
        for party in self:
            if any(designation.name.text in ("Petitioner") for designation in party.party.party_types):
                the_pets_output.append(party.party)
        the_pets_output.there_is_another = False
        the_pets_output.gathered = True
        return the_pets_output

    def the_cs(self):
        the_cs_output = SCList("the_cs_output", there_are_any=True, auto_gather=False)
        the_cs_output.clear()
        for party in self:
            if any(designation.name.text in ("Claimant") for designation in party.party.party_types):
                the_cs_output.append(party.party)
        the_cs_output.there_is_another = False
        the_cs_output.gathered = True
        return the_cs_output

    def the_as(self):
        the_as_output = SCList("the_as_output", there_are_any=True, auto_gather=False)
        the_as_output.clear()
        for party in self:
            if any(designation.name.text in ("Appellant") for designation in party.party.party_types):
                the_as_output.append(party.party)
        the_as_output.there_is_another = False
        the_as_output.gathered = True
        return the_as_output

    def the_oes(self):
        the_oes_output = SCList("the_oes_output", there_are_any=True, auto_gather=False)
        the_oes_output.clear()
        for party in self:
            if any(designation.name.text in ("Obligee") for designation in party.party.party_types):
                the_oes_output.append(party.party)
        the_oes_output.there_is_another = False
        the_oes_output.gathered = True
        return the_oes_output

    def the_ors(self):
        the_ors_output = SCList("the_ors_output", there_are_any=True, auto_gather=False)
        the_ors_output.clear()
        for party in self:
            if any(designation.name.text in ("Respondent") for designation in party.party.party_types):
                the_ors_output.append(party.party)
        the_ors_output.there_is_another = False
        the_ors_output.gathered = True
        return the_ors_output

    def the_rs(self):
        the_rs_output = SCList("the_rs_output", there_are_any=True, auto_gather=False)
        the_rs_output.clear()
        for party in self:
            if any(designation.name.text in ("Respondent") for designation in party.party.party_types):
                the_rs_output.append(party.party)
        the_rs_output.there_is_another = False
        the_rs_output.gathered = True
        return the_rs_output

    def the_ds(self):
        the_ds_output = SCList("the_ds", there_are_any=True, auto_gather=False)
        the_ds_output.clear()
        for party in self:
            if any(designation.name.text in ("Defendant", "Respondent", "Obligor") for designation in
                   party.party.party_types):
                the_ds_output.append(party.party)
        the_ds_output.there_is_another = False
        the_ds_output.gathered = True
        return the_ds_output

    def cross_ps(self):
        cross_ps_output = SCList("cross_ps_output", there_are_any=True, auto_gather=False)
        cross_ps_output.clear()
        for party in self:
            if any(designation.name.text in ("Cross-Plaintiff", "Cross-Petitioner") for designation in
                   party.party.party_types):
                cross_ps_output.append(party.party)
        cross_ps_output.there_is_another = False
        cross_ps_output.gathered = True
        return cross_ps_output

    def cross_ds(self):
        cross_ds_output = SCList("cross_ds_output", there_are_any=True, auto_gather=False)
        cross_ds_output.clear()
        for party in self:
            if any(designation.name.text in ("Cross-Defendant", "Cross-Respondent") for designation in
                   party.party.party_types):
                cross_ds_output.append(party.party)
        cross_ds_output.there_is_another = False
        cross_ds_output.gathered = True
        return cross_ds_output

    def threedp_ps(self):
        threedp_ps_output = SCList("threedp_ps_output", there_are_any=True, auto_gather=False)
        threedp_ps_output.clear()
        for party in self:
            if any(designation.name.text == "Third-Party Plaintiff" for designation in party.party.party_types):
                threedp_ps_output.append(party.party)
        threedp_ps_output.there_is_another = False
        threedp_ps_output.gathered = True
        return threedp_ps_output

    def threedp_ds(self):
        threedp_ds_output = SCList("threedp_ds_output", there_are_any=True, auto_gather=False)
        threedp_ds_output.clear()
        for party in self:
            if any(designation.name.text == "Third-Party Defendant" for designation in party.party.party_types):
                threedp_ds_output.append(party.party)
        threedp_ds_output.there_is_another = False
        threedp_ds_output.gathered = True
        return threedp_ds_output

    def intervenors(self):
        intervenors_output = SCList("intervenors_output", there_are_any=True, auto_gather=False)
        intervenors_output.clear()
        for party in self:
            if any(designation.name.text in ("Intervenor") for designation in party.party.party_types):
                intervenors_output.append(party.party)
        intervenors_output.there_is_another = False
        intervenors_output.gathered = True
        return intervenors_output

    def deceaseds(self):
        deceaseds_output = SCList("deceaseds_output", there_are_any=True, auto_gather=False)
        deceaseds_output.clear()
        for party in self:
            if any(designation.name.text in ("Deceased") for designation in party.party.party_types):
                deceaseds_output.append(party.party)
        deceaseds_output.there_is_another = False
        deceaseds_output.gathered = True
        return deceaseds_output

    def work_for(self, firm):
        output = list()
        for attorney in self:
            if attorney.firm is firm:
                output.append(attorney)
        return output

    def parse_firmz(self, parties):
        pairs = []
        for item in self:
            pairs.append(item.pair_firmz(self, parties))
        paired = [num for elem in pairs for num in elem]
        firma = self.copy()
        for x in firma:
            if x in paired:
                z = firma.index(x)
                del firma[z]
        return firma

    def parse_firms(self, case):
        pairs = []
        firmb = self.copy()
        firmc = firmb[0].pair_firms(self, case)
        del firmb[0]
        for a in firmb:
            if a in firmc:
                b = firmb.index(a)
                del firmb[b]
        firmb.reverse()
        for firm in firmb:
            pairs.clear()
            pairs = firm.pair_firms(self, case)
            for y in firmb:
                if y in pairs:
                    z = firmb.index(y)
                    del firmb[z]
        return firmb

    def caption_noun(self, **kwargs):
        language = kwargs.get("language", None)
        the_noun = str(self[0].party_type.name.text)
        the_noun = re.sub(r".*\.", "", the_noun)
        the_noun = re.sub(r"_", " ", the_noun)
        if (
                self.number() > 1
                or self.number() == 0
                or ("plural" in kwargs and kwargs["plural"])
        ) and not ("singular" in kwargs and kwargs["singular"]):
            output = noun_plural(the_noun, language=language)
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output
        else:
            output = noun_singular(the_noun)
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output

    def asnoun(self, **kwargs):
        language = kwargs.get("language", None)
        output = []
        output.clear()
        divergent = SCList("divergent", there_are_any=True, auto_gather=False, object_type=SCList)
        divergent.clear()
        strings = []
        strings.clear()
        for all_parties in self:
            if all_parties not in flatten(divergent.elements):
                log("divergent elements are" + repr(flatten(divergent.elements)))
                log("asnoun() . . . . 1")
                # The following conditionals perform identical operations, but this is intentional: the conditionals categorize the objects that should be kept in separate lists, and the operations create a new list and append the objects that were categorized by the conditionals.
                if any(party.party_types != all_parties.party_types for party in self):
                    log("asnoun() . . . . 2a")
                    divergent.appendObject()
                    divergent[-1].clear()
                    for party in [ex for ex in self if ex.party_types == all_parties.party_types]:
                        divergent[-1].append(party)
                    divergent[-1].there_is_another = False
                    divergent[-1].gathered = True
                else:
                    divergent.appendObject()
                    divergent[-1].clear()
                    for party in [ex for ex in self if ex.party_types == all_parties.party_types]:
                        divergent[-1].append(party)
                    divergent[-1].there_is_another = False
                    divergent[-1].gathered = True
        divergent.there_is_another = False
        divergent.gathered = True
        if len(divergent):
            log("asnoun() . . . . 4")
            for party_lst in divergent:
                nouns = SCList("nouns", there_are_any=True, auto_gather=False)
                nouns.clear()
                for party_type in party_lst[0].party_types:
                    if party_type.name.text not in nouns.elements:
                        nouns.append(party_type.name.text)
                nouns.there_is_another = False
                nouns.gathered = True
                designations = []
                designations.clear()
                for the_noun in nouns:
                    the_noun = re.sub(r".*\.", "", the_noun)
                    the_noun = re.sub(r"_", " ", the_noun)
                    if (party_lst.number() > 1 or party_lst.number() == 0):
                        designations.append(noun_plural(the_noun, language=language))
                    else:
                        designations.append(noun_singular(the_noun))
                the_parties = []
                the_parties.clear()
                others = 0
                log("asnoun() . . . . 5")
                for party in party_lst:
                    for different_party in self:
                        if different_party not in party_lst:
                            for list_type in designations:
                                log("asnoun() . . . . 6")
                                if any(str(party_type.name.text) == str(list_type) for party_type in
                                       different_party.party_types):
                                    log("asnoun() . . . . 7")
                                    the_parties.append(party.name.full())
                                    others += 1
                if others >= 1:
                    log("asnoun() . . . . 8a")
                    item_string = str(comma_and_list(designations) + " " + comma_and_list(the_parties))
                else:
                    log("asnoun() . . . . 8b")
                    item_string = str(comma_and_list(designations))
                strings.append(str(item_string))
        else:
            log("asnoun() . . . . 8c")
            strings = self.caption_noun()
        log("asnoun() . . . . 8d")
        return str(comma_and_list(strings))

    def asnounold(self, **kwargs):
        language = kwargs.get("language", None)
        output = []
        output.clear()
        divergent = SCList("divergent", there_are_any=True, auto_gather=False, object_type=SCList)
        divergent.clear()
        strings = []
        strings.clear()
        log("asnounold")
        for all_parties in self:
            if all_parties not in flatten(divergent.elements):
                log(repr(flatten(divergent.elements)))
                divergent.appendObject()
                divergent[-1].clear()
                # divergent[-1].append(all_parties)
                # if any(party.party.party_types != all_parties.party_types for party in self):
                for party in [ex for ex in self if ex.party_types == all_parties.party_types]:
                    divergent[-1].append(party)
                divergent[-1].there_is_another = False
                divergent[-1].gathered = True
        divergent.there_is_another = False
        divergent.gathered = True
        if len(divergent):
            for party_lst in divergent:
                # output.clear()
                nouns = SCList("nouns", there_are_any=True, auto_gather=False)
                nouns.clear()
                # for party in party_lst:
                for party_type in party_lst[0].party_types:
                    if party_type.name.text not in nouns.elements:
                        nouns.append(party_type.name.text)
                nouns.there_is_another = False
                nouns.gathered = True
                designations = []
                designations.clear()
                for the_noun in nouns:
                    the_noun = re.sub(r".*\.", "", the_noun)
                    the_noun = re.sub(r"_", " ", the_noun)
                    if (party_lst.number() > 1 or party_lst.number() == 0):
                        designations.append(noun_plural(the_noun, language=language))
                    else:
                        designations.append(noun_singular(the_noun))

                item_string = str(comma_and_list(designations))
                strings.append(str(item_string))
        return comma_and_list(strings)

    def asnounall(self, **kwargs):
        language = kwargs.get("language", None)
        output = []
        output.clear()
        divergent = SCList("divergent", there_are_any=True, auto_gather=False, object_type=SCList)
        divergent.clear()
        strings = []
        strings.clear()
        log("asnounall")
        for all_parties in self:
            if all_parties not in flatten(divergent.elements):
                log(repr(flatten(divergent.elements)))
                divergent.appendObject()
                divergent[-1].clear()
                # divergent[-1].append(all_parties)
                if any(party.party.party_types != all_parties.party_types for party in self):
                    for party in [ex for ex in self if ex.party_types == all_parties.party_types]:
                        divergent[-1].append(party)
                    divergent[-1].there_is_another = False
                    divergent[-1].gathered = True
        divergent.there_is_another = False
        divergent.gathered = True
        if len(divergent):
            for party_lst in divergent:
                # output.clear()
                nouns = SCList("nouns", there_are_any=True, auto_gather=False)
                nouns.clear()
                # for party in party_lst:
                for party_type in party_lst[0].party_types:
                    if party_type.name.text not in nouns.elements:
                        nouns.append(party_type.name.text)
                nouns.there_is_another = False
                nouns.gathered = True
                designations = []
                designations.clear()
                for the_noun in nouns:
                    the_noun = re.sub(r".*\.", "", the_noun)
                    the_noun = re.sub(r"_", " ", the_noun)
                    if (party_lst.number() > 1 or party_lst.number() == 0):
                        designations.append(noun_plural(the_noun, language=language))
                    else:
                        designations.append(noun_singular(the_noun))
                the_parties = []
                the_parties.clear()
                others = 0
                for party in party_lst:
                    for different_party in self:
                        if different_party not in party_lst:
                            for list_type in designations:
                                if any(str(party_type.name.text) == str(list_type) for party_type in
                                       different_party.party_types):
                                    the_parties.append(party.party.name.full())
                                    others += 1
                if others >= 1:
                    item_string = str(comma_and_list(designations) + " " + comma_and_list(the_parties))
                else:
                    item_string = str(comma_and_list(designations))
                strings.append(str(item_string))
        return comma_and_list(strings)

    def asnounobjects(self, **kwargs):
        language = kwargs.get("language", None)
        the_noun = str(self[0].party_type.name.text)
        the_noun = re.sub(r".*\.", "", the_noun)
        the_noun = re.sub(r"_", " ", the_noun)
        if (
                self.number() > 1
                or self.number() == 0
                or ("plural" in kwargs and kwargs["plural"])
        ) and not ("singular" in kwargs and kwargs["singular"]):
            output = noun_plural(the_noun, language=language) + " object "
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output
        else:
            output = noun_singular(the_noun) + " objects "
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output

    def asnounobjects_further(self, **kwargs):
        flow_steppers = [" also ", " further ", " ", " ", " "]
        language = kwargs.get("language", None)
        the_noun = str(self[0].party_type.name.text)
        the_noun = re.sub(r".*\.", "", the_noun)
        the_noun = re.sub(r"_", " ", the_noun)
        if (
                self.number() > 1
                or self.number() == 0
                or ("plural" in kwargs and kwargs["plural"])
        ) and not ("singular" in kwargs and kwargs["singular"]):
            output = (
                    noun_plural(the_noun, language=language)
                    + random.choice(flow_steppers)
                    + "object "
            )
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output
        else:
            output = noun_singular(the_noun) + random.choice(flow_steppers) + "objects "
            if "article" in kwargs and kwargs["article"]:
                if "some" in kwargs and kwargs["some"]:
                    output = some(output, language=language)
            elif "this" in kwargs and kwargs["this"]:
                output = these(output, language=language)
            if "capitalize" in kwargs and kwargs["capitalize"]:
                return capitalize(output)
            else:
                return output

    def caption(self, **kwargs):
        output = []
        output.clear()
        for item in self:
            output.append(
                self.fx.BOLD
                + item.name.full().strip().upper()
                + self.fx.BOLD
                + item.caption_text
            )
        return output

    def facite_rrfps(self, case):
        indexr = Counter()
        rfpval = case.prior_requests_total
        output = list()
        for item in self:
            if this_thread.evaluation_context == "docx":
                if indexr.get_incr() > 1:
                    output.append("\n\n")
                output.append(
                    self.fx.BOLD
                    + "REQUEST FOR PRODUCTION NO. "
                    + str(rfpval)
                    + ":"
                    + self.fx.BOLD
                    + fix_punctuation(single_paragraph(item.name.text))
                )
                output.append("\n\n")
                output.append(self.fx.BOLD + "RESPONSE: " + self.fx.BOLD)
            else:
                if indexr.get_incr() > 1:
                    output.append(" [NEWLINE] ")
                output.append(
                    "[FLUSHLEFT] "
                    + self.fx.BOLD
                    + "REQUEST FOR PRODUCTION NO. "
                    + str(rfpval)
                    + ":"
                    + self.fx.BOLD
                    + single_paragraph(item.name.text)
                    + " [NEWLINE] "
                )
                output.append(self.fx.BOLD + "RESPONSE:" + self.fx.BOLD)
            # if item.other_other:
            # output.append(fix_punctuation(item.other_other) + ' ')
            if item.overly_broad:
                output.append(
                    case.user_clientlist().asnounobjects()
                    + "to this request to the extent that it is overly broad"
                )
                if item.not_reasonably_calculated:
                    output.append(
                        ", unduly burdensome, and not reasonably calculated to lead to the discovery of admissible evidence. "
                    )
                else:
                    output.append(" and unduly burdensome.")
            if item.not_reasonably_calculated and not item.overly_broad:
                output.append(
                    case.user_clientlist().asnounobjects()
                    + "to this request as not reasonably calculated to lead to the discovery of admissible evidence. "
                )
            if item.attorney_client:
                if item.overly_broad or item.not_reasonably_calculated:
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request to the extent that it seeks documents protected by the attorney client privilege and/or attorney work product doctrine. "
                )
            if item.same_bodypart:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as not limited in scope as to the same or similar body parts alleged to have been injured in the Complaint, and requests documents, should any exist, that are confidential under HIPAA and the Oregon Evidence Code, including, but not limited to, the physician-patient privilege, psychotherapist-patient privilege, nurse-patient privilege and/or counselor-client privilege. "
                )
            if item.public_record:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as calling for a classification of documents that would be public record and, as such, equally available to the "
                )
                output.append(case.recipients().asnoun())
                output.append(" requesting them. ")
            if item.vauge_ambiguous:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append("to this request as vague and/or ambiguous. ")
            if item.overbroad_time:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append("to this request as overbroad as to time. ")
            if item.defendant_possesses:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking documents already within the control of the "
                )
                output.append(possessify(case.recipients().asnoun(), ""))
                output.append("requesting them. ")
            if item.interrogatory:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as constituting an impermissible interrogatory. "
                )
            if item.expert:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                        or item.interrogatory
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking impermissible expert discovery. "
                )
            if item.biz_docs:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                        or item.interrogatory
                        or item.expert
                ):
                    output.append(case.user_clientlist().asnounobjects_further())
                else:
                    output.append(case.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking highly confidential business documents. "
                )
            # if item.other:
            #  output.append(fix_punctuation(item.other) + ' ')
            if (
                    item.overly_broad
                    or item.not_reasonably_calculated
                    or item.attorney_client
                    or item.same_bodypart
                    or item.public_record
                    or item.vauge_ambiguous
                    or item.overbroad_time
                    or item.defendant_possesses
                    or item.interrogatory
                    or item.expert
                    or item.biz_docs
            ):
                output.append("Subject to")
                if item.attorney_client:
                    output.append(", and without waiving,")
                output.append(" said objections, " + item.responsive)
                if item.responsive == "see response to RFP No.":
                    output.append(" " + str(item.see_rfp) + ". ")
            else:
                output.append(capitalize(item.responsive))
                if item.responsive == "see response to RFP No.":
                    output.append(" " + str(item.see_rfp) + ". ")
            rfpval += 1
        return "".join(output)

    def facite_rrogs(self, firm, recipients, rog_val, cases):
        indexr = Counter()
        output = list()
        for item in self:
            if this_thread.evaluation_context == "docx":
                if indexr.get_incr() > 1:
                    output.append("\n\n")
                output.append(
                    self.fx.BOLD
                    + "INTERROGATORY NO. "
                    + str(rog_val)
                    + ":"
                    + self.fx.BOLD
                    + " "
                    + fix_punctuation(single_paragraph(item.name.text))
                )
                output.append("\n\n")
                output.append(self.fx.BOLD + "ANSWER: " + self.fx.BOLD)
            else:
                if indexr.get_incr() > 1:
                    output.append(" [NEWLINE] ")
                output.append(
                    "[FLUSHLEFT] "
                    + self.fx.BOLD
                    + "INTERROGATORY NO. "
                    + str(rog_val)
                    + ":"
                    + self.fx.BOLD
                    + single_paragraph(item.name.text)
                    + " [NEWLINE] "
                )
                output.append(self.fx.BOLD + "ANSWER:" + self.fx.BOLD)

            if item.overly_broad:
                output.append(
                    firm.clientlist(cases).asnounobjects()
                    + "to this request to the extent that it is overly broad"
                )
                if item.not_reasonably_calculated:
                    output.append(
                        ", unduly burdensome, and not reasonably calculated to lead to the discovery of admissible evidence. "
                    )
                else:
                    output.append(" and unduly burdensome.")
            if item.not_reasonably_calculated and not item.overly_broad:
                output.append(
                    firm.clientlist(cases).asnounobjects()
                    + "to this request as not reasonably calculated to lead to the discovery of admissible evidence. "
                )

            if item.cumulative:
                if item.overly_broad or item.not_reasonably_calculated:
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "as the information sought is unreasonably cumulative and duplicative, "
                )
                if item.obtainable_rrfp:
                    output.append(
                        "as it is obtainable from documents provided in response to "
                        + possessify(recipients.asnoun(), "")
                        + " Request for Production, "
                    )
                elif item.obtainable_others:
                    output.append(
                        "as it is obtainable from documents already in Defendant’s possession and/or which will be provided pursuant to subpoena of third parties, "
                    )
                output.append(
                    "which is a more convenient and less burdensome source of the requested information. "
                )

            if item.work_product:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append("as calling for attorney work product. ")

            if item.attorney_client:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "as calling for information protected by the attorney client privilege. "
                )

            if item.same_bodypart:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "as calling for physician-patient privileged information. "
                )

            if item.vauge_ambiguous:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.obtainable_rrfp
                        or item.work_product
                        or item.obtainable_others
                        or item.attorney_client
                        or item.same_bodypart
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append("to this request as vague and/or ambiguous. ")

            if item.public_record:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.attorney_client
                        or item.work_product
                        or item.same_bodypart
                        or item.vauge_ambiguous
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "as calling for public information equally available to Defendant. "
                )

            if item.defendant_possesses:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.attorney_client
                        or item.work_product
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "to this request as calling for information already known by requesting defendant. "
                )

            if item.not_medically_stationary:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.work_product
                        or item.cumulative
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "to the extent that "
                    + firm.clientlist(cases).pronoun_subjective()
                    + " is not medically stationary, and continues to treat for "
                    + firm.clientlist(cases).pronoun_possessive("injuries")
                )
                if item.incomplete_cns:
                    output.append(
                        ", and because "
                        + firm.clientlist(cases).pronoun_subjective()
                        + " may not have in "
                        + firm.clientlist(cases).pronoun_possessive(
                            "custody, possession, or control a complete set of all medical chart notes incurred to date. "
                        )
                    )
                if item.changing_symptoms:
                    output.append(" with symptoms that continue to change over time. ")

            if item.expert:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "as calling for expert opinion prior to the date for expert disclosures. "
                )

            if item.seeks_docs:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "to the extent that it seeks to obtain the production of documents or to request a complete recital of the contents of documents."
                )

            if item.legal_conclusion:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.obtainable_rrfp
                        or item.obtainable_others
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                ):
                    output.append(firm.clientlist(cases).asnounobjects_further())
                else:
                    output.append(firm.clientlist(cases).asnounobjects())
                output.append(
                    "to the extent that it calls for a pure legal conclusion. "
                )

            if (
                    item.overly_broad
                    or item.not_reasonably_calculated
                    or item.cumulative
                    or item.work_product
                    or item.same_bodypart
                    or item.vauge_ambiguous
                    or item.public_record
                    or item.defendant_possesses
                    or item.not_medically_stationary
                    or item.expert
                    or item.seeks_docs
                    or item.legal_conclusion
            ):
                output.append("Subject to")
                if item.attorney_client:
                    output.append(", and without waiving,")
                output.append(" said objections, " + item.answer)
            else:
                output.append(capitalize(item.answer))
        rog_val += 1
        return "".join(output)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCPersonMixin(object):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.initializeAttribute("fx", Color)

    def caption_description(self):
        item = self
        output = ""
        if item.name.caption_type.name.text == "individual":
            output = ""
        elif item.name.caption_type.name.text == "individual proceeding under a pseudonym":
            output = ", an individual proceeding under a pseudonym"
        elif item.name.caption_type.name.text == "trustee":
            output = str(
                ", in "
                + item.pronoun_possessive("capacity")
                + " as trustee of the "
                + self.fx.BOLD
                + item.trust.upper()
                + self.fx.BOLD
            )
        elif item.name.caption_type.name.text == "personal representative of an estate":
            output = str(
                ", in "
                + item.pronoun_possessive("capacity")
                + " as Personal Representative for the Estate of "
                + self.fx.BOLD
                + item.decedent.upper()
                + self.fx.BOLD
            )
        elif item.name.caption_type.name.text == "guardian ad litem":
            output = str(
                ", as guardian ad litem for "
                + self.fx.BOLD
                + item.guards.name.full().upper()
                + self.fx.BOLD
                + " (a minor)"
            )
        elif item.name.caption_type.name.text == "doing business as":
            output = str(
                ", an individual doing business under the assumed name "
                + self.fx.BOLD
                + item.dba.upper()
                + self.fx.BOLD
            )
        elif item.name.caption_type.name.text == "governmental entity":
            output = ""
        elif item.name.caption_type.name.text == "business entity":
            if item.business_entity.name.text in (
                    "Foreign Nonprofit Corporation",
                    "Foreign Business Corporation",
                    "Foreign Professional Corporation",
            ):
                output = str(
                    ", "
                    + indefinite_article(state_name(item.stofinc))
                    + " "
                    + item.business_entity.name.text[8:].lower()
                )
            else:
                output = str(", " + indefinite_article(item.business_entity.name.text.lower()))
        return output

    def relatio_scriptum(self):
        if self.is_party:
            scriputm = str(self.party_type.name.text)
        elif self.is_lawyer:
            scriputm = str("Attorney for " + comma_and_list(self.firm.clientlist(case)))

    def possessive(self, target, **kwargs):
        the_word = str(self)
        if the_word[-1] == "s" and "plural" not in kwargs:
            kwargs["plural"] = True
        return super().possessive(target, **kwargs)

    def represented_by(self, firms):
        output = []
        for item in firms:
            if self.name.full() in comma_and_list(item.firms):
                output.append(item)
        return output

    def represents(self, parties, firms):
        output = []
        for party in parties:
            if self.name.full() in comma_and_list(party.party.represented_by(firms)):
                output.append(party)
        return output

    def works_at(self, firms):
        for item in firms:
            if item.name.text == self.firm:
                return item

    def pair_firmz(self, firms, case):
        output = SCList(there_are_any=True, auto_gather=False)
        output.clear()
        for item in firms:
            if comma_and_list(case.parties.client_of(self)) in comma_and_list(
                    case.parties.client_of(item)
            ):
                output.append(item)
        output.there_is_another = False
        output.gathered = True
        return output

    def employs(self, counselors, firms):
        output = []
        for item in counselors:
            for office in self.pair_firmz(firms):
                if office.name.text in item.firm:
                    output.append(item)
        return output

    def employs_nopair(self, counselors):
        output = []
        for item in counselors:
            if self.name.text in item.firm:
                output.append(item)
        return output

    def pair_firms(self, firms, case):
        output = []
        firma = firms.copy()
        z = firma.index(self)
        del firma[z]
        for item in firma:
            if comma_and_list(case.parties.client_of(self)) in comma_and_list(
                    case.parties.client_of(item)
            ):
                output.append(item)
        return output

    def attorneys_noun(self, counselors, firms):
        output = list()
        if self.start >= 2:
            return "attorneys "
        else:
            if len(self.employs(counselors, firms)) >= 2:
                return "attorneys "
        return "attorney "

    def lawyers_as_friends(self, firms, case):
        output = list()
        for item in self.pair_firmz(firms, case):
            for attorney in case.attorneys_for(item):
                output.append(attorney)
        if len(output) >= 2:
            return "attorneys"
        else:
            return "attorney"

    def SCaddress_block(self, language=None, international=False, show_country=False):
        if this_thread.evaluation_context == "docx":
            if len(self.address.company):
                return (
                        self.name.full()
                        + '</w:t><w:br/><w:t xml:space="preserve">'
                        + self.address.company
                        + '</w:t><w:br/><w:t xml:space="preserve">'
                        + self.address.block(
                    language=language,
                    international=international,
                    show_country=show_country,
                )
                )
            else:
                return (
                        self.name.full()
                        + '</w:t><w:br/><w:t xml:space="preserve">'
                        + self.address.block(
                    language=language,
                    international=international,
                    show_country=show_country,
                )
                )
        else:
            return (
                    "[FLUSHLEFT] "
                    + self.name.full()
                    + " [NEWLINE] "
                    + self.address.block(
                language=language,
                international=international,
                show_country=show_country,
            )
            )

    def clientlist(self, cases):
        output = SCList("clientlist", there_are_any=True, auto_gather=False)
        output.clear()
        for client in cases.clients_of(self):
            output.append(client)
        output.there_is_another = False
        output.gathered = True
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCPerson(SCPersonMixin, Person):

    def last_name(self):
        return self.name.text

    def fullname(self):
        return self.name.text


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCIndividual(SCPersonMixin, Individual):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)

    def pronoun_possessive(self, target, **kwargs):
        """Given a word like "fish," returns "her fish" or "his fish," as appropriate."""
        if self == this_thread.global_vars.user and ('thirdperson' not in kwargs or not kwargs['thirdperson']):
            output = your(target, **kwargs)
        elif self.gender.name.text == 'female':
            output = her(target, **kwargs)
        elif self.gender.name.text == 'other':
            output = their(target, **kwargs)
        else:
            output = his(target, **kwargs)
        if 'capitalize' in kwargs and kwargs['capitalize']:
            return capitalize(output)
        return output

    def pronoun(self, **kwargs):
        """Returns a pronoun like "you," "her," or "him," as appropriate."""
        if self == this_thread.global_vars.user:
            output = word('you', **kwargs)
        if self.gender.name.text == 'female':
            output = word('her', **kwargs)
        elif self.gender.name.text == 'other':
            output = word('them', **kwargs)
        else:
            output = word('him', **kwargs)
        if 'capitalize' in kwargs and kwargs['capitalize']:
            return capitalize(output)
        return output

    def pronoun_objective(self, **kwargs):
        """Same as pronoun()."""
        return self.pronoun(**kwargs)

    def pronoun_subjective(self, **kwargs):
        """Returns a pronoun like "you," "she," or "he," as appropriate."""
        if self == this_thread.global_vars.user and ('thirdperson' not in kwargs or not kwargs['thirdperson']):
            output = word('you', **kwargs)
        elif self.gender.name.text == 'female':
            output = word('she', **kwargs)
        elif self.gender.name.text == 'other':
            output = word('they', **kwargs)
        else:
            output = word('he', **kwargs)
        if 'capitalize' in kwargs and kwargs['capitalize']:
            return capitalize(output)
        return output

    # function to be later adapted for class methods
    def split_title_and_name(self):
        title = person.split()[0]
        lastname = person.split()[-1]
        return '{} {}'.format(title, lastname)
        list(map(split_title_and_name, people))

    @property
    def is_doctor(self):
        if (
                hasattr(self, "name.suffix")
                and len(self.name.suffix.strip())
                and self.name.suffix in ("MD", "DO", "DC", "DMD", "PhD", "PsyD")
        ):
            return True
        else:
            return False

    def last_name(self):
        return self.name.last

    def fullname(self):
        return self.name.first + self.name.last


class SCLawfirm(SCPerson):
    def init(self, *pargs, **kwargs):
        if not hasattr(self, 'fax_number') and 'fax_number' not in kwargs:
            self.initializeAttribute('fax_number', DAObject)
        if not hasattr(self, 'phone_number') and 'phone_number' not in kwargs:
            self.initializeAttribute('phone_number', DAObject)
        super().init(*pargs, **kwargs)

    @property
    def complete_lawfirm(self):
        self.name.text
        self.address.address

    @property
    def is_userfirm(self):
        if self.name.text == "Scott M Cumming, PC":
            return True
        else:
            return False


class Provider(SCPerson, SQLObject):
    _model = ProviderModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        elif column == 'address':
            return self.address.address.strip()
        elif column == "unit":
            return self.address.unit.strip()
        elif column == "city":
            return self.address.city.strip()
        elif column == "state":
            return self.address.state.strip()
        elif column == "zip":
            return self.address.zip.strip()
        elif column == "fax_number":
            return self.fax_number
        elif column == "phone_number":
            return self.phone_number
        elif column == "email":
            return self.email
        elif column == "specialty":
            return self.specialty
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        elif column == 'address':
            self.address.address = value
        elif column == "unit":
            self.address.unit = value
        elif column == "city":
            self.address.city = value
        elif column == "state":
            self.address.state = value
        elif column == "zip":
            self.address.zip = value
        elif column == "fax_number":
            self.fax_number = value
        elif column == "phone_number":
            self.phone_number = value
        elif column == "email":
            self.email = value
        elif column == "specialty":
            self.specialty = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        elif column == 'address':
            del self.address.address
        elif column == "unit":
            del self.address.unit
        elif column == "city":
            del self.address.city
        elif column == "state":
            del self.address.state
        elif column == "zip":
            del self.address.zip
        elif column == "fax_number":
            del self.fax_number
        elif column == "phone_number":
            del self.phone_number
        elif column == "email":
            del self.email
        elif column == "specialty":
            del self.specialty
        else:
            raise Exception("Invalid column " + column)

class Lawfirm(SCLawfirm, SQLObject):
    _model = LawfirmModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        elif column == 'address':
            return self.address.address.strip()
        elif column == "unit":
            return self.address.unit.strip()
        elif column == 'city':
            return self.address.city
        elif column == 'state':
            return self.address.state
        elif column == "zip":
            return self.address.zip.strip()
        elif column == "phone_number":
            return self.phone_number.strip()
        elif column == "fax_number":
            return self.fax_number.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        elif column == 'address':
            self.address.address = value
        elif column == "unit":
            self.address.unit = value
        elif column == 'city':
            self.address.city = value
        elif column == 'state':
            self.address.state = value
        elif column == "zip":
            self.address.zip = value
        elif column == "phone_number":
            self.phone_number = value
        elif column == "fax_number":
            self.fax_number = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        elif column == 'address':
            del self.address.address
        elif column == "unit":
            del self.address.unit
        elif column == "city":
            del self.address.city
        elif column == "state":
            del self.address.state
        elif column == "zip":
            del self.address.zip
        elif column == "phone_number":
            del self.phone_number
        elif column == "fax_number":
            del self.fax_number
        else:
            raise Exception("Invalid column " + column)

    @classmethod
    def all_lawfirm_names(cls, instance_name=None):
        if 'dbcache' not in this_thread.misc:
            this_thread.misc['dbcache'] = {}
        if instance_name:
            listobj = DAList(instance_name, object_type=cls)
        else:
            listobj = DAList(object_type=cls)
            listobj.set_random_instance_name()
        for db_entry in list(cls._session.query(cls._model.name).order_by(cls._model.name).all()):
            if cls._model.__name__ in this_thread.misc['dbcache'] and db_entry.name in this_thread.misc['dbcache'][
                cls._model.__name__]:
                listobj.append(this_thread.misc['dbcache'][cls._model.__name__][db_entry.name])
            else:
                listobj.append(db_entry.name)
        listobj.gathered = True
        return listobj

    def userfirm_exists(self):
        try:
            if not self.ready():
                raise Exception("userfirm_exists: cannot retrieve data")
            db_entry = self._session.query(LawfirmModel).filter(LawfirmModel.is_userfirm == True).all()
            if db_entry is None:
                return False
            elif len(db_entry) == 1:
                return True
            else:
                return False
        except:
            self._session.rollback()

    def has_lawyer(self, lawyer):
        try:
            if not (self.ready() and lawyer.ready()):
                raise Exception("has_lawyer: cannot retrieve data")
            db_entry = self._session.query(LawfirmLawyerModel).filter(LawfirmLawyerModel.lawfirm_id == self.id,
                                                                      LawfirmLawyerModel.lawyer_id == lawyer.id).first()
            if db_entry is None:
                return False
            return True
        except:
            self._session.rollback()

    def add_lawyer(self, lawyer):
        try:
            if not self.has_lawyer(lawyer):
                db_entry = LawfirmLawyerModel(lawfirm_id=self.id, lawyer_id=lawyer.id)
                self._session.add(db_entry)
                self._session.commit()
        except:
            self._session.rollback()

    def get_lawyers(self):
        if not self.ready():
            raise Exception("get_lawyers: cannot retrieve data")
        results = SCLawyerList("get_lawyers", there_are_any=True, auto_gather=False)
        results.clear()
        for db_entry in self._session.query(LawfirmLawyerModel).filter(LawfirmLawyerModel.lawfirm_id == self.id).all():
            results.append(Lawyer.by_id(db_entry.lawyer_id))
        results.there_is_another = False
        results.gathered = True
        return results

    def del_lawyer(self, lawyer):
        if not (self.ready() and lawyer.ready()):
            raise Exception("del_lawyer: cannot retrieve data")
        self._session.query(LawfirmLawyerModel).filter(LawfirmLawyerModel.lawfirm_id == self.id,
                                                       LawfirmLawyerModel.lawyer_id == lawyer.id).delete()
        self._session.commit()

    def has_party(self, party):
        try:
            if not (self.ready() and party.ready()):
                raise Exception("has_party: cannot retrieve data")
            db_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.id,
                                                                     PartyLawfirmModel.party_id == party.id).first()
            if db_entry is None:
                return False
            return True
        except:
            self._session.rollback()

    def get_parties(self):
        try:
            if not self.ready():
                raise Exception("get_parties: cannot retrieve data")
            results = SCPartyList("get_parties", there_are_any=True, auto_gather=False)
            results.clear()
            for db_entry in self._session.query(PartyLawfirmModel).filter(
                    PartyLawfirmModel.lawfirm_id == self.id).all():
                results.append(Party.by_id(db_entry.party_id))
            results.there_is_another = False
            results.gathered = True
            return results
        except:
            self._session.rollback

    def del_party(self, party):
        if not (self.ready() and party.ready()):
            raise Exception("del_party: cannot retrieve data")
        self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.id,
                                                      PartyLawfirmModel.party_id == party.id).delete()
        self._session.commit()

    def has_party_case(self, party, case):
        try:
            if not (self.ready() and party.ready() and case.ready()):
                raise Exception("has_party: cannot retrieve data")
            db_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.id,
                                                                     PartyLawfirmModel.party_id == party.id).first()
            if db_entry is None:
                return False
            return True
        except:
            self._session.rollback()

    def add_party_case(self, party, case):
        try:
            if not self.has_party_case(party, case):
                db_entry = PartyLawfirmModel(lawfirm_id=self.id, party_id=party.id, cases_id=case.id)
                self._session.add(db_entry)
                self._session.commit()
        except:
            self._session.rollback

    def get_case_parties(self, case):
        try:
            if not self.ready() and not case.ready():
                raise Exception("get_case_parties: cannot retrieve data")
            results = SCPartyList("get_case_parties", there_are_any=True, auto_gather=False)
            results.clear()
            for db_entry in self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.id,
                                                                          PartyLawfirmModel.cases_id == case.id).all():
                results.append(Party.by_id(db_entry.party_id))
            results.there_is_another = False
            results.gathered = True
            return results
        except:
            self._session.rollback

    def del_party_case(self, party, case):
        if not (self.ready() and party.ready() and case.ready()):
            raise Exception("del_party_case: cannot retrieve data")
        self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.id,
                                                      PartyLawfirmModel.party_id == party.id,
                                                      PartyLawfirmModel.cases_id == case.id).delete()
        self._session.commit()


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLawfirmList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = CaseLawfirm
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_lawfirm"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == item.cases.id,
                                                     CaseLawfirmModel.lawfirm_id == item.lawfirm.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCClaimList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = PartyClaim
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_claim"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(PartyClaimModel).filter(PartyClaimModel.party_id == item.party.id,
                                                    PartyClaimModel.claim_id == item.claim.id,
                                                    PartyClaimModel.adverse_party_id == item.adverse_party.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLawyer(SCIndividual):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("firm", Lawfirm)
        super().init(*pargs, **kwargs)

    def bar(self):
        return str("OSB No. " + self.bar_no)
        output = ""
        if self.barred_in == "AL":
            output += "Bar Code No. "
            output += self.bar_no
        elif self.barred_in == "AZ":
            output += parens(self.bar_no)
        elif self.barred_in == "OK":
            output += "OBA #"
            output += self.bar_no
        elif self.barred_in == "OR":
            output += "OSB No. "
            output += self.bar_no
            return output
        elif state_name(self.barred_in) == "California":
            output += "SBN "
            output += self.bar_no
        elif self.barred_in == "CO":
            output += ", #"
            output += self.bar_no
        elif state_name(self.barred_in) == "Washington":
            output += "Florida Bar No. "
            output += self.bar_no
        elif state_name(self.barred_in) == "Minnesota":
            output += "#"
            output += self.bar_no
        elif state_name(self.barred_in) == "New York":
            output += "NY Bar #"
            output += self.bar_no
        elif state_name(self.barred_in) == "Nevada":
            output += "Nevada Bar No."
            output += self.bar_no
        elif state_name(self.barred_in) == "District of Columbia":
            output += "DC Bar No. "
            output += self.bar_no
        elif state_name(self.barred_in) == "Washington":
            output += "Washington State Bar No. "
            output += self.bar_no
        else:
            output += str(state_name(self.barred_in) + " Bar No. ")
            output += self.bar_no
        return output

    def pro_hac_vice(self, case):
        if self.barred_in != case.juris.state.name.text:
            return True
        else:
            return False

    @property
    def is_party(self):
        return False

    @property
    def is_lawyer(self):
        return True

    @property
    def is_lawfirm(self):
        return False

    @property
    def clio_type(self):
        return "Person"


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class Lawyer(SCLawyer, SQLObject):
    _model = LawyerModel
    _session = DBSession
    _required = ['email']
    _parent = [Lawfirm, 'lawfirm', 'lawfirm_id']
    _uid = 'email'

    @property
    def complete_lawyer(self):
        self.email
        self.name.first
        # self.firm.name.text
        self.gender.name.text

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'first_name':
            return self.name.first.strip()
        elif column == 'middle_name':
            return self.name.middle.strip()
        elif column == 'last_name':
            return self.name.last.strip()
        elif column == 'email':
            return self.email.strip()
        elif column == 'fax_number':
            return self.fax_number.strip()
        elif column == 'phone_number':
            return self.phone_number.strip()
        elif column == 'gender_id':
            return self.gender.id
        elif column == 'lawfirm_id':
            return self.firm.id
        elif column == 'state_bar_id':
            return self.barred_in.id
        elif column == 'bar_no':
            return self.bar_no
        raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'first_name':
            self.name.first = value
        elif column == 'middle_name':
            self.name.middle = value
        elif column == 'last_name':
            self.name.last = value
        elif column == 'email':
            self.email = value
        elif column == 'fax_number':
            self.fax_number = value
        elif column == 'phone_number':
            self.phone_number = value
        elif column == 'gender_id':
            self.gender = Gender.by_id(value)
        elif column == 'lawfirm_id':
            # Suppose you have a Python object plaintiff of class Party and plaintiff.pet is an object of type Pet, and you have two tables in SQL called parties and pets for tracking people and pets, respectively. A person can only have one pet in this scenario.
            # In the parties table you could have an integer column pet_id that references the id column of the pets table in the normal SQLAlchemy foreign-key way. Your db_set() for the pet_id column would set .pet to the output of Pet.by_id(value). Your db_get() for the pet_id column would return self.pet.id.
            self.firm = Lawfirm.by_id(value)
        elif column == 'state_bar_id':
            self.barred_in = State.by_id(value)
        elif column == 'bar_no':
            self.bar_no = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'first_name':
            del self.name.first
        elif column == 'middle_name':
            del self.name.middle
        elif column == 'last_name':
            del self.name.last
        elif column == 'email':
            del self.email
        elif column == 'fax_number':
            del self.fax_number
        elif column == 'phone_number':
            del self.phone_number
        elif column == "gender_id":
            del self.gender.name.text
        elif column == "lawfirm_id":
            del self.firm
        elif column == 'state_bar_id':
            del self.barred_in
        elif column == 'bar_no':
            del self.bar_no
        else:
            raise Exception("Invalid column " + column)

    def get_bar_nos(self):
        if not (self.ready()):
            raise Exception("get_bar_no: cannot retrieve data")
        get_bar_nos_output = []
        get_bar_nos_output.clear()
        log("get_bar_nos 1")
        if self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.lawyer_id == self.id).first():
            log("get_bar_nos 2a")
            for admission in self._session.query(LawyerStateBarModel).filter(
                    LawyerStateBarModel.lawyer_id == self.id).all():
                log("get_bar_nos 3")
                get_bar_nos_output.append(admission)
            return get_bar_nos_output
        else:
            log("get_bar_nos 2b")
            return None

    def get_bar_no(self, state):
        if not (self.ready() and state.ready()):
            raise Exception("get_bar_no: cannot retrieve data")
        log("get_bar_noget_bar_no 1")
        if self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.lawyer_id == self.id,
                                                           LawyerStateBarModel.state_id == state.id).first():
            log("get_bar_noget_bar_no 2a")
            admission = self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.lawyer_id == self.id,
                                                                        LawyerStateBarModel.state_id == state.id).first()
            log("get_bar_noget_bar_no 3")
            return admission.bar_no
        else:
            log("get_bar_noget_bar_no 2b")
            return None

    def get_state_admission(self):
        if not (self.ready()):
            raise Exception("get_bar_no: cannot retrieve data")
        get_state_admissions_output = SCList("get_state_admissions_output", there_are_any=True, auto_gather=True)
        get_state_admissions_output.clear()
        log("get_state_admissions_output 1")
        if self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.lawyer_id == self.id).first():
            log("get_state_admissions_output 2")
            admission = self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.lawyer_id == self.id).all()
            for admission in admissions:
                log("get_state_admissions_output 3")
                get_state_admissions_output.append(State.by_id(admission.state_id))
        get_state_admissions_output.there_is_another = False
        get_state_admissions_output.gathered = True
        return get_state_admissions_output

    def has_admission(self, state):
        if not (self.ready() and state.ready()):
            raise Exception("has_statebar_no: cannot retrieve data")
        state_bar_no_entry = self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.state_id == state.id,
                                                                             LawyerStateBarModel.lawyer_id == self.id).first()
        if state_bar_no_entry is None:
            return False
        return True

    def add_admission(self, state, bar_no):
        log("add_statebar_no1")
        if not self.has_statebar_no(state):
            log("add_statebar_no2")
            state_bar_no_entry = LawyerStateBarModel(state_id=state.id, bar_no=bar_no, lawyer_id=self.id)
            self._session.add(state_bar_no_entry)
            self._session.commit()

    def del_admission(self, state):
        if not (self.ready() and state.ready()):
            raise Exception("del_bar_no_state: cannot retrieve data")
        self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.state_id == state.id,
                                                        LawyerStateBarModel.lawyer_id == self.id).delete()
        self._session.commit()

    def db_userfirm_lawyers(self):
        self.userfirm_lawyers = SCLawyerList("self.userfirm_lawyers")
        self.userfirm_lawyers.there_are_any = True
        self.userfirm_lawyers.auto_gather = False
        self.userfirm_lawyers.clear()
        db_entries = []
        for userfirm_lawyer in self.userfirm.lawyers:
            db_entries.append(self._session.query(LawyerModel).filter(LawyerModel.id == userfirm_lawyer.id).first())
        if len(db_entries):
            for db_entry in db_entries:
                self.userfirm_lawyers.append(Lawyer.by_id({"id": db_entry.id,
                                                           "name": db_entry.first_name + " " + db_entry.last_name + ", " + db_entry.firm}).first())
            self.userfirm_lawyers.there_is_another = False
            self.userfirm_lawyers.gathered = True

    @classmethod
    def show_all(cls, lawfirms, instance_name=None):
        lawfirm_ids = []
        lawfirm_ids.clear()
        for lawfirm in lawfirms:
            lawfirm_ids.append(lawfirm.lawfirm.id)
        if 'dbcache' not in this_thread.misc:
            this_thread.misc['dbcache'] = {}
        if instance_name:
            listobj = DAList(instance_name, object_type=cls)
        else:
            listobj = DAList(object_type=cls)
            listobj.set_random_instance_name()
        for db_entry in list(
                cls._session.query(cls._model.email).filter(cls._model.lawfirm_id.in_(lawfirm_ids)).order_by(
                        cls._model.email).all()):
            if cls._model.__name__ in this_thread.misc['dbcache'] and db_entry.email in this_thread.misc['dbcache'][
                cls._model.__name__]:
                listobj.append(this_thread.misc['dbcache'][cls._model.__name__][db_entry.email])
            else:
                listobj.append(db_entry.email)
        listobj.gathered = True
        return listobj
        if len(db_entries):
            for db_entry in db_entries:
                db_lawyers.append(Lawyer.by_id({"id": db_entry.id, "email": db_entry.email}))
                db_lawyers.there_is_another = False
                db_lawyers.gathered = True
                return db_lawyers

    def db_find(self):
        try:
            db_lawyers = SCLawyerList(
                "db_lawyers", there_are_any=True, auto_gather=False
            )
            db_lawyers.clear()
            db_entries = self._session.query(LawyerModel).filter(LawyerModel.first_name == self.name.first,
                                                                 LawyerModel.last_name == self.name.last).all()
            if db_entries == None:
                return None
            elif len(db_entries):
                for db_entry in db_entries:
                    db_lawyers.append(Lawyer.by_id({"id": db_entry.id,
                                                    "name": db_entry.first_name + " " + db_entry.last_name + ", " + db_entry.firm}))
            db_lawyers.there_is_another = False
            db_lawyers.gathered = True
            return db_lawyers
        except:
            return None

    def has_lawfirm(self, lawfirm):
        try:
            if not (self.ready() and lawfirm.ready()):
                raise Exception("has_lawfirm: cannot retrieve data")
            db_entry = self._session.query(LawyerModel).filter(LawyerModel.id == self.id,
                                                               LawyerModel.lawfirm_id == lawfirm.id).first()
            if db_entry is None:
                return False
            return True
        except:
            if not (self.ready() and lawfirm.ready()):
                raise Exception("has_lawfirm: cannot retrieve data")
            lawfirm_db_entry = self._session.query(LawfirmModel).filter(LawfirmModel.name == lawfirm.name.text).first()
            result = Lawfirm.by_id(lawfirm_db_entry.id)
            db_entry = self._session.query(LawyerModel).filter(LawyerModel.id == self.id,
                                                               LawyerModel.lawfirm_id == result.id).first()
            if db_entry is None:
                return False
            return True

    def add_lawfirm(self, lawfirm):
        try:
            if not self.has_lawfirm(lawfirm):
                db_entry = LawyerModel(id=self.id, lawfirm_id=lawfirm.id)
                self._session.add(db_entry)
                self._session.commit()
        except:
            if not self.has_lawfirm(lawfirm):
                lawfirm_db_entry = self._session.query(LawfirmModel).filter(
                    LawfirmModel.name == lawfirm.name.text).first()
                result = Lawfirm.by_id(lawfirm_db_entry.id)
                db_entry = LawyerModel(id=self.id, lawfirm_id=result.id)
                self._session.add(db_entry)
                self._session.commit()

    def get_lawfirm(self):
        if not self.ready():
            raise Exception("get_lawfirm: cannot retrieve data")
        log("Lawyer . . . self.get_lawfirm()")
        if self._session.query(LawyerModel).filter(LawyerModel.id == self.id).first():
            log("Lawyer . . . if self._session.query(LawyerModel).filter(LawyerModel.id == self.id).first():")
            db_entry = self._session.query(LawyerModel).filter(LawyerModel.id == self.id).first()
            return Lawfirm.by_id(db_entry.lawfirm_id)

    def del_lawfirm(self, lawfirm):
        if not (self.ready() and lawfirm.ready()):
            raise Exception("del_lawfirm: cannot retrieve data")
        self._session.query(LawyerModel).filter(LawyerModel.id == self.id,
                                                LawyerModel.lawfirm_id == lawfirm.id).delete()
        self._session.commit()


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLawyerList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = CaseLawyer
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_lawyer"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == item.cases.id,
                                                    CaseLawyerModel.lawyer_id == item.lawyer.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")

    def list_all(self):
        results = SCLawyerList("list_all", there_are_any=True, auto_gather=False)
        results.clear()
        for db_entry in self._session.query(LawyerModel).all():
            results.append(Lawyer.by_id(db_entry.id))
        results.there_is_another = False
        results.gathered = True
        return results


class SCName(Name):
    """The name of an Individual."""

    def init(self, *pargs, **kwargs):
        if "uses_parts" not in kwargs:
            self.uses_parts = True
        return super().init(*pargs, **kwargs)

    def defined(self):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return hasattr(self, "text")
        else:
            if not self.uses_parts:
                return super().defined()
            return hasattr(self, "first")

    def familiar(self):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return self.text
        else:
            if not self.uses_parts:
                return self.full()
            return self.first

    def alt(self):
        if self.caption_type.name.text == "personal representative for an estate":
            return "Defendant Estate or " + self.decedent.full().strip()
        else:
            if not self.uses_parts:
                return self.full().strip()
            return self.first

    def full(self, middle="initial", use_suffix=True):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return self.text.strip()
        else:
            if not self.uses_parts:
                return super().full()
            names = [self.first.strip()]
            if hasattr(self, "middle"):
                if middle is False or middle is None:
                    pass
                elif middle == "initial":
                    initial = self.middle_initial().strip()
                    if initial:
                        names.append(initial)
                elif len(self.middle.strip()):
                    names.append(self.middle.strip())
            if hasattr(self, "last") and len(self.last.strip()):
                names.append(self.last.strip())
            else:
                if hasattr(self, "paternal_surname") and len(
                        self.paternal_surname.strip()
                ):
                    names.append(self.paternal_surname.strip())
                if hasattr(self, "maternal_surname") and len(
                        self.maternal_surname.strip()
                ):
                    names.append(self.maternal_surname.strip())
            if hasattr(self, "suffix") and use_suffix and len(self.suffix.strip()):
                names.append(str(", " + self.suffix.strip()))
            return " ".join(names)

    def firstlast(self):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return self.text.strip()
        else:
            if not self.uses_parts:
                return super().firstlast().strip()
            return self.first.strip() + " " + self.last.strip()

    def lastfirst(self):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return self.text.strip()
        else:
            if not self.uses_parts:
                return super().lastfirst().strip()
            output = self.last.strip()
            if hasattr(self, "suffix") and self.suffix and len(self.suffix.strip()):
                output += str(", " + self.suffix.strip())
            output += ", " + self.first.strip()
            if hasattr(self, "middle"):
                initial = self.middle_initial().strip()
                if initial:
                    output += " " + initial
            return output.strip()

    def middle_initial(self, with_period=True):
        if self.caption_type.name.text == "business entity" or self.caption_type.name.text == "governmental entity":
            return ""
        else:
            if len(self.middle.strip()) == 0:
                return ""
            if with_period:
                return self.middle[0].strip() + "."
            else:
                return self.middle[0].strip()


class SCParty(SCPerson, SCIndividual):
    NameClass = SCName

    def init(self, *pargs, **kwargs):

        super().init(*pargs, **kwargs)

    def salute_full(self):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = self.name.text.strip()
        else:
            output = self.salutation() + " " + self.name.full().strip()
        return output

    def summons_address(self, case):
        output = RichText(
            str(
                "TO:"
                + case.fx.TAB
                + self.salute_full()
                + case.fx.NLINE
                + case.fx.TAB
                + self.address.line_one()
                + case.fx.NLINE
                + case.fx.TAB
                + self.address.line_two()
            )
        )
        return output

    def summons_text(self, case):
        output = RichText(str(case.fx.TAB))
        strb(
            output,
            str(
                "IN THE NAME OF THE STATE OF "
                + case.juris.state.name.text.upper()
                + ": "
            ),
        )
        strnull(
            output,
            str(
                "You are hereby required to appear and answer the complaint filed against you in the above-entitled cause within thirty (30) days from the date of service of this summons upon you, and if you fail to so answer, for want thereof, the plaintiff(s) will take judgment against you for the relief prayed for in the complaint on file herein, a copy of which is herewith served upon you. "
                + case.fx.NLINE
                + case.fx.TAB
            ),
        )
        strb(
            output,
            str(
                "NOTICE TO "
                + self.party_type.name.text.upper()
                + ":	READ THESE PAPERS CAREFULLY!"
            ),
        )
        if self.party_type.name.text == "Defendant":
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "You must “appear” in this case or the other side will win automatically. To “appear” you must file with the court a legal document called a “motion” or “answer.” The “motion” or “answer” must be given to the court clerk or administrator within thirty (30) days along with the required filing fee. It must be in the proper form and have proof of service on the plaintiff’s attorney or, if the plaintiff does not have an attorney, proof of service on the plaintiff."
                ),
            )
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "If you have questions, you should see an attorney immediately. If you need help in finding an attorney, you may call the "
                    + case.juris.state.name.text
                    + " State Bar’s Lawyer Referral Service online at www.oregonstatebar.org or by calling (503) 684-3763 (in the Portland metropolitan area) or toll-free elsewhere in Oregon at (800) 452-7636."
                ),
            )
        else:
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "You must “appear” to protect your rights in this matter. To “appear” you must file with the court a legal document called a “motion,” a “reply” to a counterclaim, or an“answer” to a cross-claim. The “motion,” “reply,” or “answer” must be given to the court clerk or administrator within 30 days along with the required filing fee. It must be in proper form and have proof of service on the defendant’s attorney or, if the defendant does not have an attorney, proof of service on the defendant."
                ),
            )
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "If you have questions, you should see an attorney immediately. If you need help in finding an attorney, you may contact the Oregon State Bar’s Lawyer Referral Service online at http://www.oregonstatebar.org or by calling (503) 684-3763 (in the Portland metropolitan area) or toll-free elsewhere in Oregon at (800) 452-7636."
                ),
            )
        return output

    def is_defendant(self):
        if self.party_type.name.text == "Defendant":
            return True
        else:
            return False

    def pleading(self):
        if self.is_defendant(self):
            return "Answer"
        else:
            return "Complaint"

    def pname(self):
        return str(self.party_type.name.text + " " + self.name.full().strip())

    def salute_possessive(self):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            return possessify(self.name.text.strip(), "")
        else:
            return possessify(self.salutation(with_name=True).strip(), "")

    def salute(self):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            return self.name.text.strip()
        else:
            return self.salutation(with_name=True).strip()

    @property
    def complete_party(self):
        self.party_type
        log("BEFORE CAPTION_TYPE COMPLETE")
        self.name.caption_type.name.text
        log("AFTER CAPTION_TYPE COMPLETE   2")
        self.caption_type.name.text = self.name.caption_type.name.text
        # self.caption_type.db_read()
        log("AFTER CAPTION_TYPE COMPLETE   3")
        self.name.full()
        log("AFTER CAPTION_TYPE COMPLETE")
        # self.name.full()
        # self.the_one
        log("BEFORE CAPTION_text COMPLETE")
        self.caption_text
        log("AFTER CAPTION_text COMPLETE")
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            if self.name.caption_type.name.text == "business entity":
                self.business_entity.name.text
        else:
            self.gender.name.text
        if self.name.caption_type.name.text == "trustee":
            self.trust
        elif self.name.caption_type.name.text == "doing business as":
            self.dba
        log("AFTER self.firms COMPLETE")

    @property
    def is_party(self):
        return True

    @property
    def is_lawyer(self):
        return False

    @property
    def is_lawfirm(self):
        return False

    @property
    def clio_type(self):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = "Company"
        else:
            output = "Person"
        return output

    def __str__(self):
        return str(self.name.full().strip())

    def last_name(self):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            return self.name.text.strip()
        else:
            return self.name.last.strip()

    def pronoun_objective(self, **kwargs):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = word("it", **kwargs)
        else:
            if self == this_thread.global_vars.user:
                output = word("you", **kwargs)
            if self.gender.name.text == "female":
                output = word("her", **kwargs)
            elif self.gender.name.text == "other":
                output = word("them", **kwargs)
            else:
                output = word("him", **kwargs)
        if "capitalize" in kwargs and kwargs["capitalize"]:
            return capitalize(output)
        else:
            return output

    def pronoun_possessive(self, target, **kwargs):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = its(target, **kwargs)
        else:
            if self == this_thread.global_vars.user and (
                    "thirdperson" not in kwargs or not kwargs["thirdperson"]
            ):
                output = your(target, **kwargs)
            elif self.gender.name.text == "female":
                output = her(target, **kwargs)
            elif self.gender.name.text == "other":
                output = their(target, **kwargs)
            else:
                output = his(target, **kwargs)
        return output

    def possessive_pronoun(self, **kwargs):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = word("its", **kwargs)
        else:
            if self == this_thread.global_vars.user and (
                    "thirdperson" not in kwargs or not kwargs["thirdperson"]
            ):
                output = word("your", **kwargs)
            elif self.gender.name.text == "female":
                output = word("her", **kwargs)
            elif self.gender.name.text == "other":
                output = word("their", **kwargs)
            else:
                output = word("his", **kwargs)
        return output

    def pronoun_subjective(self, **kwargs):
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            output = word("it", **kwargs)
        else:
            if self == this_thread.global_vars.user and (
                    "thirdperson" not in kwargs or not kwargs["thirdperson"]
            ):
                output = word("you", **kwargs)
            elif self.gender.name.text == "female":
                output = word("she", **kwargs)
            elif self.gender.name.text == "other":
                output = word("they", **kwargs)
            else:
                output = word("he", **kwargs)
        return output

    def life_expectancy(self):
        age = self.age_in_years()
        if age < 0:
            raise Exception(
                "Error: cannot determine life expectancy of someone born in the future."
            )
        if self.race == "black":
            if self.gender.name.text == "female":
                xlsx_file = "Table12.xls"
            else:
                xlsx_file = "Table11.xls"
        elif self.race == "hispanic":
            if self.gender.name.text == "female":
                xlsx_file = "Table06.xls"
            else:
                xlsx_file = "Table05.xls"
        elif self.race == "white":
            if self.gender.name.text == "female":
                xlsx_file = "Table09.xls"
            else:
                xlsx_file = "Table08.xls"
        elif self.gender.name.text == "female":
            xlsx_file = "Table03.xls"
        elif self.gender.name.text == "male":
            xlsx_file = "Table02.xls"
        elif self.gender.name.text == "other":
            xlsx_file = "Table01.xls"
        else:
            raise Exception(
                "Cannot determine life expectancy of someone without at least a defined gender and birthdate."
            )
        path, mimetype = path_and_mimetype("data/sources/" + xlsx_file)
        wb = xlrd.open_workbook(path)
        sheet = wb.sheet_by_index(0)
        row = 0
        headers = []
        while row < sheet.nrows:
            first_cell = sheet.cell_value(row, 0)
            if first_cell.startswith("Table"):
                row += 1
                continue
            if first_cell.startswith("Age"):
                headers = get_row(sheet, row)
            if first_cell.startswith("NOTES") or first_cell.startswith("SOURCE"):
                break
            m = re.search(r"([0-9]+)[^0-9]+([0-9]+)", first_cell)
            if m:
                min_age = float(m.group(1))
                max_age = float(m.group(2))
                if age >= min_age and age <= max_age:
                    vals = get_row(sheet, row)
                    if len(headers) != len(vals):
                        raise Exception(
                            "File " + xlsx_file + " is in the wrong format:"
                        )
                    return {
                        headers[indexno]: vals[indexno] for indexno in range(len(vals))
                    }
            else:
                m = re.search(r"([0-9]+) and over", first_cell)
                if m:
                    min_age = float(m.group(1))
                    if age >= min_age:
                        vals = get_row(sheet, row)
                        if len(headers) != len(vals):
                            raise Exception(
                                "File " + xlsx_file + " is in the wrong format"
                            )
                        return {
                            headers[indexno]: vals[indexno]
                            for indexno in range(len(vals))
                        }
            row += 1
        raise Exception(
            "Unable to determine life expectancy information where age was " + str(age)
        )


class Party(SCParty, SQLObject):
    _model = PartyModel
    _session = DBSession
    _required = ['party_type_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'caption_type_id':
            return self.name.caption_type.id
        elif column == 'city':
            return self.address.city
        elif column == 'state':
            return self.address.state
        elif column == 'business_entity_id':
            return self.business_entity.id
        elif column == 'state_of_inc_id':
            return self.state_of_inc.id
        elif column == 'gender_id':
            return self.gender.id
        elif column == 'name':
            if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
                return self.name.text.strip()
            else:
                return str(self.name.first + self.name.last)
        elif column == 'first_name':
            return self.name.first.strip()
        elif column == 'middle_name':
            return self.name.middle.strip()
        elif column == 'last_name':
            return self.name.last.strip()
        elif column == 'address':
            return self.address.address.strip()
        elif column == 'unit':
            return self.address.unit.strip()
        elif column == 'zip_code':
            return self.address.zip.strip()
        elif column == "designation":
            return self.designation.strip()
        elif column == "trust":
            return self.trust
        elif column == "party_type_id":
            return self.party_type.id
        elif column == "caption_text":
            return self.caption_text
        elif column == "dba":
            return self.dba
        elif column == "is_3dp":
            return self.is_3dp
        elif column == "is_cross_p":
            return self.is_cross_p
        elif column == "is_cross_d":
            return self.is_cross_d
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'caption_type_id':
            self.name.caption_type = CaptionType.by_id(value)
        elif column == 'city':
            self.address.city = value
        elif column == 'state':
            self.address.state = value
        elif column == 'business_entity_id':
            self.business_entity = BusinessEntity.by_id(value)
        elif column == 'state_of_inc_id':
            self.state_of_inc = State.by_id(value)
        elif column == 'gender_id':
            self.gender = Gender.by_id(value)
        elif column == 'name':
            if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
                self.name.text = value
        elif column == 'first_name':
            self.name.first = value
        elif column == 'middle_name':
            self.name.middle = value
        elif column == 'last_name':
            self.name.last = value
        elif column == 'address':
            self.address.address = value
        elif column == 'unit':
            self.address.unit = value
        elif column == 'zip_code':
            self.address.zip = value
        elif column == "designation":
            self.designation = value
        elif column == 'caption_text':
            self.caption_text = value
        elif column == "trust":
            self.trust = value
        elif column == "party_type_id":
            self.party_type = Designation.by_id(value)
        elif column == "dba":
            self.dba = value
        elif column == "is_3dp":
            self.is_3dp = value
        elif column == "is_cross_p":
            self.is_cross_p = value
        elif column == "is_cross_d":
            self.is_cross_d = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == "caption_type_id":
            del self.name.caption_type.name.text
        elif column == "city":
            del self.address.city
        elif column == "state":
            del self.address.state
        elif column == "business_entity_id":
            del self.business_entity.name.text
        elif column == "state_of_inc_id":
            del self.state_of_inc
        elif column == "gender_id":
            del self.gender.name.text
        elif column == 'name':
            if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
                del self.name.text
        elif column == 'first_name':
            del self.name.first
        elif column == 'middle_name':
            del self.name.middle
        elif column == 'last_name':
            del self.name.last
        elif column == 'address':
            del self.address.address
        elif column == 'unit':
            del self.address.unit
        elif column == 'zip_code':
            del self.address.zip
        elif column == "designation":
            del self.designation
        elif column == 'caption_text':
            del self.caption_text
        elif column == "trust":
            del self.trust
        elif column == "party_type":
            del self.party_type
        elif column == "dba":
            del self.dba
        elif column == "is_3dp":
            del self.is_3dp
        elif column == "is_cross_p":
            del self.is_cross_p
        elif column == "is_cross_d":
            del self.is_cross_d
        else:
            raise Exception("Invalid column " + column)

    def get_lawfirms(self, case):
        q = Session.query(Lawfirm, ).join(CaseLawfirm).join(CaseParty).filter(CaseLawfirm.cases_id == case.id).filter(
            CaseParty.party_id == self.id, ).all()
        return q

    def db_find_all(self):
        possibilities = []
        possibilities.clear()
        if self.name.caption_type.name.text == "business entity" or self.name.caption_type.name.text == "governmental entity":
            for entry in self._session.query(PartyModel).filter(PartyModel.name == self.name.text).all():
                possibilities.append({"id": entry.id, "name": entry.name})
        else:
            for entry in self._session.query(PartyModel).filter(PartyModel.first_name == self.name.first,
                                                                PartyModel.last_name == self.name.last).all():
                possibilities.append({"id": entry.id, "name": entry.first_name + " " + entry.last_name})
                # possibilities.append(entry.first_name + " " + entry.last_name)
        return possibilities


class Designation(Thing, SQLObject):
    _model = DesignationModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)

class Claim(Thing, SQLObject):
    _model = ClaimModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class BusinessEntity(Thing, SQLObject):
    _model = BusinessEntityModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class CaptionType(Thing, SQLObject):
    _model = CaptionTypeModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class Type(Thing, SQLObject):
    _model = TypeModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class Gender(Thing, SQLObject):
    _model = GenderModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class SCCity(DAObject, SQLObject):
    _model = CityModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.strip()

    def db_set(self, column, value):
        if column == 'name':
            self = value

    def db_null(self, column):
        if column == 'name':
            del self


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCPartyList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = CaseParty
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_party"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(CasePartyModel).filter(CasePartyModel.cases_id == item.cases.id,
                                                   CasePartyModel.party_id == item.party.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")

    def pleading(self):
        return self[0].pleading()

    def salute(self):
        return comma_and_list(self)

    def salute_possessive(self):
        return possessify(comma_and_list(self), "")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCClaimList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = PartyClaim
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_partyclaim"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(PartyClaimModel).filter(PartyClaimModel.claim_id == item.claim.id,
                                                    PartyClaimModel.party_id == item.party.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCDesignationList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = PartyDesignation
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_partydesignation"
        super().init(*pargs, **kwargs)

    def hook_on_remove(self, item):
        log("hook 2")
        item._session.query(PartyDesignationModel).filter(PartyDesignationModel.designation_id == item.designation.id,
                                                          PartyDesignationModel.party_id == item.party.id).delete()
        log("hook 3")
        item._session.commit()
        log("hook 4")


def bluebook_ordinal(string_ordinal):
    if string_ordinal.lower() == "first":
        return "1st"
    elif string_ordinal.lower() == "second":
        return "2d"
    elif string_ordinal.lower() == "third":
        return "3d"
    elif string_ordinal.lower() == "fourth":
        return "4th"
    elif string_ordinal.lower() == "fifth":
        return "5th"
    elif string_ordinal.lower() == "sixth":
        return "6th"
    elif string_ordinal.lower() == "seventh":
        return "7th"
    elif string_ordinal.lower() == "eighth":
        return "8th"
    elif string_ordinal.lower() == "ninth":
        return "9th"


def jurisdiction_docx():
    document = Document()
    table = document.add_table(rows=0, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "STATE OF WISCONSIN"
    cells[1].text = "DISTRICT COURT"
    cells[2].text = "WASCO COUNTY"
    return document


def the_states():
    a = states_list()
    b = ['DC', 'GU', 'CM', 'VI', 'PR', 'UM', 'MP', 'CZ', 'AS']
    return sorted([x for x in a if x not in b])


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
CDC_VITAL_STATS = (
    "Centers for Disease Control, 69 *National Vital Statistics Report* 12 (2020)"
)
CONTINGENCY_FEE_NO_TRIAL = (
    0.3333333333333333333333333333333333333333333333333333333333333333333333333333333333
)
CONTINGENCY_FEE_AFTER_TRIAL = 0.4
APOSTROPHE = "&#8217;"

# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
ORDINANCE = {1: "First", 11: "Eleventh", 21: "Twenty-First", 31: "Thirty-First", 41: "Forty-First", 2: "Second",
             12: "Twelfth", 22: "Twenty-Second", 32: "Thirty-Second", 42: "Forty-Second", 3: "Third", 13: "Thirteenth",
             23: "Twenty-Third", 33: "Thirty-Third", 43: "Forty-Third", 4: "Fourth", 14: "Fourteenth",
             24: "Twenty-Fourth", 34: "Thirty-Fourth", 44: "Forty-Fourth", 5: "Fifth", 15: "Fifteenth",
             25: "Twenty-Fifth", 35: "Thirty-Fifth", 45: "Forty-Fifth", 6: "Sixth", 16: "Sixteenth", 26: "Twenty-Sixth",
             36: "Thirty-Sixth", 46: "Forty-Sixth", 7: "Seventh", 17: "Seventeenth", 27: "Twenty-Seventh",
             37: "Thirty-Seventh", 47: "Forty-Seventh", 8: "Eighth", 18: "Eighteenth", 28: "Twenty-Eighth",
             38: "Thirty-Eighth", 48: "Forty-Eighth", 9: "Ninth", 19: "Nineteenth", 29: "Twenty-Ninth",
             39: "Thirty-Ninth", 49: "Forty-Ninth", 10: "Tenth", 20: "Twentieth", 30: "Thirtieth", 40: "Fortieth",
             50: "Fiftieth", 51: "Fifty-First", 61: "Sixty-First", 71: "Seventy-First", 81: "Eighty-First",
             91: "Ninety-First", 52: "Fifty-Second", 62: "Sixty-Second", 72: "Seventy-Second", 82: "Eighty-Second",
             92: "Ninety-Second", 53: "Fifty-Third", 63: "Sixty-Third", 73: "Seventy-Third", 83: "Eighty-Third",
             93: "Ninety-Third", 54: "Fifty-Fourth", 64: "Sixty-Fourth", 74: "Seventy-Fourth", 84: "Eighty-Fourth",
             94: "Ninety-Fourth", 55: "Fifty-Fifth", 65: "Sixty-Fifth", 75: "Seventy-Fifth", 85: "Eighty-Fifth",
             95: "Ninety-Fifth", 56: "Fifty-Sixth", 66: "Sixty-Sixth", 76: "Seventy-Sixth", 86: "Eighty-Sixth",
             96: "Ninety-Sixth", 57: "Fifty-Seventh", 67: "Sixty-Seventh", 77: "Seventy-Seventh", 87: "Eighty-Seventh",
             97: "Ninety-Seventh", 58: "Fifty-Eighth", 68: "Sixty-Eighth", 78: "Seventy-Eighth", 88: "Eighty-Eighth",
             98: "Ninety-Eighth", 59: "Fifty-Ninth", 69: "Sixty-Ninth", 79: "Seventy-Ninth", 89: "Eighty-Ninth",
             99: "Ninety-Ninth", 60: "Sixtieth", 70: "Seventieth", 80: "Eightieth", 90: "Ninetieth", 100: "Hundredth"}


def format_date_range(dates):
    output = ""
    output += str(format_date(dates[0], format='MMMM'))
    output += " "
    output += str(format_date(dates[0], format='d'))
    if format_date(dates[0], format='YYYY') != format_date(dates[-1], format='YYYY'):
        output += format_date(dates[0], format='YYYY')
    output += "[ENDASH]"
    if format_date(dates[0], format='MMMM') == format_date(dates[-1], format='MMMM') and format_date(dates[0],
                                                                                                     format='YYYY') == format_date(
        dates[-1], format='YYYY'):
        output += str(format_date(dates[-1], format='d'))
        output += ", "
    else:
        output += str(format_date(dates[-1], format='MMMM'))
        output += str(format_date(dates[-1], format='d'))
        output += ", "
    output += str(format_date(dates[-1], format='YYYY'))
    return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def parens(text):
    output = ""
    output += "("
    output += str(text)
    output += ")"
    return output


def bld(run):
    run.bold = True
    run.italic = False
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def uline(run):
    run.bold = False
    run.italic = False
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def itx(run):
    run.bold = False
    run.italic = True
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def uline_itx(run):
    run.bold = False
    run.italic = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def bld_itx(run):
    run.bold = True
    run.italic = True
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def bld_uline(run):
    run.bold = True
    run.italic = False
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def bld_uline_itx(run):
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def nullrun(run):
    run.bold = False
    run.italic = False
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def no_indent_body(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Inches(0)
    paragraph_format.right_indent = Inches(0.06)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(24.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def h1strb(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True, style="Heading 1")


def h1stri(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, style="Heading 1")


def h1strbu(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True, underline=True, style="Heading 1")


def h1strbi(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, bold=True, style="Heading 1")


def h1strbiu(list_o_rt, rt):
    return list_o_rt.add(
        str(rt), italic=True, bold=True, underline=True, style="Heading 1"
    )


def h1strnull(list_o_rt, rt):
    return list_o_rt.add(rt, style="Normal")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def h2strb(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True, style="Heading 2")


def h2stri(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, style="Heading 2")


def h2strbu(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True, underline=True, style="Heading 2")


def h2strbi(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, bold=True, style="Heading 2")


def h2strbiu(list_o_rt, rt):
    return list_o_rt.add(
        str(rt), italic=True, bold=True, underline=True, style="Heading 2"
    )


def h2strnull(list_o_rt, rt):
    return list_o_rt.add(rt, style="Normal")


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def strb(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True)


def stri(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True)


def strbu(list_o_rt, rt):
    return list_o_rt.add(rt, bold=True, underline=True)


def strbi(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, bold=True)


def strbiu(list_o_rt, rt):
    return list_o_rt.add(str(rt), italic=True, bold=True, underline=True)


def strnull(list_o_rt, rt):
    return list_o_rt.add(rt)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def flatten(t):
    return [item for sublist in t for item in sublist]


class SCJurisdiction(DAObject):
    def init(self, *pargs, **kwargs):
        # self.initializeAttribute("jurisdiction", Jurisdiction)
        # self.initializeAttribute("state", State)
        # self.initializeAttribute("county", County)
        # self.initializeAttribute("district", District)
        # self.initializeAttribute("division", Division)
        super().init(*pargs, **kwargs)

    def civil_procedure(self):
        dictionary = CIV_PRO[str(self.state.name.text)]
        if self.jurisdiction.name.text == "federal":
            dictionary["official title"] = "Federal Rules of Civil Procedure"
            dictionary["abbreviation"] = "FRCP"
            dictionary["Duty to Disclose; General Provisions Governing Discovery"] = "26"
            dictionary["Depositions to Perpetuate Testimony"] = "27"
            dictionary["Depositions by Oral Examination"] = "30"
            dictionary["Notice or Subpoena Directed to an Organization"] = "30(b)(6)"
            dictionary["Interrogatories to Parties"] = "33"
            dictionary[
                "Producing Documents, Electronically Stored Information, and Tangible Things, or Entering onto Land, for Inspection and Other Purposes"] = "34"
            dictionary["Physical and Mental Examinations"] = "35"
            dictionary["Requests for Admission"] = "36"
        return dictionary

    def trial_court(self):
        if self.jurisdiction.name.text == "federal":
            return self.civil_procedure()["federal"]
        else:
            return self.civil_procedure()["default trial court"]

    def trial_court_dict(self):
        if self.jurisdiction.name.text != "federal":
            return self.civil_procedure()[self.trial_court()]
        else:
            return self.civil_procedure()["federal"]

    def trial_court_unit_titles(self):
        output = []
        output.clear()
        # output.append(self.county.name.text)
        if self.trial_court_has_multiple_unit_types():
            output.append(self.trial_court_dict()["unified by key value"][0])
            output.append(self.trial_court_dict()["unified by key value"][1])
        else:
            output.append(self.trial_court_dict()["divided by"])
        return output

    def county_level(self):
        return self.trial_court_unit_titles()[0]

    def district_level(self):
        if self.jurisdiction.name.text == "federal":
            return "district"
        else:
            if self.trial_court_has_multiple_unit_types():
                return self.trial_court_unit_titles()[1]
            else:
                return None

    def division_level(self):
        if self.jurisdiction.name.text == "federal":
            return "division"
        else:
            if self.trial_court_has_divisions():
                return self.trial_court_dict()["division level"]
            else:
                return None

    def trial_court_has_multiple_unit_types(self):
        if "unified by" in self.trial_court_dict().keys():
            return True
        else:
            return False

    def trial_court_unit_dictionary(self):
        if self.trial_court_has_multiple_unit_types():
            return self.trial_court_dict()["unified by"]
        else:
            return None

    def trial_court_unit_keys(self):
        if self.trial_court_has_multiple_unit_types():
            output = []
            output.clear()
            for item in self.trial_court_unit_dictionary().keys():
                output.append(item)
        else:
            output = self.trial_court_dict()[self.trial_court_dict()["divided by"]]
        return output

    def trial_court_unit_values(self):
        if self.jurisdiction.name.text == "federal":
            if len(self.trial_court().keys()) == 1:
                self.district.name.text = self.trial_court().keys()[0]
            else:
                return self.trial_court().keys()
        else:
            if self.trial_court_has_multiple_unit_types():
                return self.trial_court_unit_dictionary().values()
            else:
                return self.trial_court_dict()[self.trial_court_dict()["divided by"]]

    def trial_court_units(self):
        output = []
        output.clear()
        output.append(self.county.name.text)
        if self.trial_court_has_multiple_unit_types():
            if type(self.trial_court_unit_dictionary()[self.county.name.text]) == list:
                output.append(self.trial_court_unit_dictionary()[self.county.name.text][0])
            else:
                output.append(self.trial_court_unit_dictionary()[self.county.name.text])
        if self.trial_court_has_divisions():
            output.append(self.division.name.text)
        return output

    def trial_court_has_divisions(self):
        if "divisions" in self.trial_court_dict().keys():
            return True
        else:
            return False

    def trial_court_divisions(self):
        if self.jurisdiction.name.text == "federal":
            if len(self.trial_court()[str(self.district.name.text)]) == 1:
                self.district.name.text = self.trial_court()[str(self.district.name.text)][0]
            else:
                return self.trial_court()[str(self.district.name.text)]
        else:
            if self.trial_court_has_divisions():
                return self.trial_court_dict()["divisions"]
            else:
                return None

    def case_no_in_header(self):
        if "case no in header" in self.civil_procedure().keys():
            return True
        else:
            return False


def special_case_no_style(self):
    if "case no" in self.civil_procedure().keys():
        return True
    else:
        return False


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCRequests(Thing):
    def impetra_rrfp_objections(self):
        log("impetra_rrfp_objections")
        levl = 0
        output = []
        if self.overly_broad:
            output.append(item_label(levl, 4) + " Overly Broad & Unduly Burdensome")
            levl += 1
        if self.overbroad_time:
            output.append(item_label(levl, 4) + " Overbroad as to Time")
            levl += 1
        if self.not_reasonably_calculated:
            output.append(
                item_label(levl, 4)
                + " Not Reasonably Calculated to Lead to Discoverable Evidence"
            )
            levl += 1
        if self.attorney_client:
            output.append(item_label(levl, 4) + " Attorney Client Privelege")
            levl += 1
        if self.same_bodypart:
            output.append(item_label(levl, 4) + " Not Same or Similar Bodypart")
            levl += 1
        if self.vauge_ambiguous:
            output.append(item_label(levl, 4) + " Vague and Ambiguous")
            levl += 1
        if self.public_record:
            output.append(item_label(levl, 4) + " Public Record")
            levl += 1

        if self.defendant_possesses:
            output.append(item_label(levl, 4) + " Defendant Already Possesses")
            levl += 1

        if self.interrogatory:
            output.append(item_label(levl, 4) + " Impermissible Interrogatory")
            levl += 1
        if self.expert:
            output.append(item_label(levl, 4) + " Impermissible Expert Discovery")
            levl += 1
        if self.biz_docs:
            output.append(item_label(levl, 4) + " Confidential Business Documents")
            levl += 1
        if self.text_response:
            output.append(item_label(levl, 4) + str(" " + self.text_response))
            levl += 1
        if self.responsive:
            output.append(item_label(levl, 4) + str(" " + self.responsive + str(
                " " + str(self.see_rfp) if self.responsive == 'see response to RFP No.' else "")))
            levl += 1
        listout = ""
        for item in add_separators(output):
            listout += str(item) + "[BR][BR]"
        return listout

    def impetra_rrog_objections(self, recipients, clientlist):
        levl = 0
        output = []
        if self.overly_broad:
            output.append(item_label(levl, 4) + " Overly Broad & Unduly Burdensome")
            levl += 1
        if self.not_reasonably_calculated:
            output.append(
                item_label(levl, 4)
                + " Not Reasonably Calculated to Lead to Discoverable Evidence"
            )
            levl += 1
        if self.attorney_client:
            output.append(item_label(levl, 4) + " Attorney Client Privelege")
            levl += 1
        if self.cumulative:
            output.append(
                item_label(levl, 4) + " Unreasonably Cumulative and Duplicative"
            )
            levl += 1
        if self.same_bodypart:
            output.append(item_label(levl, 4) + " Physician-Patient Privilege")
            levl += 1
        if self.vauge_ambiguous:
            output.append(item_label(levl, 4) + " Vague and Ambiguous")
            levl += 1
        if self.public_record:
            output.append(item_label(levl, 4) + " Public Record")
            levl += 1
        if self.defendant_possesses:
            output.append(
                item_label(levl, 4) + " Info Already Known by " + recipients.asnoun()
            )
            levl += 1
        if self.not_medically_stationary:
            output.append(
                item_label(levl, 4)
                + clientlist.asnoun()
                + " Isn't Medically Stationary "
            )
            levl += 1
        if self.seeks_docs:
            output.append(
                item_label(levl, 4)
                + " Seeks Production of Documents or Requests a Complete Recital of the Contents of Documents”"
            )
            levl += 1
        if self.expert:
            output.append(item_label(levl, 4) + " Impermissible Expert Discovery")
            levl += 1
        if self.legal_conclusion:
            output.append(item_label(levl, 4) + " Calls for a Pure Legal Conclusion")
            levl += 1
        if self.annoying_embarrassing_oppressive:
            output.append(
                item_label(levl, 4) + " Annoying, Embarrassing, or Oppressive"
            )
            levl += 1
        if self.repetitive_duplicative:
            output.append(item_label(levl, 4) + " Repetitive or Duplicative")
            levl += 1
        if self.unknown_to_me:
            output.append(
                item_label(levl, 4)
                + " Not Within the Knowledge, Possession, Custody or Control of the Answering Party"
            )
            levl += 1
        if self.outside_scope_of_rules:
            output.append(
                item_label(levl, 4)
                + " Beyond the Scope/Requirements of the Rules of Civil Procedure "
            )
            levl += 1
        if self.premature:
            output.append(item_label(levl, 4) + " Premature")
            levl += 1
        if self.not_limited_in_time:
            output.append(
                item_label(levl, 4)
                + " Not Limited in Time to the Years of the Alleged Occurrence"
            )
            levl += 1
        if self.created_after_filing_of_lawsuit:
            output.append(
                item_label(levl, 4)
                + " Information Created After the Commencement of the Lawsuit"
            )
            levl += 1
        if self.seeks_witness_statement:
            output.append(
                item_label(levl, 4)
                + " Seeks Witness Statement Without the Requisite Showing of Substantial Need and Undue Hardship"
            )
            levl += 1
        if self.privacy_rights_of_others:
            output.append(item_label(levl, 4) + " Privacy Rights of Individuals")
            levl += 1
        if self.confidentiality_agreement:
            output.append(item_label(levl, 4) + " Confidentiality Agreement")
            levl += 1
        if self.protective_order:
            output.append(item_label(levl, 4) + " Court Order Restricting Disclosure")
            levl += 1
        if self.trade_secret:
            output.append(
                item_label(levl, 4)
                + " Confidential Commercial Information, Trade Secrets, or Proprietary Information"
            )
            levl += 1
        if self.not_relevant:
            output.append(item_label(levl, 4) + " Not Relevant")
            levl += 1
        listout = ""
        for item in add_separators(output):
            listout += str(item) + "[BR][BR]"
        return listout

    def impetra_rrfas_objections(self):
        levl = 0
        output = []
        # if self.intro:
        #  output.append(item_label(levl, 4) + str(self.other_other))
        #  levl +=1
        if self.text_response:
            output.append(item_label(levl, 4) + str(self.text_response))
            levl += 1
        if self.undefined_subjective:
            output.append(item_label(levl, 4) + " Undefined & Subjective Term or Phrase")
            levl += 1
        if self.reasonable_inquiry:
            output.append(item_label(levl, 4) + " Reasonable Inquiry")
            levl += 1
        if self.vauge_ambiguous:
            output.append(item_label(levl, 4) + " Vague and Ambiguous")
            levl += 1
        if self.expert:
            output.append(item_label(levl, 4) + " Impermissible Expert Discovery")
            levl += 1
        if self.attorney_client:
            output.append(item_label(levl, 4) + " Attorney Client Privelege")
            levl += 1
        listout = ""
        for item in add_separators(output):
            listout += str(item) + "[BR][BR]"
        return listout


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCCase(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("juris", SCJurisdiction)
        self.initializeAttribute("ordinal", DAObject)
        self.initializeAttribute("draft", DAObject)
        # self.initializeAttribute("parties", SCPartyList)
        # self.initializeAttribute("cross_parties", SCPartyList)
        # self.initializeAttribute("lawfirms", SCLawfirmList)
        # self.initializeAttribute("lawyers", SCLawyerList)
        self.initializeAttribute("economics", SCEconomics)
        self.initializeAttribute("noneconomics", SCNonEconomics)
        self.initializeAttribute("vehicles", SCVehicleList)
        self.initializeAttribute("fx", Color)
        self.initializeAttribute("MVA", SCMVA)
        # self.initializeAttribute("cross_claims", SCCrossClaims)
        self.initializeAttribute("UCJI", SCJuryInstruction)
        self.initializeAttribute("lien_creditors", SCLienCreditors)
        self.initializeAttribute("userfirm", DAGlobal, key="userfirm")
        self.user_title = False
        super().init(*pargs, **kwargs)

    def designations(self):
        parties = ["Intervenor", "Third-Party Defendant"]
        parties.clear()
        if any(party.party.party_type.name.text == "Plaintiff" for party in self.parties.elements) and not any(
                party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
            parties.append('Plaintiff')
            if not any(party.party.party_type.name.text == "Respondent" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Defendant')
            if not any(party.party.party_type.name.text == "Defendant" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Respondent')
        elif any(party.party.party_type.name.text == "Claimant" for party in self.parties.elements) and not any(
                party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
            parties.append('Claimant')
            if not any(party.party.party_type.name.text == "Respondent" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Defendant')
            if not any(party.party.party_type.name.text == "Defendant" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Respondent')
        elif any(party.party.party_type.name.text == "Appellant" for party in self.parties.elements) and not any(
                party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
            parties.append('Appellant')
            parties.append("Respondent")
        elif any(party.party.party_type.name.text == "Petitioner" for party in self.parties.elements) and not any(
                party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
            parties.append('Petitioner')
            if not any(party.party.party_type.name.text == "Respondent" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Defendant')
            if not any(party.party.party_type.name.text == "Defendant" for party in self.parties.elements) and not any(
                    party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
                parties.append('Respondent')
        elif any(party.party.party_type.name.text == "Obligee" for party in self.parties.elements) and not any(
                party.party.party_type.name.text == "Deceased" for party in self.parties.elements):
            parties.append('Obligee')
            parties.append("Obligor")
        else:
            parties.append('Plaintiff')
            parties.append('Defendant')
            parties.append('Respondent')
            parties.append('Petitioner')
            parties.append('Claimant')
            parties.append('Appellant')
            parties.append('Obligor')
            parties.append('Obligee')
            parties.append('Deceased')
        return parties

    def other_parties(self):
        parties = SCPartyList(
            "other_parties", there_are_any=True, auto_gather=False
        )
        parties.clear()
        for group in self.certificate_of_service_list():
            for party in group[2]:
                parties.append(party)
        parties.there_is_another = False
        parties.gathered = True
        return parties

    def other_parties_lawyers(self):
        attorneys = SCLawyerList(
            "other_parties_lawyers", there_are_any=True, auto_gather=False
        )
        attorneys.clear()
        for group in self.certificate_of_service_list():
            for firm in group[0]:
                for attorney in self.attorneys_who_work_at(firm):
                    attorneys.append(attorney)
        attorneys.there_is_another = False
        attorneys.gathered = True
        return attorneys

    def current_trial_date(self):
        dates = [self.current_trial_start_date, self.current_trial_end_date]
        output = format_date_range(dates)
        return output

    def new_trial_date(self):
        dates = [self.new_trial_start_date, self.new_trial_start_date]
        output = format_date_range(dates)
        return output

    def prior_trial_settings(self):
        output = SCList(
            "prior_trial_settingss", there_are_any=True, auto_gather=False
        )
        output.clear()
        for trial in self.prior_trial_dates:
            dates = [trial.start, trial.end]
            prior_trial = format_date_range(dates)
            output.append(prior_trial)
        output.there_is_another = False
        output.gathered = True
        return comma_and_list(output, comma_string='; ')

    def certificate_of_service_list(self):
        output = SCList(
            "output2", there_are_any=True, auto_gather=False
        )
        output.clear()
        everyfirm = SCList(
            "everyfirm", there_are_any=True, auto_gather=False
        )
        everyfirm.clear()
        for firm in self.lawfirms:
            lawfirms = SCList(
                "lawfirms", there_are_any=True, auto_gather=False
            )
            lawfirms.clear()
            attorneys = SCList(
                "attorneys", there_are_any=True, auto_gather=False
            )
            attorneys.clear()
            if not firm.lawfirm.is_userfirm and firm not in self.cofirms() and firm.lawfirm not in everyfirm.elements:
                clients = SCList(
                    "clients", there_are_any=True, auto_gather=False
                )
                clients.clear()
                lawfirms.append(firm.lawfirm)
                everyfirm.append(firm.lawfirm)
                for client in self.parties.client_of(firm.lawfirm):
                    if client not in clients.elements:
                        clients.append(client)
                clients.there_is_another = False
                clients.gathered = True
                for otherfirm in self.lawfirms:
                    if not otherfirm.lawfirm.is_userfirm and otherfirm.lawfirm not in self.cofirms() and otherfirm.lawfirm not in lawfirms.elements and otherfirm != firm and otherfirm.lawfirm not in everyfirm.elements:
                        for otherclient in self.parties.client_of(otherfirm.lawfirm):
                            if otherclient in clients and otherfirm.lawfirm not in lawfirms.elements:
                                lawfirms.append(otherfirm.lawfirm)
                                everyfirm.append(otherfirm.lawfirm)
                lawfirms.there_is_another = False
                lawfirms.gathered = True
                for allfirms in lawfirms:
                    for attorney in self.attorneys_who_work_at(allfirms):
                        attorneys.append(attorney)
                attorneys.there_is_another = False
                attorneys.gathered = True
                output.append([lawfirms, attorneys, clients])
        everyfirm.there_is_another = False
        everyfirm.gathered = True
        output.there_is_another = False
        output.gathered = True
        return output

    def discovery_rules(self):
        output = []
        output.clear()
        output.append(self.juris.civil_procedure()["Duty to Disclose; General Provisions Governing Discovery"])
        if self.rrogs:
            output.append(self.juris.civil_procedure()["Interrogatories to Parties"])
        if self.rrfps:
            output.append(self.juris.civil_procedure()[
                              "Producing Documents, Electronically Stored Information, and Tangible Things, or Entering onto Land, for Inspection and Other Purposes"])
            output.append(self.juris.civil_procedure()["Physical and Mental Examinations"])
        if self.rrfas:
            output.append(self.juris.civil_procedure()["Requests for Admission"])

        if self.juris.jurisdiction.name.text == "state" and "flattened" in self.juris.civil_procedure().keys():
            data = flatten(sorted(output))
        else:
            data = sorted(output)
        log("the_flattened data is " + repr(data))
        if self.juris.jurisdiction.name.text == "state" and self.juris.state.name.text == "CA":
            return data
        else:
            from itertools import groupby
            from operator import itemgetter
            packed = list()
            for k, g in groupby(enumerate(data), lambda ix: int(ix[0]) - int(ix[1])):
                packed.append(list(map(itemgetter(1), g)))
            cracked = list()
            for item in packed:
                the_item = ""
                if len(item) > 1:
                    first_item = str(item[0])
                    last_item = str(item[-1])
                    if len(last_item) > 2 and len(first_item) > 2:
                        if last_item[-3] != first_item[-3]:
                            last_item = str(last_item[-3:])
                        else:
                            last_item = str(last_item[-2:])
                    elif len(last_item) > 2:
                        last_item = str(last_item[-3:])
                    else:
                        last_item = str(last_item[-2:])
                    if self.juris.jurisdiction.name.text == "state" and "prefix" in self.juris.civil_procedure().keys():
                        the_item += str(self.juris.civil_procedure()["prefix"])
                        if self.juris.state.name.text == "WI":
                            if len(str(first_item)) == 1:
                                the_item += "0"
                    the_item += str(first_item)
                    the_item += self.fx.EN
                    if self.juris.jurisdiction.name.text == "state" and self.juris.state.name.text == "WI":
                        if len(str(last_item)) == 1:
                            the_item += "0"
                    the_item += last_item
                    cracked.append(the_item)
                else:
                    if self.juris.jurisdiction.name.text == "state" and "prefix" in self.juris.civil_procedure().keys():
                        the_item += str(self.juris.civil_procedure()["prefix"])
                        if self.juris.state.name.text == "WI":
                            if len(str(item[0])) == 1:
                                the_item += "0"
                    the_item += str(item[0])
                    cracked.append(the_item)
            return cracked

    def draft(self):
        if self.draft == "Discovery Response":
            output = ""
            if self.rrfps:
                try:
                    output += bluebook_ordinal(case.rrfps.request_ordinal)
                    output += "RRFP "
                except:
                    pass
            if self.rrogs:
                try:
                    output += bluebook_ordinal(case.rrogs.request_ordinal)
                    output += "RROG "
                except:
                    pass
            if self.rrfas:
                try:
                    output += bluebook_ordinal(case.rrfas.request_ordinal)
                    output += "RRFA "
                except:
                    pass
        else:
            output = ""
        return output

    def no(self):
        output = ""
        if "case no" in self.juris.civil_procedure().keys():
            output += self.juris.civil_procedure()["case no"]
        else:
            output += "Case No."
        return output

    def number(self):
        output = ""
        if "case no" in self.juris.civil_procedure().keys():
            output += self.juris.civil_procedure()["case no"]
        else:
            output += "Case No."
        output += " "
        output += self.docket_number.strip()
        return output

    @property
    def rules_of_evidence(self):
        output = []
        if self.juris.jurisdiction.name.text == "federal":
            output.append("FRE")
        elif "rules of evidence" in self.juris.civil_procedure().keys():
            output.append(self.juris.civil_procedure()["rules of evidence"])
        elif state_name(self.juris.state.name.text) == "Oregon":
            output.append("OEC")
        else:
            output.append(str(state_name(self.juris.state.name.text) + " Rules of Evidence"))
        return output

    def style_the_alldocx(self):
        # Colo. R. Civ. Proc. 310 has a bunch of really specific formatting requirements that will need to be encoded.
        alldocx = Document(self.the_alldocx.path())
        alldocx.core_properties.author = "Trial Legend"
        style = alldocx.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12.0)
        sections = alldocx.sections

        justification = WD_ALIGN_PARAGRAPH.JUSTIFY
        for section in sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
        paragraphs = alldocx.paragraphs
        for paragraph in paragraphs[10:12]:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = Pt(25.4)
            paragraph.alignment = justification

        paragraphs[12].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph12_format = paragraphs[12].paragraph_format
        paragraph12_format.line_spacing = Pt(25.4)

        for paragraph13_19 in paragraphs[13:19]:
            paragraph_format13_19 = paragraph13_19.paragraph_format
            paragraph_format13_19.line_spacing = Pt(25.4)
            paragraph13_19.alignment = justification

        paragraphs[19].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph19_format = paragraphs[19].paragraph_format
        paragraph19_format.line_spacing = Pt(25.4)

        terminal_paragraph = int(len(paragraphs) - 14)

        for paragraph20 in paragraphs[20:terminal_paragraph]:
            paragraph_format20 = paragraph20.paragraph_format
            paragraph_format20.line_spacing = Pt(25.4)
            paragraph20.alignment = justification

        paragraphs[terminal_paragraph].alignment = WD_ALIGN_PARAGRAPH.CENTER
        terminal_paragraph_format = paragraphs[terminal_paragraph].paragraph_format
        terminal_paragraph_format.line_spacing = Pt(12.7)

        for paragraph_certificate in paragraphs[int(terminal_paragraph + 1):]:
            paragraph_format_certificate = paragraph_certificate.paragraph_format
            paragraph_format_certificate.line_spacing = Pt(12.7)
            paragraph_certificate.alignment = justification

        alldocfile = DAFile("file")
        the_name = "File"
        alldocfile.initialize(filename=space_to_underscore(str(self.name() + ".docx")))
        alldocx.save(alldocfile.path())
        alldocfile.commit()
        return alldocfile

    def is_mva(self):
        if "Negligence" in self.claims:
            if "MVA" in self.negligence_type:
                return True
            else:
                return False
        else:
            return False

    def is_products_liability(self):
        if '"Products Liability"' in self.claims:
            return True
        else:
            return False

    def product_was_leased(self):
        if self.owned_or_leased == "owned":
            return False
        else:
            return True

    @property
    def fault(self):
        if "Negligence" in self.claims:
            return "negligence"
        if "Premises Liability" in self.claims:
            return "fault"
        else:
            return "conduct"

    @property
    def at_fault(self):
        if "Negligence" in self.claims:
            return "negligent"
        else:
            return "at fault"

    def name(self):
        return str(
            # self.fx.ITALIC
            self.user_clientlist()[0].last_name()
            + " v. "
            + self.recipients()[0].last_name()
            # + self.fx.ITALIC
        )

    def __str__(self):
        return str(self.name())

    def facite_complaint(self):
        para_num = 1
        output.append((str(str(para_num) + "."), para))
        para_num += 1
        for claim in self.claims:
            # if claim == 'Products Liability':
            # if "Strict Liability" in self.products_li_type:
            # if "Negligence" in self.products_li_type:
            # if "Breach of Express Warranty" in self.products_li_type:
            # if "Breach of Implied Warranty of Merchantability" in self.products_li_type:
            # if "Breach of Implied Warranty of Fitness for a Particular Purpose" in self.products_li_type:

            if claim == "Premises Liability":
                for para in self.paras_premises_liability():
                    if paras_premises_liability().index(para) == 0:
                        output.append(
                            (
                                str(
                                    str(para_num)
                                    + "."
                                    + str(
                                        self.fx.ALINE
                                        + str(
                                            str(
                                                ordinal_number(
                                                    claims.index(claim), capitalize=True
                                                )
                                            )
                                            + " Claim for Relief"
                                        ).upper()
                                        + self.fx.ALINE
                                        + str("(" + str(claim) + ")")
                                        + str(
                                            str(
                                                self.fx.ALINE
                                                + "("
                                                + self.liable_for_premises().pnameup()
                                                + ")"
                                            )
                                            if len(self.liable_for_premises())
                                               < len(self.liable_parties())
                                            else ""
                                        )
                                        if len(self.claims) > 1
                                        else ""
                                    )
                                ),
                                para,
                            )
                        )
                        para_num += 1
                    else:
                        output.append(
                            (str(str(para_num) + "."), para, False, None, None, None)
                        )
                        para_num += 1
        for party in self.complainants():
            if party.party.is_not_medically_stationary:
                output.append(
                    (
                        str(
                            "NOTICE OF POTENTIAL AMENDMENT OF CLAIMED INJURIES AND DAMAGES"
                            + self.fx.ALINE
                            + str(para_num)
                            + "."
                        ),
                        str(
                            party.party.name.full().upper()
                            + " gives notice that "
                            + party().pronoun_subjective()
                            + " is not medically stationary, and therefore continues to treat and be evaluated by "
                            + party.party.pronoun_possessive("")
                            + "medical providers. Thus, the injuries and damages alleged above are subject to amendment to conform to the proof."
                        ),
                    )
                )

    def paras_basic(self):
        output = []
        output.append(
            str(
                "Due to the amount prayed for herein this claim is "
                + str("not " if self.total_damage() < 50000 else "")
                + "subject to mandatory arbitration."
            )
        )
        if "Wrongful Death" in self.claims:
            for party in parties:
                if party.party.is_estate_rep and party.party.is_complainant:
                    output.append(
                        str(
                            party.party.name.full()
                            + " is the duly appointed Personal Representative of the Estate of "
                            + self.decedent().name.full()
                            + ", who is deceased. Decedent was a resident of "
                            + state_name(self.decedent().domicile)
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "This action is brought for the benefit of the statutory beneficiaries of the Estate of "
                            + self.decedent().name.full()
                            + " pursuant to ORS 30.020."
                        )
                    )
        else:
            for party in self.complainants():
                output.append(
                    str(
                        party.party.name.full()
                        + " is "
                        + indefinite_article(state_name(party.party.domicile))
                        + " resident."
                    )
                )
        for party in parties:
            if party.party.name.caption_type.name.text == "guardian ad litem" and party.party.is_complainant:
                output.append(
                    str(
                        party.party.name.full()
                        + " is the duly appointed guardian ad litem of "
                        + party.party.guards.name.full()
                        + ", and brings this action on "
                        + party.party.guards.pronoun_possessive("")
                        + "behalf."
                    )
                )
            elif party.party.name.caption_type.name.text == "governmental entity":
                output.append(
                    str(
                        "At all times material, "
                        + party.party.name.full()
                        + " was and is a political subdivision of the "
                        + str(
                            str("State of " + state_name(party.party.address.state))
                            if party.party.is_state_gov
                            else "United States"
                        )
                        + "."
                    )
                )
                output.append(
                    str(
                        "At all times material, "
                        + party.party.name.full()
                        + " had working for it officials, employees, agents, or others within its control, or right of control. All acts attributed to "
                        + party.party.name.full()
                        + " were performed by said persons while acting within the course and scope of said capacity."
                    )
                )
                output.append(
                    str(
                        "On "
                        + self.date_of_notice
                        + ", timely notice of this claim was given to "
                        + party.party.name.full()
                        + "."
                    )
                )
            elif party.party.name.caption_type.name.text == "doing business as":
                output.append(
                    str(
                        "At all times material, "
                        + party.party.indlpname()
                        + " was an individual doing business as “"
                        + party.party.dba
                        + ".”"
                    )
                )
                output.append(
                    str(
                        "At all times material, "
                        + party.party.indlpname()
                        + " had employed owners, officers, directors, members, managers, employees, agents, or others within "
                        + party.party.pronoun_possessive("")
                        + "control, or right of control, and all of whom were acting within the course and scope of said positions. All acts or omissions attributed to "
                        + party.party.indlpname()
                        + " were either performed by "
                        + party.party.indlpname()
                        + ", or were otherwise performed by said persons in said capacities."
                    )
                )
            elif party.party.name.caption_type.name.text == "business entity":
                output.append(
                    str(
                        "At all times material, "
                        + party.party.indlpname()
                        + " was and is "
                        + indefinite_article(state_name(party.party.stofinc))
                        + " "
                        + party.party.business_entity.name.text.lower()
                        + str(
                            str(
                                ", authorized to transact business in the state of "
                                + state_name(self.juris.state.name.text)
                                + "."
                            )
                            if party.party.orauthorized
                            else ""
                        )
                    )
                )
                output.append(
                    str(
                        "At all times material, "
                        + party.party.indlpname()
                        + " conducted and continues to conduct said business from its principal place of business/principal office located at "
                        + party.party.ppb.address.on_one_line()
                        + "."
                    )
                )
                output.append(
                    str(
                        "At all times material, "
                        + party.party.indlpname()
                        + " had working for it owners, officers, directors, members, managers, employees, agents, or others within its control, or right of control, and all of whom were acting within the course and scope of said positions.  All acts or omissions attributed to Defendant "
                        + party.party.indlpname()
                        + "  were performed by said persons in said capacities."
                    )
                )
        return output

    def premises_owners(self):
        try:
            output = SCPartyList(
                "premises_owners", there_are_any=True, auto_gather=False
            )
            output.clear()
            for party in self.parties:
                if party.party.is_owner and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w premises_owners(self)")

    def premises_tt(self):
        try:
            output = SCPartyList("premises_tt", there_are_any=True, auto_gather=False)
            output.clear()
            for party in self.parties:
                if party.party.is_tt and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w premises_tt(self)")

    def claims_premises_liability(self):
        try:
            output = SCPartyList(
                "claims_premises_liability", there_are_any=True, auto_gather=False
            )
            output.clear()
            for party in self.parties:
                if party.party.claims_premises_liability and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w claims_premises_liability(self)")

    def liable_for_premises(self):
        try:
            output = SCPartyList(
                "liable_for_premises", there_are_any=True, auto_gather=False
            )
            output.clear()
            for party in self.parties:
                if party.party.is_premises_manager or party.party.is_owner or party.party.is_tt:
                    if party not in output.elements:
                        output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w liable_for_premises(self)")

    def premises_managers(self):
        try:
            output = SCPartyList(
                "premises_managers", there_are_any=True, auto_gather=False
            )
            output.clear()
            for party in self.parties:
                if party.party.is_premises_manager and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w premises_managers(self)")

    def paras_premises_liability(self):
        output = []
        for claim in self.claims:
            if claim == "Premises Liability":
                output.append(
                    str(
                        "At all times material, "
                        + self.premises_owners().pname()
                        + " owned the real property and improvements thereon located at "
                        + self.premises.address.on_one_line()
                        + ", also known as "
                        + self.premises.county
                        + " County Map and Taxlot No. "
                        + str(self.premises.address.taxlot)
                        + " (hereinafter referred to as “The Property”)."
                    )
                )

                if self.premises.had_tt:
                    output.append(
                        str(
                            "On information and belief, at all times material, "
                            + self.premises_owners().pname()
                            + " leased The Property to "
                            + self.premises_tt().pname()
                            + ", including any buildings or improvements thereon, parking lot and sidewalks adjacent thereto, and any other systems requiring maintenance thereon."
                        )
                    )

                if self.premises.had_managers:
                    output.append(
                        str(
                            "On information and belief, at all times material, "
                            + self.premises_owners().pname()
                            + " contracted with "
                            + self.premises_managers().pname()
                            + " for asset, property and/or facility management of The Property, including any buildings or improvements thereon, parking lot and sidewalks adjacent thereto, and any other systems requiring maintenance thereon."
                        )
                    )

                if self.location_of_fall == "sidewalk":
                    output.append(
                        str(
                            "At all times material, The Property included a sidewalk intended for pedestrian use "
                            + str(
                                "that was situated between the parking spaces in the parking lot and "
                                if self.sidewalk_in_parkinglot
                                else ""
                            )
                            + str("that abbutted " if self.sidewalk_abutted else "")
                            + "the "
                            + self.building_type
                            + " located on The Property (hereinafter referred to as “"
                            + self.hazloc.name.text
                            + "”)."
                        )
                    )

                if self.location_of_fall == "parking lot":
                    output.append(
                        str(
                            "At all times material, The Property included a parking lot, intended for "
                            + self.pclass
                            + " to park their vehicles and for pedestrian use, that abbutted the "
                            + self.building_type
                            + " located on The Property (hereinafter referred to as “"
                            + self.hazloc.name.text
                            + "”)."
                        )
                    )

                if self.location_of_fall == "entryway":
                    output.append(
                        str(
                            "At all times material, abutting the "
                            + self.building_type
                            + " located on The Property, there existed an asphalt entryway that was intended for use by pedestrians for ingress and egress (hereinafter referred to as “"
                            + self.hazloc.name.text
                            + "”)."
                        )
                    )

                    output.append(
                        str(
                            "At all times material "
                            + self.hazloc.name.text
                            + " was used and intended for use by "
                            + self.liable_for_premises().pname_possessive()
                            + self.pclass
                            + "s as a means of ingress and egress."
                        )
                    )

                output.append(
                    str(
                        "At all times material, "
                        + self.liable_for_premises().pname()
                        + " had control, or the right of control, of The Property, including "
                        + self.hazloc.name.text
                        + "."
                    )
                )

                if self.it_was_overcast:
                    output.append(
                        str(
                            "At all times material, the weather conditions were overcast, reducing visibility for those, such as "
                            + comma_and_list(self.claims_premises_liability())
                            + ", using "
                            + self.hazloc.name.text
                            + ", which  "
                            + self.liable_for_premises().asnoun()
                            + " knew or should have known."
                        )
                    )
                if self.it_was_shaded:
                    output.append(
                        str(
                            "At all times material, "
                            + self.the_hazard
                            + " was covered in shade, reducing visibility for those, such as "
                            + comma_and_list(self.claims_premises_liability())
                            + ", using "
                            + self.hazloc.name.text
                            + ", which  "
                            + self.liable_for_premises().asnoun()
                            + " knew or should have known."
                        )
                    )
                if not self.claims_premises_liability().was_trespasser:
                    output.append(
                        str(
                            "At all times material, "
                            + self.liable_for_premises().pname()
                            + " had a common law duty of inspecting and maintaining "
                            + self.hazloc.name.text
                            + " in a condition to make it reasonably safe for use by "
                            + self.liable_for_premises().pronoun_possessive(self.pclass)
                            + str(
                                ", including inspecting for and making safe icy conditions."
                                if self.conditions_were_icy
                                else "."
                            )
                        )
                    )

                if self.premises.had_managers or self.premises.had_tt:
                    output.append(
                        str(
                            "On information and belief, "
                            + str(
                                self.premises_managers().pname()
                                if self.premises.had_managers
                                else str(
                                    self.premises_tt().pname()
                                    if self.premises.had_tt
                                    else ""
                                )
                            )
                            + " had a contractual duty of inspecting and maintaining "
                            + self.hazloc.name.text
                            + " in a condition to make it reasonably safe for use by "
                            + self.liable_for_premises().pronoun_possessive(self.pclass)
                            + str(
                                ", including inspecting for and making safe icy conditions."
                                if self.conditions_were_icy
                                else "."
                            )
                        )
                    )

                if self.there_was_a_foreign_substance:
                    output.append(
                        str(
                            "At all times material, there existed on "
                            + self.hazloc.name.text
                            + " "
                            + self.the_hazard.lower()
                            + " (hereinafter referred to as “"
                            + self.the_hazard
                            + "”), causing the walking surface of "
                            + self.hazloc.name.text
                            + " to become extremely slippery, which "
                            + self.liable_for_premises().asnoun()
                            + " knew or should have known."
                        )
                    )

                    if self.foreign_substance_was_clear:
                        output.append(
                            str(
                                "At all times material, "
                                + self.the_hazard
                                + " was clear, making it difficult for those using "
                                + self.hazloc.name.text
                                + " to see and appreciate its presence, which created a danger to the public using "
                                + self.hazloc.name.text
                                + str(
                                    ", particularly at night"
                                    if self.at_night
                                    else str(
                                        ", particularly during overcast weather conditions"
                                        if self.it_was_overcast
                                        else str(
                                            str(
                                                ", particularly when "
                                                + self.the_hazard
                                                + " was covered in shade"
                                            )
                                            if self.it_was_shaded
                                            else ""
                                        )
                                    )
                                )
                                + ", which "
                                + self.liable_for_premises().asnoun()
                                + " knew or should have known."
                            )
                        )

                    if any(party.party.placed_hazard for party in self.liable_for_premises()):
                        output.append(
                            str(
                                "On information and belief and at all times material, "
                                + self.liable_for_premises().asnoun()
                                + " placed "
                                + self.the_hazard
                                + " or otherwise caused "
                                + self.the_hazard
                                + " to be placed on "
                                + self.hazloc.name.text
                                + "."
                            )
                        )
                    else:
                        output.append(
                            str(
                                "On information and belief and at all times material, "
                                + self.liable_for_premises().asnoun()
                                + " knew "
                                + self.the_hazard
                                + " was on "
                                + self.hazloc.name.text
                                + " and failed to use reasonable diligence to remove it."
                            )
                        )

                        output.append(
                            str(
                                "On information and belief, in the alternative to the allegations in the Paragraph above, and at all times material, "
                                + self.the_hazard
                                + " had been on "
                                + self.hazloc.name.text
                                + " for so long that "
                                + self.liable_for_premises().asnoun()
                                + " should, in the exercise of reasonable diligence, have discovered and removed it."
                            )
                        )

                if self.elevation_difference:
                    output.append(
                        str(
                            "At all times material, "
                            + self.liable_for_premises().pname_did("was")
                            + " also responsible for inspecting and maintaining The Property, including "
                            + self.hazloc.name.text
                            + ", in a condition safe for use by the public, pursuant to the Americans with Disabilities Act."
                        )
                    )

                    output.append(
                        str(
                            "At all times material, there existed on "
                            + self.hazloc.name.text
                            + " "
                            + self.elevation_cause
                            + " (hereinafter referred to as “"
                            + self.the_hazard
                            + "”), causing an elevation difference with the surrounding walking surface of "
                            + self.hazloc.name.text
                            + " approximately 3/4 inches in height or greater."
                        )
                    )

                elif self.conditions_were_icy:
                    if self.it_snowed:
                        output.append(
                            str(
                                "On several days of the days leading up to "
                                + self.doi
                                + ", it snowed on The Property "
                                + str(
                                    str("(hereinafter “" + self.the_hazard + ")")
                                    if not self.snow_compacted and not self.freeze_thaw
                                    else ""
                                )
                                + ", including on "
                                + self.hazloc.name.text
                                + ". At all times material, "
                                + self.liable_for_premises().pname()
                                + " knew, or should have known, of this fact."
                            )
                        )

                        if self.snow_was_compacted:
                            output.append(
                                str(
                                    "On several days of the days leading up to "
                                    + self.doi
                                    + ", the snow melted and drained, or otherwise leaked, onto "
                                    + self.hazloc.name.text
                                    + ", which "
                                    + self.liable_for_premises().pname()
                                    + " knew or should have known."
                                )
                            )

                        if self.freeze_thaw:
                            output.append(
                                str(
                                    "On several days of the days leading up to "
                                    + self.doi
                                    + "some of the snow on The Property and "
                                    + self.hazloc.name.text
                                    + " refroze as a thin layer of ice (hereinafter referred to as “"
                                    + self.the_hazard
                                    + "”), causing the walking surface of "
                                    + self.hazloc.name.text
                                    + " to become extremely slippery, which "
                                    + self.liable_for_premises().pname()
                                    + " knew or should have known."
                                )
                            )

                output.append(
                    str(
                        "At all times material, "
                        + self.liable_for_premises().pname()
                        + " knew of the existence "
                        + self.the_hazard
                        + ", or should have known of its existence, because a reasonable inspection would have revealed it."
                    )
                )

                if self.inexperienced_invitee:
                    output.append(
                        str(
                            "At all times material, "
                            + comma_and_list(self.claims_premises_liability())
                            + " had insufficient prior experience or exposure to "
                            + self.the_hazard.lower()
                            + "s such as "
                            + self.the_hazard
                            + ", which "
                            + self.liable_for_premises().asnoun()
                            + " knew or should have known."
                        )
                    )

                if self.people_complained:
                    output.append(
                        str(
                            "At all times material, "
                            + self.liable_for_premises().pname()
                            + " knew of the existence of "
                            + self.the_hazard
                            + ", or otherwise should have known of the existence of "
                            + self.the_hazard
                            + ", because, prior to "
                            + self.doi
                            + ", people had complained to "
                            + self.liable_for_premises().pname()
                            + " about "
                            + self.the_hazard
                            + ", or otherwise complained about the dangerous condition of "
                            + self.hazloc.name.text
                            + "."
                        )
                    )

                if self.hazloc.name.text_not_well_lit:
                    output.append(
                        str(
                            "At all times material, the area where "
                            + self.the_hazard
                            + " was located was not well lit making it difficult for those using "
                            + self.hazloc.name.text
                            + " to see and appreciate, which created a danger to the public using "
                            + self.hazloc.name.text
                            + str(
                                ", particularly at night"
                                if self.at_night
                                else str(
                                    ", particularly during overcast weather conditions"
                                    if self.it_was_overcast
                                    else str(
                                        str(
                                            ", particularly when "
                                            + self.the_hazard
                                            + " was covered in shade"
                                        )
                                        if self.it_was_shaded
                                        else ""
                                    )
                                )
                            )
                            + "."
                        )
                    )

                if self.conditions_were_icy and self.ice_was_clear:
                    output.append(
                        str(
                            "At all timed material, "
                            + self.the_hazard
                            + " was clear, making it difficult to see and appreciate for those using "
                            + self.hazloc.name.text
                            + ", which created a danger to said persons, including "
                            + comma_and_list(self.claims_premises_liability())
                            + "."
                        )
                    )

                output.append(
                    str(
                        "At all times material, "
                        + self.liable_for_premises().pname()
                        + " knew, or should have known, that the potential danger created by "
                        + self.the_hazard
                        + " was concealed and not reasonably obvious to "
                        + str("inexperienced" if self.inexperienced_invitee else "")
                        + " persons using "
                        + self.hazloc.name.text
                        + ", including "
                        + comma_and_list(self.claims_premises_liability())
                        + "."
                    )
                )

                output.append(
                    str(
                        "At all times material, "
                        + self.liable_for_premises().pname()
                        + " knew, or should have known, that the concealed danger presented by "
                        + self.the_hazard
                        + " constituted an unreasonable risk of harm to to"
                        + str("inexperienced" if self.inexperienced_invitee else "")
                        + " persons using "
                        + self.hazloc.name.text
                        + ", including "
                        + comma_and_list(self.claims_premises_liability())
                        + "."
                    )
                )

                if any(party.party.was_invitee for party in self.claims_premises_liability()):
                    output.append(
                        str(
                            "At all times material, "
                            + self.liable_for_premises().pname()
                            + " knew, or should have known, that the concealed danger presented by "
                            + self.the_hazard
                            + " further constituted an unreasonably dangerous condition that could not be encountered with a reasonable degree of safety by to"
                            + str("inexperienced" if self.inexperienced_invitee else "")
                            + " persons using "
                            + self.hazloc.name.text
                            + ", including "
                            + comma_and_list(self.claims_premises_liability())
                            + ", even if they knew and appreciated the danger."
                        )
                    )

                if any(
                        party.party.was_licensee for party in self.claims_premises_liability()
                ):
                    output.append(
                        str(
                            "At all times material, "
                            + self.liable_for_premises().pname()
                            + " knew, or should have known, that "
                            + comma_and_list(self.claims_premises_liability())
                            + " would not discover or realize the unreasonable risk of harm created by "
                            + self.hazloc.name.text
                            + "."
                        )
                    )

                    output.append(
                        str(
                            "At all times material, "
                            + comma_and_list(self.claims_premises_liability())
                            + " did not know, nor have reason to know, of the unreasonable risk of harm created by "
                            + self.hazloc.name.text
                            + "."
                        )
                    )

                if self.claims_premises_liability().was_invitee:
                    if self.claims_premises_liability().was_binvitee:
                        for party in self.claims_premises_liability():
                            output.append(
                                str(
                                    "At all times material, "
                                    + party.party.name.full().upper()
                                    + " was "
                                    + str(
                                        indefinite_article(party.party.business_on_premises)
                                    )
                                    + " of "
                                    + str(
                                        self.premises_tt().pname()
                                        if self.premises.had_tt
                                        else self.premises_owners().pname()
                                    )
                                    + ", and was therefore on The Property at "
                                    + possessify(
                                        self.liable_for_premises().asnoun(), ""
                                    )
                                    + "invitation, and for "
                                    + possessify(
                                        self.liable_for_premises().asnoun(), ""
                                    )
                                    + "business purposes."
                                )
                            )

                        output.append(
                            str(
                                "At all times material, "
                                + comma_and_list(self.claims_premises_liability())
                                + " "
                                + self.claims_premises_liability().did_verb("was")
                                + " "
                                + self.claims_premises_liability().as_noun(
                                    "business invitee", article=True
                                )
                                + " of "
                                + self.liable_for_premises().asnoun()
                                + "."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, "
                                + self.liable_for_premises().pname()
                                + " held The Property open to business invitees, including "
                                + comma_and_list(self.claims_premises_liability())
                                + ", for "
                                + self.liable_for_premises().pronoun_possessive(
                                    "business purposes."
                                )
                            )
                        )

                    if self.claims_premises_liability().was_pinvitee:
                        output.append(
                            str(
                                "At all times material, "
                                + comma_and_list(self.claims_premises_liability())
                                + " was a public invitee of "
                                + self.liable_for_premises().pname()
                                + "."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, "
                                + self.liable_for_premises().pname()
                                + " expressly or impliedly lead the public, including "
                                + comma_and_list(self.claims_premises_liability())
                                + ", to believe that "
                                + self.hazloc.name.text
                                + " was intended for use as a public walkway."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, "
                                + self.liable_for_premises().pname()
                                + " acquiesced to the use of "
                                + self.hazloc.name.text
                                + " as a public walkway, including said use by "
                                + comma_and_list(self.claims_premises_liability())
                                + "."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, "
                                + self.liable_for_premises().pname()
                                + " knew, or should have known, that the area where "
                                + self.the_hazard
                                + " existed was available for use, and was actually used, by "
                                + +self.liable_for_premises().pronoun_possessive(
                                    self.pclass
                                )
                                + ", including "
                                + comma_and_list(self.claims_premises_liability())
                                + "."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, the use of "
                                + self.hazloc.name.text
                                + " by "
                                + comma_and_list(self.claims_premises_liability())
                                + " as a walkway was in accord with the intention or design for which the premises was adapted and prepared."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, the manner in which "
                                + comma_and_list(self.claims_premises_liability())
                                + " walked on "
                                + self.hazloc.name.text
                                + " was normal and reasonably foreseeable."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, the route in which "
                                + comma_and_list(self.claims_premises_liability())
                                + " walked on "
                                + self.hazloc.name.text
                                + " was normal and reasonably foreseeable."
                            )
                        )

                        output.append(
                            str(
                                "At all times material, the time at which "
                                + comma_and_list(self.claims_premises_liability())
                                + " walked on "
                                + self.hazloc.name.text
                                + " was normal and reasonably foreseeable."
                            )
                        )

                if self.claims_premises_liability().was_licensee:
                    output.append(
                        str(
                            "At all times material, "
                            + self.claims_premises_liability().fullname_did("was")
                            + " on The Property, at "
                            + self.liable_for_premises().asnoun_possessive()
                            + "invitation, for "
                            + self.claims_premises_liability().licensed_for
                            + "."
                        )
                    )

                    output.append(
                        str(
                            "At all times material, "
                            + self.claims_premises_liability().fullname_did("was")
                            + " "
                            + self.liable_for_premises().asnoun_possessive()
                            + "licensee."
                        )
                    )

                output.append(self.premises_li_specifications())

    def premises_li_specifications(self):
        output = []
        output.append(
            str(
                self.liable_for_premises().pname_did("was")
                + " at fault in one or more of the following particulars, each of which constituted an unreasonable and foreseeable risk of injury to "
                + comma_and_list(self.claims_premises_liability())
                + ":"
            )
        )
        output.append(
            str(
                item_label(levl, 5)
                + self.fx.TAB
                + "In causing or allowing "
                + self.the_hazard
                + " to "
                + str("accumulate or " if self.conditions_were_icy else "")
                + "exist on "
                + self.hazloc.name.text
            )
        )
        output.append(
            str(
                item_label(levl, 5)
                + self.fx.TAB
                + "In failing to inspect "
                + self.hazloc.name.text
                + str(
                    "for the existence of ice"
                    if self.conditions_were_icy
                    else str("for dangers, including " + self.the_hazard)
                )
            )
        )
        output.append(
            str(
                item_label(levl, 5)
                + self.fx.TAB
                + "In failing to discover the dangerous condition of "
                + self.the_hazard
            )
        )
        if self.conditions_were_icy:
            output.append(
                str(
                    item_label(levl, 5)
                    + self.fx.TAB
                    + "In failing to clean up, remove or melt "
                    + self.the_hazard
                )
            )
            output.append(
                str(
                    item_label(levl, 5)
                    + self.fx.TAB
                    + "In failing to place sand, gravel, cinder, or other suitable traction material on "
                    + self.the_hazard
                    + " to make it safe to walk upon"
                )
            )
        else:
            output.append(
                str(
                    item_label(levl, 5)
                    + self.fx.TAB
                    + "In failing to maintain The Property so as to remove, repair, or replace "
                    + self.the_hazard
                )
            )
        if self.hazloc.name.text_not_well_lit:
            output.append(
                str(
                    item_label(levl, 5)
                    + self.fx.TAB
                    + "In failing to illuminate "
                    + self.hazloc.name.text
                    + " so that the existence and location of "
                    + self.the_hazard
                    + " would be obvious to those using "
                    + self.hazloc.name.text
                    + str(
                        "at night"
                        if self.at_night
                        else str(
                            "during overcast weather conditions"
                            if self.it_was_overcast
                            else str(
                                str("when " + self.the_hazard + " was covered in shade")
                                if self.it_was_shaded
                                else ""
                            )
                        )
                    )
                )
            )
        if self.elevation_difference:
            output.append(
                str(
                    item_label(levl, 5)
                    + self.fx.TAB
                    + "In allowing there to exist in "
                    + self.hazloc.name.text
                    + " a misalignment in the walkway was greater than 3/8 inches (negligence per se under the Americans with Disabilities Act (ADA))"
                )
            )
        output.append(
            str(
                item_label(levl, 5)
                + self.fx.TAB
                + "In failing to warn, or otherwise make obvious to "
                + self.pclass
                + ", including "
                + comma_and_list(self.claims_premises_liability()).upper()
                + " of the existence of "
                + self.the_hazard
            )
        )
        output.append(
            str(
                item_label(levl, 5)
                + self.fx.TAB
                + "In failing to put up signs, barricades, or other devices to prevent persons from encountering "
                + self.the_hazard
            )
        )
        outout = RichText("/a".join(output))
        return outout

    def forma_sect(self, d):
        if not (d.sections):
            sect = d.add_section
        else:
            sect = d.sections[0]
        sect.left_margin = Inches(1)
        sect.right_margin = Inches(1)
        sect.top_margin = Inches(1)
        sect.bottom_margin = Inches(1)
        return d

    def forma_p_i(self, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0.5)
        paragraph_format.right_indent = Inches(0.06)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(24.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for p in paragraph.runs:
            p.font.name = "Times New Roman"
            p.font.size = Pt(12)
        return paragraph

    def forma_p_noi(self, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0)
        paragraph_format.right_indent = Inches(0.06)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(24.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for p in paragraph.runs:
            p.font.name = "Times New Roman"
            p.font.size = Pt(12)
        return paragraph

    def rfps_doc(self):
        d = docx.Document()
        indexr = Counter()
        rfp_val = self.prior_requests_total
        for item in self.rfps:
            paragraph = d.add_paragraph(
                str("REQUEST FOR PRODUCTION NO. " + str(rfp_val) + ":")
            )
            rfp_val += 1
            no_indent_body(paragraph)
            bld_uline(paragraph.runs[0])
            paragraph.add_run(
                str(" " + fix_punctuation(single_paragraph(item.name.text)))
            )
            nullrun(paragraph.runs[1])
            resp = d.add_paragraph("RESPONSE: ")
            no_indent_body(resp)
            bld_uline(resp.runs[0])
        tabula = DAFile("tabula")
        tabula.initialize(filename="tabula.docx")
        d.save(tabula.path())
        tabula.commit()
        return tabula

    @property
    def total_damage(self):
        output = 0.0
        output += self.ecomonomics.total
        output += self.nonecomonomics.amount
        return output

    def filing_fee(self):
        if self.total <= 10000:
            costs.filingfee = 170
            mandatory_arbitration = True
            rule = "ORS 21.160(a)"
        elif self.total < 50000:
            costs.filingfee = 283
            mandatory_arbitration = True
            rule = "ORS 21.160(b)"
        elif self.total < 1000000:
            costs.filingfee = 594
            mandatory_arbitration = False
            rule = "ORS 21.160(c)"
        elif self.total < 10000000:
            costs.filingfee = 884
            mandatory_arbitration = False
            rule = "ORS 21.160(d)"
        else:
            fee = 1178
            mandatory_arbitration = False
            rule = "ORS 21.160(e)"
        return (fee, mandatory_arbitration, rule)

    @property
    def sol(self):
        output = self.doi + date_interval(years=2)
        return output

    @property
    def beat_sol_by(self):
        output = date_difference(starting=self.dof, ending=self.sol)
        return output

    @property
    def incident(self):
        return "incident"
        if "Negligence" in self.claims:
            if "Animal Attack" in self.negligence_type:
                return "attack"
            elif "MVA" in self.negligence_type:
                return "collision"
            else:
                return "incident"
        elif "Premises Liability" in self.claims:
            if "Slip/Trip and Fall" in self.premises_li_type:
                return "fall"
            else:
                return "incident"
        elif "Battery" in self.claims:
            return "assault"
        elif "AVP" in self.claims:
            return "abuse"
        else:
            return "incident"

    def has_cross_parties(self):
        log("has_cross_parties")
        if self.cross_claims.there_are_any == True:
            log("has_cross_parties . . . . TRUE")
            return True
        else:
            log("has_cross_parties . . . . FALSE")
            return False

    def has_third_parties(self):
        if any(party.party.party_type.name.text == "Third-Party Defendant" for party in self.parties):
            return True
        else:
            return False

    def has_intervenors(self):
        if any(party.party.party_type.name.text == "Intervenor" for party in self.parties):
            return True
        else:
            return False

    def has_deceaseds(self):
        if any(party.party.party_type.name.text == "Deceased" for party in self.parties):
            return True
        else:
            return False

    def paras_strict_product_liability(self):
        output = []
        output.append(
            str(
                "At all times material, "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " was engaged in the business of selling and leasing "
                + self.product
                + "s, including the "
                + self.product
                + " at issue, as consumer goods from "
                + self.recipients().pronoun_possessive("")
                + self.product_seller().ppb
                + " resort identified in Paragraph 3 above."
            )
        )
        output.append(
            str(
                "The "
                + self.product
                + " at issue was "
                + self.type_of_purchase
                + " by "
                + comma_and_list(self.user_clientlist())
                + " from "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " as a consumer good."
            )
        )
        output.append(
            str(
                "At all times material, the "
                + self.product
                + " was expected to, and did, reach "
                + comma_and_list(self.user_clientlist())
                + " without a substantial change in the condition in which it was manufactured and sold."
            )
        )
        output.append(
            str(
                "Based on information and belief, the "
                + self.product
                + " was defective and unreasonably dangerous to users or consumers, including "
                + comma_and_list(self.user_clientlist())
                + ", in one or more of the following ways:"
                + self.fx.NLINE
                + item_label(levl, 4)
                + " The "
                + self.product
                + " failed to conform to the manufacturer’s specifications specific to "
                + self.defect
                + ";"
                + self.fx.NLINE
                + item_label(levl, 4)
                + " The "
                + self.product
                + " did not contain proper and adequate warnings for safe use, relative to "
                + self.defect
                + ";"
                + self.fx.NLINE
                + item_label(levl, 4)
                + " The "
                + self.product
                + " did not contain proper and adequate warnings of the dangers of use, relative to "
                + self.defect
                + "; and"
                + self.fx.NLINE
                + item_label(levl, 4)
                + " The "
                + self.product
                + " did not contain proper and adequate instructions for safe use, relative to "
                + self.defect
                + ", or otherwise concerning inspecting the "
                + self.product
                + " relative to manufacturer expectations or instructions."
            )
        )
        output.append(
            str(
                "As a result of the defective and unreasonably dangerous condition of the "
                + self.product
                + ", "
                + comma_and_list(self.user_clientlist())
                + " sustained the injuries and noneconomic damages as alleged in Paragraphs 9 and 10 above, which are realleged and incorporated by reference herein."
            )
        )
        return output

    def paras_breach_of_express_warranty(self):
        output = []
        output.append(
            str(
                "Upon "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + " "
                + self.type_of_purchase
                + " of the "
                + self.product
                + ", "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + ", in offering the product for "
                + self.type_of_purchase
                + " to consumers, made factual affirmations and promises relative to the safety of the "
                + self.product
                + " for trail use, assuring that the "
                + self.product
                + " was fit and safe for such use."
            )
        )
        output.append(
            str(
                "The factual affirmations and promises made by "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " became part of the basis of the bargain for the "
                + self.type_of_purchase
                + " of the "
                + self.product
                + " by "
                + comma_and_list(self.user_clientlist())
                + "."
            )
        )
        output.append(
            str(
                "The "
                + self.product
                + " suffered a defect, in that its seat back was not properly secured to the frame, causing users of the "
                + self.product
                + " to be exposed to sharps in the course of normal and foreseeable use, as described above. This defect constituted a breach of "
                + self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + " express warranty referenced in Paragraphs 18 and 19 above."
            )
        )
        output.append(
            str(
                "As a result of "
                + self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + " breach of express warranty as alleged herein, "
                + comma_and_list(self.user_clientlist())
                + " suffered the injuries and damages alleged in Paragraphs 9 and 10 above, which are realleged and incorporated by reference herein."
            )
        )
        output.append(
            str(
                comma_and_list(self.user_clientlist())
                + " gave timely and reasonable notice of the breach of an express warranty to "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + "."
            )
        )
        return output

    def paras_breach_of_implied_warranty_of_merchantability(self):
        output = []
        output.append(
            str(
                "At all times material, "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " was a merchant, in that it was in the business of renting/leasing and selling "
                + self.product
                + "s, including the "
                + self.product
                + " at issue."
            )
        )
        output.append(
            str(
                "At all times material, "
                + comma_and_list(self.user_clientlist())
                + " was in vertical privity with "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " as described above."
            )
        )
        output.append(
            str(
                self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " impliedly warranted to "
                + comma_and_list(self.user_clientlist())
                + " that the "
                + self.product
                + " was of a quality that would pass without objection in the trade, was fit for ordinary purposes for which such products are used, and in all other respects was of merchantable quality as further defined by "
                + self.implied_warranty_of_merchantability_rule
                + "."
            )
        )
        output.append(
            str(
                comma_and_list(self.user_clientlist())
                + " relied on this implied warranty of merchantability in the "
                + self.type_of_purchase
                + " and use of the "
                + self.product
                + "."
            )
        )
        output.append(
            str(
                self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + "​ breached this implied warranty of merchantability in that the "
                + self.product
                + " was not merchantable and was unfit for the ordinary purposes for which it was "
                + self.type_of_purchase
                + ", as it was unsafe and not reasonably fit for "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + " use given the "
                + self.defect
                + "."
            )
        )
        output.append(
            str(
                self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + "breach of the implied warranty of merchantability caused "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + "injuries and damages as alleged in Paragraphs 9 and 10 above, which are realleged and incorporated by reference herein."
            )
        )
        output.append(
            str(
                comma_and_list(self.user_clientlist())
                + " gave timely and reasonable notice of the breach of the implied warranty of merchantability to "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + "."
            )
        )
        return output

    def paras_breach_of_implied_warranty_of_fitness_for_a_particular_purpose(self):
        output = []
        output.append(
            str(
                self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " impliedly warranted to "
                + comma_and_list(self.user_clientlist())
                + " that the "
                + self.product
                + " was of a quality that would pass without objection in the trade, was fit for ordinary purposes for which such products are used, and in all other respects was of merchantable quality as further defined by "
                + self.implied_warranty_of_fitness_for_a_particular_purpose_rule
                + ", and more specifically fit for use as a safe and reasonable  "
                + self.product
                + " to be "
                + self.particular_purpose
                + " for "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + "purposes."
            )
        )

        output.append(
            str(
                comma_and_list(self.user_clientlist())
                + " relied on "
                + self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + "skill or judgment in selection of the "
                + self.product
                + " and on "
                + self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + "implied warranty of fitness for use as a safe and reasonable "
                + self.product
                + " to be "
                + self.particular_purpose
                + ". "
            )
        )

        output.append(
            str(
                "At all times material, "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + " knew of the particular purpose for which "
                + comma_and_list(self.user_clientlist())
                + " was renting/leasing the "
                + self.product
                + ", as well as "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + "reliance on "
                + self.warrantors().asnoun()
                + " "
                + possessify(comma_and_list(self.warrantors()), "")
                + "skill or judgment in selecting the appropriate pontoon boat for such purpose."
            )
        )
        output.append(
            str(
                self.recipients().asnoun()
                + " breached the implied warranty of fitness for a particular purpose in that the "
                + self.product
                + " was unsafe and not reasonably fit for use as a "
                + self.product
                + ", given the condition of the "
                + self.product
                + " relative to the "
                + self.product_defect
                + "."
            )
        )
        output.append(
            str(
                possessify(self.recipients().asnoun(), "")
                + "breach of the implied warranty of fitness for a particular purpose caused "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + "injuries and damages as alleged in Paragraphs 9 and 10 above, which are realleged and incorporated by reference herein."
            )
        )
        output.append(
            str(
                omma_and_list(self.user_clientlist())
                + " gave timely and reasonable notice of the breach of the implied warranty of fitness for a particular purpose to "
                + self.warrantors().asnoun()
                + " "
                + comma_and_list(self.warrantors())
                + "."
            )
        )
        return output

    def manufacturers(self):
        try:
            output = SCPartyList("manufacturers", there_are_any=True, auto_gather=False)
            output.clear()
            for party in self.parties:
                if party.party.is_manufacturer and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w product_seller(self)")

    def wholesalers(self):
        try:
            output = SCPartyList("wholesalers", there_are_any=True, auto_gather=False)
            output.clear()
            for party in self.parties:
                if party.party.is_wholesaler and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w product_seller(self)")

    def the_ps(self):
        log("the_ps .  . . . for party in self.parties: . . . . 1")
        the_ps_output = SCList("the_ps_output", there_are_any=True, auto_gather=False)
        the_ps_output.clear()
        for party in self.parties:
            # if any(designation.name.text in ("Plaintiff", "Petitioner", "Claimant", "Appellant", "Obligee") for designation in party.party.party_types):
            log("the_ps .  . . . for party in self.parties: . . . . 2 . . . party " + repr(party))
            for designation in party.party.party_types:
                log("the_ps .  . . . for party in self.parties: . . . . 3  . . . . designation " + repr(designation))
                if designation.name.text == "Plaintiff":
                    log("the_ps .  . . . for party in self.parties: . . . . 4")
                    the_ps_output.append(party.party)
        the_ps_output.there_is_another = False
        the_ps_output.gathered = True
        return the_ps_output

    def the_ds(self):
        output = SCPartyList("the_ds", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party.party.party_type.name.text == ("Defendant" or "Respondent"):
                output.append(party.party)
        output.there_is_another = False
        output.gathered = True
        return output

    def cross_ps(self):
        output = SCList("cross_ps", there_are_any=True, auto_gather=False)
        output.clear()
        if self.has_cross_parties():
            for party in self.parties:
                for claim in self.cross_claims:
                    if party in claim.ps:
                        output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def cross_ds(self):
        output = SCPartyList("cross_ps", there_are_any=True, auto_gather=False)
        output.clear()
        if self.has_cross_parties():
            for party in self.parties:
                for claim in self.cross_claims:
                    if party in claim.ds:
                        output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def threedp_ps(self):
        output = SCPartyList("threedp_ps", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party.party.party_type.name.text in ("Third-Party Plaintiff"):
                output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def threedp_ds(self):
        output = SCPartyList("threedp_ds", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party.party.party_type.name.text in ("Third-Party Defendant"):
                output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def all_ps(self):
        output = SCPartyList("all_ps", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party.party.party_type.name.text in (
                    "Third-Party Plaintiff",
                    "Cross-Plaintiff",
                    "Cross-Petitioner",
                    "Plaintiff",
                    "Petitioner",
                    "Claimant",
                    "Appellant",
            ):
                output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def all_ds(self):
        output = SCPartyList("all_ds", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party.party.party_type.name.text in (
                    "Third-Party Defendant",
                    "Cross-Defendant",
                    "Cross-Respondent",
                    "Defendant",
                    "Respondent",
            ):
                output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def complainants(self):
        try:
            output = SCPartyList("complainants", there_are_any=True, auto_gather=False)
            output.clear()
            for party in self.all_ds():
                party.party.is_complainant = False
            for party in self.all_ps():
                if party.party.is_complainant:
                    if party not in output.elements:
                        output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w complainants(self)")

    def warrantors(self):
        try:
            output = SCPartyList("warrantors", there_are_any=True, auto_gather=False)
            output.clear()
            for party in self.parties:
                if (
                        any(y.was_physically_injured for y in self.parties)
                        or self.product.was_a_consumer_good_to_be_sold_at_retail
                ):
                    if (
                            party.party.is_manufacturer
                            or party.party.is_wholesaler
                            or party.party.is_seller
                            and party not in output.elements
                    ):
                        output.append(party)
                else:
                    if party.party.is_seller and party not in output.elements:
                        output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w warrantors(self)")

    def product_seller(self):
        try:
            output = SCPartyList(
                "product_seller", there_are_any=True, auto_gather=False
            )
            output.clear()
            for party in self.parties:
                if party.party.is_seller and party not in output.elements:
                    output.append(party)
            output.there_is_another = False
            output.gathered = True
            return output
        except:
            log("issue w product_seller(self)")

    def adverse_parties(self):
        if any(y in self.the_ps() for y in self.user_clientlist()):
            if not self.the_ds():
                raise Exception("you got ps w no ds!")
            return self.the_ds()
        elif any(y in self.the_ds() for y in self.user_clientlist()):
            if not self.the_ps():
                raise Exception("you got ds w no ps!")
            return self.the_ps()
        elif any(y in self.threedp_ps() for y in self.user_clientlist()):
            return self.threedp_ds()
        elif any(y in self.threedp_ds() for y in self.user_clientlist()):
            return self.threedp_ps()
        elif any(y in self.cross_ps() for y in self.user_clientlist()):
            return self.cross_ds()
        elif any(y in self.cross_ds() for y in self.user_clientlist()):
            return self.parties.cross_ds()

    @property
    def rule_of_completeness(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 106; ORS 45.260")

    @property
    def rulings_on_evidence_in_jury_cases_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 103(3)")

    @property
    def waiver_from_disclosure_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 511")

    @property
    def no_evidence_of_faith(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 610")

    @property
    def expert_discretion_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 705")

    @property
    def balance_evidence_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " 401")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str("OEC 401")

    @property
    def relevancy_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + " 401, 402, and 403")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 401, 402, and 403")

    @property
    def insurance_evidence_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 411")

    @property
    def leading_questions_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.rules_of_evidence[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.rules_of_evidence[0] + " 611(3)")

    @property
    def uniform_jury_instructions(self):
        if self.juris.jurisdiction.name.text == "federal":
            return ("", "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return ("UCJI", "Oregon Uniform Civil Jury Instructions")

    @property
    def precautionary_jury_instructions(self, **kwargs):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.uniform_jury_instructions[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                if "with_title" in kwargs and kwargs["with_title"]:
                    return str(
                        self.uniform_jury_instructions[0] + " 5.01" + self.fx.EM + ""
                    )
                else:
                    return str(self.uniform_jury_instructions[0] + " 5.01")

    @property
    def ability_to_pay_instruction(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.uniform_jury_instructions[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.uniform_jury_instructions[0] + " 16.01")

    @property
    def but_for_instruction(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.uniform_jury_instructions[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.uniform_jury_instructions[0] + " 23.01")

    @property
    def substantial_factor_instruction(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.uniform_jury_instructions[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.uniform_jury_instructions[0] + " 23.02")

    @property
    def permanent_injury_instruction(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.uniform_jury_instructions[0] + "")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.uniform_jury_instructions[0] + " 74.01")

    def rfp_instructions(self):
        output = SCList("rfp_instructions", there_are_any=True, auto_gather=False)
        output.clear()
        if "Negligence" in self.claims:
            if "MVA" in self.negligence_type:
                if all(
                        party.party.mva_status != "Pedestrian" for party in self.user_clientlist()
                ):
                    output.append(
                        str(
                            "The term “collision” refers to the collision between "
                            + self.user_clientlist().pname_possessive()
                            + "vehicle and "
                            + self.recipients().pname_possessive()
                            + "vehicle on "
                            + self.doi
                            + ", which is the subject of this action."
                        )
                    )
                else:
                    output.append(
                        str(
                            "The term “collision” refers to the collision between "
                            + self.recipients().pname_possessive()
                            + "vehicle and "
                            + self.user_clientlist().pname()
                            + " on "
                            + self.doi
                            + ", which is the subject of this action."
                        )
                    )
            if "Animal Attack" in self.negligence_type:
                output.append(
                    str(
                        "The term “the "
                        + self.breed
                        + "” refers to the "
                        + self.breed
                        + ", named “"
                        + self.animal_name
                        + ",” that "
                        + self.animal_owners.asnoun()
                        + " "
                        + self.animal_owners.upper()
                        + " owned and were the possessors of on "
                        + self.doi
                        + "."
                    )
                )
                output.append(
                    str(
                        "The term “attack” refers to the "
                        + possessify(self.breed, self.doi)
                        + ", attack of "
                        + self.user_clientlist().upper()
                        + " which is the subject of this action."
                    )
                )
                output.append(
                    str(
                        "The term “animal nuisance” is used as defined by Deschutes County Code 6.08.50 and Redmond City Code 5.270."
                    )
                )
                output.append(
                    str(
                        "The term “dangerous dog” is used as defined by Deschutes County Code 6.08.020, 6.08.070 and as Redmond City Code 5.250, 5.272 define “dangerous animal.”"
                    )
                )
                output.append(
                    str(
                        "The term “unconfined” is used as defined by Deschutes County Code 6.08.045."
                    )
                )
                output.append(
                    str(
                        "The term “The Premises” refers to the real property and improvements thereon located at "
                        + self.premises_address.address
                        + ", "
                        + str(
                            str(self.premises_address.unit + ", ")
                            if self.premises_address.unit
                            else ""
                        )
                        + self.premises_address.city
                        + ", "
                        + state_name(self.premises_address.state)
                        + " "
                        + self.premises_address.zip
                        + "."
                    )
                )
        elif "Premises Liability" in self.claims:
            output.append(
                str(
                    "The term “incident” refers to "
                    + possessify(self.user_clientlist(), "")
                    + self.doi
                    + ", fall which has been made the subject of this action."
                )
            )
            output.append(
                str(
                    "The term “"
                    + self.hazloc.name.text
                    + "” refers to "
                    + self.hazloc.name.text.lower()
                    + " located on the property at "
                    + self.premises_address.address
                    + ", "
                    + str(
                        str(self.premises_address.unit + ", ")
                        if self.premises_address.unit
                        else ""
                    )
                    + self.premises_address.city
                    + ", "
                    + state_name(self.premises_address.state)
                    + " "
                    + self.premises_address.zip
                    + ", also known as "
                    + self.premises_address.county
                    + " County Map and Taxlot No. "
                    + premises
                    | upper
                    + " (hereinafter referred to as “The Property”, “"
                    + self.premises_address.city
                    + ", "
                    + state_name(self.premises_address.state)
                    + " facility”, or “"
                    + self.premises_address.city
                    + ", "
                    + state_name(self.premises_address.state)
                    + " location”)."
                )
            )
            output.append(
                str(
                    "The term “"
                    + self.the_hazard
                    + "” refers to "
                    + self.the_hazard.lower()
                    + " then existing on The Property on "
                    + self.doi
                    + " as alleged in the Complaint."
                )
            )
            output.append(
                str(
                    "The term “unreasonably dangerous condition” means a condition that cannot be encountered with reasonable safety, even if the danger is known and appreciated."
                )
            )
        else:
            output.append(
                str(
                    "The terms “incident” or “accident” when used herein refers to the factual allegations that are the subject of this litigation."
                )
            )
        output.append(
            str(
                "The words “you” and “your” refer to the "
                + self.recipients().asnoun()
                + " and "
                + self.recipients().pronoun_possessive("")
                + "agents, attorneys, and all other representatives."
            )
        )
        output.append(
            str(
                "The term “person” means the plural as well as the singular and includes any natural person, firm, corporation, association, joint venture, partnership, district, or other entity and any department, group, or section of such entity."
            )
        )
        output.append(
            str(
                "The term “documents” includes all writings, drawings, charts, drafts, notes, photographs, depictions, graphs, tapes, disks, e-mails, phonographic records, electronically stored data, and other data compilations from which information can be obtained.  Further, “documents” includes anything that is in written form or is a tangible recording of speech, sounds, pictures, words or symbols, however produced or reproduced, and including the originals (or any copies when originals are not available), and any other non-identical copies (whether different from the originals because of notes made on such copies, or because of indications that the copies were sent to different individuals than the original or because of any other reason), including, but not limited to, working papers, preliminary, intermediate or final drafts, correspondence, books, pamphlets, memoranda, interoffice memoranda, notes, reports, compilations, computer runs, drawings, charts, indexes, photographic or graphic matter, minutes and records of any sort of meetings, invoices, financial statements, financial calculations, diaries, reports of telephone or other oral conversations, desk calendars, appointment books, maps, scientific data, and any other writing and recordings of every  kind which are in your actual or constructive possession custody or control however produced or reproduced or electronically stored. This includes, *but is not limited to*, any computer generated, computer stored, or electronically stored file of any kind and description, including e-mail messages, notes, final or draft agreements, memoranda, letters or other communications as well as the tapes, cassettes, discs, diskettes, or recordings, and all transcriptions or printouts thereof. **In the case of computer generated, computer stored, or electronically stored files, the documents should be produced in hard copy as well as in an electronic medium with all metadata intact.**."
            )
        )
        output.append(
            str(
                "The term “statements” includes any statements, including but not limited to, audio, video, or written statements and including summaries or notes of statements."
            )
        )
        output.append(
            str(
                "The term “communication” means any correspondence, contact, discussion, or exchange between any two or more persons. Without limiting the foregoing, “communication” includes all documents, telephone conversations, and any means of transmitting a message, face-to-face conversations, meetings, and conferences."
            )
        )
        output.append(
            str(
                "The terms “relate,” “related,” “relating to,” “relation,” “regarding,” and “referring to” mean pertinent, relevant, material to, consisting of, summarizing, describing, mentioning concerning, evidencing, or referring in any way to the subject matter of the inquiry."
            )
        )
        output.append(
            str(
                "The terms “and” and “or” shall be individually interpreted in every instance as including both the conjunctive and disjunctive, and should not be interpreted to exclude any information within the scope of any request."
            )
        )
        output.append(
            str(
                "The terms “any,” “every,” and “all” should be interpreted to be interpreted inclusive of one another.  The term “any” includes the terms “every” and “all;” the term “every” includes the terms “any” and “all;” and the term “all” includes the terms “any” and “every.”"
            )
        )
        output.append(
            str(
                "The documents produced pursuant to this request are to be segregated and identified by the number of the request to which they are responsive."
            )
        )
        output.append(
            str(
                "All documents shall be produced in their entirety (without redaction), together with all attachments, exhibits, cover letters, and the like."
            )
        )
        output.append(
            str(
                "A draft or non-identical copy of a document is a separate document for purposes of these Requests for Production of Documents."
            )
        )
        output.append(
            str(
                "If you claim any responsive documents are not required to be produced as a result of any privilege or other protections including, but not limited to, attorney/client privilege or attorney’s work product, you are requested to identify such documents stating the author, recipients, date, and subject matter with sufficient particularity that a determination can be made with respect to the applicability of any claim or privilege or other protection against discovery."
            )
        )
        output.append(
            str(
                capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "request extends beyond all documents and information within "
                + capitalize(possessify(self.recipients(), "possession"))
                + " to include all documents and information within "
                + capitalize(possessify(self.recipients(), "possession"))
                + " and control and may, therefore, require "
                + self.recipients().asnoun()
                + " or "
                + possessify(self.recipients().asnoun(), "")
                + ""
                + self.recipients_attorneys().as_noun("attorney")
                + " to seek and obtain the specifically requested documents and information."
            )
        )
        output.there_is_another = False
        output.gathered = True
        return output

    def impetra_rrfp_objections(self):
        levl = 0
        output = []
        if self.intro:
            output.append(self.fx.NLINE + item_label(levl, 4) + str(self.other_other))
            levl += 1
        if self.rrfps.overly_broad:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Overly Broad & Unduly Burdensome"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.attorney_client:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Attorney Client Privelege"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.same_bodypart:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Not Same or Similar Bodypart"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.vauge_ambiguous:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Vague and Ambiguous"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.public_record:
            output.append(
                self.fx.NLINE + item_label(levl, 4) + " Public Record" + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.overbroad_time:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Overbroad as to Time"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.defendant_possesses:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Defendant Already Possesses"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.not_reasonably_calculated:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Not Reasonably Calculated to Lead to Discoverable Evidence"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.interrogatory:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Impermissible Interrogatory"
                + self.fx.NLINE
            )
            levl += 1
        if self.rrfps.expert:
            output.append(
                self.fx.NLINE
                + item_label(levl, 4)
                + " Impermissible Expert Discovery"
                + self.fx.NLINE
            )
            levl += 1
            if self.rrfps.additional_text:
                output.append(self.fx.NLINE + item_label(levl, 4) + str(self.rrfps.additional_text))
            levl += 1
        return "".join(output)

    def impetra_rrog_objections(self):
        levl = 0
        output = []
        if self.rrogs.overly_broad:
            output.append(
                "\n\n" + item_label(levl, 4) + " Overly Broad & Unduly Burdensome\n\n"
            )
            levl += 1
        if self.rrogs.not_reasonably_calculated:
            output.append(
                "\n\n"
                + item_label(levl, 4)
                + " Not Reasonably Calculated to Lead to Discoverable Evidence\n\n"
            )
            levl += 1
        if self.rrogs.attorney_client:
            output.append(
                "\n\n" + item_label(levl, 4) + " Attorney Client Privelege\n\n"
            )
            levl += 1
        if self.rrogs.cumulative:
            output.append(
                "\n\n"
                + item_label(levl, 4)
                + " Unreasonably Cumulative and Duplicative\n\n"
            )
            levl += 1
        if self.rrogs.same_bodypart:
            output.append(
                "\n\n" + item_label(levl, 4) + " Physician-Patient Privilege\n\n"
            )
            levl += 1
        if self.rrogs.vauge_ambiguous:
            output.append("\n\n" + item_label(levl, 4) + " Vague and Ambiguous\n\n")
            levl += 1
        if self.rrogs.public_record:
            output.append("\n\n" + item_label(levl, 4) + " Public Record\n\n")
            levl += 1
        if self.rrogs.defendant_possesses:
            output.append(
                "\n\n"
                + item_label(levl, 4)
                + " Info Already Known by "
                + self.recipients().asnoun()
                + "\n\n"
            )
            levl += 1
        if self.rrogs.not_medically_stationary:
            output.append(
                "\n\n"
                + item_label(levl, 4)
                + self.user_clientlist().asnoun()
                + " Isn't Medically Stationary \n\n"
            )
            levl += 1
        if self.rrogs.seeks_docs:
            output.append(
                "\n\n"
                + item_label(levl, 4)
                + " Seeks Production of Documents or Requests a Complete Recital of the Contents of Documents”\n\n"
            )
            levl += 1
        if self.rrogs.expert:
            output.append(
                "\n\n" + item_label(levl, 4) + " Impermissible Expert Discovery\n\n"
            )
            levl += 1
        if self.rrogs.legal_conclusion:
            output.append(
                "\n\n" + item_label(levl, 4) + "Calls for a Pure Legal Conclusion\n\n"
            )
            levl += 1
        return "".join(output)

    def facite_rrogs(self):
        indexr = Counter()
        output = list()
        listout = []
        rog_val = self.rrogs.prior_requests_total
        for item in self.rrogs:
            output.clear()
            if item.overly_broad:
                output.append(
                    self.user_clientlist().asnounobjects()
                    + "to this interrogatory to the extent that it is overly broad"
                )
                if item.not_reasonably_calculated:
                    output.append(
                        ", unduly burdensome, and not reasonably calculated to lead to the discovery of admissible evidence. "
                    )
                else:
                    output.append(" and unduly burdensome. ")
            if item.not_reasonably_calculated and not item.overly_broad:
                output.append(
                    self.user_clientlist().asnounobjects()
                    + "to this interrogatory as not reasonably calculated to lead to the discovery of admissible evidence. "
                )

            if item.cumulative:
                if item.overly_broad or item.not_reasonably_calculated:
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as the information sought is unreasonably cumulative and duplicative, "
                )
                if item.obtainable_rrfp:
                    output.append(
                        "as it is obtainable from documents provided in response to "
                        + possessify(self.recipients().asnoun(), "")
                        + " Request for Production, "
                    )
                elif item.obtainable_others:
                    output.append(
                        "as the information sought is obtainable from documents already in Defendant’s possession and/or which will be provided pursuant to subpoena of third parties, "
                    )
                output.append(
                    "which is a more convenient and less burdensome source of the requested information. "
                )

            if item.work_product:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "seeks information that is protected by the “work product” privilege. The interrogatory calls for disclosure of information or materials prepared in anticipation of litigation and/or trial preparation material without the showing required to discover this information or material. "
                )

            if item.attorney_client:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as calling for information protected by the attorney client privilege. "
                )

            if item.same_bodypart:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as calling for physician-patient privileged information. "
                )

            if item.vauge_ambiguous:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append("to this request as vague and/or ambiguous. ")

            if item.public_record:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.attorney_client
                        or item.work_product
                        or item.same_bodypart
                        or item.vauge_ambiguous
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as calling for public information equally available to "
                    + self.recipients().asnoun()
                    + ". "
                )

            if item.defendant_possesses:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.attorney_client
                        or item.work_product
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as calling for information already known by requesting defendant. "
                )

            if item.not_medically_stationary:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.work_product
                        or item.cumulative
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to the extent that "
                    + self.user_clientlist().pronoun_subjective()
                    + " is not medically stationary, and continues to treat for "
                    + self.user_clientlist().pronoun_possessive("injuries")
                )
                if item.incomplete_cns:
                    output.append(
                        ", and because "
                        + self.user_clientlist().pronoun_subjective()
                        + " may not have in "
                        + self.user_clientlist().pronoun_possessive(
                            "custody, possession, or control a complete set of all medical chart notes incurred to date. "
                        )
                    )
                elif item.changing_symptoms:
                    output.append(" with symptoms that continue to change over time. ")

            if item.expert:
                if item.no_expert_yet:
                    output.append(
                        "In response to this interrogatory requesting information about experts, no determination as to who may be called as an expert witness has yet been made. This answer will be supplemented when and if such information is known. "
                    )
                if item.expert_outside_frcp_26_b4Ai:
                    if (
                            item.overly_broad
                            or item.not_reasonably_calculated
                            or item.cumulative
                            or item.work_product
                            or item.attorney_client
                            or item.same_bodypart
                            or item.vauge_ambiguous
                            or item.public_record
                            or item.defendant_possesses
                            or item.not_medically_stationary
                    ):
                        output.append(self.user_clientlist().asnounobjects_further())
                    else:
                        output.append(self.user_clientlist().asnounobjects())
                    output.append(
                        "to those parts of the interrogatory which seek the expert’s qualifications, social or professional relationship to "
                        + self.user_clientlist().asnoun()
                        + ", a bibliography which the expert considers authoritative, a bibliography of all published works of the expert, the manner in which the expert became familiar with the facts of the case, and a list of all cases in which the expert has testified on the grounds that the information requested is outside the scope of FRCP 26(b)(4)(A)(i). "
                    )
                if item.experts_are_work_product:
                    if (
                            item.overly_broad
                            or item.not_reasonably_calculated
                            or item.cumulative
                            or item.work_product
                            or item.attorney_client
                            or item.same_bodypart
                            or item.vauge_ambiguous
                            or item.public_record
                            or item.defendant_possesses
                            or item.not_medically_stationary
                    ):
                        output.append(self.user_clientlist().asnounobjects_further())
                    else:
                        output.append(self.user_clientlist().asnounobjects())
                    output.append(
                        "to this interrogatory on the grounds that the employment of such persons, if any, occurred in anticipation of litigation, constitute work product of counsel and are not discoverable, and further that such information is specifically protected from discovery with specific exceptions, pursuant to FRCP 26(b)(4)(B). "
                    )
                if item.expert_before_expert_disclosure:
                    if (
                            item.overly_broad
                            or item.not_reasonably_calculated
                            or item.cumulative
                            or item.work_product
                            or item.attorney_client
                            or item.same_bodypart
                            or item.vauge_ambiguous
                            or item.public_record
                            or item.defendant_possesses
                            or item.not_medically_stationary
                    ):
                        output.append(self.user_clientlist().asnounobjects_further())
                    else:
                        output.append(self.user_clientlist().asnounobjects())
                    output.append(
                        "as calling for expert opinion prior to the date for expert disclosures. "
                    )

            if item.seeks_docs:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to the extent that the interrogatory seeks to obtain the production of documents or to request a complete recital of the contents of documents. "
                )

            if item.legal_conclusion:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to the extent that the interrogatory calls for a pure legal conclusion. "
                )

            if item.annoying_embarrassing_oppressive:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as the interrogatory seeks information that is annoying, embarrassing, or oppressive. "
                )

            if item.repetitive_duplicative:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as the interrogatory seeks information that is repetitive or duplicative. "
                )

            if item.unknown_to_me:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "because this information not within the knowledge, possession, custody or control of "
                    + self.user_clientlist().asnoun()
                    + ". "
                )

            if item.outside_scope_of_rules:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as this interrogatory is beyond the scope and requirements of the "
                    + self.civpro[0]
                    + ". "
                )

            if item.premature:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as this interrogatory is premature to the extent that it requests the "
                    + self.user_clientlist().asnoun()
                    + ", prior to the completion of discovery, to state all the facts supporting "
                    + self.user_clientlist().pronoun_possessive("")
                    + "current contentions and to speculate as to "
                    + self.user_clientlist().pronoun_possessive("")
                    + "future contentions. "
                )

            if item.not_limited_in_time:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it is not limited in time to the years of the alleged occurrence. "
                )

            if item.created_after_filing_of_lawsuit:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as the information sought was generated or authored subsequent to the commencement of this court action. "
                )

            if item.seeks_witness_statement:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it seeks statements made by witnesses without the requisite showing of substantial need and undue hardship. "
                )

            if item.privacy_rights_of_others:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                        or item.created_after_filing_of_lawsuit
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it seeks disclosure of information where such disclosure would violate the privacy rights of individuals. "
                )

            if item.confidentiality_agreement:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                        or item.created_after_filing_of_lawsuit
                        or item.privacy_rights_of_others
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it seeks information relating to any confidentiality agreement between "
                    + self.recipients().asnoun()
                    + " and another. "
                )

            if item.protective_order:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                        or item.created_after_filing_of_lawsuit
                        or item.privacy_rights_of_others
                        or item.confidentiality_agreement
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it seeks information which would violate a court order restricting the disclosure of information. "
                )

            if item.trade_secret:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                        or item.created_after_filing_of_lawsuit
                        or item.privacy_rights_of_others
                        or item.confidentiality_agreement
                        or item.protective_order
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this interrogatory as it seeks information which could result in the disclosure of confidential commercial information, trade secrets, or proprietary information. "
                )

            if item.not_relevant:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.cumulative
                        or item.work_product
                        or item.attorney_client
                        or item.same_bodypart
                        or item.vauge_ambiguous
                        or item.public_record
                        or item.defendant_possesses
                        or item.not_medically_stationary
                        or item.expert
                        or item.seeks_docs
                        or item.legal_conclusion
                        or item.annoying_embarrassing_oppressive
                        or item.repetitive_duplicative
                        or item.unknown_to_me
                        or item.outside_scope_of_rules
                        or item.not_limited_in_time
                        or item.created_after_filing_of_lawsuit
                        or item.created_after_filing_of_lawsuit
                        or item.privacy_rights_of_others
                        or item.confidentiality_agreement
                        or item.protective_order
                        or item.trade_secret
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "as the information sought is not relevant to any issue in this action. "
                )

            if (
                    item.overly_broad
                    or item.not_reasonably_calculated
                    or item.cumulative
                    or item.work_product
                    or item.attorney_client
                    or item.same_bodypart
                    or item.vauge_ambiguous
                    or item.public_record
                    or item.defendant_possesses
                    or item.not_medically_stationary
                    or item.expert
                    or item.seeks_docs
                    or item.legal_conclusion
                    or item.annoying_embarrassing_oppressive
                    or item.repetitive_duplicative
                    or item.unknown_to_me
                    or item.outside_scope_of_rules
                    or item.not_limited_in_time
                    or item.created_after_filing_of_lawsuit
                    or item.created_after_filing_of_lawsuit
                    or item.privacy_rights_of_others
                    or item.confidentiality_agreement
                    or item.protective_order
                    or item.trade_secret
            ):
                output.append("Subject to")
                if item.attorney_client:
                    output.append(", and without waiving,")
                output.append(" said objections, " + fix_punctuation(item.answer))
            else:
                output.append(capitalize(item.answer))
            textout = "".join(output)
            listout.append(
                (
                    str(rog_val),
                    fix_punctuation(single_paragraph(item.name.text)),
                    textout,
                )
            )
            rog_val += 1
        return listout

    def facite_rrfas(self):
        indexr = Counter()
        output = list()
        listout = []
        output.clear()
        rfpval = self.rrfas.prior_requests_total
        for item in self.rrfas:
            if item.undefined_subjective:
                output.append(
                    self.user_clientlist().asnounobjects()
                    + "to this request as the "
                    + str("phrase " if " " in item.bad_term.strip() else "term ")
                    + quote(item.bad_term)
                    + " is undefined and subjective, making this request vague and ambiguous. "
                )
            if item.attorney_client:
                if item.undefined_subjective:
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as calling for information protected under the attorney client privilege and/or attorney work product doctrine. "
                )
            if item.reasonable_inquiry:
                output.append(
                    "Reasonable inquiry has been made, and the information known or readily obtainable by "
                    + self.user_clientlist().asnoun()
                    + " is insufficient to enable "
                    + self.user_clientlist().asnoun()
                    + " to admit or deny this request. "
                )
            if item.expert:
                if item.undefined_subjective or item.attorney_client:
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as calling for an undiscoverable expert opinion. "
                )
            if item.vauge_ambiguous:
                if item.undefined_subjective or item.attorney_client or item.expert:
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as ambiguous, unclear, and confusing, precluding "
                    + self.user_clientlist().asnoun()
                    + " from sufficiently being able to admit or deny this request. "
                )
            if (
                    item.undefined_subjective
                    or item.attorney_client
                    or item.expert
                    or item.vauge_ambiguous
            ):
                output.append("Subject to")
                if item.attorney_client:
                    output.append(", and without waiving,")
                output.append(" said objections, " + item.response)
            else:
                output.append(capitalize(item.response))
            textout = "".join(output)
            listout.append(
                [
                    str(rfp_val),
                    fix_punctuation(single_paragraph(item.name.text)),
                    textout,
                ]
            )
            rfp_val += 1
        return listout

    def facite_rrfps(self):
        indexr = Counter()
        rfpval = self.rrfps.prior_requests_total
        output = list()
        listout = []
        for item in self.rrfps:
            output.clear()
            # if item.intro:
            #  output.append(fix_punctuation(item.other_other) + ' ')
            if item.overly_broad:
                output.append(
                    self.user_clientlist().asnounobjects()
                    + "to this request to the extent that it is overly broad"
                )
                if item.not_reasonably_calculated:
                    output.append(
                        ", unduly burdensome, and not reasonably calculated to lead to the discovery of admissible evidence. "
                    )
                else:
                    output.append(" and unduly burdensome.")
            if item.not_reasonably_calculated and not item.overly_broad:
                output.append(
                    self.user_clientlist().asnounobjects()
                    + "to this request as not reasonably calculated to lead to the discovery of admissible evidence. "
                )
            if item.attorney_client:
                if item.overly_broad or item.not_reasonably_calculated:
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request to the extent that it seeks documents protected by the attorney client privilege and/or attorney work product doctrine. "
                )
            if item.same_bodypart:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(str(
                    "to this request as not limited in scope as to the same or similar body parts alleged to have been injured in the Complaint, and requests documents, should any exist, that are confidential under HIPAA and the " +
                    self.rules_of_evidence[
                        0] + " including, but not limited to, the physician-patient privilege, psychotherapist-patient privilege, nurse-patient privilege and/or counselor-client privilege. "
                ))
            if item.public_record:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as calling for a classification of documents that would be public record and, as such, equally available to the "
                )
                output.append(self.recipients().asnoun())
                output.append(" requesting them. ")
            if item.vauge_ambiguous:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append("to this request as vague and/or ambiguous. ")
            if item.overbroad_time:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append("to this request as overbroad as to time. ")
            if item.defendant_possesses:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking documents already within the control of the "
                )
                output.append(possessify(self.recipients().asnoun(), ""))
                output.append("requesting them. ")
            if item.interrogatory:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as constituting an impermissible interrogatory. "
                )
            if item.expert:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                        or item.interrogatory
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking impermissible expert discovery. "
                )
            if item.biz_docs:
                if (
                        item.overly_broad
                        or item.not_reasonably_calculated
                        or item.attorney_client
                        or item.same_bodypart
                        or item.public_record
                        or item.vauge_ambiguous
                        or item.overbroad_time
                        or item.defendant_possesses
                        or item.interrogatory
                        or item.expert
                ):
                    output.append(self.user_clientlist().asnounobjects_further())
                else:
                    output.append(self.user_clientlist().asnounobjects())
                output.append(
                    "to this request as seeking highly confidential business documents. "
                )
            # if item.text_response:
            #    output.append(fix_punctuation(item.text_response) + ' ')
            if (
                    item.overly_broad
                    or item.not_reasonably_calculated
                    or item.attorney_client
                    or item.same_bodypart
                    or item.public_record
                    or item.vauge_ambiguous
                    or item.overbroad_time
                    or item.defendant_possesses
                    or item.interrogatory
                    or item.expert
                    or item.biz_docs
            ):
                output.append("Subject to")
                if item.attorney_client:
                    output.append(", and without waiving,")
                output.append(" said objections, " + item.responsive)
                if item.responsive == "see response to RFP No.":
                    output.append(" " + str(item.see_rfp) + ". ")
            else:
                output.append(capitalize(item.responsive))
                if item.responsive == "see response to RFP No.":
                    output.append(" " + str(item.see_rfp) + ". ")
            textout = "".join(output)
            listout.append(
                [
                    str(rfpval),
                    fix_punctuation(single_paragraph(item.name.text)),
                    textout,
                ]
            )
            rfpval += 1
        return listout

    def outside_cofirms(self):
        otherput = SCList("outside_cofirms", there_are_any=True, auto_gather=False)
        otherput.clear()
        for firm in self.lawfirms:
            if firm not in otherput.elements:
                output = SCList(
                    "outside_cofirms_output", there_are_any=True, auto_gather=False
                )
                output.clear()
                otherput.append(firm)
                for lawfirm in self.lawfirms:
                    if (
                            lawfirm is not firm
                            and self.clientlist(lawfirm) == self.clientlist(firm)
                            and lawfirm not in output.elements
                    ):
                        output.append(lawfirm)
                output.there_is_another = False
                output.gathered = True
                otherput.append(output)
        otherput.there_is_another = False
        otherput.gathered = True
        return otherput

    def clientlist(self, firm):
        output = SCPartyList("client_of", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if firm.name.full() in comma_and_list(party.party.firms):
                output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    @property
    def implied_warranty_of_merchantability_rule(self):
        if state_name(self.juris.state.name.text) == "Oregon":
            output = "ORS 72.3140"
        else:
            output = "ORS 72.3140"
        return output

    @property
    def implied_warranty_of_fitness_for_a_particular_purpose_rule(self):
        if state_name(self.juris.state.name.text) == "Oregon":
            output = "ORS 72.3140"
        else:
            output = "ORS 72.3140"
        return output

    # Every Jurisdiction has its own rules of civil procedure. Here, I will encode the official full name of every jurisdiction's rules, as well as the singular version of it, and the abbreviation (official abbreviations where available will likely do, but whatever abbreviation is used by practitioners in that jurisdiction is the ideal and the ultimate goal for this data set. Speaking if sets, I initially and currently set up this evaluation as a series if elif statements based on the jurisdiction's state name, which then evaluate to a set of the 3 datapoints. Possible it would be better or more efficient to set this up as a dictionary of sets where the state name is the key for each set, but I will need to research this question. Similarly, it may be better to use lists for the data rather than sets (and i suspect this may be true), but again, more research is warrnated. Finally, what follows this method are many more methods which do the same evaluation, but to bring up individual rule numbers. A person can begin to evaluate/comprehend the resultant rule structure using the federal rules as a reference point (because most states' rules are based upon the federal rules, even if they happen to be numbered differently!). Don't rely too hard on this though, as even jurisdictions that are almost entirely regurgitations of the federal rules can have important distinctions. For example, Oregon rules are almost idebtical to federal rules, but interrogatories are conspicuously missing from the Oregon Rules of Civil Procedure (because rhe Oregon legislature apparently didnt like interrogatories), so interrogatories just don't exist in oregon, so the value for that rule is None.
    @property
    def civpro(self):
        if self.juris.jurisdiction.name.text == "federal":
            output = (
                "Federal Rule of Civil Procedure",
                "FRCP",
                "Federal Rules of Civil Procedure",
            )
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                output = (
                    "Oregon Rule of Civil Procedure",
                    "ORCP",
                    "Oregon Rules of Civil Procedure",
                )
            if state_name(self.juris.state.name.text) == "Alabama":
                output = (
                    "Alabama Rule of Civil Procedure",
                    "Ala. R. Civ. P.",
                    "Alabama Rules of Civil Procedure",
                )
            elif state_name(self.juris.state.name.text) == "Alaska":
                output = ("Alaska Civil Rule", "Rule", "Alaska Civil Rules")
            elif state_name(self.juris.state.name.text) == "Arizona":
                output = (
                    "Arizona Rule of Civil Procedure",
                    "Ariz.R.Civ.P.",
                    "Arizona Rules of Civil Procedure",
                )
            elif state_name(self.juris.state.name.text) == "Arkansas":
                output = (
                    "Arkansas Rule of Civil Procedure",
                    "Ark. R. Civ. P.",
                    "Arkansas Rules of Civil Procedure",
                )

            elif state_name(self.juris.state.name.text) == "California":
                output = (
                    "Code of Civil Procedure",
                    "C.C.P.",
                    "Code of Civil Procedure",
                )

            elif state_name(self.juris.state.name.text) == "Colorado":
                output = (
                    "Colorado Rule of Civil Procedure",
                    "C.R.C.P.",
                    "Colorado Rules of Civil Procedure",
                )
            elif state_name(self.juris.state.name.text) == "Connecticut":
                output = ("Practice Book Section", "P.B. Sec.", "Practice Book")
            elif state_name(self.juris.state.name.text) == "Delaware":
                output = (
                    "Superior Court Rule of Civil Procedure",
                    "Del. R. Civ. P. Super. Ct.",
                    "Superior Court Rules of Civil Procedure",
                )
            elif state_name(self.juris.state.name.text) == "Florida":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Georgia":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Hawaii":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Idaho":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Illinois":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Indiana":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Iowa":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Kansas":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Kentucky":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Louisiana":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Maine":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Maryland":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Massachusetts":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Michigan":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Minnesota":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Mississippi":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Missouri":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Montana":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "New Hampshire":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "New Jersey":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "New Mexico":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "New York":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "North Carolina":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "North Dakota":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Ohio":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Oklahoma":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Pennsylvania":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Puerto Rico":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Rhode Island":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "South Carolina":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "South Dakota":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Tennessee":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Texas":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Utah":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Vermont":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Virginia":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Washington":
                output = (
                    "Superior Court Civil Rule",
                    "CR",
                    "Superior Court Civil Rule",
                )
            elif state_name(self.juris.state.name.text) == "West Virginia":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Wisconsin":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Wyoming":
                output = ("", "")
            elif state_name(self.juris.state.name.text) == "Nevada":
                output = ("Nevada Rules of Civil Procedure", "NRCP")
            elif state_name(self.juris.state.name.text) == "District of Columbia":
                output = ("", "")
            else:
                output = (
                    str(
                        state_name(self.juris.state.name.text)
                        + " Rules of Civil Procedure"
                    ),
                    str(state_name(self.juris.state.name.text)[0] + "RCP"),
                )
        return output

    # ------------/ Begin Civil Procedure Rules \  ------------ #

    @property
    def rfa_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 36")
            else:
                return str(self.civpro[1] + " 36")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 45")
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return str(self.civpro[1] + " 13-22–25")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 36")
            elif state_name(self.juris.state.name.text) == "Washington":
                return str(self.civpro[1] + "36")

    @property
    def rfa_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return "36"
            else:
                return "36"
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return "45"

    @property
    def rfp_intro_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 26, 34, and 35")
            else:
                return str(self.civpro[1] + " 26, 34, and 35")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 36, 43, and 44")

    @property
    def discovery_sanction_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return " "
            else:
                return " "
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return "46"
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 37")

    @property
    def rfa_sanction_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return " "
            else:
                return " "
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return "46 C"

    @property
    def rfp_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["34", "35"]
            else:
                return ["34", "35"]
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["43", "44"]
            elif state_name(self.juris.state.name.text) == "Washington":
                return ["34", "35"]
            else:
                return ["34", "35"]

    @property
    def rog_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["33"]
            else:
                return ["33"]
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["None"]
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return ["13-6–8"]
            elif state_name(self.juris.state.name.text) == "Washington":
                return "33"

    def rog_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 33")
            else:
                return str(self.civpro[1] + " 33")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return None
            elif state_name(self.juris.state.name.text) == "Washington":
                return str(self.civpro[1] + " 33")
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return str(self.civpro[1] + " 13-6–8")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 33")

    @property
    def pleading_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " ")
            else:
                return str(self.civpro[1] + " ")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 18")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 8")

    @property
    def general_discovery_rules(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["26"]
            else:
                return ["26"]
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return ["36"]
            elif state_name(self.juris.state.name.text) == "Washington":
                return ["26"]
            else:
                return ["26"]

    @property
    def general_discovery_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 26")
            else:
                return str(self.civpro[1] + " 26")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 36")
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return str(self.civpro[1] + " 13-3–5")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 26")
            elif state_name(self.juris.state.name.text) == "Washington":
                return str(self.civpro[1] + "26")

    @property
    def rfp_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 34")
            else:
                return str(self.civpro[1] + " 34")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 43")
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return str(self.civpro[1] + " 13-9–11A")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 34")

    @property
    def instructions_to_jury_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 51; LR 51.1")
            else:
                return str(self.civpro[1] + " 51")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 59 B")

    @property
    def depo_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 30")
            else:
                return str(self.civpro[1] + " 30")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 59 B")
            elif state_name(self.juris.state.name.text) == "Connecticut":
                return str(self.civpro[1] + " 13-26–31")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 30")

    @property
    def summary_jmnt_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.civpro[1] + " 56")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 47")
            elif state_name(self.juris.state.name.text) == "Washington":
                return str(self.civpro[1] + "56")

    @property
    def ime_rule(self):
        if self.juris.jurisdiction.name.text == "federal":
            return str(self.civpro[1] + " 35")
        else:
            if state_name(self.juris.state.name.text) == "Oregon":
                return str(self.civpro[1] + " 44")
            elif state_name(self.juris.state.name.text) == "Delaware":
                return str(self.civpro[1] + " 35")
            elif state_name(self.juris.state.name.text) == "Washington":
                return str(self.civpro[1] + "35")

    # ------------/ End Civil Procedure Rules \  ------------ #

    def general_objections(self):
        output = SCList("general_objections", there_are_any=True, auto_gather=False)
        output.clear()
        output.append(
            str(
                self.user_clientlist().asnounobjects()
                + "to all requests to the extent "
                + self.recipients().asnoun()
                + " "
                + self.recipients().does_verb("seek")
                + " information protected by the attorney-client privilege, work-product doctrine, or any other applicable privilege."
            )
        )
        output.append(
            str(
                self.user_clientlist().asnounobjects()
                + "to any portion of these requests that contain instructions, directions, provisions, or definitions that are inconsistent with or more onerous than the requirements of the "
                + str(
                    "Federal Rules of Civil Procedure "
                    if self.juris.jurisdiction.name.text == "federal"
                    else self.juris.civil_procedure()["official title"]
                )
            )
        )
        output.append(
            str(
                self.user_clientlist().asnounobjects()
                + "to any request to the extent it seeks the production of documents that are not currently in "
                + self.user_clientlist().pronoun_possessive("")
                + "possession, custody, or control."
            )
        )
        output.append(
            str(
                possessify(self.user_clientlist().asnoun(), "")
                + " investigation and discovery are continuing. "
                + self.user_clientlist().asnoun()
                + " therefore "
                + self.user_clientlist().does_verb("reserve")
                + "  the right to supplement "
                + self.user_clientlist().pronoun_possessive("")
                + " responses after additional discovery has taken place."
            )
        )
        output.append(
            str(
                self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("is")
                + " responding in good faith to these requests as "
                + self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("interpret")
                + " and "
                + self.user_clientlist().does_verb("understand")
                + " them. If "
                + self.recipients().asnoun()
                + " subsequently "
                + self.recipients().does_verb("assert")
                + " an interpretation of any request that differs from "
                + possessify(self.user_clientlist().asnoun(), "")
                + "understanding, "
                + self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("reserve")
                + " the right to supplement "
                + self.user_clientlist().pronoun_possessive("")
                + "responses or objections."
            )
        )
        output.append(
            str(
                "Each of these general objections is incorporated into each of "
                + possessify(self.user_clientlist().asnoun(), "")
                + "specific responses below."
            )
        )
        output.there_is_another = False
        output.gathered = True
        return output

    def attorneys_who_work_at(self, lawfirm):
        log("attorneys_who_work_at1")
        attorneys_who_work_at_output = SCList(
            "attorneys_who_work_at", there_are_any=True, auto_gather=False
        )
        log("attorneys_who_work_at2")
        attorneys_who_work_at_output.clear()
        log("attorneys_who_work_at3")
        for attorney in self.lawyers:
            log("attorneys_who_work_at4")
            if attorney.lawyer.firm.id == lawfirm.id:
                log("attorneys_who_work_at5")
                attorneys_who_work_at_output.append(attorney.lawyer)
                log("attorneys_who_work_at6")
        attorneys_who_work_at_output.there_is_another = False
        log("attorneys_who_work_at7")
        attorneys_who_work_at_output.gathered = True
        log("attorneys_who_work_at8")
        return attorneys_who_work_at_output

    @property
    def ourhouse(self):
        ourhouse_output = SCList(
            "ourhouse", there_are_any=True, auto_gather=False
        )
        ourhouse_output.clear()
        log("ourhouse1")
        for lawyer in self.lawyers:
            if lawyer.lawyer.firm.id == self.user_firm().id:
                log("ourhouse2")
                ourhouse_output.append(lawyer.lawyer)
                log("ourhouse3")
        ourhouse_output.there_is_another = False
        log("ourhouse4")
        ourhouse_output.gathered = True
        return ourhouse_output

    @property
    def is_userfirm(self):
        for lawfirm in self.lawfirms:
            if lawfirm.lawfirm.name.text == "Scott M Cumming, PC":
                lawfirm.lawfirm.is_userfirm = True
            else:
                lawfirm.lawfirm.is_userfirm = False

    def user_firm(self):
        for lawfirm in self.lawfirms:
            if lawfirm.lawfirm.name.text == "Scott M Cumming, PC":
                return lawfirm.lawfirm

    def fed_head(self):
        output = ""
        for atty in self.user_house().elements:
            output += (
                    self.fx.BOLD
                    + atty.name.full().upper()
                    + self.fx.BOLD
                    + ", "
                    + atty.bar()
                    + self.fx.NLINE)
            if atty.barred_in != self.juris.state.name.text:
                output += str(self.fx.ITALIC + "pro hac vice" + self.fx.ITALIC + self.fx.NLINE)
            output += str(
                "["
                + atty.email.lower()
                + "]("
                + atty.email.lower()
                + ")"
                + self.fx.NLINE
            )

        output += (
                self.user_firm().name.text
                + self.fx.NLINE
                + self.fx.SINGLE
                + self.user_firm().address.address
                + str(
            str(", " + self.user_firm().address.unit)
            if self.user_firm().address.unit
            else ""
        )
                + self.fx.NLINE
                + self.user_firm().address.city
                + ", "
                + state_name(self.user_firm().address.state)
                + " "
                + self.user_firm().address.zip
                + self.fx.NLINE
                + "Phone: "
                # + phone_number_formatted(self.user_firm().phone_number)
                + self.fx.NLINE
                + "Fax: "
                # + phone_number_formatted(self.user_firm().fax_number)
                + self.fx.NLINE
        )
        # if self.lawyers.gathered == True and self.parties.gathered == True and self.lawfirms.gathered == True:
        for pair in self.cofirms().elements:
            for atty in self.attorneys_who_work_at(pair).elements:
                if atty.barred_in != self.juris.state.name.text:

                    output += (
                            self.fx.BOLD
                            + atty.name.full().upper()
                            + self.fx.BOLD
                            + ", "
                            + atty.bar()
                            + self.fx.NLINE
                            + self.fx.ITALIC
                            + "pro hac vice"
                            + self.fx.ITALIC
                            + self.fx.NLINE
                            + "["
                            + atty.email.lower()
                            + "]("
                            + atty.email.lower()
                            + ")"
                            + self.fx.NLINE
                    )

                else:

                    output += (
                            self.fx.BOLD
                            + atty.name.full().upper()
                            + self.fx.BOLD
                            + ", "
                            + atty.bar()
                            + self.fx.NLINE
                            + "["
                            + atty.email.lower()
                            + "]("
                            + atty.email.lower()
                            + ")"
                            + self.fx.NLINE
                    )

            output += (
                    pair.name.text
                    + self.fx.NLINE
                    + pair.address.address
                    + str(str(", " + pair.address.unit) if pair.address.unit else "")
                    + self.fx.NLINE
                    + pair.address.city
                    + ", "
                    + state_name(pair.address.state)
                    + " "
                    + pair.address.zip
                    + self.fx.NLINE
                    + "Phone: "
                    + phone_number_formatted(pair.phone_number)
                    + self.fx.NLINE
                    + "Fax: "
                    + phone_number_formatted(pair.fax_number)
                    + self.fx.NLINE
            )
        # if self.lawyers.gathered == True and self.parties.gathered == True and self.lawfirms.gathered == True:
        output += (
                capitalize(self.colawyers().as_noun("attorney"))
                + " for "
                + self.user_clientlist().asnoun()
        )
        if this_thread.evaluation_context == "docx":
            return output
        else:
            return str(
                '<p style="font-size: 1rem">'
                + markdown_to_html(output)
                + "</p>"
                + self.fx.NLINE
            )

    def colawyers(self):
        output = SCLawyerList("colawyers", there_are_any=True, auto_gather=False)
        output.clear()
        for relationship in self.lawyer_lawfirm_rel:
            if (
                    relationship.lawfirm in self.cofirms()
                    and relationship.lawyer not in output.elements
            ):
                output.append(relationship.lawyer)
        for lawyer in self.ourhouse:
            if lawyer not in output.elements:
                output.append(lawyer)
        output.there_is_another = False
        output.gathered = True
        return output

    def user_clientlist(self):
        log("user_clientlist 1")
        output = self.get_lawfirm_parties(self.user_firm())
        log("user_clientlist 2")
        if not output:
            raise Exception(
                "user_clientlist is empty. Could be because none of case.lawfirms .is_userfirm or because there are no parties related to both the userfirm and this case. ")
        return output

    def user_house(self):
        output = SCLawyerList("user_house", there_are_any=True, auto_gather=False)
        output.clear()
        for relationship in self.lawyer_lawfirm_rel:
            if (
                    relationship.lawfirm.is_userfirm
                    and relationship.lawyer not in output.elements
            ):
                output.append(relationship.lawyer)
        output.there_is_another = False
        output.gathered = True
        return output

    def cofirms(self):
        output = SCList("cofirms", there_are_any=True, auto_gather=False)
        output.clear()
        for relationship in self.firm_client:
            if not relationship.lawfirm.is_userfirm:
                if (
                        relationship.client in self.user_clientlist()
                        and relationship.lawfirm not in output.elements
                ):
                    output.append(relationship.lawfirm)
        output.there_is_another = False
        output.gathered = True
        return output

    def recipients(self):
        recipients_output = SCList("recipients_output", there_are_any=True, auto_gather=False)
        recipients_output.clear()
        for relationship in self.firm_client:
            if relationship.client not in self.user_clientlist():
                if (
                        len(self.parties) == 2
                        and relationship.client not in recipients_output.elements
                ):
                    recipients_output.append(relationship.client)
                elif (
                        relationship.client.is_recipient
                        and relationship.client not in recipients_output.elements
                ):
                    recipients_output.append(relationship.client)
        recipients_output.there_is_another = False
        recipients_output.gathered = True
        return recipients_output

    def recipients_firms(self):
        output = SCLawfirmList(
            "recipients_firms", there_are_any=True, auto_gather=False
        )
        output.clear()
        for relationship in self.firm_client:
            if relationship.client not in self.user_clientlist():
                if (
                        len(self.parties) == 2
                        and relationship.lawfirm not in output.elements
                ):
                    output.append(relationship.lawfirm)
                elif (
                        relationship.client.is_recipient
                        and relationship.lawfirm not in output.elements
                ):
                    output.append(relationship.lawfirm)
        output.there_is_another = False
        output.gathered = True
        return output

    def recipients_attorneys(self):
        output = SCLawyerList(
            "recipients_attorneys", there_are_any=True, auto_gather=False
        )
        output.clear()
        for relationship in self.lawyer_lawfirm_rel:
            if (
                    relationship.lawfirm in self.recipients_firms()
                    and relationship.lawyer not in output.elements
            ):
                output.append(relationship.lawyer)
        output.there_is_another = False
        output.gathered = True
        return output

    def relationship_list(self, parties, lawyers):
        relationship_list = []
        for xx in self.parties:
            if not hasattr(xx, "relationship_id"):
                relationship_list.append(
                    {"description": "Party", "contact": {"id": xx.id}}
                )
        for xx in lawyers:
            if not hasattr(xx, "relationship_id"):
                relationship_list.append(
                    {"description": "Lawyer", "contact": {"id": xx.id}}
                )

    def payload(self):
        payload = {
            "data": {
                "custom_field_values": [
                    {
                        "value": str(self.doi),
                        "custom_field": {"id": self.doi_id},
                        "id": self.doi_vid,
                    },
                    {
                        "value": self.docket_number,
                        "custom_field": {"id": self.docket_number_id},
                        "id": self.docket_number_vid,
                    },
                    {
                        "value": self.juris.jurisdiction.name.text,
                        "custom_field": {"id": self.juris.jurisdiction_id},
                        "id": self.juris.jurisdiction_vid,
                    },
                    {
                        "value": self.juris.district.name.text,
                        "custom_field": {"id": self.matter_ids_dict["district"]},
                        "id": self.juris.district_vid,
                    },
                    {
                        "value": self.juris.division.name.text,
                        "custom_field": {"id": self.juris.division_id},
                        "id": self.juris.division_vid,
                    },
                    {
                        "value": self.juris.state.name.text,
                        "custom_field": {"id": self.juris.state_id},
                        "id": self.juris.state_vid,
                    },
                    {
                        "value": self.juris.county.name.text,
                        "custom_field": {"id": self.juris.county_id},
                        "id": self.juris.county_vid,
                    },
                ]
            }
        }
        return payload

    def doc_header(self):
        output = ""
        if self.juris.jurisdiction.name.text == "federal":
            output += str(
                str(
                    self.fx.CENTER + str("United States District Court"
                                         ).upper())
            )
            if len(self.juris.trial_court().keys()) > 1:
                output += str(
                    str(self.fx.NLINE + self.fx.CENTER + self.juris.district.name.text + " ").upper() + self.fx.NLINE + self.fx.CENTER)
            else:
                output += str(str(self.fx.NLINE + self.fx.CENTER))
            output += str(
                str(
                    self.fx.CENTER
                    + "District of "
                    + state_name(self.juris.state.name.text)
                ).upper()
            )
            if len(self.juris.trial_court()[str(self.juris.district.name.text)]) > 1:
                output += str(
                    str(
                        self.fx.NLINE
                        + self.fx.CENTER
                        + self.juris.division.name.text
                        + " Division"

                    ).upper()
                )
        else:
            # California and Colorado both have apparently unique caption forms, which seem to be required and may or may not be required in California. I should research this issue further.
            if state_name(self.juris.state.name.text) == "Oregon":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF THE STATE OF OREGON"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR THE COUNTY OF "
                        + self.juris.county.name.text.upper()
                    )
                )

            elif state_name(self.juris.state.name.text) == 'Alabama':
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + ", ALABAMA"
                    )
                )

            elif state_name(self.juris.state.name.text) == 'Alaska':
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR THE STATE OF ALASKA"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "AT "
                        + self.juris.county.name.text.upper()
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == 'Arizona':
                output += str(
                    str(
                        self.fx.CENTER
                        # + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF ARIZONA"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        # + "IN "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == 'Arkansas':
                output += str(
                    str(
                        self.fx.CENTER
                        + "THE CIRCUIT COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, ARKANSAS"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "CIVIL DIVISION"
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == 'California':
                output += str(
                    str(
                        self.fx.CENTER
                        + "SUPERIOR COURT OF CALIFORNIA"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + " COUNTY OF "
                        + self.juris.county.name.text.upper()
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == 'Colorado':
                output += str(
                    str(
                        self.juris.trial_court().upper()
                        + " COURT "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, COLORADO"
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == "Connecticut":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE SUPERIOR COURT OF THE STATE OF "
                        + state_name(self.juris.state.name.text).upper()
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR THE COUNTY OF "
                        + self.juris.county.name.text.upper()
                    )
                )
            elif state_name(self.juris.state.name.text) == "Delaware":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " SUPERIOR COURT OF THE STATE OF DELAWARE"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "IN AND FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Florida":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + str(" COURT" if self.juris.trial_court() not in (
                            "Court of Chancery", "Court of Common Pleas") else "")
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "OF THE "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL CIRCUIT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "IN AND FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, FLORIDA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Georgia":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + "COUNTY"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "STATE OF GEORGIA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Hawaii":
                output += str(
                    str(
                        self.fx.CENTER
                        + "STATE OF HAWAI‘I"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "CIRCUIT COURT OF THE "
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + self.juris.county.name.text.upper()
                        + " CIRCUIT"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Idaho":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR THE "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL DISTRICT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR THE STATE OF IDAHO, IN AND FOR THE COUNTY OF "
                        + self.juris.county.name.text.upper()
                        + " CIRCUIT"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Illinois":
                output += str(
                    str(
                        self.fx.CENTER
                        + "In the "
                        + self.juris.trial_court()
                        + " Court of "
                        + self.juris.county.name.text
                        + " County, Illinois"
                        + self.fx.NLINE
                        + self.juris.trial_court_units()[1]
                        + " Judicial Circuit"
                    ).upper()
                )
            elif state_name(self.juris.state.name.text) == "Indiana":
                output += str(
                    str(
                        "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Iowa":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE IOWA "
                        + self.juris.trial_court().upper()
                        + " COURT FOR "
                        + self.juris.county.name.text.upper()
                        + "COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Kansas":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL DISTRICT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "DISTRICT COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, KANSAS"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "[SITTING AT (Name of city, if more than one court location in county) ]"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Kentucky":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR THE STATE OF KENTUCKY"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR THE "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL DISTRICT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "COUNTY OF "
                        + self.juris.county.name.text.upper()
                    )
                )
            elif state_name(self.juris.state.name.text) == "Louisiana":
                output += str(
                    str(
                        self.fx.CENTER
                        + str(ordinal(self.juris.trial_court_units()[1]).upper() if self.juris.trial_court_units()[
                                                                                        1] != "Orleans" else
                              self.juris.trial_court_units()[1].upper())
                        + " JUDICIAL DISTRICT COURT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + self.juris.county.name.text.upper()
                        + " PARISH, LOUISIANA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Maine":
                output += str(
                    str(
                        self.fx.CENTER
                        + "MAINE JUDICIAL BRANCH"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + str(str(self.fx.NLINE + self.fx.CENTER + "IN " + self.juris.trial_court_units()[
                            1]).upper() if self.juris.county.name.text == "Aroostook" else "")
                    )
                )
            elif state_name(self.juris.state.name.text) == "Maryland":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court_units()[1].upper()
                        + " "
                        + self.juris.trial_court().upper()
                        + " COURT OF MARYLAND FOR "
                        + self.juris.county.name.text.upper()
                        + str(" CITY" if self.juris.county.name.text == "Baltimore" else " COUNTY")
                    )
                )

            elif state_name(self.juris.state.name.text) == "Massachusetts":
                output += str(
                    str(
                        self.fx.CENTER
                        + "COMMONWEALTH OF MASSACHUSETTS"
                        + self.fx.NLINE
                        + self.fx.NLINE

                        + self.fx.CENTER
                        + self.juris.county.name.text.upper()
                        + ", ss"
                        + self.fx.TAB
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + self.fx.NLINE
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + str(str(self.juris.trial_court_units()[
                                      1].upper() + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER) if self.juris.county.name.text in (
                        "Bristol", "Essex", "Middlesex", "Plymouth") else "")
                    )
                )
            elif state_name(self.juris.state.name.text) == "Michigan":
                output += str(
                    str(
                        self.fx.CENTER
                        + "STATE OF MICHIGAN"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL CIRCUIT"
                        + str(str(self.fx.NLINE + self.fx.CENTER + self.juris.county.name.text.upper() + " COUNTY") if
                              self.juris.trial_court_units()[1] in (
                                  "eighth", "eleventh", "twelfth", "thriteenth", "nineteenth", "twenty-third",
                                  "twenty-sixth", "twenty-seventh", "twenty-eighth", "twenty-ninth", "thirty-second",
                                  "thirty-fourth", "fourty-first", "fourty-sixth", "fourty-ninth", "fifty-first",
                                  "fifty-third", "fifty-fifth") else "")
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "THE "
                        + self.juris.trial_court().upper()
                        + " COURT"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Minnesota":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE DISTRICT COURT FOR THE STATE OF MINNESOTA"
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL DISTRICT"
                        + str(
                            str(self.fx.NLINE + self.fx.CENTER + self.juris.county.name.text.upper() + " COUNTY") if not
                            self.juris.trial_court_units()[1] in ("second", "fourth") else "")
                    )
                )
            # The Mississippi CIVPRO dictionary is keyed such as to first divide the Circuit court by district (e.g., fifth,  sixth, etc.), and then further divide by county, whereas according to the form of the caption it may be of better use the other way around (and which judicial district a court is sitting in seems to be of questionable use). But I am leaving it for now pending further research because this form that I found and used for Mississippi wasn't the greatest, but it was the only form with a caption I could find.
            elif state_name(self.juris.state.name.text) == "Mississippi":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + ", MISSISSIPPI"
                    )
                )

            elif state_name(self.juris.state.name.text) == "Missouri":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " CIRCUIT COURT"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + ", MISSOURI"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Montana":
                output += str(
                    str(
                        self.fx.CENTER
                        + "MONTANA "
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL DISTRICT COURT"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Nebraska":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "NEBRASKA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Nevada":
                output += str(
                    str(
                        self.fx.CENTER
                        + self.fx.BOLD
                        + str(str(ordinal(self.juris.trial_court_units()[1]).upper() + " JUDICIAL ") if ordinal(
                            self.juris.trial_court_units()[1]) in ("First", "Fifth", "Seventh", "Eleventh") else "")
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + self.fx.BOLD
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY, NEVADA "
                    )
                )
            elif state_name(self.juris.state.name.text) == "New Hampshire":
                output += str(
                    str(
                        self.fx.CENTER
                        + self.juris.county.name.text.upper()
                        + " "
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + str(str(" " + self.juris.trial_court_units()[
                            1].upper()) if self.juris.county.name.text == "Hillsborough" and self.juris.trial_court() == "Superior" else "")
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "STATE OF NEW HAMPSHIRE"
                    )
                )
            elif state_name(self.juris.state.name.text) == "New Jersey":
                output += str(
                    str(
                        self.fx.CENTER
                        + self.juris.trial_court().upper()
                        + " COURT OF NEW JERSEY"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.division.name.text.upper()
                        + " DIVISION"
                    )
                )

            elif state_name(self.juris.state.name.text) == "New Mexico":
                output += str(
                    str(
                        self.fx.FLUSHLEFT
                        + "STATE OF NEW MEXICO"
                        + self.fx.FLUSHLEFT
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + self.fx.FLUSHLEFT
                        + self.fx.NLINE
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL "
                        + self.juris.trial_court().upper()
                        + " COURT"
                    )
                )
            elif state_name(self.juris.state.name.text) == "New York":
                output += str(
                    str(
                        self.fx.FLUSHLEFT
                        + self.juris.trial_court().upper()
                        + " COURT OF THE STATE OF NEW YORK  "
                        # + self.fx.FLUSHLEFT
                        + self.fx.NLINE
                        + "COUNTY OF "
                        + self.juris.county.name.text.upper()
                    )
                )
            # Has a set up that will require a row of two cells for the header.
            elif state_name(self.juris.state.name.text) == "North Carolina":
                output += str(
                    [
                        str("NORTH CAROLINA"
                            + self.fx.NLINE
                            + self.fx.NLINE
                            + self.juris.county.name.text.upper()
                            + " COUNTY "),
                        str(self.fx.CENTER
                            + self.fx.NLINE
                            + "NORTH CAROLINA"
                            + self.fx.CENTER
                            + "IN THE GENERAL COURT OF JUSTICE"
                            + self.fx.CENTER
                            + self.fx.NLINE
                            + self.juris.trial_court().upper()
                            + " COURT DIVISION"
                            + self.fx.CENTER
                            + self.fx.NLINE
                            + str(self.docket_number)
                            )

                    ]
                )
            elif state_name(self.juris.state.name.text) == "North Dakota":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT, COUNTY OF "
                        + self.juris.county.name.text.upper()
                        + ", STATE OF NORTH DAKOTA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Ohio":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + str(" COURT, " if self.juris.trial_court().upper() in (
                            "Municipal", "County", "Mayor’s") else ", ")
                        + self.juris.county.name.text.upper()
                        + "COUNTY, OHIO"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Oklahoma":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "STATE OF OKLAHOMA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Pennsylvania":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + str(" COURT " if self.juris.trial_court() != "Court of Common Pleas" else "")
                        + " OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, PENNSYLVANIA"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "CIVIL DIVISION"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Rhode Island":
                output += str(
                    str(
                        self.fx.CENTER
                        + "STATE OF RHODE ISLAND"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + ", S.C."
                        # CELL2
                        + self.juris.trial_court().upper()
                        + str(" COURT" if self.juris.trial_court() != "Traffic Tribunal" else "")
                    ))
            elif state_name(self.juris.state.name.text) == "South Carolina":
                output += str(
                    str(
                        "STATE OF SOUTH CAROLINA"
                        + self.fx.TAB
                        + ")"
                        + self.fx.TAB
                        + "IN THE "
                        + str("COURT OF COMMON PLEAS" if self.juris.trial_court() == "Circuit" else str(
                            self.juris.trial_court().upper() + " COURT"))
                        + " COURT"
                        + self.fx.NLINE
                        + "COUNTY OF "
                        + self.juris.county.name.text.upper()
                        + self.fx.TAB
                        + ")"
                        + self.fx.TAB
                        + ordinal(self.juris.trial_court_units()[1]).upper()
                        + " JUDICIAL CIRCUIT"
                    )
                )

            # South Dakota may require the two cell header, but I tried to make it work with the tab configuration here.
            elif state_name(self.juris.state.name.text) == "South Dakota":
                output += str(
                    str(
                        "STATE OF SOUTH DAKOTA"
                        + self.fx.TAB
                        + ")"
                        + self.fx.TAB
                        + "IN "
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + self.fx.NLINE
                        + "COUNTY OF "
                        + self.juris.county.name.text.upper()
                        + self.fx.TAB
                        + ")"
                        + self.fx.TAB
                        + self.juris.trial_court_units()[1].upper()
                        + " JUDICIAL CIRCUIT"
                    )
                )

            elif state_name(self.juris.state.name.text) == "Tennessee":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, TENNESSEE"
                    )
                )
            # Texas has distinct formatting with the case number where the header usually is, the header where the case number and document title usually is, and the document title beneath the caption and centered. I will need to return and update for format, but this is the header (which goes where the case number and document title usually is when its time to format).
            elif state_name(self.juris.state.name.text) == "Texas":
                output += str(
                    str(
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + str(" COURT" if self.juris.trial_court() != "County Court At Law" else "")
                        + self.fx.NLINE
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY, TEXAS"
                        + self.fx.NLINE
                        + self.fx.NLINE
                        + str(str(ordinal(self.juris.trial_court_units()[1]).upper() + " JUDICIAL DISTRICT") if
                              self.juris.trial_court_units()[1] != "1 - A" else "JUDICIAL DISTRICT 1 - A")
                    )
                )
            elif state_name(self.juris.state.name.text) == "Utah":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + str(str(self.juris.trial_court_units()[
                                      1].upper() + " JUDICIAL DISTRICT ") if self.juris.trial_court() != "District" else "")
                        + "COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "STATE OF UTAH"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Vermont":
                output += str(
                    str(
                        self.fx.CENTER
                        + "STATE OF VERMONT"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.juris.county.name.text.upper()
                        + " COUNTY, SS.")
                )

            elif state_name(self.juris.state.name.text) == "Virginia":
                output += str(
                    str(
                        self.fx.CENTER
                        + "VIRGINIA: IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Washington":
                output += str(
                    str(
                        self.fx.CENTER
                        + self.juris.trial_court().upper()
                        + " COURT OF THE STATE OF WASHINGTON"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "FOR "
                        + self.juris.county.name.text.upper()
                        + " COUNTY"
                    )
                )
            elif state_name(self.juris.state.name.text) == "West Virginia":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + self.juris.trial_court().upper()
                        + " COURT OF "
                        + self.juris.county.name.text.upper()
                        + " COUNTY, WEST VIRGINIA"
                    )
                )
            elif state_name(self.juris.state.name.text) == "Wisconsin":
                output += str(
                    str(
                        "STATE OF WISCONSIN "
                        + self.juris.trial_court().upper()
                        + " COURT "
                        # CELL3
                        + self.juris.county.name.text.upper()
                        + " COUNTY"

                    ))
            elif state_name(self.juris.state.name.text) == "Wyoming":
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE "
                        + str(str(ordinal(self.juris.trial_court_units()[
                                              1]).upper() + " JUDICIAL ") if self.juris.trial_court() == "District" else "")
                        + self.juris.trial_court().upper()
                        + " COURT"
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + "THE STATE OF WYOMING, COUNTY OF "
                        + self.juris.county.name.text.upper()
                        + self.fx.CENTER
                        + self.fx.NLINE
                        + self.number()
                    )
                )


            elif state_name(self.juris.state.name.text) == "District of Columbia":
                output = ("", "")
            else:
                output += str(
                    str(
                        self.fx.CENTER
                        + "IN THE CIRCUIT COURT OF THE STATE OF "
                        + state_name(self.juris.state.name.text.upper())
                        + self.fx.NLINE
                        + self.fx.CENTER
                        + "FOR THE COUNTY OF "
                        + self.juris.county.name.text.upper()
                    )
                )
        return output

    def minnesota(self):
        output = ""
        log("starting minnesota")
        output += str(str("STATE OF MINNESOTA" + self.fx.NLINE + self.fx.NLINE + str(
            str(self.juris.county.name.text.upper() + " COUNTY") if not
            self.juris.trial_court_units()[1] in ("second", "fourth") else "")))
        output += str(
            str(self.juris.trial_court().upper() + " COURT" + self.fx.NLINE + self.fx.NLINE + ordinal(
                self.juris.trial_court_units()[1]).upper() + " JUDICIAL DISTRICT"))
        log(str(output))
        return output

    def jurisdiction_docx(self):
        log("starting make_caption")

        fifth_cell = str(self.fx.CENTER
                         + self.fx.BOLD
                         + self.document_title().upper()
                         + self.fx.BOLD)

        if self.juris.jurisdiction.name.text == 'federal':
            return self.second_column(), None

        elif self.juris.state.name.text == "CO":
            first_cell = str(
                self.fx.SINGLE + self.doc_header() + self.fx.NLINE + self.juris.courthouse.line_one() + self.fx.NLINE + self.juris.courthouse.line_two())
            third_cell = self.fed_head()
            fourth_cell = str(
                self.fx.CENTER + self.fx.SINGLE + "**⮝COURT USE ONLY⮝**" + self.fx.NLINE + self.fx.SINGLE + self.number() + self.fx.NLINE + "Courtroom: ")
            fifth_cell = str(self.fx.CENTER
                             + self.fx.BOLD
                             + self.document_title().upper()
                             + self.fx.BOLD)
            return first_cell, third_cell, fourth_cell, fifth_cell

        elif self.juris.state.name.text == "FL":
            return self.second_column(), None

        elif self.juris.state.name.text == "IN":
            first_cell = str("STATE OF INDIANA")
            second_cell = str(self.doc_header())
            third_cell = "SS:"
            fourth_cell = str(self.juris.county.name.text.upper() + " COUNTY")
            sixth_cell = self.number()
            fifth_cell = self.document_title()
            return first_cell, second_cell, third_cell, fourth_cell, fifth_cell, sixth_cell

        elif self.juris.state.name.text == "MN":
            log(repr(self.juris.trial_court_units()[1]))
            first_cell = str(str("STATE OF MINNESOTA" + self.fx.NLINE + self.fx.NLINE + str(
                str(self.juris.county.name.text.upper() + " COUNTY") if not
                self.juris.trial_court_units()[1] in ("second", "fourth") else "")))
            second_cell = str(
                str(self.juris.trial_court().upper() + " COURT" + self.fx.NLINE + self.fx.NLINE + ordinal(
                    self.juris.trial_court_units()[1]).upper() + " JUDICIAL DISTRICT"))
            return first_cell, second_cell, self.second_column()

        elif self.juris.state.name.text == "NC":
            first_cell = str("NORTH CAROLINA" + self.fx.NLINE + self.fx.NLINE + self.juris.county.name.text + " COUNTY")
            second_cell = str(self.fx.CENTER + self.doc_header() + self.fx.NLINE + self.number())
            return first_cell, second_cell, self.second_column()

        elif self.juris.state.name.text == "NM":
            first_cell = self.doc_header()
            return first_cell, self.second_column(), fifth_cell

        elif self.juris.state.name.text == "NY":
            first_cell = self.doc_header()
            return first_cell, self.second_column()

        elif self.juris.state.name.text == "RI":
            first_cell = str(
                "STATE OF RHODE ISLAND" + self.fx.NLINE + self.juris.county.name.text.upper() + ", S.C.")
            second_cell = str(self.juris.trial_court().upper() + str(
                " COURT" if self.juris.trial_court() != "Traffic Tribunal" else ""))
            return first_cell, second_cell, self.second_column()

        elif self.juris.state.name.text == "SC":
            first_cell = str(
                self.fx.FLUSHLEFT + "STATE OF SOUTH CAROLINA" + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + "COUNTY OF " + self.juris.county.name.text.upper() + self.fx.NLINE)
            second_cell = str(self.fx.CENTER + "IN THE " + str(
                "COURT OF COMMON PLEAS" if self.juris.trial_court() == "Circuit" else str(
                    self.juris.trial_court().upper() + " COURT")) + self.fx.NLINE + self.fx.CENTER + ordinal(
                self.juris.trial_court_units()[1]).upper() + " JUDICIAL CIRCUIT")
            return first_cell, second_cell, self.second_column()


        elif self.juris.state.name.text == "SD":
            first_cell = str(
                self.fx.FLUSHLEFT + "STATE OF SOUTH DAKOTA" + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + "COUNTY OF " + self.juris.county.name.text.upper() + self.fx.NLINE)
            second_cell = str(self.fx.CENTER + "IN "
                              + self.juris.trial_court().upper()
                              + " COURT"
                              + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + ordinal(
                self.juris.trial_court_units()[1]).upper() + " JUDICIAL CIRCUIT")
            return first_cell, second_cell, self.second_column()

        elif self.juris.state.name.text == "TX":
            return self.second_column(), fifth_cell

        elif self.juris.state.name.text == "UT":
            first_cell = self.doc_header()
            return first_cell, self.second_column()

        elif self.juris.state.name.text == "WI":
            first_cell = "STATE OF WISCONSIN"
            second_cell = str(self.juris.trial_court().upper() + " COURT")
            third_cell = str(self.juris.county.name.text.upper() + " COUNTY")
            return first_cell, second_cell, third_cell, self.second_column()

        elif self.juris.state.name.text == "WY":
            first_cell = self.doc_header()
            return first_cell, fifth_cell

        elif self.juris.state.name.text == "WV":
            first_cell = self.doc_header()
            return first_cell, fifth_cell, self.second_column()

        elif self.juris.state.name.text in (
                "AL", "AR", "AK", "FL", "OH", "GA", "IL", "IN", "KS", "MD", "MO", "NM", "OH", "OK", "PA", "RI", "TX",
                "TN",
                "WV", "WY", "VA", "VT"):
            return self.second_column(), fifth_cell

        else:
            return self.second_column(), None

    def party_captions(self):
        output_party_captions = []
        output_party_captions.clear()
        base_party_column = ""
        if any(party.party.party_type.name.text in (
        "Plaintiff", "Petitioner", "Appellant" "Respondent", "Defendant", "Claimant") for party in self.parties):
            for party in add_separators(
                    self.parties.the_ps().caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                base_party_column += party
            base_party_column += (
                    self.fx.NLINE
                    + self.fx.NLINE
                    + " [TAB] "
                    + self.parties.the_ps().caption_noun()
                    + ",  "
                    # + self.fx.NLINE2
                    + self.fx.NLINE2
                    + " [NBSP] "
                    + " [NBSP] "
                    + " [NBSP] "
                    + "vs. "
                    + self.fx.NLINE2
                # + self.fx.NLINE2
                # + self.fx.NLINE2
                # + self.fx.NLINE
            )
            for party in add_separators(
                    self.parties.the_ds().caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                base_party_column += party
            base_party_column += (
                    self.fx.NLINE2
                    # + self.fx.NLINE
                    + self.fx.NLINE2
                    + " [TAB] "
                    + self.parties.the_ds().caption_noun()
                    + "."
                    + self.fx.NLINE
            )
            log("made base party column")
            output_party_captions.append(str(base_party_column))
        log("AFTER made base party column . . . . BEFORE has_cross_parties")
        if self.has_cross_parties():
            log("if self.has_cross_parties():")
            cross_party_column = ""
            for claim in self.cross_claims:
                for party in add_separators(
                        claim.ps.caption(),
                        separator=", ",
                        last_separator=", and ",
                        end_mark=", ",
                ):
                    cross_party_column += party
                cross_party_column += (
                        self.fx.NLINE
                        + self.fx.NLINE
                        + " [TAB] Cross-"
                        + claim.ds.caption_noun()
                        + ","
                        + self.fx.NLINE
                        + " [NBSP] "
                        + " [NBSP] "
                        + " [NBSP] "
                        + "vs."
                        + self.fx.NLINE
                        + self.fx.NLINE
                )
                for party in add_separators(
                        claim.ds.caption(),
                        separator=", ",
                        last_separator=", and ",
                        end_mark=", ",
                ):
                    cross_party_column += party
                cross_party_column += (
                        self.fx.NLINE
                        + self.fx.NLINE
                        + " [TAB] Cross-"
                        + claim.ps.caption_noun()
                        + "."
                )
            fourth_column = str(self.fx.NLINE)
            log("made cross party columns")
            output_party_captions.append(cross_party_column)
        log("passed cross party columns")
        if self.has_third_parties():
            log("has_third_parties")
            third_party_column = ""
            for party in add_separators(
                    self.parties.threedp_ps().caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                log("third_party_column += party")
                third_party_column += party
            third_party_column += (
                    self.fx.NLINE
                    + self.fx.NLINE
                    + " [TAB] "
                    + "Third-Party "
                    + self.parties.threedp_ps().as_noun("Plaintiff")
                    + ","
                    + self.fx.NLINE
                    + " [NBSP] "
                    + " [NBSP] "
                    + " [NBSP] "
                    + "vs."
                    + self.fx.NLINE
                    + self.fx.NLINE
            )
            for party in add_separators(
                    self.parties.threedp_ds().caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                third_party_column += party
            third_party_column += (
                    self.fx.NLINE
                    + self.fx.NLINE
                    + " [TAB] "
                    + self.parties.threedp_ds().caption_noun()
                    + "."
            )
            sixth_column = str(self.fx.NLINE)
            output_party_captions.append(third_party_column)
            log("made sixth column")
        log("passed 3dp column")

        if self.has_intervenors():
            intervenor_column = ""
            for party in add_separators(
                    self.parties.intervenors().caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                intervenor_column += party
            intervenor_column += (
                    self.fx.NLINE
                    + self.fx.NLINE
                    + " [TAB] "
                    + self.parties.intervenors().caption_noun()
            )
            log("made intervenor column")
            output_party_captions.append(intervenor_column)

        if self.has_deceaseds():
            deceased_column = "In the Matter of the Estate of "
            for party in add_separators(
                    claim.deceaseds.caption(),
                    separator=", ",
                    last_separator=", and ",
                    end_mark=", ",
            ):
                deceased_column += party
            deceased_column += (
                    self.fx.NLINE
                    + self.fx.NLINE
                    + " [TAB]"
                    + self.parties.deceaseds().caption_noun()
                    + "."

            )
            output_party_captions.append(deceased_column)
            log("made deceased column")
        return output_party_captions

    def second_column(self):
        if self.juris.jurisdiction.name.text == "federal":
            second_column_output = self.document_title()

        elif self.juris.state.name.text == "FL":
            second_column_output = str(
                self.doc_header() + self.fx.NLINE + self.fx.NLINE + self.number().upper() + self.fx.NLINE + "DIVISION:" + self.juris.division.name.text)

        elif self.juris.state.name.text == "MN":
            second_column_output = str(
                self.number() + self.fx.NLINE
                # + self.type()
                + self.fx.NLINE + self.fx.BOLD + self.document_title().upper() + self.fx.BOLD)

        elif self.juris.state.name.text == "NC":
            second_column_output = str(self.fx.CENTER
                                       + self.fx.BOLD
                                       + self.document_title().upper()
                                       + self.fx.BOLD)

        elif self.juris.state.name.text == "NV":
            second_column_output = str(self.number()
                                       + self.fx.NLINE
                                       + self.fx.NLINE
                                       + "Dept. No" + self.dept_no + self.fx.NLINE + self.fx.NLINE)

        elif self.juris.state.name.text == "TX":
            second_column_output = str(
                "IN THE "
                + self.juris.trial_court().upper()
                + str(" COURT" if self.juris.trial_court() != "County Court At Law" else "")
                + self.fx.NLINE
                + self.fx.NLINE
                + self.juris.county.name.text.upper()
                + " COUNTY, TEXAS"
                + self.fx.NLINE
                + self.fx.NLINE
                + str(str(ordinal(self.juris.trial_court_units()[1]).upper() + " JUDICIAL DISTRICT") if
                      self.juris.trial_court_units()[1] != "1 - A" else "JUDICIAL DISTRICT 1 - A")
            )

        elif self.juris.state.name.text == "VT":
            second_column_output = str(self.juris.county.name.text.upper()
                                       + " "
                                       + self.juris.trial_court().upper()
                                       + str(" COURT" if self.juris.trial_court() != "Judicial Bureau" else "")
                                       + self.fx.NLINE
                                       + self.fx.NLINE
                                       + self.number())

        elif self.juris.state.name.text == "WI":
            second_column_output = str(self.fx.BOLD + self.number() + self.fx.BOLD)

        elif self.juris.state.name.text in (
                "AL", "AR", "AK", "OH", "GA", "IL", "IN", "KS", "MD", "MO", "NM", "OH", "OK", "PA", "RI", "TN", "WV",
                "WY",
                "VA"):
            second_column_output = self.number()
        else:
            second_column_output = str(self.number()
                                       + self.fx.NLINE
                                       + self.fx.NLINE
                                       + self.fx.BOLD
                                       + self.document_title().upper()
                                       + self.fx.BOLD
                                       + self.fx.NLINE
                                       )
        return second_column_output

    def make_caption(self):
        log("starting make_caption")
        output = ""

        if self.juris.jurisdiction.name.text == 'federal':
            log("making federal cap . . . . 1")
            output += str(self.fx.FLUSHLEFT + self.fed_head() + self.fx.NLINE)
            output += self.doc_header()

        elif self.juris.state.name.text in ("CA", "AZ", "UT", "NJ"):
            output += str(self.fx.FLUSHLEFT + self.fed_head() + self.fx.NLINE)

        elif self.juris.state.name.text == "TX":
            output += str(self.fx.CENTER + self.number() + self.fx.NLINE + self.fx.NLINE)

        if self.juris.jurisdiction.name.text == 'federal':
            log("making federal cap . . . . 2")
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            log("making federal cap . . . . 3")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>"
                    )
            log("making federal cap . . . . 4")


        elif self.juris.state.name.text == "CO":
            first_cell = str(
                self.fx.SINGLE + self.doc_header() + self.fx.NLINE + self.juris.courthouse.line_one() + self.fx.NLINE + self.juris.courthouse.line_two())
            third_cell = self.fed_head()
            fourth_cell = str(
                self.fx.CENTER + self.fx.SINGLE + "**⮝COURT USE ONLY⮝**" + self.fx.NLINE + self.fx.SINGLE + self.number() + self.fx.NLINE + "Courtroom: ")
            fifth_cell = str(self.fx.CENTER
                             + self.fx.BOLD
                             + self.document_title().upper()
                             + self.fx.BOLD)
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(first_cell)
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: middle; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 1px; border-bottom-width: 0px">'
                + ""
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em;  border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px vertical-align: top">'
                + ""
                + "</td></tr></table>"
            )
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em;  border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px vertical-align: top">'
                        + ""
                        + "</td></tr></table>"
                    )

            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em;  border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(third_cell)
                + '</td><td style="padding-left: 1em; width: 50%; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px vertical-align: top">'
                + markdown_to_html(fourth_cell)
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 100%; vertical-align: middle; text-align: center; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(fifth_cell)
                + '</td></tr></table>')

        elif self.juris.state.name.text == "FL":
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            # bottom border on left cell
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "IN":
            output += str(
                '<table style="width: 100%"><tr><td style="width: 35%; text-align: left; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html("STATE OF INDIANA")
                + '</td><td style="padding-left: 1em; width: 65%; text-align: right; vertical-align: top">'
                + markdown_to_html(self.doc_header())
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 35%; text-align: left; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.fx.NLINE)
                + '</td><td style="padding-left: 1em; width: 65%; text-align: left; vertical-align: top">'
                + markdown_to_html("SS:")
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 35%; text-align: left; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.juris.county.name.text.upper() + " COUNTY")
                + '</td><td style="padding-left: 1em; width: 65%; text-align: right; vertical-align: top">'
                + markdown_to_html(self.number())
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "MN":
            output = str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; text-align: left; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(str("STATE OF MINNESOTA" + self.fx.NLINE + self.fx.NLINE + str(
                    str(self.juris.county.name.text.upper() + " COUNTY") if not
                    self.juris.trial_court_units()[1] in ("second", "fourth") else "")))
                + '</td><td style="width: 50%; vertical-align: bottom; text-align: right;  border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px"">'
                + markdown_to_html(
                    str(self.fx.NLINE + self.juris.trial_court().upper() + " COURT" + self.fx.NLINE + self.fx.NLINE + ordinal(
                        self.juris.trial_court_units()[1]).upper() + " JUDICIAL DISTRICT"))
                + "</td></tr></table>"
            )

            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "NC":
            first_cell = str("NORTH CAROLINA" + self.fx.NLINE + self.fx.NLINE + self.juris.county.name.text + " COUNTY")
            second_cell = str(self.fx.CENTER + self.doc_header() + self.fx.NLINE + self.number())
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(first_cell)
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(second_cell)
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: middle">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: middle">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "NM":
            output += self.doc_header()
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "NY":
            output += str(
                '<table style="width: 100%"><tr><td style="width: 75%; vertical-align: top; border-style: dashed; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.doc_header())
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(str())
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: dashed; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: dashed; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>"
                    )

        elif self.juris.state.name.text == "RI":
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: text-align: left; top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(
                    str("STATE OF RHODE ISLAND" + self.fx.NLINE + self.juris.county.name.text.upper() + ", S.C."))
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; text-align: right;  border-style: solid; border-right-width: 0px; padding-right: 0em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px"">'
                + markdown_to_html(str(self.juris.trial_court().upper() + str(
                    " COURT" if self.juris.trial_court() != "Traffic Tribunal" else "")))
                + "</td></tr></table>")
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "SC":
            first_cell = str(
                self.fx.FLUSHLEFT + "STATE OF SOUTH CAROLINA" + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + "COUNTY OF " + self.juris.county.name.text.upper() + self.fx.NLINE)
            second_cell = str(self.fx.CENTER + "IN THE " + str(
                "COURT OF COMMON PLEAS" if self.juris.trial_court() == "Circuit" else str(
                    self.juris.trial_court().upper() + " COURT")) + self.fx.NLINE + self.fx.CENTER + ordinal(
                self.juris.trial_court_units()[1]).upper() + " JUDICIAL CIRCUIT")
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 1px; border-bottom-width: 0px">'
                + markdown_to_html(first_cell)
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: topborder-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 1px; border-bottom-width: 0px">'
                + markdown_to_html(second_cell)
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "SD":
            first_cell = str(
                self.fx.FLUSHLEFT + "STATE OF SOUTH DAKOTA" + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + "COUNTY OF " + self.juris.county.name.text.upper() + self.fx.NLINE)
            second_cell = str(self.fx.CENTER + "IN "
                              + self.juris.trial_court().upper()
                              + " COURT"
                              + self.fx.NLINE + self.fx.NLINE + self.fx.CENTER + ordinal(
                self.juris.trial_court_units()[1]).upper() + " JUDICIAL CIRCUIT")
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(first_cell)
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: topborder-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(second_cell)
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "UT":
            output += self.doc_header()
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 2px; border-bottom-width: 2px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 2px; border-bottom-width: 2px">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 2px; border-bottom-width: 2px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 1px; border-top-width: 2px; border-bottom-width: 2px">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>"
                    )

        elif self.juris.state.name.text == "WI":
            first_cell = "STATE OF WISCONSIN"
            second_cell = str(self.juris.trial_court().upper() + " COURT")
            third_cell = str(self.juris.county.name.text.upper() + " COUNTY")
            output += str(
                '<table style="width: 100%"><tr><td style="width: 33%; vertical-align: top; border-style: dashed; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(first_cell)
                + '</td><td style="padding-left: 1em; width: 33%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(second_cell)
                + '</td><td style="padding-left: 1em; width: 33%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(third_cell)
                + "</td></tr></table>"
            )
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 1px; border-bottom-width: 1px">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>"
                    )

        elif self.juris.state.name.text == "WY":
            output += self.doc_header()
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: dashed; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 0em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: dashed; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 0em; border-left-width: 1px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(self.fx.NLINE + self.fx.NLINE)
                        + "</td></tr></table>")

        elif self.juris.state.name.text == "WV":
            output += self.doc_header()
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: middle">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: middle">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>")

        else:
            log("making caption 1 . . . .")
            if self.juris.state.name.text != "TX":
                output += self.doc_header()
            output += str(
                '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                + markdown_to_html(self.party_captions()[0])
                + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                + markdown_to_html(self.second_column())
                + "</td></tr></table>"
            )
            log("making caption 2 . . . .")
            if len(self.party_captions()) > 1:
                for party_caption in self.party_captions()[1:]:
                    output += str(
                        '<table style="width: 100%"><tr><td style="width: 50%; vertical-align: top; border-style: solid; border-right-width: 1px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 1px">'
                        + markdown_to_html(party_caption)
                        + '</td><td style="padding-left: 1em; width: 50%; vertical-align: top">'
                        + markdown_to_html(self.fx.NLINE)
                        + "</td></tr></table>"
                    )
            log("making caption 3 . . . .")
        if self.juris.jurisdiction.name.text != 'federal':
            if self.juris.state.name.text in (
                    "AL", "AR", "AK", "FL", "OH", "GA", "IL", "IN", "KS", "MD", "MO", "NM", "OH", "OK", "PA", "RI",
                    "TX", "TN",
                    "WV", "WY", "VA", "VT"):
                fifth_cell = str(self.fx.CENTER
                                 + self.fx.BOLD
                                 + self.document_title().upper()
                                 + self.fx.BOLD)
                output += str(
                    '<table style="width: 100%"><tr><td style="width: 100%; vertical-align: middle; text-align: center; border-style: solid; border-right-width: 0px; padding-right: 1em; border-left-width: 0px; border-top-width: 0px; border-bottom-width: 0px">'
                    + markdown_to_html(fifth_cell)
                    + '</td></tr></table>')
        log("made it to output")
        return output

    def destination_base(self):
        if self.draft == "RRFP" or self.draft == "RFP":
            return "RFP"

        elif self.draft == "RROG" or self.draft == "ROG":
            return "ROG"

        elif self.draft == "RRFA" or self.draft == "RFA":
            return "RFA"
        elif self.draft == "MIL":
            return output.append(
                "Motions" + self.fx.ITALIC + " in limine" + self.fx.ITALIC
            )
        elif self.draft == "Witness List":
            return "Witness List"
        elif self.draft == "Trial Memo":
            return "Trial Memo"
        elif self.draft == "Depo Notice":
            return "Notice of Deposition"

    def document_title(self):
        output = ""
        if self.user_title == False:
            if self.draft == "RRFP":
                output = str(
                    possessify(self.user_clientlist().asdesignation(), "Response")
                    + " to "
                    + possessify(self.recipients().asnoun(), str(self.request_ordinal))
                    + " Request for Production of Documents"
                )
            if self.draft == "Notice of New Firm":
                output += str("Notice of New Address of Counsel for " + self.user_clientlist().asnoun())
            elif self.draft == "ProposeJuryInstructions":
              output = str(
                    possessify(self.user_clientlist().asnoun(), "Requested Jury Instructions")
                )
            elif self.draft == "Motion to Amend to Allege Punitive Damages":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "Motion to Amend to Allege Punitive Damages")
                )
            elif self.draft == "RFP":
                output = str(
                    possessify(
                        self.user_clientlist().asnoun(), str(self.request_ordinal)
                    )
                    + " Request for Production of Documents to "
                    + self.recipients().asnoun()
                )
            elif self.draft == "RROG":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "Response")
                    + " to "
                    + possessify(self.recipients().asnoun(), str(self.request_ordinal))
                    + " Set of Interrogatories"
                )
            elif self.draft == "ROG":
                output = str(
                    possessify(
                        self.user_clientlist().asnoun(), str(self.request_ordinal)
                    )
                    + " Set of Interrogatories to "
                    + self.recipients().asnoun()
                )
            elif self.draft == "RRFA":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "Response")
                    + " to "
                    + possessify(self.recipients().asnoun(), str(self.request_ordinal))
                    + " Set of Requests for Admissions"
                )
            elif self.draft == "RFA":
                output = str(
                    possessify(
                        self.user_clientlist().asnoun(), str(self.request_ordinal)
                    )
                    + " Set of Requests for Admissions to "
                    + self.recipients().asnoun()
                )
            elif self.draft == "MIL":

                output = str(
                    possessify(self.user_clientlist().asnoun(), "")
                    + "Motions"
                    + self.fx.ITALIC
                    + " in limine"
                    + self.fx.ITALIC
                )
            elif self.draft == "Witness List":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "") + "Witness List"
                )
            elif self.draft == "Trial Memo":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "") + "Trial Memorandum"
                )
            elif self.draft == "Jury Instructions":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "")
                    + "Proposed Jury Instructions"
                )
            elif self.draft == "Depo Notice":
                output = str(
                    possessify(self.user_clientlist().asnoun(), "")
                    + "Notice of Deposition of "
                    + self.recipients().asnoun()
                    + " "
                    + comma_and_list(self.recipients())
                )
            elif self.draft == "Complaint":
                output = "Complaint"
            elif self.draft == "Summons":
                output = "Summons"
            elif self.draft == 'Motion to Postpone':
                output = 'Motion to Postpone'
            elif self.draft == "Motion":
                output = ""
                output += "Motion for Order to Postpone"
            elif self.draft == "Order":
                output = ""
                output += "Order to Postpone"
            elif self.draft == "Declaration":
                output = ""
                output += "Declaration in Support of Motion for Order to Postpone"
            elif self.draft == "Discovery Response":
                output = ""
                responses = []
                output += str(
                    possessify(self.user_clientlist().asnoun(), "Response")
                    + " to "
                    + possessify(self.recipients().asnoun(), "")
                )
                if len(self.rrfps):
                    responses.append(
                        str(
                            self.rrfps.request_ordinal
                            + " Request for Production of Documents"
                        )
                    )
                if len(self.rrogs):
                    responses.append(
                        str(self.rrogs.request_ordinal + " Set of Interrogatories")
                    )
                if len(self.rrfas):
                    responses.append(
                        str(
                            self.rrfas.request_ordinal
                            + " Set of Requests for Admissions"
                        )
                    )
                output += comma_and_list(responses)
        else:
            output = str(self.utitle)
        return output

    def witness_list_outro(self):
        output = str(
            self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("reserve")
            + " the right to call additional rebuttal witnesses who may or may not be named above. However, as it stands, "
            + self.user_clientlist().asnoun()
            + self.user_clientlist().does_verb("does")
            + " not plan to do so."
        )
        return output

    def negligent_employee(self):
        return self.recipients()

    def response_intro(self):
        log("response_intro . . . . 1")
        if self.draft == "Discovery Response":
            output = "Pursuant to "
            responses = []
            rules = []
            rules.clear()
            output += str(
                self.juris.civil_procedure()["abbreviation"] + " " + comma_and_list(self.discovery_rules()) + ", ")
            output += (
                    self.user_clientlist().asnoun()
                    + " "
                    + self.user_clientlist().does_verb("respond")
                    + " to "
                    + self.recipients().asnoun()
                    + " "
                    + possessify(comma_and_list(self.recipients()), "")
            )
            if self.rrfps:
                responses.append(
                    str(
                        self.rrfps.request_ordinal
                        + " Request for Production of Documents"
                    )
                )
            if self.rrogs:
                responses.append(
                    str(self.rrogs.request_ordinal + " Set of Interrogatories")
                )
            if self.rrfas:
                responses.append(
                    str(self.rrfas.request_ordinal + " Set of Requests for Admissions")
                )
            output += comma_and_list(responses)
            output += " as follows: "

        if self.draft == "RRFP" or self.draft == "RROG" or self.draft == "RRFA":
            output = ""
            output += (
                    self.user_clientlist().asnoun()
                    + " responds to "
                    + self.recipients().asnoun()
                    + " "
                    + possessify(comma_and_list(self.recipients()), self.request_ordinal)
            )
            if self.draft == "RRFP":
                output += " Request for Production of Documents "
            elif self.draft == "RROG":
                output += " Set of Interrogatories "
            elif self.draft == "RRFA":
                output += " Set of Requests for Admissions "
            output += "as follows: "
        elif self.draft in ('Motion', 'Order', 'Declaration'):
            output = ""
        elif self.draft == "RFP":
            output = str(
                "Pursuant to "
                + self.rfp_intro_rules
                + ", "
                + self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("request")
                + " that "
                + self.recipients().asnoun()
                + " produce and make available for inspection and copying all of the documents described herein which are within the custody and control of "
                + self.recipients().asnoun()
                + ". "
                + self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("request")
                + " that said production take place within 30 days from the date of service of this request at the offices of "
                + self.user_firm().name.text
                + ", "
                + self.user_firm().address.address
                + ", "
                + self.user_firm().address.city
                + ", "
                + self.user_firm().address.state
                + ", "
                + self.user_firm().address.zip
                + "."
            )
        elif self.draft == "RFA":
            output = str(
                "Pursuant to "
                + self.rfa_rule
                + ", "
                + self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("request")
                + " the admission by "
                + self.recipients().asnoun()
                + " of the truth of the following facts or opinions of facts, application of law to facts, or the genuineness of the documents referred to herein and attached to this request:"
            )
        elif self.draft == "Jury Instructions":
            output = str(
                self.user_clientlist().asnoun()
                + self.user_clientlist().does_verb("request")
                + " that the court address to the jury the following approved "
                + self.uniform_jury_instructions[1]
                + " and the special jury instructions that follow, separately stated and consecutively numbered:"
            )
        elif self.draft == "MIL":
            output = str(
                self.user_clientlist().asnoun()
                + "  moves for an Order prohibiting "
                + possessify(self.recipients().asnoun(), "")
                + "reference to, offering evidence of, or argument regarding the following matters during the trial of this action:"
                + self.fx.NLINE
            )
        elif self.draft == "Witness List":
            output = str(
                self.user_clientlist().asnoun()
                + comma_and_list(self.user_clientlist())
                + ", by and through "
                + self.user_clientlist().pronoun_possessive("")
                + self.colawyers().as_noun("attorney")
                + ", "
                + self.colawyers()
                + +", "
                + self.user_clientlist().does_verb("give")
                + " notice that "
                + self.user_clientlist().pronoun_subjective()
                + self.user_clientlist().does_verb("intend")
                + " to call the following witnesses to testify at trial in this matter:"
            )
        elif self.draft == "Complaint":
            if any(
                    party.party.name.caption_type.name.text == "personal representative for an estate"
                    for party in self.user_clientlist()
            ):
                output = str(
                    self.estate_reps().pname
                    + ", in "
                    + self.estate_reps().pronoun_possessive(
                        "capacity as Personal Representative of the Estate of "
                    )
                    + self.decedent().name.full().upper()
                    + ", "
                    + self.estate_reps().does_verb("allege")
                    + ":"
                )
            else:
                output = str(self.user_clientlist().pname_does("allege") + ":")

        elif self.draft == "Depo Notice":
            if self.remote_depo:
                output = str(
                    self.fx.BOLD
                    + "PLEASE TAKE NOTICE"
                    + self.fx.BOLD
                    + " that pursuant to "
                    + self.depo_rule
                    + ", on the date and at the location indicated below, or on an alternative date and location agreed by the parties and confirmed in writing, "
                    + self.user_clientlist().asnoun()
                    + " will take the deposition of "
                    + self.recipients().asnoun()
                    + " "
                    + comma_and_list(self.recipients())
                    + ". The deposition will be taken by remote audio/video conference and will be taken upon oral examination, before a court reporter who is duly authorized to administer oaths by the laws of the state. "
                )
            else:
                output = str(
                    self.fx.BOLD
                    + "PLEASE TAKE NOTICE"
                    + self.fx.BOLD
                    + "that pursuant to "
                    + self.depo_rule
                    + ", on the date and at the location indicated below, or on an alternative date and location agreed by the parties and confirmed in writing, "
                    + self.user_clientlist().asnoun()
                    + " will take the deposition of "
                    + self.recipients().asnoun()
                    + " "
                    + comma_and_list(self.recipients())
                    + " at the offices of "
                    + self.user_firm().name.text
                    + self.user_firm().address.line_one()
                    + ", "
                    + self.user_firm().address.line_two()
                    + ", before a court reporter who is duly authorized to administer oaths by the laws of the state."
                )
            if (
                    not this_thread.evaluation_context == "docx"
                    and not self.pmk
                    and not self.staff_witness_depo
            ):
                output += str(
                    " [BEGIN_TWOCOL] "
                    + self.fx.BOLD
                    + "PERSON TO BE EXAMINED:"
                    + self.fx.BOLD
                    + " [BREAK] "
                    + comma_and_list(self.recipients())
                    + " [END_TWOCOL] "
                )
                output += str(
                    " [BEGIN_TWOCOL] "
                    + self.fx.BOLD
                    + "DATE AND TIME OF DEPOSITION:"
                    + self.fx.BOLD
                    + " [BREAK] "
                    + format_date(self.proposed_depo_date)
                    + str(
                        str(self.fx.NLINE + format_time(self.proposed_depo_time))
                        if self.proposed_depo_time
                        else ""
                    )
                    + " [END_TWOCOL] "
                )
                output += str(
                    " [BEGIN_TWOCOL] "
                    + self.fx.BOLD
                    + "PLACE OF DEPOSITION:"
                    + self.fx.BOLD
                    + " [BREAK] "
                    + self.depo_location.name.text
                    + self.fx.NLINE
                    + self.depo_location.address.block()
                    + " [END_TWOCOL] "
                )
                output += str(
                    " [BEGIN_TWOCOL] "
                    + self.fx.BOLD
                    + "COURT REPORTER:"
                    + self.fx.BOLD
                    + " [BREAK] "
                    + self.court_reporter.name.text
                    + str(
                        str(
                            self.fx.NLINE
                            + phone_number_formatted(self.court_reporter.phone_number)
                        )
                        if self.court_reporter.phone_number
                        else ""
                    )
                    + " [END_TWOCOL] "
                )
            if self.staff_witness_depo:
                staffdepo = ""
                staffdepo += str(
                    self.recipients().pname_possessive()
                    + " staff, employees, or agents who witnessed the "
                    + self.incident
                    + " at issue in this matter, the aftermath, or who were present or on duty at the location at the time of the "
                    + self.incident
                    + ", on "
                    + format_date(self.doi)
                    + " at "
                    + format_time(self.toi)
                    + "."
                )
            if self.pmk:
                output += str(
                    " Pursuant to "
                    + self.juris.civil_procedure()["Notice or Subpoena Directed to an Organization"]
                    + ", "
                    + self.user_clientlist().asnoun()
                    + " "
                    + self.user_clientlist().does_verb("advise")
                    + " "
                    + self.recipients().asnoun()
                    + " of "
                    + self.recipients().pronoun_possessive("")
                    + "duty to designate such officers, directors, managing agents, or other persons who consent to testify on "
                    + possessify(self.recipients().asnoun(), "")
                    + "behalf with regard to the following matters, and to set forth for each person designated the matters on which such person must testify:"
                )
        log("response_intro . . . . 2")
        return output

    def pmk_bullets(self):
        bullets = []
        bullets.clear()
        bullets.append(
            "The "
            + self.incident
            + ", including witness statements, employee statements, observations of the incident or of "
            + self.user_clientlist().salute()
            + "."
        )
        bullets.append(
            "Other "
            + self.incident
            + "s similar to the "
            + self.incident
            + " at issue."
        )
        if "Dram Shop" in self.claims:
            bullets.append(
                str(
                    "Background information of your company, including its organizational structure, the services it provides, and its compliance with any applicable state licensure requirements."
                )
            )
            bullets.append(
                str(
                    "Rules, policies, procedures, training and supervision concerning service of alcohol, and specifically including detecting intoxication and how to respond to visible intoxication. This includes any OLCC server education designees."
                )
            )
            bullets.append(
                str(
                    "Knowledge of and past interactions with "
                    + self.adverse_individuals().asnoun()
                    + " "
                    + comma_and_list(self.adverse_individuals())
                    + "."
                )
            )
            bullets.append(
                str("Layout of your bar as of " + format_date(self.doi) + ".")
            )
            bullets.append(
                str(
                    "Alcohol and food available at your bar as of  "
                    + format_date(self.doi)
                    + "."
                )
            )
            bullets.append(
                str(
                    "Identity and last known contact information for staff on duty at your bar referenced in the operative complaint the night of  "
                    + format_date(self.doi)
                    + "."
                )
            )
            bullets.append(
                str(
                    "The incident alleged in the operative complaint, including investigation, witnesses, statements, incident reports, surveillance video, photographs, receipts, and bar logs."
                )
            )
            bullets.append(
                str(
                    "Your first notice of the "
                    + self.incident
                    + " that is the subject of this lawsuit."
                )
            )
            bullets.append(
                str(
                    "You first notice of any claim or allegation that "
                    + self.adverse_individuals().asnoun()
                    + " "
                    + comma_and_list(self.adverse_individuals())
                    + " "
                    + self.adverse_individuals().did_verb("was")
                    + " at your bar on  "
                    + format_date(self.doi)
                    + "."
                )
            )
            bullets.append(
                str(
                    "Any complaints or allegations of alleged service of alcohol to others who were visibility intoxicated from "
                    + format_date(current_datetime() - date_interval(years5))
                    + " to the present."
                )
            )
            bullets.append(
                str(
                    "Allegations made in the Answer and Affirmative Defense to the "
                    + self.operative_complaint
                    + "."
                )
            )

        if "Breach of PIP Contract" in self.claims:
            bullets.append("The claim file regarding the matter at issue.")
            bullets.append(
                "Information regarding the payment or nonpayment for any portion of "
                + self.user_clientlist().pname_possessive()
                + " claim for PIP benefits."
            )
            bullets.append(
                "Policy and procedure relative to "
                + self.recipients().pname_possessive()
                + "investigations into claims for PIP benefits that were in effect from "
                + format_date(self.doi)
                + " to the present."
            )
            bullets.append(
                self.recipients().pname_possessive()
                + "investigation into "
                + possessify(self.user_clientlist().asnoun(), "")
                + "claim for PIP benefits."
            )
        if "Premises Liability" in self.claims:
            bullets.append(
                "Training rules, policies, procedures, schedules, and supervision of "
                + possessify(self.recipients(), "")
                + " employees’ duties to clean, inspect, and/or maintian the "
                + self.the_hazard
                + " and the "
                + self.hazloc.name.text
                + "."
            )
            bullets.append(
                "Investigation and retention policies and procedures regarding customer complaints."
            )
            bullets.append(
                "Investigation and retention policies and procedures regarding incident reports."
            )
            bullets.append(
                "Investigation and retention policies and procedures regarding reports of injuries on the premises."
            )
            bullets.append(
                "Rules, policies, procedures, schedules, and supervision concerning capturing and retention of security cameras at the premises where the "
                + self.incident
                + " occured."
            )

        if (
                self.recipients() in self.the_ds()
        ):
            if self.recipients().has_answered:
                bullets.append(
                    "The factual allegations denied in "
                    + possessify(self.recipients(), "")
                    + "Answer."
                )
                if len(self.recipients().affds):
                    bullets.append(
                        "The facts underlying the Affirmative Defenses asserted in "
                        + possessify(self.recipients(), "")
                        + "Answer."
                    )
                if self.recipients().answer_denied_negligence:
                    bullets.append(
                        "The specifications of negligence denied in "
                        + possessify(self.recipients(), "")
                        + "Answer."
                    )
        if any(recipient.has_responded_to_rfa for recipient in self.recipients()):
            bullets.append(
                "The facts denied in "
                + possessify(self.recipients(), "")
                + "Response to "
                + possessify(comma_and_list(self.user_clientlist()), "")
                + "Request for Admissions."
            )
        return bullets

    def pmk_instructions(self):
        output = ""
        output += "Pursuant to ORCP 39 C(6), " + self.recipients().asnoun() + " " + self.recipients().does_verb(
            "is") + " required to designate one or more officers, directors, managing agents, or other persons who consent to testify on behalf of " + self.recipients().asnoun() + ", and whom " + self.recipients().asnoun() + " will fully prepare to testify regarding all information that is known or reasonably available to " + self.recipients().asnoun() + " regarding the designated matters set forth below, and whose testimony shall be binding on " + self.recipients().asnoun() + " in this litigation. In addition, " + self.recipients().asnoun() + " " + self.recipients().does_verb(
            "is") + " is required to provide notice of no fewer than three (3) days before the scheduled deposition, absent good cause or agreement of the parties and the deponent, designating the name(s) of the officers, directors, managing agents, or other persons who consent to testify on " + possessify(
            self.recipients().asnoun(),
            "") + "behalf and setting forth, for each person designated, the matters on which such person will testify. To the extent the persons designated to testify on " + possessify(
            self.recipients().asnoun(),
            "") + "behalf regarding the matters below have additional or other discoverable personal knowledge or information beyond the scope of the matters below but related to the parties’ claims and defenses, " + self.user_clientlist().salute() + " hereby provides notice of her intent to further depose such persons at the same time and place as fact witnesses, not as corporate designees, pursuant to ORCP 39 C(1)."
        return output

    def remote_depo_notices(self):
        return (
            "The deposition will be conducted remotely, using audio-visual conference technology.",
            "The court reporter will report the deposition from a remote location separate from the witness.",
            "Counsel for the parties and their clients will be participating from various, separate locations.",
            "The court reporter will administer the oath to the witness remotely.",
            "The witness will be required to provide government-issued identification satisfactory to the court reporter and this identification must be legible on camera.",
            "The court reporter will record the testimony.",
        )

    def depo_outro(self):
        output = str(
            "You are invited to attend and cross-examine. The oral examination will continue from day to day until completed, unless you are sooner discharged, before a court reporter duly authorized to administer oaths by the laws of the State of "
            + state_name(self.juris.state.name.text)
            + "."
        )
        return output

    def dear_recipients(self):
        log("dear_recipients . . . 1")
        dear_recipients_output = str(
            "TO: "
            + self.recipients().asnoun()
            + " "
            + comma_and_list(self.recipients())
            + " and "
            + self.recipients().pronoun_possessive("")
            + self.recipients_attorneys().as_noun("attorney")
            + ", "
            + comma_and_list(self.recipients_attorneys())
            + "."
        )
        log("dear_recipients . . . 1")
        return dear_recipients_output

    # Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
    def impetra_rfps(self):
        output = SCList("impetra_rfps", there_are_any=True, auto_gather=False)
        output.clear()
        output.append(
            str(
                "Any and all photographs, either laser prints, copies made from negatives, (not photocopies of photographs), or original digital files in a commonly used format (such as .JPEG or .TIFF), videotapes, or other visual media (including digital video files), of the "
                + str(
                    str(
                        "vehicles, "
                        if "MVA" in self.negligence_type
                           and all(
                            party.party.mva_status != "Pedestrian"
                            for party in self.user_clientlist()
                        )
                        else str(
                            "vehicle, "
                            if "MVA" in self.negligence_type
                               and any(
                                party.party.mva_status == "Pedestrian"
                                for party in self.user_clientlist()
                            )
                            else str(
                                self.breed + ", The Premises, "
                                if "Animal Attack" in self.negligence_type
                                else ""
                            )
                        )
                    )
                    if "Negligence" in self.claims
                    else ""
                )
                + "parties, witnesses, "
                + str("premises, " if "Premises Liability" in self.claims else "")
                + self.incident
                + str(
                    str(
                        ", "
                        + self.drunk().asnoun()
                        + self.drunk()
                        + "on the day of the incident, "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "vehicle on "
                        + self.doi
                        + ", "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "purchase of alcohol on "
                        + self.doi
                        + ", any other individuals’ purchase of alcohol for "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "on "
                        + self.doi
                        + ", any alcohol that was provided to "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "on "
                        + self.doi
                        + ", for free, or otherwise on a bar tab, "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "consumption of alcohol on "
                        + self.doi
                        + ", "
                        + self.drunk().asnoun()
                        + self.drunk().possessive("").upper()
                        + "purchase of food on "
                        + self.doi
                        + ", "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "consumption of food "
                        + self.doi
                        + ", "
                        + "the scene of the subject "
                        + self.incident
                        + ", "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "use of drugs, including prescription medications or recreational drugs, on "
                        + self.doi
                        + ", witnesses to "
                        + self.drunk().asnoun()
                        + possessify(self.drunk().name.full().upper(), "")
                        + "consumption of alcohol "
                        + self.doi
                        + ", or witnesses to the subject "
                        + self.incident
                        + " or its immediate aftermath, including observation of "
                        + self.drunk().asnoun()
                        + self.drunk().name.full().upper()
                        + " through the time he was booked in jail."
                    )
                    if "statutory dram shop" in self.claims
                       or "social host liability" in self.claims
                    else ", or scene in question."
                )
                + " The copies are to be in color, if possible."
            )
        )
        output.append(
            str(
                "Any and all drawings, maps, or sketches of the scene of the "
                + self.incident
                + "."
            )
        )
        output.append(
            str(
                "A copy of any surveillance movies or photographs of "
                + self.user_clientlist().asnoun()
                + "."
            )
        )
        output.append(
            str(
                "Any and all statements previously made by "
                + self.user_clientlist().asnoun()
                + " related to the subject matter of this lawsuit, including any written statement signed or otherwise adopted or approved by "
                + self.user_clientlist().asnoun()
                + " and any stenographic, mechanical, electrical, or other type of recording or any transcription thereof made by "
                + self.user_clientlist().asnoun()
                + " about this lawsuit and contemporaneously recorded."
            )
        )
        output.append(
            str(
                "Any and all statements previously made by "
                + self.recipients().asnoun()
                + " related to the subject matter of this lawsuit, including any written statement signed or otherwise adopted or approved by "
                + self.recipients().asnoun()
                + " and any stenographic, mechanical, electrical, or other type of recording or any transcription thereof made by "
                + self.recipients().asnoun()
                + " about this lawsuit and contemporaneously recorded."
            )
        )
        output.append(
            str(
                "Copies of any and all documents produced pursuant to a medical release signed by "
                + self.user_clientlist().asnoun()
                + ", other than those documents produced by this office."
            )
        )
        if len(self.recipients()) >= 2:
            output.append(
                str(
                    "All documents exchanged between "
                    + self.recipients().asnoun()
                    + " regarding the "
                    + self.incident
                    + ", including those exchanged by and through agents, insurers or attorneys of the same, and including documents that were produced pursuant to any prior request for production."
                )
            )
        output.append(
            str(
                "Any and all videos or photographs of "
                + self.user_clientlist().asnoun()
                + "."
            )
        )
        output.append(
            str(
                "Certified copies of all insurance agreements or policies, including but not limited to, any liability, business, blanket, motor vehicle, personal injury protection, homeowners, commercial, business, farm, or umbrella policies with the accompanying declaration pages (not copies of computer printouts), under which a person transacting insurance may be liable to satisfy part or all of a judgment which may be entered in this action, or to indemnify or reimburse any payment made to satisfy a judgment."
            )
        )
        output.append(
            str(
                "Documents sent to "
                + self.user_clientlist().asnoun()
                + " or "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "agents from "
                + self.recipients().asnoun()
                + " or "
                + possessify(self.recipients().asnoun(), "")
                + "agents regarding the "
                + self.incident
                + ", "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "injuries, or "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "claim for damages."
            )
        )
        output.append(
            str(
                "Documents signed by "
                + self.user_clientlist().asnoun()
                + " or by others on "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "behalf."
            )
        )
        output.append(
            str(
                "Any and all notes of conversations with "
                + self.user_clientlist().asnoun()
                + " or with any others on "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "behalf."
            )
        )
        output.append(
            str(
                "Documents regarding the facts of the "
                + self.incident
                + " at issue, including, but not limited to, incident reports, witness reports, witness statements, employee statements, notes, letters, memoranda, e-mails, and texts."
            )
        )
        output.append(
            str(
                "Documents relating to all postings to or by "
                + self.recipients().asnoun()
                + " on any social media site or other website (including, but not limited to: Facebook, Twitter, Instagram, Pinterest, YouTube, or any other social media platform) regarding the subject "
                + self.incident
                + " or lawsuit (including printouts of related historical content including messages, posts, comments, and photographs)."
            )
        )
        output.append(
            str(
                "Documents containing the names, addresses, or phone numbers of individuals with knowledge of discoverable material regarding the "
                + self.incident
                + str(
                    str(
                        ", "
                        + self.recipients().asnoun()
                        + " consumption of alcohol on "
                        + self.doi
                        + ","
                    )
                    if "Negligence" in self.claims
                       and "MVA" in self.negligence_type
                       and self.dui
                    else ""
                )
                + " or "
                + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                + "claim for damages."
            )
        )
        output.append(
            str(
                "Any and all notes, diaries, journals, or other documents concerning the "
                + self.incident
                + " at issue, or of any statements or observations of the parties or witnesses."
            )
        )
        output.append(
            str(
                "Documents subpoenaed in this case, including documents that were produced pursuant to the subpoena, but not specifically subpoenaed."
            )
        )
        output.append(
            str(
                "The names and contact information for any witness to the "
                + self.incident
                + "."
            )
        )
        output.append(
            str(
                "Documents related to or otherwise indicating the identity of any person, other than the parties to this case, believed by "
                + self.recipients().asnoun()
                + " to be at fault for the "
                + self.incident
                + "."
            )
        )
        for defendant in self.recipients():
            if defendant.name.caption_type.name.text == "individual":
                output.append(
                    str(
                        "Documents relating to any criminal charge of "
                        + defendant.pname()
                        + " at any time, including documents indicating the case name, case number, and venue, as well as any convictions, even if they were subsequently set aside, expunged, or otherwise reduced."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, a copy of any movies or photographs of "
                        + defendant.pname()
                        + " on "
                        + self.doi
                        + "."
                    )
                )
                output.append(
                    str(
                        "Documents relating to any civil lawsuit "
                        + defendant.pname()
                        + " has been involved in at any time, including, but not limited to, documents indicating the case name, case number, and venue; copies of any pleadings filed in those lawsuits, and copies of any deposition transcripts relative to those lawsuits."
                    )
                )
            if defendant.name.caption_type.name.text == "business entity":
                output.append(
                    str(
                        "Documents indicating whether and when "
                        + possessify(self.user_clientlist().asnoun(), self.incident)
                        + " was discussed at any time, including the substance of any such discussion. This request does not seek attorney client privileged information, but does seek safety committee meetings, personnel meetings, trainings, and any other discussion or meeting where the "
                        + self.incident
                        + " at issue was either discussed or was made a factual basis for any such discussion."
                    )
                )
                if "Corporation" in defendant.business_entity.name.text:
                    output.append(
                        str(
                            "Documents relating to "
                            + defendant.party_type
                            + " "
                            + possessify(defendant.name.full().upper(), "")
                            + "corporate structure on "
                            + self.doi
                            + "."
                        )
                    )
        if (
                "statutory dram shop" in self.claims
                or "social host liability" in self.claims
        ):
            output.append(
                str(
                    "To the extent not otherwise requested above, copy of any surveillance movies or photographs of  "
                    + self.drunk().asnoun()
                    + self.drunk().name.full().upper()
                    + " on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Certified copies of all insurance agreements or policies, including but not limited to, any auto, liability, business, dram shop, blanket, or umbrella policies with the accompanying declaration pages (not copies of computer printouts) under which a person transacting insurance may be liable to satisfy part or all of a judgment which may be entered in this action, or to indemnify or reimburse any payment made to satisfy a judgment."
                )
            )
            output.append(
                str(
                    "A copy of "
                    + self.dram_shops().asnoun().possessive("")
                    + " business licenses."
                )
            )
            output.append(
                str(
                    "Any and all documents containing the names, addresses, or phone numbers of individuals with knowledge of discoverable material regarding the incident and/or "
                    + self.parties.the_ps().asnoun().possessive("")
                    + " claim for damages."
                )
            )
            output.append(
                str(
                    "Documents indicating ownership of the subject premises at the time of the incident."
                )
            )
            output.append(
                str(
                    "A copy of the liquor license "
                    + self.dram_shops().asnoun()
                    + self.dram_shops().did_verb("was")
                    + "operating under at the time of the incident."
                )
            )
            output.append(
                str(
                    "A copy of all applications made for liquor licenses or renewals thereof submitted by or on behalf of "
                    + self.dram_shops().asnoun()
                    + "from "
                    + format_date(current_datetime() - date_interval(years=10))
                    + " to the present."
                )
            )
            output.append(
                str(
                    "Complete copies of all OLCC permits and certifications held by Defendants’ employees who worked at the subject premises on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "A copy of the till records from "
                    + format_date(self.doi - date_interval(days=1))
                    + ", and "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents referencing, evidencing, or reflecting the different types of alcoholic beverages sold at the subject premises on the date of the incident, including documents indicating the price for sale of the beverages to the general public, and the containers in which the beverages were sold, i.e., size of glass or other container."
                )
            )
            output.append(
                str(
                    "Documents referencing, evidencing, or reflecting the different types of menu items, other than beverages, sold at the subject premises on the date of the incident, including documents indicating the items available, price for sale of those items, and the serving sizes in which the items were sold, and the hours during which those items could be purchased."
                )
            )
            output.append(
                str(
                    "Documents referencing, evidencing, or reflecting the names, last known addresses, last known telephone numbers, and last known e-mail addresses of employees on duty on "
                    + self.doi
                    + ", at the subject premises, as well as said individuals’ records indicating the days and hours said individuals worked and, if possible, in what capacity they did so. This request is limited to those employees working in the bar or kitchen of each premises."
                )
            )
            output.append(
                str(
                    "Documents, files, or memoranda maintained by "
                    + self.dram_shops().asnoun()
                    + " concerning "
                    + self.doi
                    + ", including any accident reports, incident reports, internal and memoranda, staff schedules, external communications, photographs, witness statements, and surveillance video recordings."
                )
            )
            output.append(
                str(
                    "Transcripts or other recordings of testimony from "
                    + format_date(current_datetime() - date_interval(years=5))
                    + ", to the present, whether given during a trial, hearing, or deposition, of any person employed by Defendants in any case involving the sale of alcoholic beverages at the subject premises or any other establishment owned and/or operated by Defendants."
                )
            )
            output.append(
                str(
                    "Copies of any and all liability insurance applications made by "
                    + self.dram_shops().asnoun()
                    + " from "
                    + format_date(current_datetime() - date_interval(years=10))
                    + ", to the present, including any documents submitted with said applications, or otherwise made a supplement of the same."
                )
            )
            output.append(
                str(
                    "Documents regarding the policies, protocols, and procedures effective on "
                    + self.doi
                    + ", concerning 1) the provision of alcoholic beverages to customers, 2) the recognition of visible intoxication, 3) how to respond to visible intoxication, 4) the prevention of injuries caused by alcohol consumption, 5) documenting visible intoxication of patrons, 6) documenting actions taken in response to visible intoxication of patrons, and 6) the method of removal of intoxicated patrons from the premises."
                )
            )
            output.append(
                str(
                    "To the extent not otherwise requested above, a complete copy of all documentation, including handbooks, safety manuals, videotapes or other records which "
                    + self.dram_shops().asnoun()
                    + " relied upon to train its employees regarding 1) the provision of alcoholic beverages to customers, 2) the recognition of visible intoxication, 3) how to respond to visible intoxication, 4) the prevention of injuries caused by alcohol consumption, 5) documenting visible intoxication of patrons, 6) documenting actions taken in response to visible intoxication of patrons, and 6) the method of removal of intoxicated patrons from the premises."
                )
            )
            output.append(
                str(
                    "To the extent not otherwise requested above, a complete copy of all documents concerning any changes to policies or procedures of "
                    + self.dram_shops().asnoun()
                    + " made since the incident concerning 1) the provision of alcoholic beverages to customers, 2) the recognition of visible intoxication, 3) how to respond to visible intoxication, 4) the prevention of injuries caused by alcohol consumption, 5) documenting visible intoxication of patrons, 6) documenting actions taken in response to visible intoxication of patrons, and 6) the method of removal of intoxicated patrons from the premises."
                )
            )
            output.append(
                str(
                    "A complete copy of all training materials and documentation provided to "
                    + self.dram_shops().asnoun()
                    + "  by the Oregon Liquor Control Commission or any other entity which relate in any way to the service of alcohol of visibly intoxicated patrons."
                )
            )
            output.append(
                str(
                    "Complete copies of any and all logs maintained by "
                    + self.dram_shops().asnoun()
                    + " related to refusal to serve alcohol to patrons due to visible intoxication."
                )
            )
            output.append(
                str(
                    "Complaints or citations issued by the Oregon Liquor Control Commission or any other governmental entity relating to the sale of alcoholic beverages by "
                    + self.dram_shops().asnoun()
                    + "."
                )
            )
            output.append(
                str(
                    "Documents concerning any claims of service, or actual service, of visibly intoxicated patrons by "
                    + self.dram_shops().asnoun()
                    + " at any time, including, but not limited to, reprimand of any employee for the same."
                )
            )
            output.append(
                str(
                    "Documents describing, identifying, and providing contact information for individuals who may have witnessed "
                    + self.drunk().asnoun()
                    + self.drunk().possessive("").upper()
                    + " presence at "
                    + self.dram_shops().asnoun().possessive("")
                    + "premises on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents describing or pertaining to measures taken by "
                    + self.dram_shops().asnoun()
                    + " to prevent the intoxication of  "
                    + self.drunk().asnoun()
                    + self.drunk().name.full().upper()
                    + ", or any other patron at the involved premises on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Any and all credit card, debit card, checks or other receipts or documents for purchased made by, or for, "
                    + self.drunk().asnoun()
                    + self.drunk().upper()
                    + " or anybody else in his party on "
                    + self.doi
                    + ". This includes purchases of food, alcoholic drinks and nonalcoholic drinks."
                )
            )
            output.append(
                str(
                    "Documents identifying the type and number of alcoholic beverages sold to, served to, or consumed by  "
                    + self.drunk().asnoun()
                    + self.drunk().name.full().upper()
                    + " on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents identifying the type and number of nonalcoholic beverages sold to, served to, or consumed by "
                    + self.drunk().asnoun()
                    + self.drunk().upper()
                    + " on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents identifying the type and amount of food sold to, served to, or consumed by "
                    + self.drunk().asnoun()
                    + self.drunk().name.full().upper()
                    + " on "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents indicating "
                    + self.drunk().asnoun()
                    + possessify(self.drunk().name.full().upper(), "")
                    + "work schedule from "
                    + format_date(self.doi - date_interval(days=7))
                    + ", through "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents indicating "
                    + self.drunk().asnoun()
                    + possessify(self.drunk().name.full().upper(), "")
                    + "cell phone carrier and cell phone number as of "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Documents indicating "
                    + self.drunk().asnoun()
                    + possessify(self.drunk().name.full().upper(), "")
                    + "incoming and outgoing phone calls on the day of the incident, including the time of any such calls, the phone numbers involved, and the identity of the persons involved in any such calls."
                )
            )
            output.append(
                str(
                    "Documents indicating "
                    + self.drunk().asnoun()
                    + possessify(self.drunk().name.full().upper(), "")
                    + "texts received and sent on "
                    + self.doi
                    + ", along with the text messages themselves, including the time of any such texts, the phone numbers involved, and the identity of the persons involved in any such texts."
                )
            )
            output.append(
                str(
                    "Copies of any and all documents where "
                    + self.dram_shops().asnoun()
                    + " discuss, reference, or otherwise comment on  "
                    + self.drunk().asnoun()
                    + possessify(self.drunk().name.full().upper(), "")
                    + " level of intoxication at any time on the day of "
                    + self.doi
                    + "."
                )
            )
            output.append(
                str(
                    "Non-privileged statements made by anyone, whether written, transcribed, summarized or maintained in any electronic format, concerning any allegation in this lawsuit."
                )
            )
            output.append(
                str(
                    "All documents concerning any potential or actual reprimand of any employee of "
                    + self.dram_shops().asnoun()
                    + " concerning the incident."
                )
            )
            output.append(
                str(
                    "Copies of any and all documents relating to any civil lawsuit "
                    + self.dram_shops().asnoun()
                    + " have been involved in at any time for liquor liability."
                )
            )
            for defendant in self.dram_shops():
                output.append(
                    str(
                        "The name of any individual who served  "
                        + self.drunk().asnoun()
                        + self.drunk().upper()
                        + " alcohol at "
                        + defendant.name.full()
                        + ", as well as her last known contact information, including address, e-mail address and phone number."
                    )
                )
        if "Negligence" in self.claims:
            if "MVA" in self.negligence_type:
                output.append(
                    str(
                        "A copy of front and back of "
                        + possessify(self.recipients().asnoun(), "")
                        + "driver license."
                    )
                )
                output.append(
                    str(
                        "A copy of the title to the vehicle "
                        + self.recipients().asnoun()
                        + " was driving at the time in question."
                    )
                )
                output.append(
                    str(
                        "Documents reflecting repair work performed on the vehicle after the collision, including any and all damage appraisals, estimates, work orders, and invoices for such repair work or appraisal of the vehicles following the collision."
                    )
                )
                output.append(
                    str(
                        "Any and all documentation indicating the speed "
                        + self.recipients().asnoun()
                        + " was travelling for the 30 seconds immediately preceding the "
                        + self.incident
                        + ", including, but not limited to, GPS tracking data, cellular phone tracking data, and vehicle sensor data. If any such data exists but was not recorded in the 30 seconds immediately preceding the collision, the last data available immediately prior to the collision is requested."
                    )
                )
                if self.dui:
                    output.append(
                        str(
                            "Documents reflecting the name and contact information of any business or person who served "
                            + self.recipients().asnoun()
                            + " on "
                            + self.doi
                            + "."
                        )
                    )
                if (
                        self.dui
                        or "statutory dram shop" in self.claims
                        or "social host liability" in self.claims
                ):
                    output.append(
                        str(
                            "Any and all credit card, debit card, checks or other receipts or documents for purchased made by, or for, "
                            + self.recipients().asnoun()
                            + " or anybody else in his party on "
                            + self.doi
                            + ".  This includes purchases of food, alcoholic drinks and nonalcoholic drinks."
                        )
                    )
                    output.append(
                        str(
                            "Documents identifying the type and amount of alcoholic beverages sold to, served to, or consumed by "
                            + self.recipients().asnoun()
                            + " on "
                            + self.doi
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "Documents identifying the type and amount of nonalcoholic beverages sold to, served to, or consumed by "
                            + self.recipients().asnoun()
                            + " on "
                            + self.doi
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "Documents identifying the type and amount of food sold to, served to, or consumed by "
                            + self.recipients().asnoun()
                            + " on "
                            + self.doi
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "Documents indicating "
                            + possessify(self.recipients().asnoun(), "")
                            + "work schedule during the seven days immediately preceding "
                            + self.doi
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "Documents indicating "
                            + possessify(self.recipients().asnoun(), "")
                            + "cell phone carrier and cell phone number as of "
                            + self.doi
                            + "."
                        )
                    )
                    output.append(
                        str(
                            "Documents indicating "
                            + possessify(self.recipients().asnoun(), "")
                            + "incoming and outgoing phone calls on "
                            + self.doi
                            + ", including the time of any such calls, the phone numbers involved, and the identity of the persons involved in any such calls."
                        )
                    )
                    output.append(
                        str(
                            "Documents indicating "
                            + possessify(self.recipients().asnoun(), "")
                            + "texts received and sent on "
                            + self.doi
                            + ", along with the text messages themselves, including the time of any such texts, the phone numbers involved, and the identity of the persons involved in any such texts."
                        )
                    )
                    output.append(
                        str(
                            "Copies of any and all documents where "
                            + self.recipients().asnoun()
                            + " "
                            + comma_and_list(self.recipients()).does_verb("discuss")
                            + ", reference, or otherwise comment on "
                            + possessify(self.recipients().asnoun(), "")
                            + "level of intoxication at any on "
                            + self.doi
                            + "."
                        )
                    )
        if "Animal Attack" in self.negligence_type:
            output.append(
                str("Any and all veterinary records for the " + self.breed + ".")
            )
            output.append(
                str(
                    "Documents relating to ownership of the "
                    + self.breed
                    + " at the time of the "
                    + self.incident
                    + ", including but not limited to settlement agreements, certificate of title, text messages, emails, diaries, notes, reports, receipts, or journals."
                )
            )
            output.append(
                str(
                    "Documents relating to the purchase or sale of the "
                    + self.breed
                    + "."
                )
            )
            output.append(
                str("Any and all vaccination certificates for the " + self.breed + ".")
            )
            output.append(
                str(
                    "Documents relating to the import of the "
                    + self.breed
                    + " to the United States, including, but not limited to, any documents required by any federal or state agencies."
                )
            )
            output.append(
                str(
                    "Documents relating to the names of any medications the "
                    + self.breed
                    + " was ever administered."
                )
            )
            output.append(
                str(
                    "Documents relating to the diagnoses for which any medications were administered to the "
                    + self.breed
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to another dog biting, attacking, or behaving aggressively toward the "
                    + self.breed
                    + ", including, but not limited to, text messages, emails, social media posts, letters, agreements, veterinary records, medical records, reports, notes, journals, or diaries."
                )
            )
            output.append(
                str(
                    "Documents relating to the "
                    + self.breed
                    + " biting, attacking, or behaving aggressively toward another dog, including, but not limited to, text messages, emails, social media posts, letters, agreements, veterinary records, medical records, reports, notes, journals, or diaries."
                )
            )
            output.append(
                str(
                    "Documents relating to the "
                    + self.breed
                    + " biting, attacking, or behaving aggressively toward any person, including, but not limited to, text messages, emails, social media posts, letters, agreements, veterinary records, medical records, reports, notes, journals, or diaries."
                )
            )
            output.append(
                str(
                    "Any complaints relating to the "
                    + self.breed
                    + " for any reason, including, but not limited to, for barking, growling, behaving aggressively, biting, attacking, roaming, being unconfined, constituting an animal nuisance, or being a dangerous dog."
                )
            )
            output.append(
                str(
                    "Documents reflecting the purchase or ownership of any equipment used or intended for use with the "
                    + self.breed
                    + ", including, but not limited to, any kennel, restraint, barrier, or muzzle systems."
                )
            )
            output.append(
                str(
                    "Documents reflecting the "
                    + possessify(self.breed, "")
                    + "physical attributes on or about "
                    + self.doi
                    + ", including, but not limited to, the "
                    + possessify(self.breed, "")
                    + "weight in pounds, age, gender, and whether or not it was spayed/neutered."
                )
            )
            output.append(
                str(
                    "Documents relating to ownership, control, or possession of “The Premises” on "
                    + self.doi
                    + ", including but not limited to, title, deeds, leases, text messages, emails, social media posts, letters, agreements, records, reports, notes, journals, DMV records, registered mailing addresses with any company or organization, or diaries."
                )
            )
            output.append(
                str(
                    "Documents relating to any training or need to train the "
                    + self.breed
                    + ", including, but not limited to, text messages, emails, social media posts, letters, notices, agreements, veterinary records, medical records, reports, notes, journals, or diaries."
                )
            )
            output.append(
                str(
                    "Any and all notices received from any government agency relating to the "
                    + self.breed
                    + " received at any time."
                )
            )
        if "Products Liability" in self.claims:
            output.append(
                str(
                    "Documents relating to "
                    + possessify(self.recipients().asnoun(), "purchase")
                    + "of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to the year, make, model and size of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to the components of the [PRODUCT] at the time at issue."
                )
            )
            output.append(
                str(
                    "Documents related to warnings or instructions relating to the [PRODUCT] by the manufacturer or person/entity from whom "
                    + self.recipients().asnoun()
                    + " purchased the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to the design according to which the [PRODUCT] or any component part thereof was manufactured, including but not limited to, the design, blueprint, or any other manufacturing procedure."
                )
            )
            output.append(
                str(
                    "Documents relating to any changes made after the "
                    + self.incident
                    + " to the design according to which the [PRODUCT] or any component part thereof was manufactured, including but not limited to, the design, blueprint, or any other manufacturing procedure. "
                    + self.user_clientlist().asnoun()
                    + +" acknowledges that "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407 generally makes such evidence inadmissible; however, pursuant to "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407, it may be admissible for another purpose."
                )
            )
            output.append(
                str(
                    "Documents relating to the replacement of [COMPONENT]s of [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to the construction, assembly or maintenance of the [PRODUCT], including documents created or kept by "
                    + self.recipients().asnoun()
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to any design or process involved in the production or manufacture of the [PRODUCT], specifically including any portion of the [PRODUCT] designed to prevent [INJURIOUS RESULT]."
                )
            )
            output.append(
                str(
                    "Documents relating to any change made in the design or process involved in the production or manufacture of the [PRODUCT]. "
                    + self.user_clientlist().asnoun()
                    + +" acknowledges that "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407 generally makes such evidence inadmissible; however, pursuant to "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407, it may be admissible for another purpose."
                )
            )
            output.append(
                str(
                    "Documents relating to the design, according to which products similar to the [PRODUCT] or any component thereof are currently manufactured, including, but not limited to, the design, blueprint, or any other manufacturing procedure."
                )
            )
            output.append(
                str(
                    "Documents relating to any statement, printed or graphic representation, catalogue, circular, manual, brochure, report, memorandum, transcript, communication, letter, label, or other document which in any way mentions, describes, or otherwise refers to the design according to which products similar to the [PRODUCT] or any component thereof are currently manufactured."
                )
            )
            output.append(
                str("Documents relating to any component part of the [PRODUCT].")
            )
            output.append(
                str(
                    "Documents relating to any component part of the [PRODUCT] that was purchased, leased or otherwise obtained by "
                    + self.recipients().asnoun()
                    + " from a third party.party."
                )
            )
            output.append(
                str(
                    "Documents relating to any test, procedure, inspection, or examination which "
                    + self.recipients().asnoun()
                    + " performed on each component part of the [PRODUCT] that was purchased, leased or otherwise obtained by "
                    + self.recipients().asnoun()
                    + " from a third party.party."
                )
            )
            output.append(
                str(
                    "Documents included or otherwise packaged with the [PRODUCT] or "
                    + self.recipients().asnoun()
                    + " otherwise believes "
                    + self.user_clientlist().asnoun()
                    + " should have received upon purchase of the [PRODUCT], including, but not limited to, promotional materials, manuals, instructions, and warranties."
                )
            )
            output.append(
                str(
                    "Documents relating to the documents included or otherwise packaged with the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to any change in the Product's condition between the time it was sold by [DEFENDANT MANUFACTURER] to [DEFENDANT SELLER] and the time it was sold by [DEFNDANT SELLER] to "
                    + self.user_clientlist().asnoun()
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to any change in the Product's condition between the time it was sold to "
                    + self.user_clientlist().asnoun()
                    + " and the time of its failure."
                )
            )
            output.append(
                str(
                    "Documents relating to the intended or ordinary use of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to any representations made by [DEFENDANT MANUFACTURER] to [DEFENDANT SELLER] relating to the safety, fitness, or capacity of the [PRODUCT] for consumer use."
                )
            )
            output.append(
                str(
                    "Documents relating to any representations made by [DEFENDANT MANUFACTURER] to [DEFENDANT SELLER] relating to the ordinary or intended use of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents not otherwise requested which relates to the design or manufacture of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents relating to any testing or inspection performed by "
                    + self.recipients().asnoun()
                    + ", or any agent or employee thereof, on the [PRODUCT] or similar products, including any component parts of the [PRODUCT] or similar products, prior to the date of purchase alleged in the complaint."
                )
            )
            output.append(
                str(
                    "Documents relating to any testing or inspection performed by "
                    + self.recipients().asnoun()
                    + ", or any agent or employee thereof, on the [PRODUCT] or similar products, including any component parts of the [PRODUCT] or similar products, subsequent to the date of purchase alleged in the complaint."
                )
            )
            output.append(
                str(
                    "Any and all patents obtained by "
                    + self.recipients().asnoun()
                    + " with respect to the [PRODUCT], other similar products, and any component parts thereof."
                )
            )
            output.append(
                str(
                    "Each and every law, rule, regulation, standard, statute, ordinance, or other requirement or recommendation established by any local, state, or federal government body or officer, whether legislative, executive, or administrative in character, which deals with, defines, limits, or specifies any aspect of the design, manufacture, composition, distribution, or use of the [PRODUCT] or similar products."
                )
            )
            output.append(
                str(
                    "Each and every law, rule, regulation, standard, statute, ordinance, or other requirement or recommendation which deals with, defines, limits, or specifies any aspect of the design, manufacture, composition, distribution, or use of the [PRODUCT] or similar products, and which was established, published, or promulgated by any professional association, trade association, industry association, or private group."
                )
            )
            output.append(
                str(
                    "Each and every certificate, commendation, seal of approval, or other award which has been awarded or granted to the [PRODUCT] or similar products, or which has been awarded or granted to "
                    + self.recipients().asnoun()
                    + " due to the [PRODUCT] or similar products."
                )
            )
            output.append(
                str(
                    "Documents relating to repair work performed on the Product before and after the accident, including any and all damage appraisals, estimates, work orders, and invoices for such repair work or appraisal of the Product before and following the accident."
                )
            )
            output.append(
                str(
                    "Documents containing the following information relating to the Product: \n\t(a) The date of manufacture;\n\t(b)The location where the Product was manufactured;\n\t(c)The shipping date of the Product to the first purchaser;\n\t(d)The identity of the individual or entity that first purchased the Product;\n\t(e)If the Product was manufactured outside of the United States of America and imported into the US as a finished product, identify the port of entry and which State the Product was imported into;\n\t(f)Identify the [COMPONENT] originally included with the Product when sold;\n\t(g)Identify the [COMPONENT 2] originally included with the Product when sold."
                )
            )
            output.append(
                str(
                    "Documents related to any recall (including any non-CPCS recall) that the Product was subject to, or potentially subject to, including but not limited to customer complaints, internal communications regarding the recall, communications with the Consumers Product Safety Commission, internal investigation or testing related to the recall issue, and public communications concerning the recall. For the purpose of this request, the Product is potentially subject to a recall, if, based on the information known to the "
                    + self.recipients().asnoun()
                    + " about the Product, the Product might have been subject to a specific recall but more information is needed in order to definitively determine if the Product was actually subject to a specific recall."
                )
            )
            output.append(
                str(
                    "Documents related to any recall, warranty work, repair, or service work involving the Product or any component part, such as [COMPONENT]."
                )
            )
            output.append(
                str(
                    "Documents related to the names of all parties involved in the repair, maintenance, or inspection of the [COMPONENT] of the [PRODUCT] in question before and after the "
                    + self.incident
                    + "."
                )
            )
            output.append(
                str(
                    "Copies of Documents in any way related to or evidencing the names of all parties involved in the repair, maintenance, or inspection of the [PRODUCT] in question before and after the "
                    + self.incident
                    + "."
                )
            )
            output.append(
                str(
                    "Documents related to the names, roles, and time periods of involvement of all persons involved in the development, design and testing of the [COMPONENT] of the [PRODUCT] in question before and after the "
                    + self.incident
                    + "."
                )
            )
            output.append(
                str(
                    "Documents related to the names, roles, and time periods of involvement of all persons involved in the development, design and testing of the [PRODUCT] in question before and after the "
                    + self.incident
                    + "."
                )
            )
            output.append(
                str(
                    "Documents related to complaints or reports involving the [PRODUCT] model line for [ISSUES]."
                )
            )
            output.append(
                str(
                    "For the specific version/generation/ manufacturing period of the [PRODUCT] model line which includes the Product: Documents related to the design and specifications of the [COMPONENT] including any testing conducted on any aspect of the [COMPONENT]."
                )
            )
            output.append(
                str(
                    "For the initial version/generation/ manufacturing period of the [PRODUCT] model line: Documents related to the design and specifications of the [COMPONENT] including any testing conducted on any aspect of the [COMPONENT]."
                )
            )
            output.append(
                str(
                    "For the specific version/generation/ manufacturing period of the [PRODUCT] model line which includes the Product: Documents related to difficulties, problems or concerns about the [PRODUCT]'s [COMPONENT] related to or involving the manufacturing of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents related to the names, position, and contact information of the employees who sold or assisted in the sale of the [PRODUCT] to the "
                    + self.user_clientlist().asnoun()
                    + "."
                )
            )
            output.append(
                str(
                    "Documents related to instructions, warnings, labeling or other information relating to the [PRODUCT] or any component thereof provided to "
                    + self.recipients().asnoun()
                    + "."
                )
            )
            output.append(
                str(
                    "Documents related to any defect in the manufacture of the [PRODUCT]."
                )
            )
            output.append(
                str(
                    "Documents related to any representation of the load capacity of the [PRODUCT] made by, or on behalf of, the manufacturer or seller."
                )
            )
            output.append(
                str(
                    "Any and all communications between [DEFENDANT MANUFACTURER] and [DEFENDANT SELLER] relating to the return of or complaints about office chairs of the same or similar model as the Product."
                )
            )
        if "Premises Liability" in self.claims:
            output.append(
                str(
                    "Documents that were handed out, or otherwise made available, to employees at "
                    + possessify(self.recipients().asnoun(), "")
                    + ""
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " facility concerning their job duties that were in use at the time of the "
                    + self.incident
                    + ", including but not limited to employee handbooks, policies, procedures, safety manuals, or the like."
                )
            )
            output.append(
                str(
                    "Documents that were handed out, or otherwise made available, to managers at "
                    + possessify(self.recipients().asnoun(), "")
                    + ""
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " facility concerning their job duties that were in use at the time of the "
                    + self.incident
                    + ", including, but not limited to, manager handbooks, policies, procedures, safety manuals, supervisory guidelines, or the like."
                )
            )
            output.append(
                str(
                    "Documents reflecting the names, position, and contact information of all employees on duty at "
                    + possessify(self.recipients().asnoun(), "")
                    + ""
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " facility  at the time of the "
                    + self.incident
                    + "."
                )
            )
            output.append(
                str(
                    "All documents related to "
                    + possessify(self.recipients().asnoun(), "")
                    + "policies, practices, or procedures for creating or maintaining warnings for any potential unreasonably dangerous condition on The Property."
                )
            )
            output.append(
                str(
                    ""
                    + other_counter.reset()
                    + "Documents relating to changes made since the date of the "
                    + self.incident
                    + " to "
                    + possessify(self.recipients().asnoun(), "")
                    + "policies, practices, or procedures for "
                    + comma_and_list(self.nonnegligence)
                    + ". "
                    + self.user_clientlist().asnoun()
                    + " acknowledges that "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407 generally makes such evidence inadmissible; however, pursuant to "
                    + (
                        "FRE "
                        if self.juris.jurisdiction.name.text == "federal"
                        else "OEC "
                    )
                    + "407, it may be admissible for another purpose."
                )
            )
            output.append(
                str(
                    "Documents relating to any safety committee meetings, or any other meeting of "
                    + self.recipients().asnoun()
                    + " or their employees concerning the use, inspection, or maintenance of "
                    + self.hazloc.name.text
                    + ", limits of "
                    + self.hazloc.name.text
                    + ", using alternative means for the same purposes for which "
                    + self.hazloc.name.text
                    + " was being used by "
                    + self.recipients().asnoun()
                    + " at the time of the "
                    + self.incident
                    + ", or any other topic concerning the use or nonuse of "
                    + self.hazloc.name.text
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to any representation of the capacity of "
                    + self.hazloc.name.text
                    + " made by, or on behalf of, the manufacturer."
                )
            )
            output.append(
                str(
                    "Documents relating to "
                    + possessify(self.recipients().asnoun(), "")
                    + "policies, customs, or procedures concerning the use, inspection, maintenance, or repair of "
                    + self.hazloc.name.text
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to the policies, customs or procedures generally used in "
                    + possessify(self.recipients().asnoun(), "")
                    + "industry concerning the use, inspection, or maintenance of "
                    + self.negligent_implement
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to the purchase or use of any substances, machines, or devices used at the Property to mitigate "
                    + self.the_hazard
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to any agreements regarding the inspection, maintenance, servicing, or repair of "
                    + self.hazloc.name.text
                    + "."
                )
            )
            output.append(
                str(
                    "Documents relating to any warnings or complaints made to or by any person or entity relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities."
                )
            )
            output.append(
                str(
                    "Documents relating to the custody, control, ownership, title, rental, or lease of The Property at the time of the self.incident."
                )
            )
            output.append(
                str(
                    "Documents relating to any persons other than "
                    + self.user_clientlist().asnoun()
                    + ", including agents or employees of "
                    + self.recipients().asnoun()
                    + ", who had been injured in or near "
                    + self.hazloc.name.text
                    + ", or who had complained about conditions in The Property and referencing any failure to inspect or maintain The Property at any time from "
                    + format_date(self.doi - date_interval(years=10))
                    + ", to the present."
                )
            )
            output.append(
                str(
                    "Documents relating to the warning, instruction, or reprimand of any employee or agent of "
                    + self.recipients().asnoun()
                    + ", or person or entity otherwise employed or contracted by "
                    + self.recipients().asnoun()
                    + ", relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities made at any time from "
                    + format_date(self.doi - date_interval(years=10))
                    + " to the present"
                )
            )
            output.append(
                str(
                    "Documents relating to complaints by anyone, including any employee or agent of "
                    + self.recipients().asnoun()
                    + ", or person or entity otherwise employed or contracted by "
                    + self.recipients().asnoun()
                    + ", relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities on The Property made at any time from "
                    + format_date(self.doi - date_interval(years=10))
                    + " to the present."
                )
            )
            output.append(
                str(
                    "Documents transmitted to "
                    + self.recipients().asnoun()
                    + " from any governmental or regulatory agency relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities on The Property at any time from "
                    + format_date(self.doi - date_interval(years=10))
                    + " to the present."
                )
            )
            output.append(
                str(
                    "Documents relating to safety, condition, or inspection of "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities on The Property."
                )
            )
            output.append(
                str(
                    "Documents relating to injuries or claims of injury made by any person at any time relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities on The Property."
                )
            )
            output.append(
                str(
                    "Documents relating to inspection, maintenance, servicing, or repair records relating to "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities on The Property at any time from "
                    + format_date(self.doi - date_interval(years=10))
                    + " to the present."
                )
            )
            output.append(
                str(
                    "Documents reflecting the names, position, and contact information of all individuals or entities who, at the time of the "
                    + self.incident
                    + ", were responsible for the inspection, maintenance, servicing, or repair of "
                    + self.hazloc.name.text
                    + "; or any related devices, tools, appliances, machinery, equipment, implements, or instrumentalities."
                )
            )
            output.append(
                str(
                    "Documents relating to the facts of the "
                    + self.incident
                    + " at issue, including but not limited to, self.incident reports, witness reports, witness statements, employee statements, notes, letters, memoranda, e-mails, and texts."
                )
            )
            output.append(
                str(
                    "Any and all communications sent from "
                    + self.user_clientlist().asnoun()
                    + " or those on "
                    + self.user_clientlist().pronoun_possessive("")
                    + "behalf. "
                    + self.recipients().asnoun()
                    + " need not respond with any correspondence previously sent to or from this office."
                )
            )
            output.append(
                str(
                    "Any and all notes of communications with "
                    + self.user_clientlist().asnoun()
                    + " or anyone else on "
                    + self.user_clientlist().pronoun_possessive("")
                    + "behalf."
                )
            )
            output.append(
                str(
                    "Documents relating instructions, advise, recommendations, or directions regarding the inspection, maintenance, servicing, or repair of "
                    + self.hazloc.name.text
                    + "."
                )
            )
            output.append(
                str(
                    "Documents indicating whether and to what extent "
                    + possessify(self.recipients().asnoun(), "")
                    + "processes for inspecting, maintaining, or repairing The Property were different at any of its facilities as compared to its "
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " location at issue."
                )
            )
            output.append(
                str(
                    "Documents identifying any safety committee members as of the time of the "
                    + self.incident
                    + ". This request concerns "
                    + possessify(self.recipients().asnoun(), "")
                    + ""
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " facility only."
                )
            )
            output.append(
                str(
                    "Documents indicating the dates of any safety committee meetings, the subjects discussed, and the persons in attendance, from "
                    + format_date(self.doi - date_interval(years=10))
                    + " to the present. This request concerns "
                    + possessify(self.recipients().asnoun(), "")
                    + ""
                    + self.premises.address.city
                    + ", "
                    + state_name(self.premises.address.state)
                    + " facility only."
                )
            )
        if "Breach of PIP Contract" in self.claims:
            output.append(
                str(
                    "Documents relating to any investigation conducted by "
                    + self.recipients().asnoun()
                    + " in response to, or otherwise relating to, "
                    + possessify(self.user_clientlist().asnoun(), "")
                    + "claim for PIP benefits."
                )
            )
            output.append(
                str(
                    "Documents showing any actions taken pursuant to or in response to any investigation conducted by "
                    + self.recipients().asnoun()
                    + " in response to, or otherwise relating to, "
                    + possessify(self.user_clientlist().asnoun(), "")
                    + "claim for PIP benefits, and the dates said actions were taken."
                )
            )
            output.append(
                str(
                    "Documents relating to or otherwise reflecting the information upon which "
                    + self.recipients().asnoun()
                    + " has based any investigation conducted by it in response to, or otherwise relating to, "
                    + possessify(self.user_clientlist().asnoun(), "")
                    + "claim for PIP benefits, including, but not limited to, notes of conversations, notes of meetings, investigations, interviews, medical reviews, wage loss reviews, notes to the claim file, correspondence, e-mails, instant messages, claim analysis, opinions as to causation of claimed injuries or damages."
                )
            )
            output.append(
                str(
                    "To the extent not otherwise requested above, a full and complete copy of the claim file at issue in this matter, including any analysis, recommendation or instruction generated by computer software."
                )
            )
            output.append(
                str(
                    "Documents identifying, to the extent possible, the name, position and role of each person who has had involvement in the handling of the claim file that is the subject of this litigation, whether or not they were an employee of "
                    + self.recipients().asnoun()
                    + " at the time, including the dates they were so involved."
                )
            )
        for party in self.recipients():
            if party.party.is_therapeutic_boarding_school:
                output.append(
                    str(
                        "Proof of any license of issued by the State of Oregon to "
                        + self.recipients().asnoun()
                        + " that was active as of "
                        + self.doi
                        + ", including, but not limited to, any license required for "
                        + self.recipients().asnoun()
                        + " to meet the requirements of OAR 413-215-0011."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any denial for a license to "
                        + self.recipients().asnoun()
                        + " from the State of Oregon at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to, any such denial under OAR 413-215-0121."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any suspension of "
                        + possessify(self.recipients().asnoun(), "")
                        + "license from the State of Oregon to operate as a therapeutic boarding school at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to, any such suspension under OAR 413-215-0121."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any revocation of "
                        + possessify(self.recipients().asnoun(), "")
                        + "license from the State of Oregon to operate as a therapeutic boarding school at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to, any such revocation under OAR 413-215-0121."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any conditions placed on "
                        + possessify(self.recipients().asnoun(), "")
                        + "license from the State of Oregon to operate as a therapeutic boarding school at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to, any such conditions placed under OAR 413-215-0121."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any inspections of "
                        + self.recipients().asnoun()
                        + " by the Oregon Department of Human Services or any employee, agent, or agency thereof, at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to any inspections conducted pursuant to OAR 413-215-0101."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any investigations of "
                        + self.recipients().asnoun()
                        + " by the Oregon Department of Human Services or any employee, agent, or agency thereof, at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to any investigations conducted pursuant to OAR 413-215-0106 and any documents reflecting the substance of any underlying complaints that prompted the investigation."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any corrective actions taken by the Oregon Department of Human Services any or employee, agent, or agency thereof, against "
                        + self.recipients().asnoun()
                        + " at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to any corrective actions taken pursuant to OAR 413-215-0111."
                    )
                )
                output.append(
                    str(
                        "Any and all civil penalties assessed against "
                        + self.recipients().asnoun()
                        + " by the State of Oregon or any employee, agent, or agency thereof, at any time from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present, including, but not limited to any civil penalties assessed under OAR 413-215-0116."
                    )
                )
                output.append(
                    str(
                        "A copy of any and all correspondence, notes of correspondence, call logs, audio recordings, video recordings, or any other record documenting any communication between "
                        + self.user_clientlist().asnoun()
                        + " or "
                        + capitalize(possessify(self.user_clientlist().asnoun(), ""))
                        + "family and "
                        + self.recipients().asnoun()
                        + " at any time."
                    )
                )
                output.append(
                    str(
                        "Documents relating to the training of any employee or agent of "
                        + self.recipients().asnoun()
                        + " on duty at the time of the self.incident regarding "
                        + possessify(self.recipients().asnoun(), "")
                        + "discipline and behavior management protocols including de-escalation skills training, crisis prevention skills, positive behavior management, and disciplinary techniques, including, but not limited to, documents reflecting the substance of any such training as well as dated proofs of completion for any employee or agent of "
                        + self.recipients().asnoun()
                        + " on duty at the time of the self.incident."
                    )
                )
                output.append(
                    str(
                        "A copy of any document reflecting the number of employees or agents of "
                        + self.recipients().asnoun()
                        + " on duty at the time of the self.incident."
                    )
                )
                output.append(
                    str(
                        "A copy of any document reflecting the number of students in "
                        + possessify(self.recipients().asnoun(), "")
                        + "care at the time of the self.incident."
                    )
                )
                output.append(
                    str(
                        "Proof of any annual child abuse reporting training, staff orientation, or background and reference checks for any for any employee or agent of "
                        + self.recipients().asnoun()
                        + " on duty at the time of the self.incident that were completed and current as of "
                        + self.doi
                        + ", including, but not limited to, any such tasks required to have been completed by "
                        + self.recipients().asnoun()
                        + " per OAR 413-215-0061."
                    )
                )
                output.append(
                    str(
                        "Proof of any annual child abuse reporting training, staff orientation, and background and reference check completion for any for any employee or agent of "
                        + self.recipients().asnoun()
                        + " involved in the management, reporting, review, or discipline relating to the self.incident were completed and current as of "
                        + self.doi
                        + ", including, but not limited to, any such tasks required to have been completed by "
                        + self.recipients().asnoun()
                        + " per OAR 413-215-0061."
                    )
                )
                output.append(
                    str(
                        "The complete personnel file, job title, written job description, staff development plan, and written annual evaluations for any employee or agent of "
                        + self.recipients().asnoun()
                        + " on duty at the time of the self.incident, including, but not limited to, any documents required by OAR 413-215-0061."
                    )
                )
                output.append(
                    str(
                        "The complete personnel file, job title, written job description, staff development plan, and written annual evaluations for any employee or agent of "
                        + self.recipients().asnoun()
                        + " involved in the management, reporting, review, or discipline relating to the self.incident, including, but not limited to, any documents required by OAR 413-215-0061."
                    )
                )
                output.append(
                    str(
                        "A copy of documents indicating the identity of any employee or agent of "
                        + self.recipients().asnoun()
                        + " who handled any disputes or altercations, whether physical or verbal, between "
                        + self.user_clientlist().asnoun()
                        + " and STUDENT X, including the self.incident."
                    )
                )
                output.append(
                    str(
                        "A copy of any reports relating to the self.incident, including, but not limited to, any such reports made to the Oregon Department of Human Services."
                    )
                )
                output.append(
                    str(
                        "A copy of any written policies or procedures or procedures that were in effect on "
                        + self.doi
                        + ", that "
                        + self.recipients().asnoun()
                        + " was required to have and adhere to under OAR 413-215-0056."
                    )
                )
                output.append(
                    str(
                        "Any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", that address the use or nonuse of physical restraints or seclusion as a means of discipline, behavior management, or suicide, including, but not limited to, any such policies adopted pursuant to OAR 413-215-0076."
                    )
                )
                output.append(
                    str(
                        "Any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", that address the process by which students are referred to the boarding school, including, but not limited to, policies reflecting (1) from whom referrals are accepted; (2) on what basis new students are accepted by "
                        + self.recipients().asnoun()
                        + "; and (3) how information necessary to provide for the safety and care of the students in "
                        + possessify(self.recipients().asnoun(), "")
                        + "care was provided to the appropriate employees or agents of "
                        + self.recipients().asnoun()
                        + " as of "
                        + self.doi
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all documents reflecting the initial evaluation of STUDENT X, including, but not limited to, any and all documents relating to the initial evaluation of STUDENT X required by OAR 413-215-0671."
                    )
                )
                output.append(
                    str(
                        "Any and all documents reflecting the initial evaluation of "
                        + self.user_clientlist().asnoun()
                        + ", including, but not limited to, any and all documents relating to the initial evaluation of "
                        + self.user_clientlist().asnoun()
                        + " required by OAR 413-215-0671."
                    )
                )
                output.append(
                    str(
                        "A copy of any consents, disclosures, or authorizations obtained by "
                        + self.recipients().asnoun()
                        + " relating to "
                        + self.user_clientlist().asnoun()
                        + "’s enrollment, including, but not limited to, any such consents, disclosures, or authorizations required by OAR-215-0676."
                    )
                )
                output.append(
                    str(
                        "A copy of any consents, disclosures, or authorizations obtained by "
                        + self.recipients().asnoun()
                        + " relating to STUDENT X’s enrollment, including, but not limited to, any such consents, disclosures, or authorizations required by OAR-215-0676."
                    )
                )
                output.append(
                    str(
                        "The case file for STUDENT X that "
                        + self.recipients().asnoun()
                        + " was required to maintain under OAR 413-215-0681."
                    )
                )
                output.append(
                    str(
                        "The case file for "
                        + self.user_clientlist().asnoun()
                        + " that "
                        + self.recipients().asnoun()
                        + " was required to maintain under OAR 413-215-0681."
                    )
                )
                output.append(
                    str(
                        "Any and all documents reflecting the dates that STUDENT X was under the custody, care, and control of "
                        + self.recipients().asnoun()
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any property damage alleged to have occurred during the period of time when STUDENT X was under the custody, care, and control of "
                        + self.recipients().asnoun()
                        + ", including, but not limited to invoices, estimates for the repair of any furniture, walls, or doors in the facility."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any thefts alleged to have occurred during the period of time when STUDENT X was under the custody, care, and control of "
                        + self.recipients().asnoun()
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to reports or complaints made by students or staff indicating that they were scared of STUDENT X or otherwise feared for their safety because of STUDENT X."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any threats of bodily harm alleged to have been made by STUDENT X while she was under the custody, care, and control of "
                        + self.recipients().asnoun()
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any bodily harm alleged to have been caused by STUDENT X while she was under the custody, care, and control of "
                        + self.recipients().asnoun()
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to STUDENT X’s enrollment, including, but not limited to, any applications, orders, disciplinary records, counseling or therapy records, medical records, or administrative records."
                    )
                )
                output.append(
                    str(
                        "Any and all documents relating to any request by "
                        + self.user_clientlist().asnoun()
                        + " or "
                        + self.user_clientlist().asnoun()
                        + "’s family for physical health care services from "
                        + self.doi
                        + ", through June 7, 2019, including but not limited to documents reflecting the date upon any such request was made by "
                        + self.user_clientlist().asnoun()
                        + " and the date upon which any such health care services were rendered or otherwise made accessible to "
                        + self.user_clientlist().asnoun()
                        + "."
                    )
                )
                output.append(
                    str(
                        "Any and all policies or procedures relating to discipline, behavior management, or suicide prevention that were in effect on "
                        + self.doi
                        + ", including, but not limited to, any such policy required by OAR 413-215-0076."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", regarding violence, including, but not limited to, policies relating to whether and how staff are supposed to handle violence among students, including, but not limited to, any policy that did not allow or would have otherwise purported to restrict staff’s ability to intervene if a physical or verbal altercation were to occur between students."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", relating to self-defense or defense of others."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", relating to intervention in disputes or altercations, whether physical or verbal, between students by any agent or employee of "
                        + possessify(self.recipients().asnoun(), "")
                        + "who was on duty at the time of the "
                        + self.incident
                        + "."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all documents relating to any training or instruction that any agent or employee of "
                        + possessify(self.recipients().asnoun(), "")
                        + "who was on duty at the time of the "
                        + self.incident
                        + " received prior to "
                        + self.doi
                        + ", relating to staff intervention in disputes or altercations, whether physical or verbal, between students."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", relating to medical treatment for students, including, but not limited to, the manner in which student requests for medical attention are to be handled by staff."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all policies or procedures that were in effect on "
                        + self.doi
                        + ", relating to records of student misconduct, including, but not limited to the manner in which staff were to record student misconduct, the types of student misconduct that staff were to record, and the form and location in which any such records were to be maintained."
                    )
                )
                output.append(
                    str(
                        possessify(self.recipients().asnoun(), "")
                        + "complete budget information from "
                        + format_date(self.doi - date_interval(years=10))
                        + ", to the present."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any documents that we effective as of "
                        + self.doi
                        + ", that instructed or otherwise addressed whether or how students were to respond when they were being threatened, assaulted, or battered by another student."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, any and all accident and/or incident reports and investigations prepared by any person at any time concerning the "
                        + self.incident
                        + "."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, documents concerning or referencing complaints or concerns regarding any of "
                        + possessify(self.recipients().asnoun(), "")
                        + "students from "
                        + format_date(self.doi - date_interval(years=10))
                        + " to the present."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, documents or communication received from STUDENT X at any time."
                    )
                )
                output.append(
                    str(
                        "To the extent not otherwise requested above, documents relating to any remedial steps taken since the self.incident so as to prevent any similar "
                        + self.incident
                        + "s in the future."
                    )
                )
                output.append(
                    str(
                        "A complete copy of any correspondence, including email, between "
                        + self.recipients().asnoun()
                        + " and any law enforcement or emergency services agency regarding the "
                        + self.incident
                        + " or the claims made in the subject lawsuit."
                    )
                )
                output.append(
                    str(
                        "Copies of any and all documents relating to any public relations firm or personnel employed, hired or engaged by "
                        + self.recipients().asnoun()
                        + " concerning the "
                        + self.incident
                        + ", including suggestions made and actions taken by that firm, public relations person(s), or "
                        + self.recipients().asnoun()
                        + " as a result."
                    )
                )
        output.there_is_another = False
        output.gathered = True
        return output

    def jury_issues(self):
        output = [
            str(
                "the nature and extent of "
                + self.parties.the_ps().asnoun()
                + "injuries"
            ),
            str("liability"),
        ]
        return output

    # @property
    # def liminator(self):
    # return [str(possessify(self.recipients().asnoun(), '') + "financial condition or ability to pay a judgment."), str(possessify(self.recipients().asnoun(), '') + "remorse."), str("Requiring that expert witnesses bring their entire file to trial; disclose and allow review of the same prior to testifying."), str(possessify(self.recipients().asnoun(), '') + "expert medical opinions not included in its " + self.ime_rule + " report."), str("Comments concerning the truthfulness or credibility of other parties or witnesses."), str("Liability opinion of any witness."), str("Argument or evidence that " + self.user_clientlist().asnoun() + " had health insurance, or could have submitted " + self.user_clientlist().pronoun_possessive('') + "medical expenses to Medicare or any other health insurer."), str("Any reference to “insurance discounts” or words of similar impact in an attempt to argue that billed medical expenses are *per se* unreasonable in amount."), str("Failure to seek additional treatment—failure to mitigate."), str("The fact that a judgment would not be taxed."), str("Any suggestion that " + self.user_clientlist().asnoun() + " is seeking a windfall, rolling the dice, that this action is equivalent to a lottery, that there is no quality control over the filing of lawsuits, or other words of similar impact."), str("Argument that the jury should apportion liability between " + self.recipients().asnoun() + ", the driver of " + possessify(self.user_clientlist().asnoun(),'') + "vehicle and/or the driver of the phantom vehicle."), str("Argument or evidence that " + self.user_clientlist().asnoun() + " could have filed an uninsured motorist claim."), str("Any prior " + self.incident() + " involving " + self.user_clientlist().asnoun() + " where " + self.user_clientlist().pronoun_possessive('') + self.user_clientlist().does_verb('was') "not injured, or otherwise sustained no permanent injuries."), str("Fact that " + self.recipients().asnoun() + " was riding as a part of a church group") str(self.recipients().asnoun() + " should not be allowed to testify as a representative of the estate."), str(possessify(self.user_clientlist().asnoun(), '') + "state of intoxication at the time of the " + self.incident() + "."), str("Prior worker’s compensation claims."), str("Any prior motor vehicle collisions involving " + self.user_clientlist().asnoun() + " where he was not injured, or otherwise sustained no permanent injuries.")]

    def facite_motions_in_limine(self):
        output = [["", str(
            "To the extent that any of these motions are granted, "
            + self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("request")
            + " that the Court require opposing counsel to advise "
            + self.recipients().pronoun_possessive("")
            + self.recipients().as_noun("client")
            + " and witnesses, in advance and outside the presence of the jury, of any rulings that apply to their anticipated testimony. In support of "
            + self.user_clientlist().pronoun_possessive("")
            + "motions, "
            + self.user_clientlist().asnoun()
            + "  relies on OEC 103(3) and the following points and authorities."
        )]]

        if self.mil_exclude_remorse:
            head = str(
                "The Court should exclude all reference, mention, argument, or evidence of remorse."
            )
            output.append([str(head), str(self.exclude_remorse)])

        if self.mil_exclude_ability_to_pay:
            head = str(
                " "
                + " "
                + str(
                    possessify(self.recipients().asnoun(), "")
                    + "financial condition or ability to pay a judgment."
                )

            )
            output.append([str(head), self.exclude_ability_to_pay()])

        if self.mil_require_expert_homework:
            head = str(
                "The Court Should Require Expert Witnesses to Bring Their Entire File to Trial and to Disclose and Allow Review of the Same Prior to Testifying."
            )
            output.append([str(head), str(self.require_expert_homework)])

        if self.mil_witness_on_witness_violence:
            head = str(
                "The Court Should Prevent Parties and Witnesses from Commenting on the Truthfulness or Credibility of Other Parties or Witnesses."
            )
            output.append([str(head), str(self.witness_on_witness_violence)])

        if self.mil_exclude_health_insurance_counterfactual:
            head = str(
                "The Court Should Exclude Argument and Evidence that "
                + self.user_clientlist().asnoun()
                + " Had Health Insurance."
            )
            output.append([str(head), str(self.exclude_health_insurance_counterfactual)])

        if self.mil_exclude_insurance_discounts:
            head = str(
                "The Court Should Exclude Any reference to “insurance discounts” or words of similar impact in an attempt to argue that billed medical expenses are *per se* unreasonable in amount."
            )
            output.append([str(head), str(self.exclude_insurance_discounts)])

        if self.mil_exclude_mitigation:
            head = str(
                "The Court Should Exclude Any Evidence of Failure to Seek Additional Treatment—Failure to Mitigate."
            )
            output.append([str(head), str(self.exclude_mitigation)])

        if self.mil_exclude_tax_free_jmnts:
            head = str(
                "The Court Should Exclude Any Evidence That a Judgment Would Not Be Taxed."
            )
            output.append([str(head), str(self.exclude_tax_free_jmnts)])

        if self.mil_exclude_tort_reform:
            head = str(
                "The Court Should Exclude Any suggestion that "
                + self.user_clientlist().asnoun()
                + " is seeking a windfall, rolling the dice, that this action is equivalent to a lottery, that there is no quality control over the filing of lawsuits, or other words of similar impact.")
            output.append([str(head), str(self.exclude_tort_reform)])

        if self.mil_exclude_phantom_vehicle:
            head = str(
                self.recipients().asnoun()
                + " Should Not Be Allowed to Argue That The Phantom Vehicle "
                + str(
                    str(
                        "or the Driver of "
                        + possessify(self.user_clientlist().asnoun(), "")
                        + "Vehicle"
                    )
                    if self.user_clientlist().mva_status == "Pedestrian"
                    else ""
                )
                + " Share Liability for "
                + possessify(self.user_clientlist().asnoun(), "Injuries.")
            )
            output.append([str(head), str(self.exclude_phantom_vehicle)])

        if self.mil_exclude_prior_workers_comp:
            head = str(
                " "
                + "Prior worker’s compensation "
                # + self.workers_comp.claims().as_noun("claim")
                + "."

            )
            output.append([str(head), str(self.exclude_prior_workers_comp)])

        if self.mil_exclude_prior_mvas:
            head = str(
                " "
                + "Any prior motor vehicle collisions involving "
                + self.user_clientlist().asnoun()
                + " where he was not injured, or otherwise sustained no permanent injuries."

            )
            output.append([str(head), str(self.exclude_prior_mvas)])

        if self.mil_exclude_clients_intoxication:
            head = str(
                " "
                + possessify(self.user_clientlist().asnoun(), "")
                + "state of intoxication at the time of the "
                + self.incident()
                + "."
            )
            output.append([str(head), str(self.exclude_clients_intoxication)])

        if self.mil_exclude_religious_group:
            head = str(
                " "
                + possessify(self.recipients().asnoun(), "")
                + "Religion and Facts Related to It Should Be Excluded as Irrelevant."
            )
            output.append([str(head), str(self.exclude_religious_group)])

        if self.mil_exclude_estate:
            head = str(
                " "
                + self.recipients().trustees()
                + " Should Not Be Allowed to Testify As a Representative of the "
                + self.recipients().asnoun()
                + " estate."
            )
            output.append([str(head), str(self.exclude_estate)])

        return output

    @property
    def exclude_phantom_vehicle(self):
        output = str(
            "Argument that the jury should apportion liability between "
            + self.recipients().asnoun()
            + ", the driver of "
            + possessify(self.user_clientlist().asnoun(), "")
            + "vehicle and/or the driver of the phantom vehicle."
            + self.fx.NLINE
        )
        levl = 0
        # for item in self.workers_comp_claims:
        #    output += str(
        #        self.fx.NLINE
        #        + self.fx.TAB
        #        + item_label(levl, 4)
        #        + " "
        #        + item.year
        #        + " – "
        #        + fix_punctuation(item.name.text)
        #        + self.fx.NLINE
        #    )
        #    levl += 1
        output += str(
            "While "
            + self.recipients().asnoun()
            + " "
            + self.recipients().does_verb("is")
            + " free to deny that "
            + self.recipients().pronoun_subjective()
            + " "
            + self.recipients().did_verb("was")
            + " at fault, "
            + self.recipients().pronoun_subjective()
            + " and "
            + self.recipients().pronoun_possessive("")
            + "counsel should not be allowed to argue or suggest to the jury that it should apportion liability among "
            + comma_and_list(
                self.recipients().asnoun(),
                str(
                    str(
                        "the driver of "
                        + possessify(self.user_clientlist().asnoun(), "")
                        + "vehicle"
                    )
                    if any(
                        party.party.mva_status == "Pedestrian"
                        for party in self.user_clientlist()
                    )
                    else self.user_clientlist().asnoun()
                ),
                "the phantom vehicle",
            )
            + ", or that any such apportionment should reduce "
            + possessify(self.user_clientlist().asnoun(), "")
            + "recovery. Ultimately the jury will be asked to determine whether or not "
            + self.recipients().pronoun_subjective()
            + " "
            + self.recipients().did_verb("was")
            + " negligent. The jury cannot compare "
            + self.recipients().pronoun_possessive("")
            + "negligence with the phantom driver"
            + str(
                str(
                    ", neither with the driver of "
                    + possessify(self.user_clientlist().asnoun(), "")
                    + "vehicle, as neither is"
                )
                if any(
                    party.party.mva_status == "Pedestrian" for party in self.user_clientlist()
                )
                else ", as he is not"
            )
            + " a party to the case nor a settled party.party. ORS 31.600(2). Moreover, comparative negligence has not been pled as an affirmative defense, nor is it an issue for the jury on the verdict form."
        )
        return output

    @property
    def exclude_prior_mvas(self):
        output = str(
            self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("has")
            + " been involved in "
            + self.prior_mvas.as_noun("collision")
            + " prior to the injury at issue, but "
            + self.user_clientlist().pronoun_possessive("")
            + " was never permanently injured prior to this "
            + self.incident
            + ". The prior "
            + self.prior_mvas.as_noun("collision")
            + " include:"
        )
        levl = 0
        for item in self.prior_mvas:
            output += str(
                self.fx.NLINE
                + RichText.add(str(self.fx.TAB))
                + item_label(levl, 4)
                + item.year
                + " – "
                + fix_punctuation(item.name.text)
                + self.fx.NLINE
            )
            levl += 1
        output += str(
            self.recipients().asnoun()
            + " should be precluded from introducing "
            + self.prior_mvas.pronoun_possessive("")
            + self.prior_mvas.as_noun("collision")
            + " into evidence, as "
            + self.prior_mvas.pronoun_possessive("")
            + self.prior_mvas.does_verb("is")
            + " not relevant to the issues involved in the case, invite juror speculation, and are unduly prejudicial to "
            + self.user_clientlist().asnoun()
            + ". "
            + self.relevancy_rules
            + "; "
            + self.precautionary_jury_instructions
            + ". "
            + "This would also necessarily include evidence of any insurance claims arising out of the prior "
            + self.prior_mvas.as_noun("collision")
            + ", to the extend there were any."
        )
        return output

    @property
    def exclude_prior_workers_comp(self):
        output = str(
            self.user_clientlist().asnoun()
            + " has had prior worker’s compensation "
            # + self.workers_comp_claims.as_noun("claim")
            + ", but never suffered permanent disability. These include:"
            + self.fx.NLINE
        )
        levl = 0
        # for item in self.workers_comp_claims:
        #    output += str(
        #        self.fx.NLINE
        #        + self.fx.TAB
        #        + item_label(levl, 4)
        #        + " "
        #        + item.year
        #        + " – "
        #        + fix_punctuation(item.name.text)
        #        + self.fx.NLINE
        #    )
        #    levl += 1
        output += str(
            "Defendant should be precluded from introducing the  "
            # + self.workers_comp_claims.as_noun("claim")
            + " into evidence, as "
            # + self.workers_comp_claims.pronoun_possessive("")
            + "are not relevant to the issues involved in the case, invite juror speculation, and are unduly prejudicial to "
            + self.user_clientlist().asnoun()
            + ". "
            + self.relevancy_rules
            + "; "
            + self.precautionary_jury_instructions
            + ". This would also necessarily include the underlying incidents that brought about the claims themselves. "
        )
        return output

    @property
    def exclude_clients_intoxication(self):
        output = str(
            self.user_clientlist().asnoun()
            + " was intoxicated at the time of the "
            + self.incident
            + ". In any case, the only issues for the jury are "
            + comma_and_list(self.jury_issues)
            + ". The fact that "
            + self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("was")
            + " intoxicated at the time of the "
            + self.incident
            + " is meaningless to these issues, making any such testimony irrelevant under "
            + self.balance_evidence_rule
            + ", as it would not tend to make any allegation more or less likely to be true. Moreover, evidence of "
            + possessify(self.user_clientlist().asnoun(), "")
            + "intoxication should be excluded as unduly prejudicial and misleading, as it would tent to invoke sympathy for "
            + self.recipients().asnoun()
            + ". "
            + self.relevancy_rules
            + "; "
            + self.precautionary_jury_instructions
            + ". "
        )
        return output

    @property
    def exclude_estate(self):
        output = str(
            "The only issues before the jury in this case are "
            + comma_and_list(self.jury_issues)
            + ". Any representative of the defendant Estate has nothing meaningful to add on those issues, making the testimony irrelevant, as it would not tend to make said allegations more or less likely to be true. "
            + self.balance_evidence_rule
            + ". Moreover, such evidence should be excluded as unduly prejudicial and misleading, as it would tent to invoke sympathy on the part of the defendant Estate. "
            + self.relevancy_rules
            + "; "
            + self.precautionary_jury_instructions
            + ". "
        )
        return output

    @property
    def exclude_religious_group(self):
        output = str(
            "The police report indicates that "
            + self.recipients().asnoun()
            + " of the "
            + self.incident
            + " was a part of a church group. Again, the only issues before the jury in this matter "
            + self.jury_issues
            + ". Thus, the fact that the decedent was reportedly part of a church group adds nothing meaningful on those issues, making any such testimony irrelevant, as it would not tend to make said allegations more or less likely to be true. "
            + self.balance_evidence_rule
            + ". Moreover, such evidence should be excluded as unduly prejudicial and misleading, as it would tent to invoke sympathy on the part of "
            + self.recipients().asnoun()
            + ". "
            + self.relevancy_rules
            + "; "
            + self.precautionary_jury_instructions
            + ". "
            + self.fx.ITALIC
            + "See, also"
            + self.fx.ITALIC
            + ", "
            + self.no_evidence_of_faith
            + " (“Evidence of beliefs or opinions of a witness on matters of religion not admissible for purpose of showing that by reason of their nature the credibility of the witness is impaired or enhanced.”)."
        )
        return output

    # @property
    def exclude_ability_to_pay(self):
        output = ""
        output += self.recipients().asnoun()
        output += ", "
        output += possessify(self.recipients().asnoun(), "")
        output += "attorney, and "
        output += possessify(self.recipients().asnoun(), "")
        output += "witnesses should not be allowed to offer any evidence or argument regarding "
        output += possessify(self.recipients().asnoun(), "")
        output += "financial condition or ability to pay for any liability, loss, damage, or injury. Whether a defendant "
        output += "has the ability to pay has no bearing on the issues of the case, and this evidence should be excluded from trial. "
        output += self.fx.ITALIC
        output += "See"
        output += self.fx.ITALIC
        output += ", "
        output += self.ability_to_pay_instruction
        output += "; "
        output += self.fx.ITALIC
        output += "Brooks v. Bergholm"
        output += self.fx.ITALIC
        output += ", 256 Or 1, 4-6 (1970); "
        output += self.fx.ITALIC
        output += "Benton v. Johnson"
        output += self.fx.ITALIC
        output += ", 45 Or App 959, 963 (1980). Moreover, these topics tend to invoke sympathy or bias on the part of the "
        output += self.recipients().asnoun()
        output += ", which the jury is forbidden from considering. "
        output += self.precautionary_jury_instructions
        output += ". "
        output += self.fx.NLINE
        output += "Such testimony is also irrelevant to the sole issues before this jury—that of "
        # output += self.jury_issues
        output += ". "
        output += self.relevancy_rules
        output += ". Moreover, it should be excluded because any probative value is substantially outweighed by the danger of unfair prejudice, since such testimony would tend to arouse the jury’s sympathy. "
        output += self.relevancy_rules
        output += "; "
        output += self.precautionary_jury_instructions
        output += ". "
        output += self.fx.NLINE
        output += "If "
        output += self.recipients().asnoun()
        output += " alludes to, or argues, this issue, "
        output += self.user_clientlist().asnoun()
        output += " should be entitled to introduce evidence that "
        output += self.recipients().asnoun()
        output += "  is insured, as well as the amount of insurance available. At that point, "
        output += possessify(self.recipients().asnoun(), "")
        output += "insured status would not go to the issue of whether "
        output += self.recipients().asnoun()
        output += " acted negligently—the only reason for which it is excluded under "
        output += self.insurance_evidence_rule
        output += "—but rather to whether "
        output += self.recipients().pronoun_subjective()
        output += " can pay damages."
        return output

    @property
    def exclude_remorse(self):
        output = str(
            "The Court should exclude all reference, mention, argument, or evidence of "
            + possessify(self.recipients().asnoun(), "")
            + "remorse, or that he, or any of "
            + self.user_clientlist().pronoun_possessive("")
            + "agents or representatives, felt or still feels sorry about "
            + possessify(self.user_clientlist().asnoun(), "")
            + "injuries. Oregon has consistently rejected the admission of a "
            + possessify(self.recipients().asnoun(), "")
            + "post-injury behavior or state of mind. In "
            + self.fx.ITALIC
            + "Byers v. Santiam Ford"
            + self.fx.ITALIC
            + ", Inc., 281 Or 411, 415 (1978), the court held that subsequent acts of contrition and a conciliatory attitude were not relevant and not admissible regarding punitive damages. In *Mason v. Householder*, 58 Or App 192 (1982), the court held that admission of testimony that "
            + self.recipients().asnoun()
            + " was “rehabilitated” after "
            + possessify(self.user_clientlist().asnoun(), "")
            + "injury was reversible error, and that matters occurring after the happening of the bad act, and unrelated to the "
            + possessify(self.recipients().asnoun(), "")
            + "state of mind at the time of the events, are inadmissible on the issue of punitive damages. "
            + self.fx.ITALIC
            + "Id."
            + self.fx.ITALIC
            + " at 195. "
            + self.fx.NLINE
            + "Thus, the court should exclude any such evidence on the ground that it is not relevant to any person’s state of mind at the time of the "
            + self.incident
            + ". "
            + self.balance_evidence_rule
            + ". If such evidence is not admissible as to punitive damages (where deterrence is a major component of the award), it is  not relevant to compensation for ordinary negligence. This type of evidence should be excluded as irrelevant, unduly prejudicial, misleading, and capable of confusing the issues before the jury. "
            + self.relevancy_rules
            + ". "
        )
        return output

    @property
    def require_expert_homework(self):
        output = str(
            "This case involves expert testimony. There is no expert discovery under the Oregon Rules of Civil Procedure and, therefore, the only opportunity to know the identity and confront the expert witnesses is when they appear in the courtroom to testify. Under these circumstances, there is no prior opportunity to learn the factual basis for the opinions that will be rendered, nor to see what the expert considered, or was able to consider, prior to rendering an "
            + self.ime_rule
            + " report or otherwise testifying at trial. "
            + self.fx.NLINE
            + "It is accepted that once an expert witness testifies, all privileges protecting communications and documents considered in developing expert opinions are waived. "
            + self.waiver_from_disclosure_rule
            + "; "
            + self.fx.ITALIC
            + "State ex rel Grimm v. Ashmanskas"
            + self.fx.ITALIC
            + ", 298 Or 206, 690 P2d 1063 (1984); "
            + self.fx.ITALIC
            + "Marriage of Boon"
            + self.fx.ITALIC
            + ", 100 Or App 354, 357-58, 786 P2d 215 (1990). In custom and practice, attorneys instruct their expert witnesses to bring their entire file on the case with them to allow opposing counsel the opportunity to review it. When trial courts have confronted situations where an expert witness has failed to bring the complete file to the trial, the practice has been enforced by requiring the witness to go get it before testimony is given or at least before cross examination begins. "
            + self.fx.NLINE
            + "Under the provisions of "
            + self.expert_discretion_rule
            + ", the court has the discretion to manage the presentation of expert opinions. This includes the manner and timing of disclosure of the underlying facts or evidence upon which an expert will rely in rendering opinions. Given this, "
            + self.user_clientlist().asnoun()
            + "  moves that "
            + possessify(self.recipients().asnoun(), "")
            + "expert witnesses be required to bring their entire file to trial. Further, "
            + self.user_clientlist().asnoun()
            + " moves that"
            + self.user_clientlist().pronoun_possessive("")
            + "counsel be allowed to review any testifying expert witness’ file prior to that expert giving testimony. Being allowed to review the expert’s complete file prior to testimony being given will insure an orderly presentation of expert testimony, will enable counsel to anticipate objectionable testimony, and will minimize the likelihood of a mistrial because of inadmissible evidence getting to the jury."
        )
        return output

    @property
    def witness_on_witness_violence(self):
        output = str(
            self.recipients().asnoun()
            + "  should be precluded from questioning "
            + possessify(self.user_clientlist().asnoun(), "")
            + "truthfulness or credibility through"
            + self.user_clientlist().pronoun_possessive("")
            + "witnesses, including"
            + self.user_clientlist().pronoun_possessive("")
            + "experts. "
            + self.fx.ITALIC
            + "See State v. Lupoli"
            + self.fx.ITALIC
            + ", 348 Or 346, 357 (2010). (“This court has long held that one witness may not give an opinion on whether he or she believes another witness is telling the truth.”); "
            + self.fx.ITALIC
            + "State v. Middleton"
            + self.fx.ITALIC
            + ", 294 Or 427, 438 (1983) (“We expressly hold that in Oregon a witness, expert or otherwise, may not give an opinion on whether [the witness] believes a[nother] witness is telling the truth.”). Testimony of one witness, expert or lay, concerning the truthfulness of another witness is wholly inadmissible and improper.")

        return output

    @property
    def exclude_health_insurance_counterfactual(self):
        output = str(
            self.user_clientlist().asnoun()
            + " did not submit "
            + self.user_clientlist().pronoun_possessive("")
            + "related medical expenses to Medicare or "
            + self.user_clientlist().pronoun_possessive("")
            + "supplemental health insurer. Such evidence should be excluded as irrelevant, unduly prejudicial, misleading, and capable of confusing the issues before the jury. "
            + self.relevancy_rules
            + ". Moreover, it relates to a collateral source and is inadmissible as such—particularly where it would naturally cause the jury to question whether or not any such benefits would have paid at a lower rate than that which was billed, or would otherwise be subject to reimbursement out of the proceeds of this case. Neither issue is admissible.  ORS 31.580; "
            + self.fx.ITALIC
            + "White v. Jubitz Corp."
            + self.fx.ITALIC
            + ", 219 Or App 62 (2008); "
            + self.fx.ITALIC
            + "Coshens v. McGee"
            + self.fx.ITALIC
            + ", 219 Or App 78 (2008);"
            + self.precautionary_jury_instructions
            + "."
        )
        return output

    @property
    def exclude_insurance_discounts(self):
        output = str(
            self.recipients().asnoun()
            + " , "
            + self.recipients().pronoun_possessive("")
            + " witnesses, and attorneys should not be permitted to introduce evidence of the fact that some people who obtain medical treatment are insured, and that some of those insurers pay the medical provider a contractual discounted amount of the bill to prove or argue that the amount billed for "
            + possessify(self.user_clientlist().asnoun(), "")
            + "treatment was unreasonable. "
            + self.fx.NLINE
            + "ORS 31.580 (Collateral Benefits) and 31.550, et seq. (Advanced Payments) specifically preclude the admission of collateral payments and specify that such payment or related write-offs cannot be entered into evidence at trial. Whatever form of compensation a plaintiff received from third parties is not relevant to any question facing the jury, including insurance, payment of medical or living expenses, the Oregon Health Plan, Medicaid, Medicare, educational benefits, etc. Indeed, the Oregon Supreme Court has held that the collateral source rule “permits a plaintiff to recover from a tortfeasor the reasonable value of the medical treatment that he or she receives whether plaintiff is liable to pay or pays the medical providers’ charges for that treatment, the providers waive those charges, or a third party pays or otherwise satisfies those charges.”"
            + self.fx.ITALIC
            + "White v. Jubitz Corp."
            + self.fx.ITALIC
            + ", 347 Or 212, 236 (2009).  Simply put, evidence that a plaintiff’s insurance company or that anyone else paid or wrote off medical charges “is not admissible at trial.” ORS 31.580(2); "
            + self.fx.ITALIC
            + "accord White"
            + self.fx.ITALIC
            + ", 347 Or at 243; "
            + self.fx.ITALIC
            + "see also Cohens v. McGee"
            + self.fx.ITALIC
            + ", 219 Or App 78, (2008). "
            + self.fx.NLINE
            + "Some defense attorneys have recently tried to make an end-run around the inadmissibility of such evidence by attempting to introduce evidence that medical charges, as a whole, are unreasonable simply because some medical providers sometimes accept some unspecified discounted payment secondary to insurance contracts. Any such evidence is inadmissible and improper for two primary reasons: "
            + self.fx.NLINE
            + "First, it is an attempt at an end run around the collateral source rule and "
            + self.fx.ITALIC
            + "White v. Jubitz Corp."
            + self.fx.ITALIC
            + " under the guise of challenging the reasonableness of medical bills in general. While a defendant has the right to challenge the amount charged for medical services as not reasonable and/or necessary, they don’t get to rely on the presence of collateral source benefits to do so. Any testimony to this effect necessarily relies on collateral source payments to determine the reasonableness of bills. This type of evidence is contrary to Oregon’s strict prohibition against it. "
            + self.fx.NLINE
            + "Second, any such argument is circular and nonsensical. If the theory is that the amount charged for a medical bill is unreasonable simply because the provider would accept a percentage of that amount as payment from an insurer, then any amount charged would be per se unreasonable as a matter of law. Insurers pay a percentage of any amount charged, pursuant to their contract with the provider. A percentage of any positive integer will always be the lesser. Thus, insurers always pay at a discount, even with the most reasonable of charges. "
            + self.fx.NLINE
            + "For example, if the theory is that a $100 medical bill is unreasonable because the provider accepts 60% of that charge for payment from a particular health insurer, then $60 would also be unreasonable in the event the provider reduced the amount charged to that figure, as the insurer would still only pay 60% of that bill. Additionally, different insurers pay at different rates, with Medicare often being the lowest. Who is to say at what percentage a bill becomes “reasonable?”  Indeed, how could they? It would be wholly provider dependent and would necessarily require each and every provider to come to trial to testify as to what their average insurance payments are, from whom, over what time, for what services, and so on. "
            + self.fx.NLINE
            + "Moreover, folks without health insurance are billed the very same rate, and are expected to pay the full amount charged, and get turned over to collections if they don’t do so. If the court is to look at “reasonableness” as to what each provider expects to get paid for their services, are uninsured folks excluded from that conversation? Surely not. Thus, under any such theory of “reasonableness,” the jury would necessarily need to know what insurer covers the plaintiff, how much they paid for each service, what their deductible was for each bill, what their co-pay was for each bill, and how much was written off.  In other words, it would require evidence that is specifically excluded under Oregon’s collateral source rule discussed above. "
            + self.fx.NLINE
            + "Finally, insurers can negotiate these discounted rates thanks to a few special characteristics that no individual would have. For example, insurers have unparalleled bargaining power in that their insureds will actively avoid out of network providers. Further, insurers employ expert negotiators to get these special rates. A plaintiff in a personal injury lawsuit can’t honestly be held to the standard of an expert negotiator. "
            + self.fx.NLINE
            + "This line of examination or evidence of the same is wholly speculative, meritless, and should be denied. "
            + self.user_clientlist().asnoun()
            + "is not aware of any appellate authority in Oregon addressing this issue (likely because this novel argument often advanced by the defense bar is such a new formulation), although appellate courts in other jurisdictions, as well as numerous Oregon trial courts, have rejected similar evidence. "
            + self.fx.NLINE
            + "Given the above, this Court should order that "
            + self.recipients().asnoun()
            + " be prohibited from introducing similar evidence, or from arguing along these lines during closing argument. ORS 31.580, ORS 31.550, "
            + self.fx.ITALIC
            + "et seq."
            + self.fx.ITALIC
            + ", "
            + self.fx.ITALIC
            + "White v. Jubitz Corp."
            + self.fx.ITALIC
            + ", 347 Or 212 (2009), "
            + self.fx.ITALIC
            + "Cohens v. McGee"
            + self.fx.ITALIC
            + ", 219 Or App 78 (2008). Additionally, it should be prohibited as irrelevant to the issues before the court, unduly prejudicial to the "
            + self.user_clientlist().asnoun()
            + " , capable of misleading the jury, collateral, and a waste of time. "
            + self.relevancy_rules
            + ". It also runs afoul of Oregon’s prohibition of offering evidence of insurance or the ability to pay damages. "
            + self.ability_to_pay_instruction
        )
        return output

    @property
    def exclude_mitigation(self):
        output = str(
            self.recipients().asnoun()
            + ","
            + self.recipients().pronoun_possessive("")
            + "witnesses, and attorneys should be prevented from opining or otherwise arguing that "
            + self.user_clientlist().asnoun()
            + " failed to mitigate "
            + self.user_clientlist().pronoun_possessive("")
            + "damages, or that "
            + self.user_clientlist().pronoun_subjective()
            + " should have otherwise sought additional medical treatment.  Failure to mitigate has not been raised as an affirmative defense, and is thus outside the scope of the pleadings. "
            + self.fx.NLINE
            + "Moreover, if such evidence does not come in during trial, "
            + self.recipients().asnoun()
            + " should be precluded from arguing or suggesting the same in closing argument. In the event the Court allows such evidence or argument, the jury should be instructed that "
            + self.recipients().asnoun()
            + " carries the burden of proof on that issue, as it is an affirmative defense."
        )
        return output

    @property
    def exclude_tax_free_jmnts(self):
        output = str(
            "Argument or evidence of the fact that personal injury judgments are not taxed is wholly irrelevant to the issues before the court, is unduly prejudicial to a plaintiff, and is capable of misleading the jury. "
            + self.relevancy_rules
            + "."
        )
        return output

    @property
    def exclude_tort_reform(self):
        output = str(
            "Defense counsel in cases of this type are sometimes tempted to attack plaintiffs using popular inflammatory political language for its demeaning impact on the plaintiffs, their counsel, and the process of resolving civil disputes by jury trial.  Examples include argument that “anybody can go to the courthouse and file a lawsuit - all it takes is a few hundred dollars and a piece of paper,” or “the courts have no quality control over lawsuits” or that civil lawsuits amount to nothing more than “litigation lotto.”  No such statements or innuendos should be permitted during trial of this case, as they are highly likely to persuade the jury to decide the case on passion and prejudice, contrary to "
            + self.precautionary_jury_instructions
            + ".  As stated by the South Dakota Supreme Court: “Defense counsel’s statement that plaintiff was trying to hit the lottery by"
            + self.user_clientlist().pronoun_possessive("")
            + "lawsuit demeaned not only the plaintiff but also the judicial system itself, and impugned the trial court’s judgment of allowing the punitive damage claim to proceed.  The comments denigrated the fairness, integrity and public perception of the judicial system. ”"
            + self.fx.NLINE
            + "Counsel’s reference to playing lotto or powerball, or rolling the dice, were only meant to inflame the jury, and were beyond the bounds of proper final argument . . . . Interposing such remarks . . . can only be meant to persuade the jury to decide the case based on passion and prejudice. "
            + self.fx.ITALIC
            + "Schoon v. Looby"
            + self.fx.ITALIC
            + ", 2003 SD 123 (2003). Such statements are irrelevant to the issues before the court, are unduly prejudicial to the plaintiff, and are capable of misleading the jury. "
            + self.relevancy_rules
            + ".  They are demeaning, amount to nothing more than a slap to the face to the civil justice system, and tend to mock our courts."
            + self.fx.NLINE
            + "In the event "
            + self.recipients().asnoun()
            + ", "
            + self.user_clientlist().asnoun()
            + " witnesses, or "
            + self.recipients().asnoun()
            + " attorney mentions or argue that there is no “quality control” over lawsuits, or words to that effect, "
            + self.user_clientlist().asnoun()
            + "should be allowed to point out the availability of motions to strike, motions to dismiss, motions to make more definite and certain, motions for summary judgment, and so-on, to show that there is, in fact, “quality control” within the civil justice system, or, at the very least, to have the court give such instruction to the jury.  Moreover, "
            + self.user_clientlist().asnoun()
            + " would request a correcting instruction be given, reminding the jury that some injuries has been admitted by the  "
            + possessify(self.recipients().asnoun(), "")
            + "medical expert in this case."
        )
        return output

    def facite_trial_memo(self):
        output = []
        if self.tm_entitled_to_meds_as_billed:
            head = str(
                comma_and_list(self.user_clientlist())
                + " "
                + self.user_clientlist().does_verb("Is")
                + " Entitled to Plead and Recover "
                + capitalize(self.user_clientlist().pronoun_possessive(""))
                + "Billed Medical Expenses."
            )

            output.append([str(head), self.entitled_to_meds_as_billed_tm])

        if self.tm_all_meds_attributable_to_injury:
            head = str(
                "All of "
                + possessify(self.user_clientlist(), "")
                + "Medical Treatment Is Attributable to the "
                + self.incident
                + "."
            )
            output.append([str(head), self.all_meds_attributable_to_injury_tm])

        if self.tm_ability_to_pay:
            head = str(
                self.ability_to_pay_instruction
                + " (Ability to Pay) Should be Given "
            )
            output.append([str(head), self.ability_to_pay_tm])

        if self.tm_permanent_injury_instruction:
            head = str(
                self.permanent_injury_instruction
                + " Should be Given Because "
                + self.user_clientlist().salute()
                + " Has a Permanent Injury."
            )

            output.append([str(head), self.permanent_injury_instruction_tm])

        if self.tm_future_complications_compensible:
            head = str("Future Possible Complications are Compressible.")
            output.append([str(head), self.future_complications_compensible_tm])

        if self.tm_lay_witnesses_competent_for_pain_and_suffering:
            head = str(
                self.user_clientlist().asnoun()
                + " and Lay Witnesses are Competent to Testify as to "
                + possessify(self.user_clientlist().asnoun(), "")
                + "“Pain and Suffering.”"
            )

            output.append([str(head), self.lay_witnesses_competent_for_pain_and_suffering_tm])

        if self.tm_must_produces_recorded_statements:
            head = str(
                "Recorded Statements of "
                + possessify(self.recipients().asnoun(), "")
                + "Witnesses Must be Produced."
            )
            output.append([str(head), self.must_produces_recorded_statements_tm])

        if self.tm_rule_of_completeness:
            head = str("“Rule of Completeness”")
            output.append([str(head), self.rule_of_completeness_tm])

        if self.tm_ability_to_pay:
            head = str(
                self.ability_to_pay_instruction
                + " (Ability to Pay) Should be Given."
            )
            output.append([str(head), self.ability_to_pay_tm])

        if self.tm_written_jury_instruction:
            head = str(
                self.user_clientlist().asnoun()
                + " "
                + self.user_clientlist().does_verb("Requests")
                + " that the Jury Receive a Written Copy of the Jury Instructions."
            )
            output.append([str(head), self.written_jury_instruction_tm])

        if self.tm_phantom_motorist:
            head = "The Jury Cannot Compare Negligence with the Phantom Motorist."
            output.append([str(head), self.phantom_motorist_tm])

        if self.tm_videoconference_testimony:
            head = "Some Witnesses Will be Testifying by Videoconference."
            output.append([str(head), self.videoconference_testimony_tm])
        if self.tm_thin_skull:
            head = "Previous Infirm Conditions and Exacerbations are Compensate."
            output.append([str(head), self.thin_skull_tm])

        if self.tm_leading_questions:
            head = str(
                "Counsel May Ask Leading Questions When Examining Adverse Witnesses."
            )
            output.append([str(head), self.leading_questions_tm])
        return output

    @property
    def phantom_motorist_tm(self):
        output = str(
            "Ultimately the jury will be asked to determine whether or not "
            + self.recipients().asnoun()
            + self.recipients().did_verb("was")
            + " negligent. They cannot compare "
            + self.recipients().pronoun_possessive("")
            + "negligence to the phantom driver, nor to the driver of "
            + self.user_clientlist().salute_possessive()
            + " SUV, as neither is a party to the case nor a settled party. ORS 31.600(2). Thus, if a jury concludes that "
            + self.recipients().asnoun()
            + self.recipients().did_verb("was")
            + " negligent, even if only slightly, "
            + self.recipients().pronoun_possessive("")
            + " will become liable for all of "
            + self.user_clientlist().salute_possessive()
            + " injuries and damages."
        )
        return output

    @property
    def videoconference_testimony_tm(self):
        output = str(
            self.user_clientlist().salute()
            + " "
            + self.user_clientlist().does_verb("live")
            + ", and "
            + self.user_clientlist().did_verb("was")
            + " primarily treated by medical professionals, in West Palm Beach, Florida. Two of his treating physicians are expected to testify in the matter. Given the geographic distance, and pursuant to ORS 45.400(2), on November 4, 2019, "
            + self.user_clientlist().salute_possessive()
            + " counsel gave written notice to "
            + possessify(self.recipients().asnoun(), "")
            + " counsel that he wished to have those two physicians testify at trial via videoconference. On November 6, 2019, "
            + possessify(self.recipients().asnoun(), "")
            + " counsel responded in writing that she did not oppose said experts testifying at trial by videoconference. On December 2, 2019, "
            + self.user_clientlist().salute_possessive()
            + " counsel called Presiding Judge Bushong’s judicial assistant to inquire as to whether, under the circumstances, a motion needed to be filed via OJD File and Serve, or otherwise presented at ex parte, concerning the videoconference testimony. A short while later, Judge Bushong’s judicial assistant called back and stated she had spoken with Judge Bushong, who asked "
            + self.recipients().pronoun_possessive("")
            + "to relay to "
            + self.user_clientlist().salute_possessive()
            + " counsel that no motion was needed, and that he simply needed to make the assigned trial judge aware of the videoconference testimony."
        )
        return output

    @property
    def thin_skull_tm(self):
        output = str(
            "In this case, there is medical evidence that "
            + self.user_clientlist().asnoun()
            + " had a prior history of neck degeneration, with occasional neck pain. However, he had been quite asymptomatic with only intermittent cervical symptoms in the three years prior to the subject collision. "
            + self.user_clientlist().salute()
            + " also experienced some level of intermittent low back pain prior to the collision. However, immediately previous to the collision at issue, "
            + self.user_clientlist().salute_possessive()
            + " only had occasional symptoms related to these conditions. Conversely, immediately following the collision at issue "
            + self.user_clientlist().salute_possessive()
            + " pain and limitations were severe and constant."
            + self.fx.NLINE
            + "If the jury finds that "
            + self.user_clientlist().salute()
            + " was predisposed to injury due to his age, body habitus, or pre-existing bodily conditions, he is nonetheless entitled to full compensation for his injuries. The Prior Infirm Condition jury instruction, UCJI 70.06, states:"
            + self.fx.NLINE
            + self.fx.BLOCK
            + " If you find that the "
            + self.user_clientlist().asnoun()
            + " had a bodily condition that predisposed him to be more subject to injury than a person in normal health, nevertheless the "
            + self.recipients().asnoun()
            + " would be liable for any and all injuries and damage that may have been suffered by the "
            + self.user_clientlist().asnoun()
            + " as the result of the "
            + possessify(self.recipients().asnoun(), "")
            + " negligence, even though those injuries, due to the prior condition, may have been greater than those that would have been suffered by another person under the same circumstances."
            + self.fx.NLINE
            + "Thus, if the jury concludes that some underlying condition was symptomatically activated as a result of this collision, "
            + self.user_clientlist().salute()
            + " is nonetheless entitled to "
            + self.user_clientlist().pronoun_possessive("")
            + "full measure of damages on account of his symptoms. "
            + self.fx.NLINE
            + "Likewise, if the jury finds that "
            + self.user_clientlist().salute()
            + " was, in some way, symptomatic prior to this collision, he is nonetheless entitled to compensation for any worsening of that condition. The Aggravation of Preexisting Injury or Disability jury instruction, UCJI 70.07, states "
            + self.fx.NLINE
            + " > In the present case "
            + self.user_clientlist().asnoun()
            + " has alleged that the injury which he sustained as the result of the negligence of the "
            + self.recipients().asnoun()
            + " aggravated a preexisting injury or disability of his."
            + self.fx.NLINE
            + "In determining the amount of damages, if any, to be awarded to the "
            + self.user_clientlist().asnoun()
            + " in this case, you will allow him reasonable compensation for the consequences of any such aggravation that you find to have taken place as the result the "
            + possessify(self.recipients().asnoun(), "")
            + " negligence."
            + self.fx.NLINE
            + "The recovery should not include damages for the earlier injury or disability but only those that are due to its enhancement or aggravation."
            + self.fx.NLINE
            + self.user_clientlist().salute_possessive()
            + " experts are expected to testify that "
            + self.user_clientlist().salute_possessive()
            + " neck pain, back pain, and radicular symptoms were either new or otherwise greatly exacerbated by the collision. Ultimately, "
            + self.user_clientlist().salute()
            + " received several rounds of injections in his neck and back, and has been recommended cervical surgery. "
            + self.user_clientlist().salute_possessive()
            + " expert neurosurgeon surgeon, Dr. Aaron Smith, is expected to opine that the treatment and need for surgery was made necessary by the collision at issue."
            + self.fx.NLINE
            + "Given "
            + self.user_clientlist().salute_possessive()
            + " injuries, the impact of the collision, and the force applied to the occupants of the vehicle, it is clear that "
            + self.user_clientlist().salute_possessive()
            + " subsequent neck, shoulder, and back pain are a result of the collision at issue. Nonetheless, should the jury find that "
            + self.user_clientlist().salute_possessive()
            + " injuries are an aggravation of a preexisting injury or disability of his; Oregon law clearly dictates that "
            + self.user_clientlist().salute()
            + " is still entitled to compensation for any worsening of those conditions, including treatment required to address the same. "
        )
        return output

    @property
    def issues_bf_jury_tm(self):
        output = str(
            self.recipients().asnoun()
            + " admits liability but disputes the seriousness of "
            + possessify(self.user_clientlist().asnoun(), "")
            + " injuries. Therefore, the issues to be decided by the jury are those of causation and damages, including medical expenses (past and future) and noneconomic damages. "
            + self.fx.NLINE
        )
        return output

    @property
    def evidentiary_issues_tm(self):
        output = str(
            self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("does")
            + " not expect any significant evidentiary issues at trial. In addition to the above, "
            + self.user_clientlist().asnoun()
            + " also "
            + self.user_clientlist().does_verb("offer")
            + " the following points and authorities for the Court’s convenience: "
            + self.fx.NLINE
        )
        return output

    @property
    def entitled_to_meds_as_billed_tm(self):
        output = RichText(
            str(
                "The related medical bills "
                + self.user_clientlist().asnoun()
                + " has incurred to date total "
                + "$12,307.39. "
                + "A portion of these bills were written off, secondary to "
                + self.user_clientlist().pronoun_possessive("")
                + "personal injury protection insurance benefits or health insurance benefits. However, any write-offs on account of insurance discounts are a collateral source benefit, as "
                + self.user_clientlist().asnoun()
                + " paid the premiums for said insurance coverage. "
                + self.user_clientlist().asnoun()
                + " is therefore entitled to also recover those write-offs, and is thus entitled to an award for the gross amount of "
                + self.user_clientlist().pronoun_possessive("")
                + "medical bills. ORS 31.580; "
            )
        )
        stri(output, "White v. Jubitz Corp.")
        strnull(output, ", 219 Or App 62 (2008); ")
        stri(output, "Cohens v. McGee")
        strnull(output, str(", 219 Or App 78 (2008). " + self.fx.NLINE))
        return output

    @property
    def all_meds_attributable_to_injury_tm(self):
        output = RichText(
            str(
                self.recipients().asnoun()
                + " "
                + self.recipients().does_verb("is")
                + " not entitled to an offset for any visits that "
                + self.user_clientlist().salute()
                + " might have otherwise participated in had "
                + self.user_clientlist().pronoun_subjective()
                + " never been injured. "
                + self.fx.NLINE
                + " First, the jury is not allowed to speculate. UCJI 5.01. "
                # + self.precautionary_jury_instructions
                + "The question of how much medical treatment "
                + self.user_clientlist().asnoun()
                + " would have had, if any at all, if he were never injured invites the jury to speculate. Accordingly, it is not a proper question for the jury to consider, and the Court should exclude any such evidence or conjecture. "
                + self.fx.NLINE
                + "Second, the incident"

                # + self.incident
                + " is the legal cause of "
                + self.user_clientlist().salute_possessive()
                + " subsequent treatment. UCJI 23.01. "
                # + self.but_for_instruction
                + "(Causation – “But For”) explains: "
                + self.fx.NLINE
                + " > The defendant's conduct is a cause of the plaintiff's [harm / injury] if the [harm / injury] would not have occurred but for that conduct; conversely, the defendant's conduct is not a cause of the plaintiff's [harm / injury] if that [harm / injury] would have occurred without that conduct. UCJI 23.02. "
                + self.fx.NLINE
                + self.fx.NLINE
                # + self.substantial_factor_instruction
                + "(Causation – “Substantial Factor”) goes on to state: "
                + self.fx.NLINE
                + self.fx.NLINE
                + " > Many factors [or things] may operate either independently or together to cause [harm / injury]. In such a case, each may be a cause of the [harm / injury] even though the others by themselves would have been sufficient to cause the same [harm / injury]. "
                + self.fx.NLINE
                + "If you find that the defendant's act or omission was a substantial factor in causing the [harm / injury] to the plaintiff, you may find that the defendant's conduct caused the [harm / injury] even though it was not the only cause. [A substantial factor is an important factor and not one that is insignificant.] "
                + self.fx.NLINE
                + "Because the incident clearly caused injuries to "
                + self.user_clientlist().salute()
                + " (and resultant pain), and because those injuries necessitated the subsequent treatment, those injuries are a substantial factor in bringing about the medical expenses (economic damages) "
                + self.user_clientlist().salute()
                + " incurred. Thus, "
                + self.recipients().asnoun()
                + " is liable for those medical expenses, regardless as to whether "
                + self.user_clientlist().salute()
                + " might have incurred some of them had "
                + self.user_clientlist().pronoun_subjective()
                + " never been shot. Thus, "
                + self.recipients().asnoun()
                + " is not entitled to an offset for any medical visits that "
                + self.user_clientlist().salute()
                + " would have otherwise made had "
                + self.user_clientlist().pronoun_subjective()
                + " never been injured. "
                + self.fx.NLINE
                + " Moreover, it is anticipated that "
                + possessify(self.recipients().asnoun(), "")
                + " medical expert also will appear, and admit that "
                + self.user_clientlist().asnoun()
                + " suffered some injury as a result of being injured and, presumably, that some of "
                + self.user_clientlist().pronoun_possessive("")
                + "treatment was also reasonable and necessitated by being injured. "
                + self.fx.NLINE
            )
        )
        return output

    @property
    def permanent_injury_instruction_tm(self):
        output = RichText(str(
            "Mortality tables are admissible into evidence once evidence of permanent injury has been shown. "
            + self.fx.ITALIC
            + "Frangos v. Edmunds"
            + self.fx.ITALIC
            + ", 179 Or 577, 604 (1946). "
            + str(
                str(
                    "Here, "
                    + possessify(self.user_clientlist().asnoun(), "")
                    + str(
                        str(
                            "injuries caused "
                            + self.user_clientlist().pronoun_objective()
                            + " to require surgery to be made whole. Those "
                            + str(
                                "future "
                                if self.user_clientlist()[0].future_surgery
                                else ""
                            )
                            + "surgical changes are permanent. "
                            + str(
                                str(
                                    "Moreover, "
                                    + self.user_clientlist().pronoun_possessive("")
                                    + self.user_clientlist()[0].surgeon.type
                                    + " is expected to testify that "
                                    + self.user_clientlist().pronoun_subjective()
                                    + " has suffered a permenent injury. "
                                )
                            )
                            if self.user_clientlist()[
                                0
                            ].surgeon_expected_to_testify_permanency
                            else ""
                        )
                    )
                    if self.user_clientlist()[0].surgery_required
                    else self.permanency_statement
                )
            )
            + "Thus, the jury should receive the "
            + self.permanent_injury_instruction
            + " jury instruction. According to "
            + CDC_VITAL_STATS
            + ", "
            + self.user_clientlist().asnoun()
            + ", now age "
            + str(self.user_clientlist()[0].age_in_years())
            + ", can expect to live an additional "
            + str(self.user_clientlist()[0].life_expectancy())
            + " years with these changes. (See Exhibit 1)."
            + self.fx.NLINE
        ))
        return output

    @property
    def future_complications_compensible_tm(self):
        output = str(
            "The Oregon Supreme Court has acknowledged that a jury may consider future possible complications in determining damages. "
            + self.fx.ITALIC
            + "Feist v. Sears Roebuck & Co."
            + self.fx.ITALIC
            + ", 267 Or 402, 410 (1973) (holding that a mere possibility of future meningitis stemming from a head injury was properly submitted to the jury where a medical probability as to causation of initial head injury had been shown). The Oregon Court of Appeals has reiterated this rule. "
            + self.fx.ITALIC
            + "Pelcha v. United Amusement Co."
            + self.fx.ITALIC
            + ", 44 Or App 675, 678, rev. den., 289 Or 275 (1980) (finding that the likelihood of future complications was a proper subject of evidence and a proper question for the jury testimony where plaintiff’s doctor testified that plaintiff’s chance of future physical complications and surgery was between 30% and 45%); "
            + self.fx.ITALIC
            + "Henderson v. Hercules, Inc."
            + self.fx.ITALIC
            + ", 57 Or App 791, 796-797 (1982) (holding that a “less than 50% possibility” of future physical problems and surgery was properly submitted to the jury). "
            + self.fx.NLINE
            + "Here, "
            + possessify(self.user_clientlist().asnoun(), "")
            + "medical experts will opine that, based on a reasonable medical probability, "
            + self.user_clientlist().salute_possessive()
            + "injuries were exacerbated and/or caused by the "
            + self.incident
            + ". "
            + capitalize(self.user_clientlist().pronoun_possessive(""))
            + "medical experts will also testify that, because of the "
            + self.incident
            + ", "
            + self.user_clientlist().asnoun()
            + " is left susceptible to future complications. Thus, the issue of "
            + possessify(self.user_clientlist().asnoun(), "")
            + "possible future complications should be submitted to the jury for consideration. "
            + self.fx.NLINE
        )
        return output

    @property
    def lay_witnesses_competent_for_pain_and_suffering_tm(self):
        output = str(
            "A plaintiff's testimony about "
            + self.user_clientlist().pronoun_possessive("")
            + "own discomfort is always competent evidence on the issue of past and future pain and suffering. "
            + self.fx.ITALIC
            + "Skeeters v. Skeeters"
            + self.fx.ITALIC
            + ", 237 Or 204, 231 (1964). Lay witnesses may also testify about the plaintiff’s present complaints of pain or disability or about the witness’ own observation of the plaintiff’s limited activity or pain behavior. "
            + self.fx.ITALIC
            + "Frangos v. Edmunds"
            + self.fx.ITALIC
            + ", 179 Or 577, 593, (1946) (holding testimony of plaintiff’s wife about plaintiff’s physical condition and restrictions was sufficient to allow the issue of plaintiff’s physical condition and restrictions to go to the jury). "
            + self.fx.NLINE
        )
        return output

    @property
    def must_produces_recorded_statements_tm(self):
        output = str(
            "After a witness testifies, opposing counsel is entitled to any statements a witness gave, even if to his or her own insurance company. "
            + self.fx.ITALIC
            + "Rigelman v. Gilligan"
            + self.fx.ITALIC
            + ", 265 Or 109, 115 (1973), citing with approval "
            + self.fx.ITALIC
            + "Pacific N.W. Bell v. Century Home"
            + self.fx.ITALIC
            + ", 261 Or 333, 339-340 (1972) (holding that the transcript of a recorded statement of an adverse testifying witness must be produced prior to cross examination, even though the transcript was work product). "
            + self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("request")
            + " that all recorded statements of "
            + possessify(self.recipients().asnoun(), "")
            + " witnesses be produced prior to cross examination of said witnesses. "
            + self.fx.NLINE
        )
        return output

    @property
    def rule_of_completeness_tm(self):
        output = str(
            "If "
            + self.recipients().asnoun()
            + " reads, or otherwise offers, portions of deposition transcripts at trial, "
            + self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("reserve")
            + " "
            + self.user_clientlist().pronoun_possessive("")
            + "right to have additional portions read at that time. The “Rule of Completeness” requires that, if one side reads portions of a deposition at trial, opposing counsel has the right to have other portions of the deposition on the same subject read at that time. "
            + self.rule_of_completeness
            + ". Fairness requires that opposing counsel should be able to introduce other portions of the deposition that qualify, or provide context for, the proffered testimony. It is reversible error not to permit the introduction of other relevant portions of the deposition, at least at the conclusion of the reading of the deposition. "
            + self.fx.ITALIC
            + "Westinghouse Elec. Corp. v. Wray Equip. Corp."
            + self.fx.ITALIC
            + ", 286 F2d 491, 494 (1st Cir.), "
            + self.fx.ITALIC
            + "cert den"
            + self.fx.ITALIC
            + " 366 US 929 (1961). "
            + self.fx.NLINE
        )
        return output

    @property
    def ability_to_pay_tm(self):
        output = str(
            "The jury will not be entitled to know about the insurance status of "
            + self.recipients().asnoun()
            + ", or the fact that there is a very large insurance policy available to "
            + self.recipients().asnoun()
            + ". They may well wonder who will be paying any verdict in this matter. Because the issue of liability insurance is inadmissible for most purposes, the jury should receive the Ability to Pay instruction, UCJI 16.01"

            # + self.ability_to_pay_instruction
            + ". "
            + self.fx.NLINE
            + " Truth be told, "
            + self.recipients().asnoun()
            + " is a very nice person who doesn't appear to be a person of great means. Therefore, the jury may well feel sorry for "
            + self.recipients().pronoun_objective()
            + " and "
            + self.recipients().pronoun_possessive("")
            + "family, particularly when they are not allowed to know about the insured status or applicable insurance limits relative to this claim. This natural sympathy towards the defendant runs the substantial risk of diminishing the amount of damages to which "
            + self.user_clientlist().salute()
            + " is lawfully entitled. It is for these reasons that "
            + self.user_clientlist().asnoun()
            + " therefore respectfully "
            + self.user_clientlist().does_verb("requests")
            + " that this instruction be given. Failure to do so would cause "
            + self.user_clientlist().asnoun()
            + " undue prejudice, as the jury would be inclined to allow sympathy for the "
            + self.recipients().asnoun()
            + " to creep into the case, contrary to UCJI 5.01."
            # + self.precautionary_jury_instructions
            + self.fx.NLINE
        )
        return output

    @property
    def written_jury_instruction_tm(self):
        output = str(
            self.instructions_to_jury_rule
            + " B provides that “In charging the jury . . . the court shall reduce, or require a party to reduce, the charge to writing. However, if the preparation of written instructions is not feasible, the court may record the instructions electronically during the charging of the jury.” "
            + self.user_clientlist().asnoun()
            + " "
            + self.user_clientlist().does_verb("request")
            + " that the jury receive a written copy of the jury instructions charged and offers to reduce those instructions to writing. "
            + self.fx.NLINE
            + possessify(self.user_clientlist().asnoun(), "")
        )
        return output

    @property
    def leading_questions_tm(self):
        output = str(
            "In "
            + self.user_clientlist().pronoun_possessive("")
            + "case in chief, "
            + possessify(self.user_clientlist().asnoun(), "")
            + " counsel expects to call "
            + self.recipients().asnoun()
            + " as a witness. "
            + self.leading_questions_rule
            + " permits an attorney to use leading questions when examining the adverse party or those associated with the adverse party. Therefore, "
            + possessify(self.user_clientlist().asnoun(), "")
            + " counsel should be allowed to ask leading questions on "
            + possessify(self.recipients().asnoun(), "")
            + " direct examination. "
        )
        return output

    def dram_shops(self):
        output = SCPartyList("dram_shops", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text in ("business entity")
                    and party not in output.elements
            ):
                if party.party.is_dram_shop:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def drunk(self):
        output = SCPartyList("drunk", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text not in ("business entity", "governmental entity")
                    and party not in output.elements
            ):
                if party.party.is_drunkard:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def battered(self):
        output = SCPartyList("battered", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text not in ("business entity", "governmental entity")
                    and party not in output.elements
            ):
                if party.party.was_battered:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def batteror(self):
        output = SCPartyList("batteror", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text not in ("business entity", "governmental entity")
                    and party not in output.elements
            ):
                if party.party.is_batteror:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def hurt_by_hockey(self):
        output = SCPartyList("hurt_by_hockey", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text not in ("business entity", "governmental entity")
                    and party not in output.elements
            ):
                if party.party.was_hurt_by_hockey:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def animal_attacked(self):
        output = SCPartyList("animal_attacked", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if (
                    party.party.name.caption_type.name.text not in ("business entity", "governmental entity")
                    and party not in output.elements
            ):
                if party.party.was_attacked_by_animal:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def animal_owners(self):
        output = SCPartyList("animal_owners", there_are_any=True, auto_gather=False)
        output.clear()
        for party in self.parties:
            if party not in output.elements:
                if party.party.owned_animal:
                    output.append(party)
        output.there_is_another = False
        output.gathered = True
        return output

    def facite_rfps(self):
        indexr = Counter()
        request = list()
        rfp_val = self.prior_requests_total
        for item in self.rfps:
            request.append(
                (str(rfp_val), fix_punctuation(single_paragraph(item.name.text)))
            )
            rfp_val += 1
        return request

    def friends_of(self, lawfirm):
        output = []
        for relationship in self.lawyer_firm_rel:
            if self.clients_of(relationship.lawfirm) == self.clients_of(lawfirm):
                output.append(firm)
        return output

    def firm_of(self, attorney):
        for relationship in self.lawyer_lawfirm_rel:
            if relationship.lawyer == attorney:
                return relationship.lawfirm

    def attorneys_for(self, firm):
        output = SCLawyerList("attorneys_for", there_are_any=True, auto_gather=False)
        output.clear()
        for relationship in self.lawyer_lawfirm_rel:
            if relationship.lawfirm is firm:
                output.append(relationship.lawyer)
        output.there_is_another = False
        output.gathered = True
        return output

    def clients_of(self, firm):
        output = SCList(
            "clients_of",
            object_type=SCIndividual,
            there_are_any=True,
            auto_gather=False,
        )
        output.clear()
        for relationship in self.firm_client:
            if relationship.lawfirm == firm:
                output.append(relationship.client)
        output.there_is_another = False
        output.gathered = True
        return output

    def pair_firmz(self, firm):
        output = SCList(object_type=SCIndividual, there_are_any=True, auto_gather=False)
        output.clear()
        for relationship in self.firm_client:
            if self.clients_of(relationship.lawfirm) in self.clients_of(firm):
                put.append(item)
        output = put.copy()
        for x in output:
            if x == firm:
                z = output.index(x)
                del output[z]
        output.there_is_another = False
        output.gathered = True
        return output

    def attorneys_of_list(self, parties):
        try:
            output = SCList(
                "attorneys_of_list",
                object_type=SCIndividual,
                there_are_any=True,
                auto_gather=False,
            )
            for party in parties:
                for relationship in self.lawyer_lawfirm_rel:
                    if relationship.lawfirm in party.party.firms:
                        if relationship.lawyer not in output.elements:
                            output.append(relationship.lawyer)
            output.there_is_another = False
            output.gathered = True
            return output
        except Exception as err:
            raise Exception(err.__class__.__name__ + ": " + str(err))


class Jurisdiction(Thing, SQLObject):
    _model = JurisdictionModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class State(Thing, SQLObject):
    _model = StateModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class County(Thing, SQLObject):
    _model = CountyModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()
        
    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        elif column == 'address':
            return self.address.address.strip()
        elif column == "unit":
            return self.address.unit.strip()
        elif column == 'city':
            return self.address.city.strip()
        elif column == 'state':
            return self.address.state.strip()
        elif column == "zip":
          return self.address.zip.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        elif column == 'address':
            self.address.address = value
        elif column == "unit":
            self.address.unit = value
        elif column == 'city':
            self.address.city = value
        elif column == 'state':
            self.address.state = value
        elif column == "zip":
            self.address.zip = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        elif column == 'address':
            del self.address.address
        elif column == "unit":
            del self.address.unit
        elif column == "city":
            del self.address.city
        elif column == "state":
            del self.address.state
        elif column == "zip":
            del self.address.zip
        else:
            raise Exception("Invalid column " + column)


class District(Thing, SQLObject):
    _model = DistrictModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class Division(Thing, SQLObject):
    _model = DivisionModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class Specialty(Thing, SQLObject):
    _model = SpecialtyModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'name':
            return self.name.text.strip()
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'name':
            self.name.text = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'name':
            del self.name.text
        else:
            raise Exception("Invalid column " + column)


class Case(SCCase, SQLObject):
    _model = CaseModel
    _session = DBSession
    _required = ['casename']
    _uid = 'casename'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()


    def db_get(self, column):
        if column == 'casename':
            return self.casename.strip()
        elif column == 'number':
            return self.docket_number.strip()
        elif column == "jurisdiction_id":
            return self.juris.jurisdiction.id
        elif column == "state_id":
            return self.juris.state.id
        elif column == "county_id":
            return self.juris.county.id
        elif column == "district_id":
            return self.juris.district.id
        elif column == "division_id":
            return self.juris.division.id
        elif column == "type":
            return self.type.strip()
        elif column == "dol":
            return self.dol
        elif column == "tol":
            return self.tol
        elif column == "sol":
            return self.sol
        elif column == "filed":
            return self.filed
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'casename':
            self.casename = value
        elif column == 'number':
            self.docket_number = value
        elif column == "jurisdiction_id":
            self.juris.jurisdiction = Jurisdiction.by_id(value)
        elif column == "state_id":
            self.juris.state = State.by_id(value)
        elif column == "county_id":
            self.juris.county = County.by_id(value)
        elif column == "district_id":
            self.juris.district = District.by_id(value)
        elif column == "division_id":
            self.juris.division = Division.by_id(value)
        elif column == "type":
            self.type = value
        elif column == "dol":
            self.dol = value
        elif column == "tol":
            self.tol = value
        elif column == "sol":
            self.sol = value
        elif column == "filed":
            self.filed = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'casename':
            del self.casename
        elif column == 'number':
            del self.docket_number
        elif column == "jurisdiction_id":
            del self.juris.jurisdiction
        elif column == "state_id":
            del self.juris.state
        elif column == "county_id":
            del self.juris.county
        elif column == "district_id":
            del self.juris.district
        elif column == "division_id":
            del self.juris.division
        elif column == "type":
            del self.type
        elif column == "dol":
            del self.dol
        elif column == "tol":
            del self.tol
        elif column == "sol":
            del self.sol
        elif column == "filed":
            del self.filed
        else:
            raise Exception("Invalid column " + column)

    @classmethod
    def all_case_names(cls, instance_name=None):
        if 'dbcache' not in this_thread.misc:
            this_thread.misc['dbcache'] = {}
        if instance_name:
            listobj = DAList(instance_name, object_type=cls)
        else:
            listobj = DAList(object_type=cls)
            listobj.set_random_instance_name()
        if cls._session.query(cls._model.casename).order_by(cls._model.casename).first():
            for db_entry in list(cls._session.query(cls._model.casename).order_by(cls._model.casename).all()):
                if cls._model.__name__ in this_thread.misc['dbcache'] and db_entry.casename in \
                        this_thread.misc['dbcache'][cls._model.__name__]:
                    listobj.append(this_thread.misc['dbcache'][cls._model.__name__][db_entry.casename])
                else:
                    listobj.append(db_entry.casename)
        listobj.gathered = True
        return listobj

    def get_3dp_candidates(self):
        output_3dp_candidates = SCList("output_3dp_candidates")
        output_3dp_candidates.there_are_any = True
        output_3dp_candidates.auto_gather = False
        output_3dp_candidates.clear()
        for party in self.parties:
            db_entry = self._session.query(PartyModel).filter(PartyModel.id == party.party.id).first()
            if db_entry.party_type != "Third-Party Defendant":
                output_3dp_candidates.append(Party.by_id(db_entry.id))
        output_3dp_candidates.there_is_another = False
        output_3dp_candidates.gathered = True
        return output_3dp_candidates

    def has_lawyer(self, lawyer):
        if not (self.ready() and lawyer.ready()):
            raise Exception("has_lawyer: cannot retrieve data")
        log("Case . . . self.has_lawyer()")
        db_entry = self._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == self.id,
                                                               CaseLawyerModel.lawyer_id == lawyer.id).first()
        if db_entry is None:
            return False
        return True

    def add_lawyer(self, lawyer):
        log("Case . . . self.add_lawyer()")
        if not self.has_lawyer(lawyer):
            log("Case . . . self.add_lawyer()2")
            db_entry = CaseLawyerModel(cases_id=self.id, lawyer_id=lawyer.id)
            self._session.add(db_entry)
            self._session.commit()

    def add_lawyers(self):
        log("Case . . . self.add_lawyers()")
        for lawyer in self.lawyers.elements:
            log("Case . . . self.add_lawyers()2")
            self.add_lawyer(lawyer.lawyer)

    def get_lawyers(self):
        if not self.ready():
            raise Exception("get_lawfirms: cannot retrieve data")
        log("Case . . . self.get_lawyers()")
        if self._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == self.id).first():
            log("Case . . . self.get_lawyers()2")
            self.lawyers.clear()
            for db_entry in self._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == self.id).all():
                log("Case . . . self.get_lawyers()3")
                self.lawyers.appendObject()
                self.lawyers[-1].lawyer = Lawyer.by_id(db_entry.lawyer_id)
                # entry = self._session.query(LawyerModel).filter(LawyerModel.id == self.lawyers[-1].id)
            log("Case . . . self.get_lawyers()4")

    def del_lawyer(self, lawyer):
        if not (self.ready() and lawyer.ready()):
            raise Exception("del_lawyer: cannot retrieve data")
        self._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == self.id,
                                                    CaseLawyerModel.lawyer_id == lawyer.id).delete()
        self._session.commit()

    def add_party_designation(self, party, designation):
        try:
            if not self.has_party_designation(party, designation):
                db_entry = PartyDesignationModel(party_id=party.id, designation_id=designation.id)
                self._session.add(db_entry)
                self._session.commit()
        except:
            self._session.rollback

    def has_party_designation(self, party, designation):
        try:
            if not (self.ready() and party.ready() and designation.ready()):
                raise Exception("has_party_designation: cannot retrieve data")
            db_entry = self._session.query(PartyDesignationModel).filter(
                PartyDesignationModel.designation_id == designation.id,
                PartyDesignationModel.party_id == party.id).first()
            if db_entry is None:
                return False
            return True
        except:
            self._session.rollback()

    def get_party_designations(self, party):
        if not (self.ready() and party.ready()):
            raise Exception("get_party_designations: cannot retrieve data")
        log("self.get_party_designations()")
        if self._session.query(PartyDesignationModel).filter(PartyDesignationModel.party_id == party.id).first():
            log("self.get_party_designations.auto_gather=False")
            party.party_types.clear()
            for db_entry in self._session.query(PartyDesignationModel).filter(
                    PartyDesignationModel.party_id == party.id).all():
                # party.party_types.appendObject()
                log("designation_id " + repr(db_entry.designation_id) + " party_id " + repr(db_entry.party_id))
                party.party_types.append(Designation.by_id(db_entry.designation_id))
                # party.party_types[-1].db_read()
            party.party_types.there_is_another = False
            party.party_types.gathered = True

    def del_party_designation(self, party, designation):
        if not (self.ready() and party.ready() and designation.ready()):
            raise Exception("del_party_designation: cannot retrieve data")
        self._session.query(PartyDesignationModel).filter(PartyDesignationModel.designation_id == designation.id,
                                                          PartyDesignationModel.party_id == party.id).delete()
        self._session.commit()

    def has_party(self, party):
        try:
            if not (self.ready() and party.ready()):
                raise Exception("has_party: cannot retrieve data")
            db_entry = self._session.query(CasePartyModel).filter(CasePartyModel.cases_id == self.id,
                                                                  CasePartyModel.party_id == party.id).first()
            if db_entry is None:
                return False
            return True
        except:
            self._session.rollback()

    def add_party(self, party):
        if not self.has_party(party):
            db_entry = CasePartyModel(cases_id=self.id, party_id=party.id)
            self._session.add(db_entry)
            self._session.commit()

    def add_parties(self):
        for party in self.parties.elements:
            self.add_party(party.party)

    def get_parties(self):
        if not self.ready():
            raise Exception("get_parties: cannot retrieve data")
        log("self.get_parties()")
        if self._session.query(CasePartyModel).filter(CasePartyModel.cases_id == self.id).first():
            log("self.parties.auto_gather=False")
            self.parties.clear()
            for db_entry in self._session.query(CasePartyModel).filter(CasePartyModel.cases_id == self.id).all():
                self.parties.appendObject()
                self.parties[-1].party = Party.by_id(db_entry.party_id)
            self.parties.there_is_another = False

    def del_party(self, party):
        if not (self.ready() and party.ready()):
            raise Exception("del_party: cannot retrieve data")
        self._session.query(CasePartyModel).filter(CasePartyModel.cases_id == self.id,
                                                   CasePartyModel.party_id == party.id).delete()
        self._session.commit()

    def has_lawfirm(self, lawfirm):
        if not (self.ready() and lawfirm.ready()):
            raise Exception("has_lawfirm: cannot retrieve data")
        db_entry = self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id,
                                                                CaseLawfirmModel.lawfirm_id == lawfirm.id).first()
        if db_entry is None:
            return False
        return True

    def add_lawfirm(self, lawfirm):
        if not self.has_lawfirm(lawfirm):
            db_entry = CaseLawfirmModel(cases_id=self.id, lawfirm_id=lawfirm.id)
            self._session.add(db_entry)
            self._session.commit()

    def add_lawfirms(self):
        for lawfirm in self.lawfirms.elements:
            self.add_lawfirm(lawfirm.lawfirm)

    def get_lawfirms(self):
        if not self.ready():
            raise Exception("get_lawfirms: cannot retrieve data")
        if self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id).first():
            self.lawfirms.clear()
            for db_entry in self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id).all():
                self.lawfirms.appendObject()
                self.lawfirms[-1].lawfirm = Lawfirm.by_id(db_entry.lawfirm_id)
            self.lawfirms.there_is_another = False

    def del_lawfirm(self, lawfirm):
        if not (self.ready() and lawfirm.ready()):
            raise Exception("del_lawfirm: cannot retrieve data")
        self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id,
                                                     CaseLawfirmModel.lawfirm_id == lawfirm.id).delete()
        self._session.commit()

    def list_lawfirms(self):
        if not self.ready():
            raise Exception("get_lawfirms: cannot retrieve data")
        lawfirms = []
        if self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id).first():
            lawfirms.clear()
            for db_entry in self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.id).all():
                lawfirms.append(Lawfirm.by_id(db_entry.lawfirm_id))
            return lawfirms

    def get_lawfirm_parties(self, lawfirm):
        if not (self.ready() and lawfirm.ready()):
            raise Exception("get_lawfirm_parties: cannot retrieve data")
        get_lawfirm_parties_results = SCList("get_lawfirm_parties_results", there_are_any=True, auto_gather=True)
        get_lawfirm_parties_results.clear()
        for party in self.parties:
            log("get_lawfirm_parties_results 1")
            if self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.party.id,
                                                             PartyLawfirmModel.lawfirm_id == lawfirm.id).first():
                log("get_lawfirm_parties_results 2")
                partylawfirm_db = self._session.query(PartyLawfirmModel).filter(
                    PartyLawfirmModel.party_id == party.party.id, PartyLawfirmModel.lawfirm_id == lawfirm.id).first()
                if self._session.query(CasePartyLawfirmModel).filter(CasePartyLawfirmModel.cases_id == self.id,
                                                                     CasePartyLawfirmModel.party_lawfirm_id == partylawfirm_db.id).first():
                    log("get_lawfirm_parties_results 3")
                    get_lawfirm_parties_results.append(Party.by_id(party.party.id))
        get_lawfirm_parties_results.there_is_another = False
        get_lawfirm_parties_results.gathered = True
        return get_lawfirm_parties_results

    def has_lawfirm_party(self, party, lawfirm):
        if not (self.ready() and party.ready() and lawfirm.ready()):
            raise Exception("has_lawfirm_party: cannot retrieve data")
        if self.has_partylawfirm(party, lawfirm):
            party_lawfirm_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.id,
                                                                                PartyLawfirmModel.lawfirm_id == lawfirm.id).first()
            db_entry = self._session.query(CasePartyLawfirmModel).filter(CasePartyLawfirmModel.cases_id == self.id,
                                                                         CasePartyLawfirmModel.party_lawfirm_id == party_lawfirm_entry.id).first()
            if db_entry is None:
                return False
            return True
        else:
            return False

    def add_lawfirm_party(self, party, lawfirm):
        if not self.has_lawfirm_party(party, lawfirm):
            party_lawfirm_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.id,
                                                                                PartyLawfirmModel.lawfirm_id == lawfirm.id).first()
            db_entry = CasePartyLawfirmModel(cases_id=self.id, party_lawfirm_id=party_lawfirm_entry.id)
            self._session.add(db_entry)
            self._session.commit()

    def del_lawfirm_party(self, party, lawfirm):
        if not (self.ready() and party.ready() and lawfirm.ready()):
            raise Exception("del_lawfirm_party: cannot retrieve data")
        party_lawfirm_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.id,
                                                                            PartyLawfirmModel.lawfirm_id == lawfirm.id).first()
        self._session.query(CasePartyLawfirmModel).filter(CasePartyLawfirmModel).filter(
            CasePartyLawfirmModel.cases_id == self.id,
            CasePartyLawfirmModel.party_lawfirm_id == party_lawfirm_entry.id).delete()
        self._session.commit()

    def has_partylawfirm(self, party, lawfirm):
        if not (self.ready() and party.ready() and lawfirm.ready()):
            raise Exception("has_partylawfirm: cannot retrieve data")
        party_lawfirm_entry = self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.id,
                                                                            PartyLawfirmModel.lawfirm_id == lawfirm.id).first()
        if party_lawfirm_entry is None:
            return False
        return True

    def add_partylawfirm(self, party, lawfirm):
        log("add_partylawfirm1")
        if not self.has_partylawfirm(party, lawfirm):
            log("add_partylawfirm2")
            party_lawfirm_entry = PartyLawfirmModel(party_id=party.id, lawfirm_id=lawfirm.id)
            self._session.add(party_lawfirm_entry)
            self._session.commit()

    def del_partylawfirm(self, party, lawfirm):
        if not (self.ready() and party.ready() and lawfirm.ready()):
            raise Exception("del_lawfirm_party: cannot retrieve data")
        self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.party_id == party.id,
                                                      PartyLawfirmModel.lawfirm_id == lawfirm.id).delete()
        self._session.commit()

    def relate_lawfirms_to_parties(self, party):
        log("relate_lawfirms_to_parties")
        if self._session.query(CasePartyLawfirmModel).filter(CasePartyLawfirmModel.cases_id == self.id).first():
            for db_entry in self._session.query(CasePartyLawfirmModel).filter(
                    CasePartyLawfirmModel.cases_id == self.id).all():
                if self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.id == db_entry.party_lawfirm_id,
                                                                 PartyLawfirmModel.party_id == party.id).first():
                    party.firms.auto_gather = False
                    party.firms.clear()
                    for party_lawfirm_entry in self._session.query(PartyLawfirmModel).filter(
                            PartyLawfirmModel.id == db_entry.party_lawfirm_id,
                            PartyLawfirmModel.party_id == party.id).all():
                        party.firms.append(Lawfirm.by_id(party_lawfirm_entry.lawfirm_id))
                    party.firms.there_is_another = False
                    party.firms.gathered = True


def quotefull(text):
    words = text.strip()
    output = ""
    left = "&#8220;"
    right = "&#8221;"
    if not words[0].isupper():
        output += str(left + "[" + words[0].upper() + "]" + words[1:] + right)
    else:
        output += str(left + words + right)
    return output


def quote(text):
    words = str(text.strip())
    output = ""
    left = "&#8220;"
    right = "&#8221;"
    output += str(left + words + right)
    return output


def possessify(a, b, **kwargs):
    the_word = str(a)
    # apostrophe = "’"
    apostrophe = "&#8217;"
    if the_word[-1] == "s" and "plural" not in kwargs:
        kwargs["plural"] = True
    return possessify_en(a, b, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class Color(DAObject):
    @property
    def BOLD(self):
        return "**"

    @property
    def ITALIC(self):
        return "*"

    @property
    def ALINE(self):
        if this_thread.evaluation_context == "docx":
            return "\n"
        else:
            return "\n\n"

    @property
    def NLINE(self):
        if this_thread.evaluation_context == "docx":
            return "\n\n"
        else:
            return "\n\n"

    @property
    def NLINE2(self):
        if this_thread.evaluation_context == "docx":
            return "\a"
        else:
            return "\n"

    @property
    def TAB(self):
        if this_thread.evaluation_context == "docx":
            return "\t"
        else:
            return "[TAB]"

    @property
    def CENTER(self):
        return "[CENTER]"

    @property
    def BEGIN_TWOCOL(self):
        return "[BEGIN_TWOCOL]"

    @property
    def BREAK(self):
        return "[BREAK]"

    @property
    def END_TWOCOL(self):
        return "[END_TWOCOL]"

    @property
    def BORDER(self):
        return "[BORDER]"

    @property
    def NEWPAR(self):
        return "[NEWPAR]"

    @property
    def BLANK(self):
        return "[BLANK]"

    @property
    def FLUSHLEFT(self):
        return "[FLUSHLEFT]"

    @property
    def FLUSHRIGHT(self):
        return "[FLUSHRIGHT]"

    @property
    def CENTER(self):
        return "[CENTER]"

    @property
    def BOLDCENTER(self):
        return "[BOLDCENTER]"

    @property
    def INDENTBY5(self):
        return "[INDENTBY 0.5in]"

    @property
    def INDENTBY25(self):
        return "[INDENTBY 0.25in]"

    @property
    def BLANKFILL(self):
        return "[BLANKFILL]"

    @property
    def BOLDCENTER(self):
        return "[BOLDCENTER]"

    @property
    def SINGLE(self):
        return "[SINGLESPACING]"

    @property
    def DOUBLE(self):
        return "[ONEANDAHALFSPACING]"

    @property
    def EM(self):
        return "[EMDASH]"

    @property
    def EN(self):
        return "[ENDASH]"

    @property
    def PAGEBREAK(self):
        return "[PAGEBREAK]"

    @property
    def INDENT_START(self):
        return "[START_INDENTATION]"

    @property
    def INDENT_STOP(self):
        return "[STOP_INDENTATION]"

    @property
    def BLOCK(self):
        return "[INDENTBY 0.5in 0.5in]"


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCDocumentList(DAList):
    def init(self, *pargs, **kwargs):
        self.object_type = DAObject
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
def get_row(sheet, row):
    output = []
    col = 1
    while col < sheet.ncols:
        val = sheet.cell_value(row, col)
        if not val:
            break
        output.append(val)
        col += 1
    return output


class SCUnsortedRequestsList(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCUnsortedRequests
        # self.there_are_any = True
        self.complete_attribute = "complete_unsorted"
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCUnsortedRequests(SCRequests):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)

    @property
    def complete_unsorted(self):
        self.name.text
        self.default
        self.sorted


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCVehicleList(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCVehicle
        self.there_are_any = True
        self.complete_attribute = "vehicle_complete"
        super().init(*pargs, **kwargs)

    def w_driver_possessive(self):
        output = []
        output.clear()
        for item in self:
            output.append(item.w_driver_possessive())
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCVehicle(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("passengers", SCList)
        self.initializeAttribute("driver", SCIndividual)
        self.initializeAttribute("address", Address)
        super().init(*pargs, **kwargs)

    @property
    def vehicle_complete(self):
        self.make
        self.model
        self.year
        self.type
        self.driver
        self.passengers
        self.address.address
        self.dot
        self.lot
        self.move

    def __unicode__(self):
        return self.make + " " + self.model

    def ymm(self):
        return self.year + " " + self.make + " " + self.model

    def w_driver_possessive(self):
        return possessify(self.driver.name.full(), str(self.make + " " + self.model))


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCMVA(Address):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCWitnessList(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCWitness
        self.there_are_any = True
        self.complete_attribute = "complete_witness"
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCWitness(SCIndividual):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("relationship", Thing)
        self.initializeAttribute("activity_changes", SCList)
        super().init(*pargs, **kwargs)

    @property
    def complete_witness(self):
        self.name.first
        self.address.address
        self.activity_changes
        self.phone_number
        self.activity_changes
        self.education

    @property
    def is_party(self):
        return False

    @property
    def is_lawyer(self):
        return False

    @property
    def is_lawfirm(self):
        return False

    @property
    def is_witness(self):
        return True

    @property
    def clio_type(self):
        return "Person"


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class Counter:
    def __init__(self, *pargs, **kwargs):
        self.val = 1

    def set(self, val):
        self.val = val
        return ""

    def get(self):
        return self.val

    def get_incr(self):
        to_return = self.val
        self.increment()
        return to_return

    def incr_get(self):
        self.increment()
        return self.val

    def increment(self):
        self.val += 1
        return ""

    def decrement(self):
        self.val -= 1
        return ""

    def reset(self):
        self.val = 1
        return ""


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCNonEconomics(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        self.initializeAttribute("wrongful_death", SCWrongfulDeath)
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCEconomics(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("meds", SCMeds)
        self.initializeAttribute("future_meds", SCFutureMeds)
        self.initializeAttribute("lost_wages", SCLostWages)
        self.initializeAttribute("lost_profits", SCLostProfits)
        self.initializeAttribute("rehab", SCRehab)
        self.initializeAttribute("reduced_earning_capacity", SCReducedEarningCapacity)
        self.initializeAttribute("pecuniary_losses", SCPecuniaryLosses)
        self.initializeAttribute("mortuary", SCMortuary)
        super().init(*pargs, **kwargs)

    def total(self, case):
        output = 0.0
        output += self.meds.total
        output += self.future_meds.total
        output += self.lost_wages.total(case)
        output += self.lost_profits.total
        output += self.reduced_earning_capacity.total
        return currency(output)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLostProfits(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Lost profits in an amount to be determined at the time of trial, but for the purposes of ORCP 18, estimated to be "
        output += str(self.fx.BOLD + currency(self.amount) + self.fx.BOLD)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCWrongfulDeath(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self, decedent):
        output = ""
        output += str(
            decedent.name.full()
            + " is survived by "
            + comma_and_list(decedent.survivors())
            + ". As a result of "
            + decedent.possessive("")
            + "death, "
            + decedent.pronoun_possessive("family")
            + " has and will suffer noneconomic damages for loss of society and companionship in an amount to be determined by the jury to be fair and reasonable, but not to exceed the sum of "
        )
        output += str(self.fx.BOLD + currency(self.amount) + self.fx.BOLD)
        output += ", all of which was reasonably foreseeable."
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCRehab(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self, decedent):
        output = ""
        output += "Reasonable and necessary medical, hospital, doctor, therapy, and vocational rehabilitation expenses in the future in an amount that is not known at this time and will be amended prior to trial, but for the purposes of ORCP 18, is estimated to be "
        output += str(self.fx.BOLD + currency(self.amount) + self.fx.BOLD)
        output += ", all of which was reasonably foreseeable."
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCMortuary(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Mortuary, funeral and burial expenses in the approximate sum of "
        output += str(self.fx.BOLD + currency(self.amount) + self.fx.BOLD)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCPecuniaryLosses(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("amount", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Pecuniary loss in the approximate sum of "
        output += str(self.fx.BOLD + currency(self.amount) + self.fx.BOLD)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCFutureMed(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("estimated_cost", Value)
        super().init(*pargs, **kwargs)

    @property
    def complete_future_med(self):
        self.anticipated_med
        self.estimated_cost


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCReducedEarningCapacity(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("reduction", Value)
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Future impairment of earning capacity in an amount to be determined at the time of trial, but for the purposes of ORCP 18, estimated to be "
        output += str(self.fx.BOLD + currency(self.reduction) + self.fx.BOLD)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCWageLoss(DAObject):
    def loss(self, case):
        if self.has_worked:
            if self.tried_work and self.has_returned:
                timesince = date_difference(
                    starting=self.datestopped, ending=self.datereturned
                )
            if self.tried_work and not self.has_returned:
                timesince = date_difference(starting=self.datestopped)
            if self.has_returned and not self.tried_work:
                timesince = date_difference(starting=case.doi, ending=self.datereturned)
        else:
            timesince = date_difference(starting=case.doi)
        if self.earns == "Hourly":
            interval = timesince.weeks * self.hours_per_week
            output = interval * self.wage
        elif self.earns == "Salary":
            interval = timesince.weeks * 4
            output = interval * self.wage
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLostWages(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCWageLoss
        self.complete_attribute = "complete_wages"
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Lost income to date in the approximate sum of "
        output += str(self.fx.BOLD + currency(self.total) + self.fx.BOLD)
        return output

    @property
    def total(self, case):
        output = 0.0
        for item in self:
            output += item.loss(case)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCFutureMeds(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCFutureMed
        self.complete_attribute = "complete_future_med"
        super().init(*pargs, **kwargs)

    @property
    def complaint_text(self):
        output = ""
        output += "Future reasonable and necessary medical expenses in an amount to be determined at the time of trial, but for the purposes of ORCP 18, estimated to be "
        output += str(self.fx.BOLD + currency(self.total) + self.fx.BOLD)
        return output

    @property
    def total(self):
        output = 0.0
        for item in self:
            output += item.estimated_cost
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCMeds(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCMedsVisit
        self.complete_attribute = "complete_provider"
        super().init(*pargs, **kwargs)

    @property
    def total(self):
        output = 0.0
        for item in self:
            output += item.total_bill
        return output

    @property
    def complaint_text(self):
        output = ""
        output += "Reasonable and necessary medical expenses to date in the approximate sum of "
        output += str(self.fx.BOLD + currency(self.total) + self.fx.BOLD)
        return output


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCMedsVisit(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("chart_note", Thing)
        self.initializeAttribute("clinician", SCIndividual)
        self.initializeAttribute("clinician", SCPerson)
        self.initializeAttribute("amount_billed", Value)
        self.initializeAttribute("pip", Value)
        self.initializeAttribute("primary_health_insurance", Value)
        self.initializeAttribute("secondary_health_insurance", Value)
        self.initializeAttribute("out_of_pocket", Value)
        self.initializeAttribute("written_off", Value)
        super().init(*pargs, **kwargs)

    def balance(self):
        return float(
            self.amount_billed
            - (
                    self.pip
                    + self.primary_health_insurance
                    + self.secondary_health_insurance
                    + self.out_of_pocket
                    + self.written_off
            )
        )


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCProvider(SCPerson):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("visits", SCMedsVisitsList)
        super().init(*pargs, **kwargs)

    @property
    def complete_provider(self):
        self.name.text
        self.address

    @property
    def total_bill(self):
        output = 0.0
        for visit in self.visits:
            output += visit.medsbill.amount_billed
        return output

    def subpoena_address(self, case):
        output = RichText(
            str(
                "TO:"
                + case.fx.TAB
                + "Records Custodian"
                + case.fx.NLINE
                + case.fx.TAB
                + self.name.full()
                + case.fx.NLINE
                + case.fx.TAB
                + self.address.line_one()
                + case.fx.NLINE
                + case.fx.TAB
                + self.address.line_two()
            )
        )
        return output

    def subpoena_text(self, case):
        output = RichText(str(case.fx.TAB))
        strb(
            output,
            str(
                "IN THE NAME OF THE STATE OF "
                + case.juris.state.name.text.upper()
                + ": "
            ),
        )
        strnull(
            output,
            str(
                "You are hereby required to appear and answer the complaint filed against you in the above-entitled cause within thirty (30) days from the date of service of this summons upon you, and if you fail to so answer, for want thereof, the plaintiff(s) will take judgment against you for the relief prayed for in the complaint on file herein, a copy of which is herewith served upon you. "
                + case.fx.NLINE
                + case.fx.TAB
            ),
        )
        strb(
            output,
            str(
                "NOTICE TO "
                + self.party_type.upper()
                + ":	READ THESE PAPERS CAREFULLY!"
            ),
        )
        if self.party_type.name.text == "Defendant":
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "You must “appear” in this case or the other side will win automatically. To “appear” you must file with the court a legal document called a “motion” or “answer.” The “motion” or “answer” must be given to the court clerk or administrator within thirty (30) days along with the required filing fee. It must be in the proper form and have proof of service on the plaintiff’s attorney or, if the plaintiff does not have an attorney, proof of service on the plaintiff."
                ),
            )
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "If you have questions, you should see an attorney immediately. If you need help in finding an attorney, you may call the "
                    + case.juris.state.name.text
                    + " State Bar’s Lawyer Referral Service online at www.oregonstatebar.org or by calling (503) 684-3763 (in the Portland metropolitan area) or toll-free elsewhere in Oregon at (800) 452-7636."
                ),
            )
        else:
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "You must “appear” to protect your rights in this matter. To “appear” you must file with the court a legal document called a “motion,” a “reply” to a counterclaim, or an“answer” to a cross-claim. The “motion,” “reply,” or “answer” must be given to the court clerk or administrator within 30 days along with the required filing fee. It must be in proper form and have proof of service on the defendant’s attorney or, if the defendant does not have an attorney, proof of service on the defendant."
                ),
            )
            strnull(
                output,
                str(
                    case.fx.NLINE
                    + case.fx.TAB
                    + "If you have questions, you should see an attorney immediately. If you need help in finding an attorney, you may contact the Oregon State Bar’s Lawyer Referral Service online at http://www.oregonstatebar.org or by calling (503) 684-3763 (in the Portland metropolitan area) or toll-free elsewhere in Oregon at (800) 452-7636."
                ),
            )
        return output

# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class MedsVisit(SCMedsVisit, SQLObject):
    _model = MedsVisitModel
    _session = DBSession

    @property
    def complete_meds_visit(self):
        self.date
        self.provider


    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'chart_note':
            return self.chart_note.name.text.strip()
        elif column == 'date':
            return self.date.strip()
        elif column == 'clinician_id':
            return self.clinician.id
        elif column == 'lawfirm_id':
            return self.firm.id
        elif column == 'amount_billed':
            return self.amount_billed.value.strip()
        elif column == 'pip':
            return self.pip.value.strip()
        elif column == 'primary_health_insurance':
            return self.primary_health_insurance.value.strip()
        elif column == 'secondary_health_insurance':
            return self.secondary_health_insurance.value.strip()
        elif column == 'written_off':
            return self.written_off.value.strip()
        elif column == 'out_of_pocket':
            return self.out_of_pocket.value.strip()
        raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'chart_note':
            self.chart_note.name.text = value
        elif column == 'date':
            self.date = value
        elif column == 'clinician_id':
            self.clinician = Gender.by_id(value)
        elif column == 'lawfirm_id':
            self.firm = Lawfirm.by_id(value)
        elif column == 'amount_billed':
            self.amount_billed.value = value
        elif column == 'pip':
            self.pip.value = value
        elif column == 'primary_health_insurance':
            self.primary_health_insurance.value = value
        elif column == 'secondary_health_insurance':
            self.secondary_health_insurance.value = value
        elif column == 'secondary_health_insurance':
            self.secondary_health_insurance.value = value
        elif column == 'written_off':
            self.written_off.value = value
        elif column == 'out_of_pocket':
            self.out_of_pocket.value = value
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'chart_note':
            del self.chart_note.name.text
        elif column == 'date':
            del self.date
        elif column == 'clinician_id':
            del self.clinician
        elif column == "lawfirm_id":
            del self.firm
        elif column == 'amount_billed':
            del self.amount_billed.value
        elif column == 'pip':
            del self.pip.value
        elif column == 'primary_health_insurance':
            del self.primary_health_insurance.value
        elif column == 'secondary_health_insurance':
            del self.secondary_health_insurance.value
        elif column == 'written_off':
            del self.written_off.value
        elif column == 'out_of_pocket':
            del self.out_of_pocket.value
        else:
            raise Exception("Invalid column " + column)

class SCExhibitsList(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCExhibit
        self.complete_attribute = "complete_exhibit"
        super().init(*pargs, **kwargs)


class SCExhibit(DAFile):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("nom", DAObject)
        self.initializeAttribute("author", DAObject)
        super().init(*pargs, **kwargs)

    def stamp(self):
        self.FirstFooterLeft = "Exhibit " + self.nom + "Page [PAGENUM] of [TOTALPAGES]"


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCCrossClaims(SCList):
    def init(self, *pargs, **kwargs):
        self.complete_attribute = "complete_cross_claim"
        self.object_type = SCCrossClaim
        super().init(*pargs, **kwargs)


class SCLienCreditors(SCList):
    def init(self, *pargs, **kwargs):
        self.object_type = SCLienCreditor
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCLienCreditor(SCPerson):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("attn", SCIndividual)
        super().init(*pargs, **kwargs)


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCCrossParties(SCList):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.there_are_any = True


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCCrossClaim(DAObject):
    def init(self, *pargs, **kwargs):
        self.initializeAttribute("ds", SCCrossParties)
        self.initializeAttribute("ps", SCCrossParties)
        super().init(*pargs, **kwargs)

    @property
    def complete_cross_claim(self):
        self.ps
        self.ds


# Trial Legend ©2021 Scott Cumming - ALL RIGHTS RESEVERED
class SCJuryInstruction(Thing):
    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)

class JuryInstruction(SCJuryInstruction, SQLObject):
    _model = JuryInstructionModel
    _session = DBSession
    _required = ['name']
    _uid = 'name'

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    def db_get(self, column):
        if column == 'no':
            return self.no.strip()
        elif column == 'name':
            return self.name.text.strip()
        elif column == 'instruction':
            return self.instruction.strip()
        elif column == 'comment_title':
            return self.comment_title.strip()
        elif column == 'comment':
            return self.comment.strip()
        elif column == 'title':
            return self.title.strip()
        elif column == 'update':
            return as_datetime(self.update.strip())
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'no':
            self.no = value
        elif column == 'name':
            self.name.text = value
        elif column == 'instruction':
            self.instruction = replace_underscores(replace_newlines(single_to_double_newlines(value)))
        elif column == 'comment_title':
            self.comment_title = value
        elif column == 'comment':
            self.comment = case_name_italicizer(replace_newlines(single_to_double_newlines(value)))
        elif column == 'title':
            self.title = value
        elif column == 'update':
            self.update = as_datetime(value)
        else:
            raise Exception("Invalid column " + column)

    def db_null(self, column):
        if column == 'no':
            del self.no
        elif column == 'name':
            del self.name.text
        elif column == 'comment_title':
            del self.comment_title
        elif column == 'instruction':
            del self.instruction
        elif column == 'comment':
            del self.comment
        elif column == 'title':
            del self.title
        elif column == 'update':
            del self.update
        else:
            raise Exception("Invalid column " + column)
            
    @property
    def complete_jury_instruction(self):
        self.name.text
        self.title
        self.no
        self.instruction
        self.comment
        #self.update

    @classmethod
    def all_instruction_titles(cls, instance_name=None):
        if 'dbcache' not in this_thread.misc:
            this_thread.misc['dbcache'] = {}
        if instance_name:
            listobj = DAList(instance_name, object_type=cls)
        else:
            listobj = DAList(object_type=cls)
            listobj.set_random_instance_name()
        for db_entry in list(cls._session.query(cls._model.name).order_by(cls._model.name).all()):
            if cls._model.__name__ in this_thread.misc['dbcache'] and db_entry.name in this_thread.misc['dbcache'][
                cls._model.__name__]:
                listobj.append(this_thread.misc['dbcache'][cls._model.__name__][db_entry.name])
            else:
                listobj.append(db_entry.name)
        listobj.gathered = True
        return listobj

class SCJuryInstructionList(SCList):
    def init(self, *pargs, **kwargs):
        if 'object_type' not in kwargs:
            kwargs['object_type'] = JuryInstruction
        if 'there_are_any' not in kwargs:
            kwargs['there_are_any'] = True
        if 'auto_gather' not in kwargs:
            kwargs['auto_gather'] = False
        if 'complete_attribute' not in kwargs:
            kwargs['complete_attribute'] = "complete_jury_instruction"
        super().init(*pargs, **kwargs)

class PartyLawfirm(DAObject, SQLObjectRelationship):
    _model = PartyLawfirmModel
    _session = DBSession
    _parent = [Lawfirm, 'lawfirm', 'lawfirm_id']
    _child = [Party, 'party', 'party_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'lawfirm_id':
            return self.lawfirm.id
        elif column == 'party_id':
            return self.party.id

    def db_set(self, column, value):
        if column == 'lawfirm_id':
            self.lawfirm = Lawfirm.by_id(value)
        elif column == 'party_id':
            self.party = Party.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(PartyLawfirmModel).filter(PartyLawfirmModel.lawfirm_id == self.lawfirm.id,
                                                                 PartyLawfirmModel.party_id == self.party.id).first()
        except:
            return None


class CaseLawyer(DAObject, SQLObjectRelationship):
    _model = CaseLawyerModel
    _session = DBSession
    _parent = [Case, 'cases', 'cases_id']
    _child = [Lawyer, 'lawyer', 'lawyer_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    @property
    def complete_lawyer(self):
        self.case
        self.lawyer

    def db_get(self, column):
        if column == 'cases_id':
            return self.cases.id
        elif column == 'lawyer_id':
            return self.lawyer.id

    def db_set(self, column, value):
        if column == 'cases_id':
            self.case = Case.by_id(value)
        elif column == 'lawyer_id':
            self.lawyer = Lawyer.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(CaseLawyerModel).filter(CaseLawyerModel.cases_id == self.cases.id,
                                                               CaseLawyerModel.lawyer_id == self.lawyer.id).first()
        except:
            return None


class CaseType(DAObject, SQLObjectRelationship):
    _model = CaseTypeModel
    _session = DBSession
    _parent = [Case, 'cases', 'cases_id']
    _child = [Type, 'type', 'type_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'cases_id':
            return self.cases.id
        elif column == 'type_id':
            return self.type.id

    def db_set(self, column, value):
        if column == 'cases_id':
            self.case = Case.by_id(value)
        elif column == 'type_id':
            self.type = Type.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(CaseTypeModel).filter(CaseTypeModel.cases_id == self.cases.id,
                                                             CaseTypeModel.type_id == self.type.id).first()
        except:
            return None


class CaseParty(DAObject, SQLObjectRelationship):
    _model = CasePartyModel
    _session = DBSession
    _parent = [Case, 'cases', 'cases_id']
    _child = [Party, 'party', 'party_id']

    @property
    def complete_party(self):
        self.case
        self.party

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'cases_id':
            return self.cases.id
        elif column == 'party_id':
            return self.party.id

    def db_set(self, column, value):
        if column == 'cases_id':
            self.case = Case.by_id(value)
        elif column == 'party_id':
            self.party = Party.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(CasePartyModel).filter(CasePartyModel.cases_id == self.cases.id,
                                                              CasePartyModel.party_id == self.party.id).first()
        except:
            return None


class PartyClaim(DAObject, SQLObject):
    _model = PartyClaimModel
    _session = DBSession
    _required = ['aggreived_party_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.sql_init()

    @property
    def complete_claim(self):
        self.claim
        self.party
        self.aggreived_party

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'claim_id':
            return self.claim.id
        elif column == 'party_id':
            return self.party.id
        elif column == 'aggreived_party_id':
            return self.aggreived_party.id
        elif column == 'third_party_claim':
            return self.third_party_claim
        elif column == 'cross_claim':
            return self.cross_claim
        else:
            raise Exception("Invalid column " + column)
    def db_set(self, column, value):
        if column == 'claim_id':
            self.claim = Claim.by_id(value)
        elif column == 'party_id':
            self.party = Party.by_id(value)
        elif column == 'aggreived_party_id':
            self.aggreived_party = Party.by_id(value)
        elif column == 'third_party_claim':
            self.third_party_claim = value
        elif column == 'cross_claim':
            self.third_party_claim = value
        else:
            raise Exception("Invalid column " + column)
    def db_null(self, column):
        if column == 'claim_id':
            del self.claim.name.text
        elif column == 'party_id':
            del self.party
        elif column == 'aggreived_party_id':
            del self.aggreived_party
        elif column == 'third_party_claim':
            del self.third_party_claim
        elif column == 'cross_claim':
            del self.cross_claim
        else:
            raise Exception("Invalid column " + column)

    def db_find_existing(self):
        try:
            return self._session.query(PartyClaimModel).filter(PartyClaimModel.claim_id == self.claim.id,
                                                               PartyClaimModel.party_id == self.party.id,
                                                               PartyClaimModel.aggreived_party_id == self.aggreived_party.id).first()
        except:
            return None


class PartyDesignation(DAObject, SQLObjectRelationship):
    _model = PartyDesignationModel
    _session = DBSession
    _parent = [Party, 'party', 'party_id']
    _child = [Designation, 'designation', 'designation_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'designation_id':
            return self.designation.id
        elif column == 'party_id':
            return self.party.id

    def db_set(self, column, value):
        if column == 'designation_id':
            self.designation = Designation.by_id(value)
        elif column == 'party_id':
            self.party = Party.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(PartyDesignationModel).filter(
                PartyDesignationModel.designation_id == self.designation.id,
                PartyDesignationModel.party_id == self.party.id).first()
        except:
            return None


class CaseLawfirm(DAObject, SQLObjectRelationship):
    _model = CaseLawfirmModel
    _session = DBSession
    _parent = [Case, 'cases', 'cases_id']
    _child = [Lawfirm, 'lawfirm', 'lawfirm_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    @property
    def complete_lawfirm(self):
        self.case
        self.lawfirm

    def db_get(self, column):
        if column == 'cases_id':
            return self.cases.id
        elif column == 'lawfirm_id':
            return self.lawfirm.id

    def db_set(self, column, value):
        if column == 'cases_id':
            self.case = Case.by_id(value)
        elif column == 'lawfirm_id':
            self.lawfirm = Lawfirm.by_id(value)

    def db_find_existing(self):
        try:

            return self._session.query(CaseLawfirmModel).filter(CaseLawfirmModel.cases_id == self.cases.id,
                                                                CaseLawfirmModel.lawfirm_id == self.lawfirm.id).first()
        except:
            return None


class LawfirmLawyer(DAObject, SQLObjectRelationship):
    _model = LawfirmLawyerModel
    _session = DBSession
    _parent = [Lawfirm, 'lawfirm', 'lawfirm_id']
    _child = [Lawyer, 'lawyer', 'lawyer_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'lawfirm_id':
            return self.lawfirm.id
        elif column == 'lawyer_id':
            return self.lawyer.id

    def db_set(self, column, value):
        if column == 'lawfirm_id':
            self.lawfirm = Lawfirm.by_id(value)
        elif column == 'lawyer_id':
            self.lawyer = Lawyer.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(LawfirmLawyerModel).filter(LawfirmLawyerModel.lawfirm_id == self.lawfirm.id,
                                                                  LawfirmLawyerModel.lawyer_id == self.lawyer.id).first()
        except:
            return None


class CasePartyLawfirm(DAObject, SQLObjectRelationship):
    _model = CasePartyLawfirmModel
    _session = DBSession
    _parent = [Case, 'cases', 'cases_id']
    _child = [PartyLawfirm, 'party_lawfirm', 'party_lawfirm_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'cases_id':
            return self.cases.id
        elif column == 'party_lawfirm_id':
            return self.party_lawfirm.id

    def db_set(self, column, value):
        if column == 'cases_id':
            self.case = Case.by_id(value)
        elif column == 'party_lawfirm_id':
            self.party_lawfirm = PartyLawfirm.by_id(value)

    def db_find_existing(self):
        try:
            return self._session.query(CasePartyLawfirmModel).filter(CasePartyLawfirmModel.cases_id == self.cases.id,
                                                                     CasePartyLawfirmModel.party_lawfirm_id == self.party_lawfirm.id).first()
        except:
            return None


class LawyerStateBar(DAObject, SQLObjectRelationship):
    _model = LawyerStateBarModel
    _session = DBSession
    _parent = [State, 'state', 'state_id']
    _child = [Lawyer, 'lawyer', 'lawyer_id']

    def init(self, *pargs, **kwargs):
        super().init(*pargs, **kwargs)
        self.rel_init(*pargs, **kwargs)

    def db_get(self, column):
        if column == 'state_id':
            return self.state.id
        elif column == 'lawyer_id':
            return self.lawyer.id
        elif column == 'bar_no':
            return self.bar_no
        else:
            raise Exception("Invalid column " + column)

    def db_set(self, column, value):
        if column == 'state_id':
            self.state = State.by_id(value)
        elif column == 'lawyer_id':
            self.lawyer = Lawyer.by_id(value)
        elif column == 'bar_no':
            self.bar_no = value
        else:
            raise Exception("Invalid column " + column)

    def db_find_existing(self):
        try:
            return self._session.query(LawyerStateBarModel).filter(LawyerStateBarModel.state_id == self.state.id,
                                                                   LawyerStateBarModel.lawyer_id == self.lawyer.id).first()
        except:
            return None

def db_find_all(party):
    possibilities = []
    possibilities.clear()
    if party.name.caption_type.name.text == "business entity" or party.name.caption_type.name.text == "governmental entity":
        for entry in party._session.query(PartyModel).filter(PartyModel.name == party.name.text).all():
            possibilities.append({"id": entry.id, "name": entry.name})
    else:
        for entry in party._session.query(PartyModel).filter(PartyModel.first_name == party.name.first,
                                                             PartyModel.last_name == party.name.last).all():
            possibilities.append({"id": entry.id, "name": entry.first_name + " " + entry.last_name})
    return possibilities